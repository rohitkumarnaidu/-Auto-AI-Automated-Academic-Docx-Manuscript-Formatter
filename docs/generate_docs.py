# -*- coding: utf-8 -*-
"""
ScholarForm AI -- Production Application Documentation Generator
Generates a company-grade .docx documentation like Stripe/Vercel/Notion docs.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()
section = doc.sections[0]
section.page_height  = Cm(29.7)
section.page_width   = Cm(21.0)
section.left_margin  = section.right_margin  = Cm(2.54)
section.top_margin   = section.bottom_margin = Cm(2.54)

BRAND  = RGBColor(0x1A, 0x56, 0xDB)
DARK   = RGBColor(0x11, 0x18, 0x27)
GRAY   = RGBColor(0x6B, 0x72, 0x80)
GREEN  = RGBColor(0x05, 0x7A, 0x55)
RED    = RGBColor(0xE0, 0x2A, 0x2A)
AMBER  = RGBColor(0xB4, 0x5D, 0x09)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LBLUE  = RGBColor(0xEB, 0xF5, 0xFF)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def pb(): doc.add_page_break()

def h1(t):
    p = doc.add_heading(t, 1); p.runs[0].font.color.rgb = BRAND; return p

def h2(t):
    p = doc.add_heading(t, 2); p.runs[0].font.color.rgb = DARK; return p

def h3(t):
    p = doc.add_heading(t, 3); p.runs[0].font.color.rgb = GRAY; return p

def body(t):  return doc.add_paragraph(t)

def note(t):
    p = doc.add_paragraph(); r = p.add_run(f"  NOTE  {t}")
    r.font.color.rgb = BRAND; r.italic = True; return p

def tip(t):
    p = doc.add_paragraph(); r = p.add_run(f"  TIP   {t}")
    r.font.color.rgb = GREEN; r.italic = True; return p

def warn(t):
    p = doc.add_paragraph(); r = p.add_run(f"  WARN  {t}")
    r.font.color.rgb = AMBER; r.bold = True; return p

def danger(t):
    p = doc.add_paragraph(); r = p.add_run(f"  IMPORTANT  {t}")
    r.font.color.rgb = RED; r.bold = True; return p

def bullet(t, lvl=0):
    return doc.add_paragraph(t, style="List Bullet")

def numbered(t):
    return doc.add_paragraph(t, style="List Number")

def code(t):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    return p

def tbl(headers, rows, fill="1A56DB"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = WHITE
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        c._tc.get_or_add_tcPr().append(shd)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
    doc.add_paragraph()
    return t

# ---------------------------------------------------------------------------
# COVER
# ---------------------------------------------------------------------------
def cover():
    for _ in range(4): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ScholarForm AI"); r.bold = True
    r.font.size = Pt(40); r.font.color.rgb = BRAND

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Automated Academic Manuscript Formatter")
    r.font.size = Pt(16); r.font.color.rgb = DARK

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Product Documentation")
    r.font.size = Pt(13); r.font.color.rgb = GRAY

    for _ in range(3): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Version 1.0")
    r.font.size = Pt(11); r.font.color.rgb = GRAY

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(datetime.date.today().strftime("%B %Y"))
    r.font.size = Pt(11); r.font.color.rgb = GRAY

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CONFIDENTIAL -- For Internal Use Only")
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RED
    pb()

# ---------------------------------------------------------------------------
# 1. INTRODUCTION
# ---------------------------------------------------------------------------
def introduction():
    h1("1. Introduction")
    h2("1.1 About ScholarForm AI")
    body(
        "ScholarForm AI is a cloud-based intelligent document formatting platform built "
        "for academic researchers, graduate students, and research institutions. It "
        "automatically transforms raw manuscripts into publication-ready documents that "
        "conform to the formatting guidelines of major academic journals and publishers."
    )
    body(
        "Unlike traditional word processors or manual formatting services, ScholarForm AI "
        "uses a multi-layer artificial intelligence pipeline -- combining large language "
        "models (LLMs), scientific NLP models, and retrieval-augmented generation (RAG) -- "
        "to understand the semantic structure of your document and apply the correct "
        "formatting rules intelligently."
    )

    h2("1.2 What ScholarForm AI Does")
    tbl(
        ["Capability", "Description"],
        [
            ["Multi-format input", "Accepts DOCX, PDF, LaTeX (.tex), HTML, Markdown, plain text, and legacy .doc files"],
            ["15 journal templates", "IEEE, APA, ACM, Springer, Elsevier, Nature, Harvard, Chicago, MLA, Vancouver, Numeric, Modern Blue, Modern Gold, Modern Red, and a passthrough (None) mode"],
            ["AI block classification", "Identifies each paragraph as title, abstract, heading, body, reference, figure caption, table caption, equation, methodology, or conclusion"],
            ["Citation formatting", "Validates DOIs via CrossRef; applies 10,000+ CSL citation styles from the UI"],
            ["Neural PDF parsing", "Uses Meta's Nougat deep learning model to extract text from complex academic PDFs"],
            ["GROBID metadata extraction", "Extracts authors, title, affiliations, and bibliography from PDF files using ML"],
            ["Export formats", "Downloads the formatted document as DOCX, PDF, or JATS XML (open-access journals)"],
            ["Real-time progress", "Streams processing progress in real time using Server-Sent Events (SSE)"],
            ["Document comparison", "Side-by-side HTML diff of original vs. formatted document"],
            ["Custom templates", "Create and edit your own journal template contracts from the browser"],
        ]
    )

    h2("1.3 How It Works (Overview)")
    body("When you upload a document, ScholarForm AI runs it through an 8-stage pipeline:")
    numbered("Ingest -- detect file type and route to the appropriate parser")
    numbered("Parse -- extract raw blocks (paragraphs, headings, tables, figures, equations, references)")
    numbered("Structure Detection -- infer heading levels, section boundaries, and document hierarchy")
    numbered("Semantic Classification -- use SciBERT and heuristics to label each block by type")
    numbered("RAG Retrieval -- pull the exact formatting rules for your chosen journal template from the vector knowledge base")
    numbered("LLM Reasoning -- NVIDIA Llama 3.3 70B generates a block-by-block formatting instruction set")
    numbered("Formatting -- apply all instructions to produce a formatted DOCX using python-docx")
    numbered("Export -- convert and deliver as DOCX, PDF, or JATS XML")
    note("If the primary LLM is unavailable, the system automatically falls back to DeepSeek R1 (local), then to deterministic rule-based formatting. You always get output.")
    pb()

# ---------------------------------------------------------------------------
# 2. GETTING STARTED
# ---------------------------------------------------------------------------
def getting_started():
    h1("2. Getting Started")

    h2("2.1 Prerequisites")
    h3("System Requirements")
    tbl(
        ["Requirement", "Minimum", "Recommended"],
        [
            ["OS", "Windows 10 / Ubuntu 20.04 / macOS 12", "Ubuntu 22.04 LTS"],
            ["RAM", "8 GB", "16 GB (32 GB for local Nougat)"],
            ["CPU", "4 cores", "8+ cores"],
            ["GPU", "Not required", "NVIDIA GPU (CUDA 11.8+) for Nougat/SciBERT speed"],
            ["Python", "3.10", "3.11"],
            ["Node.js", "18.x", "20.x"],
            ["Docker", "Required for GROBID", "Docker 24+"],
            ["Disk", "10 GB free", "50 GB (for ML model cache)"],
        ]
    )
    h3("Required Accounts & API Keys")
    tbl(
        ["Service", "Purpose", "Where to Get"],
        [
            ["Supabase", "Authentication + PostgreSQL database", "supabase.com -- free tier available"],
            ["NVIDIA NIM", "Primary LLM (Llama 3.3 70B) cloud inference", "build.nvidia.com -- API key needed"],
            ["Ollama (optional)", "Local LLM fallback (DeepSeek R1)", "ollama.com -- local install, no account needed"],
        ]
    )

    h2("2.2 Installation")
    h3("Step 1: Clone the Repository")
    code("git clone https://github.com/rohitkumarnaidu/-Auto-AI-Automated-Academic-Docx-Manuscript-Formatter.git")
    code("cd automated-manuscript-formatter")

    h3("Step 2: Set Up the Backend")
    code("cd backend")
    code("python -m venv .venv")
    code("# Windows:")
    code(".venv\\Scripts\\activate")
    code("# Linux/macOS:")
    code("source .venv/bin/activate")
    code("pip install -r requirements.txt")

    h3("Step 3: Configure Backend Environment")
    body("Create a file named .env inside the backend/ folder:")
    code("# backend/.env")
    code("SUPABASE_URL=https://your-project.supabase.co")
    code("SUPABASE_ANON_KEY=your-anon-key")
    code("SUPABASE_SERVICE_KEY=your-service-role-key")
    code("NVIDIA_API_KEY=your-nvidia-nim-key")
    code("REDIS_URL=redis://localhost:6379")
    code("OLLAMA_BASE_URL=http://localhost:11434")
    code("SECRET_KEY=your-random-secret-key-32chars")

    h3("Step 4: Run Database Migrations")
    code("alembic upgrade head")

    h3("Step 5: Seed the AI Knowledge Base (ChromaDB)")
    body("This step loads the formatting guidelines into the RAG vector store. Run once.")
    code("python scripts/ingest_guidelines.py")

    h3("Step 6: Start GROBID (Docker)")
    body("GROBID is required for PDF metadata extraction. Run it in a separate terminal.")
    code("docker run --rm -p 8070:8070 lfoppiano/grobid:0.7.0")

    h3("Step 7: Start the Backend Server")
    code("uvicorn app.main:app --reload --host 0.0.0.0 --port 8000")
    body("The API will be available at http://localhost:8000")
    body("Interactive API docs: http://localhost:8000/docs")

    h3("Step 8: Set Up the Frontend")
    code("cd ../frontend")
    code("npm install")
    code("# Create frontend/.env:")
    code("VITE_SUPABASE_URL=https://your-project.supabase.co")
    code("VITE_SUPABASE_ANON_KEY=your-anon-key")
    code("VITE_API_BASE_URL=http://localhost:8000")

    h3("Step 9: Start the Frontend")
    code("npm run dev")
    body("The application will be available at http://localhost:5173")

    h3("Step 10 (Optional): Start Ollama for Local LLM Fallback")
    code("ollama serve")
    code("ollama pull deepseek-r1")
    tip("Ollama is optional. Without it, the system uses NVIDIA NIM for LLM calls. If NVIDIA is also unavailable, rule-based formatting is used automatically.")

    h2("2.3 Quick Start -- First Document")
    numbered("Open http://localhost:5173 in your browser")
    numbered("Create an account or log in")
    numbered("Click 'Upload Document' on the Dashboard")
    numbered("Select your manuscript file (DOCX, PDF, TEX, HTML, MD, TXT, or DOC)")
    numbered("Choose a journal template from the dropdown (e.g. IEEE, APA)")
    numbered("Click 'Format Document'")
    numbered("Watch real-time progress on the Processing page")
    numbered("Download your formatted document as DOCX, PDF, or JATS XML")
    pb()

# ---------------------------------------------------------------------------
# 3. USER GUIDE
# ---------------------------------------------------------------------------
def user_guide():
    h1("3. User Guide")

    h2("3.1 Dashboard")
    body(
        "The Dashboard is your central document management hub. It displays all your "
        "uploaded and formatted documents with their current status, creation date, "
        "and template used."
    )
    h3("Document Statuses")
    tbl(
        ["Status", "Meaning"],
        [
            ["Pending", "Document uploaded, waiting to begin processing"],
            ["Processing", "AI pipeline is actively formatting the document"],
            ["Done", "Formatting complete -- ready to download"],
            ["Failed", "Processing encountered an error -- see validation results for details"],
        ]
    )
    h3("Dashboard Actions")
    bullet("Click a document row to open its detail page")
    bullet("Use the search bar to filter by filename or template")
    bullet("Delete a document using the trash icon (irreversible)")

    h2("3.2 Uploading a Document")
    h3("Single Upload")
    body("Navigate to Upload > choose your file > select a template > click 'Format Document'. Accepted formats: .docx, .pdf, .tex, .html, .md, .txt, .doc")
    note("Files larger than 10 MB are automatically uploaded in 5 MB chunks and reassembled on the server.")

    h3("Batch Upload")
    body("Use the Batch Upload panel to upload up to 10 documents at once with the same template settings. Each file is processed independently and appears as a separate entry in the Dashboard.")

    h3("Processing Options")
    tbl(
        ["Option", "Description"],
        [
            ["Template", "Select the target journal style (15 options available)"],
            ["Fast Mode", "Skips the Nougat neural parser for PDFs -- faster but less accurate for complex layouts"],
            ["CSL Style", "Override the template's default citation style with any of 10,000+ CSL styles"],
        ]
    )

    h2("3.3 Choosing a Template")
    body("ScholarForm AI includes 15 built-in journal templates:")
    tbl(
        ["Template", "Best For", "Citation Style"],
        [
            ["IEEE", "Electrical eng., computer science, electronics", "Numbered [1]"],
            ["APA (7th)", "Psychology, social sciences, education", "Author-year (Smith, 2023)"],
            ["ACM", "Computer science, information technology", "ACM numbered"],
            ["Springer", "Physics, chemistry, biology, engineering", "Author-year"],
            ["Elsevier", "Medicine, pharmacology, engineering", "Numbered"],
            ["Nature", "Multidisciplinary science, biology", "Numbered superscript"],
            ["Harvard", "General academic use", "Author-year"],
            ["Chicago (17th)", "Humanities, history, arts", "Notes-bibliography or author-date"],
            ["MLA (9th)", "Literature, language, humanities", "In-text (Author page)"],
            ["Vancouver", "Medicine, nursing, health sciences", "Numbered"],
            ["Numeric", "General numbered reference style", "Numbered [1]"],
            ["Modern Blue", "Custom professional blue theme", "IEEE-based"],
            ["Modern Gold", "Custom professional gold theme", "APA-based"],
            ["Modern Red", "Custom professional red theme", "APA-based"],
            ["None", "No formatting changes -- passthrough mode", "Preserved from original"],
        ]
    )
    tip("Use the CSL Importer in the Templates page to search and import any of 10,000+ Citation Style Language styles by journal name.")

    h2("3.4 Tracking Processing Progress")
    body(
        "After submitting a document, you are automatically redirected to the Processing page. "
        "This page streams real-time progress updates from the server using Server-Sent Events (SSE). "
        "You will see each pipeline stage complete in sequence: parsing, classification, RAG retrieval, "
        "LLM reasoning, and formatting."
    )
    note("Do not close this tab while processing. If you navigate away, the processing continues in the background and you can find the result in the Dashboard.")

    h2("3.5 Comparing Documents")
    body(
        "The Compare page shows a side-by-side HTML diff of your original document versus "
        "the formatted output. Additions are highlighted in green, deletions in red. "
        "Use this to verify that formatting was applied correctly before downloading."
    )

    h2("3.6 Downloading Your Document")
    body("From the Download page, choose your preferred export format:")
    tbl(
        ["Format", "Use Case", "Notes"],
        [
            ["DOCX", "Submit to journal, continue editing in Word", "Fully formatted Word document"],
            ["PDF", "Share, print, or submit to portals accepting PDF", "Generated from DOCX via conversion"],
            ["JATS XML", "Open-access journals, PubMed Central submissions", "Journal Article Tag Suite standard XML"],
        ]
    )

    h2("3.7 Document History & Versions")
    body("The History page shows all versions of each document. You can re-download any previous version of a formatted document.")

    h2("3.8 Template Editor")
    body(
        "Advanced users can create and edit custom template contracts from the Template Editor page. "
        "A template contract is a YAML file that defines all formatting rules: "
        "font family, font sizes, line spacing, margins, heading styles, section ordering, "
        "citation style, page numbering, and cover page settings."
    )
    note("The Template Editor auto-saves your changes 2 seconds after you stop typing (debounced autosave).")

    h2("3.9 Validation Results")
    body(
        "After formatting, ScholarForm AI validates the output against the template rules "
        "and academic writing standards. The Validation Results page shows a summary of "
        "passed checks, warnings, and errors -- each with an AI-generated explanation and "
        "a suggested fix."
    )

    h2("3.10 Account & Profile")
    tbl(
        ["Feature", "Location"],
        [
            ["Update profile name and email", "Profile page"],
            ["Change password", "Profile > Security"],
            ["Dark / Light mode toggle", "Navbar (top right)"],
            ["Notification preferences", "Notifications page"],
            ["Sign out", "Navbar > Avatar > Sign Out"],
        ]
    )
    pb()

# ---------------------------------------------------------------------------
# 4. ADMIN GUIDE
# ---------------------------------------------------------------------------
def admin_guide():
    h1("4. Administrator Guide")

    h2("4.1 Accessing the Admin Dashboard")
    body("The Admin Dashboard is accessible only to accounts with the admin role. Navigate to /admin from the main navigation or visit http://your-domain/admin directly.")

    h2("4.2 User Management")
    body("From the Admin Dashboard you can:")
    bullet("View all registered users with their registration date and last active date")
    bullet("Search users by email address")
    bullet("Deactivate or delete user accounts")
    bullet("View per-user document counts and processing history")

    h2("4.3 System Metrics")
    body("The Metrics section shows real-time and historical performance data:")
    tbl(
        ["Metric", "Description"],
        [
            ["Total documents processed", "Cumulative count since deployment"],
            ["Processing success rate", "Percentage of jobs that completed without error"],
            ["Average processing time", "Mean time from upload to formatted output"],
            ["LLM tier distribution", "Which AI tier (NVIDIA / DeepSeek / Rules) served each request"],
            ["Token usage", "LLM tokens consumed per day (for cost tracking)"],
            ["Model latency", "P50, P95, P99 response times per LLM model"],
            ["Prometheus endpoint", "Raw metrics at /metrics for Grafana/Prometheus integration"],
        ]
    )

    h2("4.4 System Health")
    body("The Health Status Indicator (visible in the navigation bar) shows the status of all backend services:")
    tbl(
        ["Service", "Port", "Health Check"],
        [
            ["FastAPI Backend", "8000", "GET /health"],
            ["GROBID", "8070", "GET /api/isalive"],
            ["Ollama", "11434", "GET /api/tags"],
            ["Redis", "6379", "PING command"],
            ["Supabase", "Cloud", "Connection pool check"],
        ]
    )

    h2("4.5 Rate Limiting")
    body("The API enforces rate limiting to protect the system from abuse:")
    tbl(
        ["Limit", "Value", "Scope"],
        [
            ["Request rate", "100 requests per minute", "Per IP address"],
            ["Upload size", "50 MB per file", "Per upload request"],
            ["Batch size", "10 files per batch", "Per batch upload request"],
            ["Feedback length", "500 characters", "Per feedback submission"],
        ]
    )
    note("Rate limit headers (X-RateLimit-Limit, X-RateLimit-Remaining, X-RateLimit-Reset) are included in every API response.")
    pb()

# ---------------------------------------------------------------------------
# 5. API REFERENCE
# ---------------------------------------------------------------------------
def api_reference():
    h1("5. API Reference")
    body("ScholarForm AI exposes a RESTful JSON API. Base URL: http://your-domain:8000")
    body("Interactive documentation: http://your-domain:8000/docs (Swagger UI)")

    h2("5.1 Authentication")
    body("All API endpoints (except /health and /templates list) require a valid JWT Bearer token.")
    body("Include the token in the Authorization header:")
    code("Authorization: Bearer eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9...")
    body("Obtain a token by logging in via the frontend or directly through Supabase Auth.")

    h3("Authentication Error Codes")
    tbl(
        ["HTTP Status", "Meaning"],
        [
            ["401 (login endpoint)", "Invalid email or password"],
            ["401 (other endpoints)", "Session expired -- re-authenticate"],
            ["403", "Insufficient permissions (admin-only route)"],
        ]
    )

    h2("5.2 Document Endpoints")

    h3("POST /documents/upload")
    body("Upload a single document for formatting.")
    body("Request (multipart/form-data):")
    code('curl -X POST http://localhost:8000/documents/upload \\')
    code('  -H "Authorization: Bearer TOKEN" \\')
    code('  -F "file=@my_paper.pdf" \\')
    code('  -F "template=ieee" \\')
    code('  -F "fast_mode=false"')
    body("Response (201 Created):")
    code('{')
    code('  "id": "3fa85f64-5717-4562-b3fc-2c963f66afa6",')
    code('  "filename": "my_paper.pdf",')
    code('  "template": "ieee",')
    code('  "status": "pending",')
    code('  "created_at": "2025-02-28T13:45:00Z"')
    code('}')

    h3("POST /documents/upload/chunked/init")
    body("Initialize a chunked upload session for files larger than 10 MB.")
    code('curl -X POST http://localhost:8000/documents/upload/chunked/init \\')
    code('  -H "Authorization: Bearer TOKEN" \\')
    code('  -H "Content-Type: application/json" \\')
    code('  -d \'{"filename": "large_paper.pdf", "total_size": 52428800}\'')
    body("Response:")
    code('{ "upload_id": "abc123-def456", "chunk_size": 5242880 }')

    h3("GET /documents")
    body("List all documents for the authenticated user.")
    code('curl http://localhost:8000/documents -H "Authorization: Bearer TOKEN"')
    body("Response (200 OK):")
    code('[')
    code('  {')
    code('    "id": "3fa85f64-...",')
    code('    "filename": "my_paper.pdf",')
    code('    "template": "ieee",')
    code('    "status": "done",')
    code('    "created_at": "2025-02-28T13:45:00Z"')
    code('  }')
    code(']')

    h3("GET /documents/{doc_id}")
    body("Get the full details and current status of a document.")
    code('curl http://localhost:8000/documents/3fa85f64-5717-4562-b3fc-2c963f66afa6 \\')
    code('  -H "Authorization: Bearer TOKEN"')

    h3("DELETE /documents/{doc_id}")
    body("Permanently delete a document and its associated files.")
    code('curl -X DELETE http://localhost:8000/documents/3fa85f64-... \\')
    code('  -H "Authorization: Bearer TOKEN"')
    body("Response: 204 No Content")

    h3("GET /documents/{doc_id}/download")
    body("Download the formatted DOCX file.")
    code('curl -OJ http://localhost:8000/documents/3fa85f64-.../download \\')
    code('  -H "Authorization: Bearer TOKEN"')

    h3("GET /documents/{doc_id}/download/pdf")
    body("Download the formatted PDF file.")

    h3("GET /documents/{doc_id}/download/jats")
    body("Download the JATS XML file (for open-access journal submissions).")

    h3("GET /documents/{doc_id}/compare")
    body("Returns the HTML diff comparing original vs formatted document.")
    body("Response: { \"diff_html\": \"<div class='diff'>...</div>\" }")

    h3("GET /documents/{doc_id}/validation")
    body("Returns validation results with pass/fail checks and AI explanations.")
    code('{')
    code('  "score": 87,')
    code('  "checks": [')
    code('    { "name": "Font consistency", "status": "pass" },')
    code('    { "name": "Citation format", "status": "fail",')
    code('      "message": "3 citations do not follow IEEE style",')
    code('      "suggestion": "Use [Author, Year] format for references 5, 12, 17" }')
    code('  ]')
    code('}')

    h2("5.3 Template Endpoints")
    h3("GET /templates")
    body("List all available built-in templates.")
    code('curl http://localhost:8000/templates')
    body("Response:")
    code('{ "templates": [{ "id": "ieee", "name": "IEEE", "category": "Engineering" }, ...] }')
    warn("Frontend expects a flat array but backend returns { templates: [...] }. This mismatch is a known bug being fixed.")

    h3("GET /templates/csl/search?q={query}")
    body("Search for CSL citation styles by journal name.")
    code('curl "http://localhost:8000/templates/csl/search?q=nature" \\')
    code('  -H "Authorization: Bearer TOKEN"')
    body("Response: { \"query\": \"nature\", \"results\": [{\"id\": \"nature\", \"title\": \"Nature\"}] }")

    h2("5.4 Stream Endpoint")
    h3("GET /stream/{job_id}")
    body("Subscribe to real-time processing progress via Server-Sent Events (SSE).")
    code('curl -N http://localhost:8000/stream/JOB_ID -H "Authorization: Bearer TOKEN"')
    body("Event stream example:")
    code('data: {"stage": "parsing", "progress": 10, "message": "Extracting document structure"}')
    code('data: {"stage": "classifying", "progress": 30, "message": "AI classifying blocks"}')
    code('data: {"stage": "formatting", "progress": 80, "message": "Applying IEEE template"}')
    code('data: {"stage": "done", "progress": 100, "message": "Formatting complete"}')

    h2("5.5 Other Endpoints")
    tbl(
        ["Method", "Endpoint", "Auth", "Description"],
        [
            ["GET", "/health", "None", "Health check -- returns { status: ok }"],
            ["GET", "/metrics", "Admin JWT", "Prometheus metrics (text/plain format)"],
            ["GET", "/metrics/models", "JWT", "LLM model performance stats JSON"],
            ["POST", "/auth/login", "None", "Login with email + password"],
            ["POST", "/auth/logout", "JWT", "Invalidate current session"],
            ["POST", "/feedback", "JWT", "Submit feedback (body: {content, rating})"],
            ["POST", "/documents/batch-upload", "JWT", "Upload up to 10 files at once"],
        ]
    )

    h2("5.6 Error Response Format")
    body("All API errors return a consistent JSON structure:")
    code('{')
    code('  "error": "DOCUMENT_NOT_FOUND",')
    code('  "message": "No document found with the given ID",')
    code('  "status": 404')
    code('}')
    tbl(
        ["Error Code", "HTTP Status", "Meaning"],
        [
            ["UNAUTHORIZED", "401", "Missing or invalid JWT token"],
            ["FORBIDDEN", "403", "Valid token but insufficient permissions"],
            ["NOT_FOUND", "404", "Requested resource does not exist"],
            ["VALIDATION_ERROR", "422", "Request body failed schema validation"],
            ["RATE_LIMITED", "429", "Too many requests -- slow down"],
            ["PROCESSING_ERROR", "500", "Internal pipeline error"],
            ["SERVICE_UNAVAILABLE", "503", "External dependency (LLM/GROBID) is down"],
        ]
    )
    pb()

# ---------------------------------------------------------------------------
# 6. ARCHITECTURE
# ---------------------------------------------------------------------------
def architecture():
    h1("6. System Architecture")

    h2("6.1 High-Level Overview")
    body(
        "ScholarForm AI is a three-tier web application: a React single-page application (SPA) "
        "frontend, a FastAPI REST API backend, and a multi-service AI/ML layer. "
        "Supabase provides PostgreSQL as the primary database and handles authentication. "
        "Redis handles caching. GROBID runs as a Docker sidecar."
    )
    body("Architecture layers (top to bottom):")
    tbl(
        ["Layer", "Technology", "Responsibility"],
        [
            ["Browser (Client)", "React 18 + Vite + Tailwind CSS", "User interface, routing, real-time progress display"],
            ["API Gateway", "FastAPI (Python 3.10+)", "REST endpoints, JWT auth, rate limiting, SSE streaming"],
            ["Business Logic", "Pipeline modules (Python)", "Document parsing, AI classification, formatting, export"],
            ["AI Layer", "NVIDIA NIM / Ollama / HuggingFace", "LLM reasoning, embedding, semantic parsing"],
            ["Vector Store", "ChromaDB", "RAG knowledge base of formatting guidelines"],
            ["Database", "Supabase (PostgreSQL)", "User data, document records, version history"],
            ["Cache", "Redis", "LLM response caching, session data"],
            ["Sidecar", "GROBID (Docker)", "PDF bibliographic metadata extraction"],
        ]
    )

    h2("6.2 AI Pipeline Architecture")
    h3("LLM Tier System (3 tiers)")
    body("The system never fails due to LLM unavailability. It cascades through three tiers:")
    tbl(
        ["Tier", "Model", "Access Method", "When Used"],
        [
            ["Tier 1 (Primary)", "NVIDIA Llama 3.3 70B Instruct", "NVIDIA NIM cloud API via LiteLLM", "Default -- always tried first"],
            ["Tier 2 (Fallback)", "DeepSeek R1 (8B parameters)", "Ollama local server via LangChain", "When NVIDIA API fails or circuit is open"],
            ["Tier 3 (Last Resort)", "Rule-based heuristic engine", "Pure Python pattern matching", "When all LLMs are unavailable"],
        ]
    )
    h3("RAG Retrieval Tier System (4 tiers)")
    body("Formatting guidelines are retrieved using embedding similarity search:")
    tbl(
        ["Tier", "Embedding Model", "Dimensions", "Vector Store"],
        [
            ["Tier 1", "BAAI/bge-m3", "1024d, 8192 token window", "ChromaDB persistent"],
            ["Tier 2", "BAAI/bge-small-en-v1.5", "384d", "ChromaDB persistent"],
            ["Tier 3", "Any SentenceTransformer", "Variable", "Native JSON cosine similarity"],
            ["Tier 4", "Deterministic hash vector", "256d", "Native JSON (no ML needed)"],
        ]
    )
    h3("Safety Layer")
    body("Every LLM call is protected by a four-component safety system:")
    tbl(
        ["Component", "Behavior"],
        [
            ["Circuit Breaker", "After 3 consecutive LLM failures: opens circuit for 60 seconds. Requests bypass LLM to next tier during this window. Automatically attempts recovery."],
            ["Retry Guard", "Retries each LLM call up to 3 times with exponential backoff (0.5s, 1s, 2s delays)"],
            ["LLM Validator", "Validates every LLM JSON response against a Pydantic schema before use. Malformed responses trigger a retry."],
            ["Safe Function", "Every model method is wrapped in a @safe_function decorator that catches all exceptions and returns a pre-defined fallback value instead of crashing."],
        ]
    )

    h2("6.3 Data Flow: Document Upload to Download")
    numbered("User uploads file via browser")
    numbered("FastAPI router receives multipart/form-data, saves file with UUID name, creates DB record")
    numbered("Background task starts the 8-stage AI pipeline")
    numbered("SSE events are pushed to client at each stage completion")
    numbered("Formatted DOCX is saved to storage, DB record updated to 'done'")
    numbered("User downloads result via signed download URL")

    h2("6.4 Frontend Architecture")
    body("The frontend is a React 18 single-page application (SPA) built with Vite.")
    tbl(
        ["Module", "Technology", "Purpose"],
        [
            ["Routing", "React Router v6", "Client-side navigation between 27 pages"],
            ["Server State", "TanStack React Query v5", "API calls, background refetch, optimistic updates, caching"],
            ["Auth State", "React Context + Supabase JS", "Session management, JWT refresh, OAuth callbacks"],
            ["Styling", "Tailwind CSS v3.4", "Utility-first responsive styling, dark mode"],
            ["Real-time", "Native EventSource (SSE)", "Processing progress streaming"],
            ["Comparison", "diff npm package", "HTML side-by-side document diff"],
            ["API Client", "Axios + api.js", "Centralized API calls with auth interceptors"],
        ]
    )

    h2("6.5 Security Architecture")
    tbl(
        ["Control", "Implementation"],
        [
            ["Authentication", "Supabase JWT -- RS256 signed tokens, 1-hour expiry with auto-refresh"],
            ["Authorization", "Row-Level Security (RLS) in PostgreSQL -- users access only their own data"],
            ["Transport Security", "HTTPS enforced in production -- HSTS header with 1-year max-age"],
            ["Content Security Policy", "CSP headers restrict script and resource origins"],
            ["Clickjacking protection", "X-Frame-Options: DENY header on all responses"],
            ["Rate limiting", "100 requests/minute per IP -- 429 status with Retry-After header"],
            ["File validation", "MIME type and extension whitelist; size limit enforced before storage"],
            ["Secrets management", "All keys in environment variables -- never hardcoded"],
        ]
    )
    pb()

# ---------------------------------------------------------------------------
# 7. CONFIGURATION REFERENCE
# ---------------------------------------------------------------------------
def configuration():
    h1("7. Configuration Reference")

    h2("7.1 Backend Environment Variables")
    tbl(
        ["Variable", "Required", "Default", "Description"],
        [
            ["SUPABASE_URL", "Yes", "--", "Full URL of your Supabase project (https://xxx.supabase.co)"],
            ["SUPABASE_ANON_KEY", "Yes", "--", "Supabase anonymous (public) API key"],
            ["SUPABASE_SERVICE_KEY", "Yes", "--", "Supabase service role key (never expose to frontend)"],
            ["NVIDIA_API_KEY", "Yes*", "--", "NVIDIA NIM API key for Llama 3.3 70B. *Required unless Ollama is configured."],
            ["REDIS_URL", "No", "redis://localhost:6379", "Redis connection string"],
            ["OLLAMA_BASE_URL", "No", "http://localhost:11434", "Ollama server URL for DeepSeek fallback"],
            ["OLLAMA_MODEL", "No", "deepseek-r1", "Ollama model name to use as LLM Tier 2"],
            ["SECRET_KEY", "Yes", "--", "Random secret for JWT signing (minimum 32 chars)"],
            ["GROBID_URL", "No", "http://localhost:8070", "GROBID service URL"],
            ["CHROMA_PERSIST_DIR", "No", "./chroma_db", "ChromaDB persistence directory"],
            ["MAX_UPLOAD_SIZE_MB", "No", "50", "Maximum file upload size in megabytes"],
            ["RATE_LIMIT_PER_MINUTE", "No", "100", "Max API requests per IP per minute"],
            ["LOG_LEVEL", "No", "INFO", "Logging level: DEBUG, INFO, WARNING, ERROR"],
            ["ENVIRONMENT", "No", "development", "Environment name: development, staging, production"],
        ]
    )

    h2("7.2 Frontend Environment Variables")
    tbl(
        ["Variable", "Required", "Description"],
        [
            ["VITE_SUPABASE_URL", "Yes", "Supabase project URL (same as backend)"],
            ["VITE_SUPABASE_ANON_KEY", "Yes", "Supabase anonymous key (public, safe to expose)"],
            ["VITE_API_BASE_URL", "Yes", "Backend API base URL (e.g. http://localhost:8000)"],
        ]
    )

    h2("7.3 Template Contract Configuration")
    body("Each template contract (contract.yaml) supports the following configuration keys:")
    tbl(
        ["Key", "Type", "Example", "Description"],
        [
            ["name", "string", "IEEE", "Template display name"],
            ["font_family", "string", "Times New Roman", "Body text font"],
            ["font_size_body", "number", "10", "Body text size in points"],
            ["font_size_h1", "number", "14", "Heading 1 size in points"],
            ["line_spacing", "number", "1.0", "Line spacing multiplier"],
            ["columns", "number", "2", "Number of columns (1 or 2)"],
            ["margin_top_cm", "number", "2.54", "Top margin in centimeters"],
            ["margin_bottom_cm", "number", "2.54", "Bottom margin in centimeters"],
            ["margin_left_cm", "number", "2.0", "Left margin in centimeters"],
            ["margin_right_cm", "number", "2.0", "Right margin in centimeters"],
            ["citation_style", "string", "ieee", "CSL style identifier"],
            ["section_order", "list", "[title, abstract, intro, method, ...]", "Ordered list of section types"],
            ["page_numbers", "boolean", "true", "Include page numbers"],
            ["cover_page", "boolean", "false", "Generate cover page"],
            ["table_of_contents", "boolean", "false", "Generate table of contents"],
            ["layout", "string", "two-column", "Page layout descriptor"],
        ]
    )
    pb()

# ---------------------------------------------------------------------------
# 8. DEPLOYMENT GUIDE
# ---------------------------------------------------------------------------
def deployment():
    h1("8. Deployment Guide")

    h2("8.1 Production Checklist")
    tbl(
        ["Item", "Action"],
        [
            ["Environment variables", "Set all required vars in your server's environment manager (not .env files in production)"],
            ["HTTPS", "Configure a reverse proxy (Nginx/Caddy) with SSL certificate (Let's Encrypt)"],
            ["SECRET_KEY", "Generate a cryptographically secure random key: python -c \"import secrets; print(secrets.token_hex(32))\""],
            ["Debug mode", "Set ENVIRONMENT=production -- disables /docs Swagger UI and debug error details"],
            ["GROBID", "Run GROBID as a managed Docker container with restart policy (--restart always)"],
            ["Redis", "Use Redis with password authentication in production"],
            ["Database", "Supabase handles this -- ensure your project plan supports your expected load"],
            ["Logs", "Configure log aggregation (CloudWatch, Datadog, or Loki)"],
        ]
    )

    h2("8.2 Nginx Configuration (Reference)")
    code("server {")
    code("    listen 443 ssl;")
    code("    server_name your-domain.com;")
    code("")
    code("    location /api/ {")
    code("        proxy_pass http://127.0.0.1:8000/;")
    code("        proxy_set_header Authorization $http_authorization;")
    code("    }")
    code("")
    code("    location / {")
    code("        root /var/www/scholarform/dist;")
    code("        try_files $uri $uri/ /index.html;")
    code("    }")
    code("}")

    h2("8.3 Building the Frontend for Production")
    code("cd frontend")
    code("npm run build")
    body("The build output is in frontend/dist/ -- serve this directory with your web server.")

    h2("8.4 Running the Backend with Gunicorn (Production)")
    code("pip install gunicorn")
    code("gunicorn app.main:app -w 4 -k uvicorn.workers.UvicornWorker --bind 0.0.0.0:8000")
    note("Use 4 workers for a 4-core server. Adjust -w based on available CPU cores.")

    h2("8.5 Managing GROBID with Docker Compose")
    code("# docker-compose.yml")
    code("version: '3.8'")
    code("services:")
    code("  grobid:")
    code("    image: lfoppiano/grobid:0.7.0")
    code("    ports:")
    code('      - "8070:8070"')
    code("    restart: always")
    code("    mem_limit: 4g")
    pb()

# ---------------------------------------------------------------------------
# 9. TROUBLESHOOTING
# ---------------------------------------------------------------------------
def troubleshooting():
    h1("9. Troubleshooting")

    h2("9.1 Common Issues and Solutions")
    tbl(
        ["Problem", "Likely Cause", "Solution"],
        [
            ["Template library not loading", "API contract mismatch (known bug)", "Ensure frontend uses response.templates not response directly"],
            ["Processing stuck at 'Pending'", "GROBID not running", "Start GROBID: docker run --rm -p 8070:8070 lfoppiano/grobid:0.7.0"],
            ["LLM timeout errors", "NVIDIA API key expired or quota exceeded", "Check your NVIDIA NIM dashboard; ensure OLLAMA_BASE_URL is set as fallback"],
            ["ChromaDB errors on startup", "Guidelines not seeded", "Run: python scripts/ingest_guidelines.py"],
            ["401 on every request", "JWT expired or wrong SUPABASE_URL/KEY", "Verify env vars and re-login"],
            ["403 on upload", "File type not in whitelist", "Supported types: .docx .pdf .tex .html .md .txt .doc"],
            ["'Alembic migration failed'", "DB not reachable or schema mismatch", "Check SUPABASE_URL and run: alembic upgrade head"],
            ["Nougat model download very slow", "'Fast Mode' can skip Nougat for PDFs", "Toggle Fast Mode on in upload options, or pre-download with: transformers-cli download facebook/nougat-base"],
            ["Redis connection refused", "Redis not started", "Start Redis: redis-server (or use docker run -p 6379:6379 redis)"],
            ["CORS errors in browser", "Frontend and backend on different origins", "Verify VITE_API_BASE_URL matches backend URL exactly"],
        ]
    )

    h2("9.2 Log Locations")
    tbl(
        ["Component", "Log Location", "How to View"],
        [
            ["FastAPI backend", "stdout / LOG_LEVEL", "uvicorn console or systemd journal"],
            ["GROBID", "Docker container logs", "docker logs grobid-container-name"],
            ["Ollama", "~/.ollama/logs/", "cat ~/.ollama/logs/server.log"],
            ["Frontend (dev)", "Browser console", "F12 > Console tab"],
        ]
    )

    h2("9.3 Checking AI Model Status")
    code("# Check NVIDIA NIM connectivity:")
    code("curl -H \"Authorization: Bearer $NVIDIA_API_KEY\" https://integrate.api.nvidia.com/v1/models")
    code("")
    code("# Check Ollama:")
    code("curl http://localhost:11434/api/tags")
    code("")
    code("# Check GROBID:")
    code("curl http://localhost:8070/api/isalive")
    code("")
    code("# Check Redis:")
    code("redis-cli ping  # should return PONG")
    pb()

# ---------------------------------------------------------------------------
# 10. DEVELOPER GUIDE
# ---------------------------------------------------------------------------
def developer_guide():
    h1("10. Developer Guide")

    h2("10.1 Project Structure")
    tbl(
        ["Directory", "Purpose"],
        [
            ["backend/app/routers/", "FastAPI route handlers (documents, templates, metrics, stream, auth, feedback)"],
            ["backend/app/pipeline/intelligence/", "AI core: RAG engine, LLM reasoning engine, SciBERT semantic parser"],
            ["backend/app/pipeline/parsing/", "File parsers: DOCX, PDF, Nougat, LaTeX, HTML, Markdown, TXT, converter"],
            ["backend/app/pipeline/formatting/", "DOCX output generation: formatter, template renderer, reference formatter"],
            ["backend/app/pipeline/safety/", "AI safety layer: circuit breaker, retry guard, LLM validator"],
            ["backend/app/pipeline/agents/", "Agent system: main orchestrator (document_agent.py) + future agents"],
            ["backend/app/pipeline/contracts/", "Journal template contracts (15 YAML files)"],
            ["backend/app/templates/", "Base DOCX template files (one per journal style)"],
            ["backend/app/services/", "External service clients: Supabase, NVIDIA NIM, CrossRef, model metrics, A/B testing"],
            ["backend/app/middleware/", "HTTP middleware: rate limiting, security headers, Prometheus metrics"],
            ["backend/tests/", "Test suite: 39 unit tests + integration/safety/stress subdirectories"],
            ["frontend/src/pages/", "27 application pages/screens"],
            ["frontend/src/components/", "21 reusable UI components"],
            ["frontend/src/services/api.js", "Centralized API client with auth interceptors"],
        ]
    )

    h2("10.2 How to Add a New Journal Template")
    numbered("Create a new folder: backend/app/pipeline/contracts/{template_name}/")
    numbered("Add a contract.yaml file that defines all formatting rules (see Section 7.3 for keys)")
    numbered("Create a base DOCX template: backend/app/templates/{template_name}/template.docx")
    numbered("Add Jinja2 markers to the DOCX ({{ content }}, {% for block in blocks %} etc.)")
    numbered("Add the template name to the frontend template gallery in frontend/src/pages/Templates.jsx")
    numbered("Add it to the upload template dropdown in frontend/src/components/upload/TemplateSelector.jsx")
    numbered("Test with: python -m pytest tests/test_templates.py -v")
    tip("Run python scripts/check_template_markers.py after creating the template DOCX to verify Jinja markers are correctly embedded.")

    h2("10.3 How to Add a New File Parser")
    numbered("Create a new parser file in backend/app/pipeline/parsing/ (e.g. docbook_parser.py)")
    numbered("Inherit from BaseParser (base_parser.py) and implement the parse() method")
    numbered("The parse() method must return List[Block]")
    numbered("Register the new parser in parser_factory.py with the appropriate file extension mapping")
    numbered("Add tests in backend/tests/test_pipeline.py")

    h2("10.4 How to Extend the AI Pipeline")
    body("The main orchestration logic lives in document_agent.py. To add a new processing step:")
    numbered("Add your new function/class in a new file under pipeline/agents/ or pipeline/intelligence/")
    numbered("Import and call it from document_agent.py in the appropriate pipeline stage")
    numbered("Wrap any external calls with @safe_function from safety/safe_execution.py")
    numbered("Add circuit breaker protection if the step makes external API calls")
    numbered("Write tests covering success path, failure path, and fallback behavior")

    h2("10.5 Running Tests")
    code("# All backend tests:")
    code("cd backend && python -m pytest tests/ -v --tb=short")
    code("")
    code("# Specific test file:")
    code("python -m pytest tests/test_rag_engine.py -v")
    code("")
    code("# With coverage report:")
    code("python -m pytest tests/ --cov=app --cov-report=html")
    code("# Open htmlcov/index.html in browser")
    code("")
    code("# Frontend tests:")
    code("cd frontend && npm run test")
    code("")
    code("# Frontend tests with UI:")
    code("npm run test -- --ui")

    h2("10.6 Code Style Guidelines")
    bullet("Python: Follow PEP 8; use type hints on all function signatures")
    bullet("JavaScript/JSX: ESLint with react-hooks rules; functional components only")
    bullet("Commits: Conventional Commits format (feat:, fix:, docs:, refactor:, test:)")
    bullet("New backend routes: always require JWT auth (use get_current_user dependency injection)")
    bullet("New agent code: always wrap with @safe_function and provide fallback_value")
    pb()

# ---------------------------------------------------------------------------
# 11. CHANGELOG
# ---------------------------------------------------------------------------
def changelog():
    h1("11. Changelog")

    h2("Version 1.0.0 -- February 2025 (Current)")
    h3("New Features")
    bullet("Initial production release of ScholarForm AI")
    bullet("15 journal templates: IEEE, APA, ACM, Springer, Elsevier, Nature, Harvard, Chicago, MLA, Vancouver, Numeric, None, Modern Blue, Modern Gold, Modern Red")
    bullet("Multi-format input: DOCX, PDF, LaTeX, HTML, Markdown, plain text, and legacy .doc (LibreOffice conversion)")
    bullet("3-tier LLM pipeline: NVIDIA Llama 3.3 70B -- DeepSeek R1 (Ollama) -- Rule-based heuristics")
    bullet("4-tier RAG pipeline: BGE-M3 + ChromaDB -- BGE-Small -- SentenceTransformers -- Deterministic hash")
    bullet("27-page React frontend with dark mode, real-time SSE progress, and HTML diff comparison")
    bullet("AI safety layer: circuit breaker, retry guard, Pydantic LLM validator, @safe_function decorator")
    bullet("Export to DOCX, PDF, and JATS XML")
    bullet("CSL citation style importer (10,000+ styles)")
    bullet("CrossRef DOI validation for citations")
    bullet("GROBID integration for PDF metadata extraction")
    bullet("Meta Nougat integration for neural PDF parsing")
    bullet("Batch upload (up to 10 files)")
    bullet("Chunked upload for files > 10 MB")
    bullet("Custom template editor with debounced autosave")
    bullet("Admin dashboard with model performance metrics and Prometheus integration")
    bullet("A/B testing framework for LLM model comparison")
    bullet("Full test suite: 39 backend unit tests + 19 frontend tests")

    h3("Known Issues in v1.0.0")
    bullet("Template library API: 4 contract mismatches between frontend and backend (fix in progress)")
    bullet("Frontend template gallery: shows 8 of 15 available templates (fix in progress)")
    bullet("OCR for scanned PDFs: Tesseract stub implemented, binary integration pending")
    bullet("Modern Blue template: filenames corrupted in git (rename fix required)")

    h2("Upcoming: Version 1.1.0 (Planned)")
    h3("Fixes")
    bullet("Fix all 4 API contract mismatches for Template Library")
    bullet("Show all 15 templates in frontend gallery")
    bullet("Fix Modern Blue corrupted template filenames")
    bullet("Wire Tesseract OCR into the parsing pipeline")

    h3("New Features")
    bullet("Document Generator: create academic documents from scratch using templates")
    bullet("Resume and Portfolio template contracts")
    bullet("Version History UI for browsing and restoring previous formatted versions")
    bullet("Dashboard pagination with virtual scrolling for 100+ documents")
    bullet("WCAG 2.1 AA accessibility improvements")
    bullet("KeyBERT and YAKE! keyword extraction integration")
    bullet("Groq API as LLM Tier 2 alternative (faster and cheaper)")
    pb()

# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------
def main():
    cover()
    introduction()
    getting_started()
    user_guide()
    admin_guide()
    api_reference()
    architecture()
    configuration()
    deployment()
    troubleshooting()
    developer_guide()
    changelog()

    out = "ScholarForm_AI_Product_Documentation.docx"
    doc.save(out)
    print(f"SUCCESS: {out} generated.")

if __name__ == "__main__":
    main()
