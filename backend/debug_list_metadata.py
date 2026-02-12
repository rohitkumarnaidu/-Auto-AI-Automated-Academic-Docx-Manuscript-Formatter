
import os
import sys
from docx import Document
from pathlib import Path

# Add backend to path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../')))

from app.pipeline.parsing.parser import DocxParser

def create_list_docx(path):
    doc = Document()
    doc.add_paragraph("Top Level Paragraph")
    
    # Add lists
    p1 = doc.add_paragraph("List 1 Item 1", style='List Bullet')
    p2 = doc.add_paragraph("List 1 Item 2", style='List Bullet')
    
    # Add nested list item (manually manipulating XML if necessary, but List Bullet 2 usually works)
    try:
        p3 = doc.add_paragraph("List 1 Nested Item", style='List Bullet 2')
    except:
        # Fallback to manual ilvl if style missing
        p3 = doc.add_paragraph("List 1 Nested Item (Manual)")
        # Simple ilvl injection is complex, but python-docx sometimes supports it
        pass

    doc.save(path)
    print(f"Created test document at {path}")

def verify_list_extraction(path):
    parser = DocxParser()
    doc_obj = parser.parse(str(path), "test_list")
    
    list_items = [b for b in doc_obj.blocks if b.metadata.get("is_list_item")]
    print(f"Parsed {len(list_items)} list items")
    
    if not list_items:
         print("❌ FAILED: No list items detected")
         sys.exit(1)

    all_passed = True
    for i, li in enumerate(list_items):
        lvl = li.metadata.get("list_level")
        lid = li.metadata.get("list_id")
        print(f"Item {i}: '{li.text[:20]}...' -> Level: {lvl}, ID: {lid}")
        if lvl is None or lid is None:
            print(f"⚠️ Warning: Missing level/ID on item {i}")
            # If style is 'List Bullet', Word might not add numPr unless it's a "real" list
            # But usually it does.
    
    print("✅ Extraction attempt completed")

if __name__ == "__main__":
    test_path = Path("test_list_gen.docx")
    try:
        create_list_docx(test_path)
        verify_list_extraction(test_path)
    finally:
        if test_path.exists():
            os.remove(test_path)
