"""
Template Renderer - Jinja2/docxtpl rendering for manuscript output.
"""

from __future__ import annotations

import logging
import re
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional
from zipfile import ZipFile

logger = logging.getLogger(__name__)

try:
    from docxtpl import DocxTemplate
    _DOCXTPL_AVAILABLE = True
except ImportError:
    DocxTemplate = None  # type: ignore[assignment,misc]
    _DOCXTPL_AVAILABLE = False

from docx import Document as WordDocument

from app.models import Block, PipelineDocument as Document


class TemplateRenderer:
    """
    Renders pipeline documents into DOCX using Jinja2 tags in docxtpl templates.
    """

    def __init__(self, templates_dir: str = "app/templates"):
        self.templates_dir = Path(templates_dir)
        self._template_marker_cache: Dict[Path, bool] = {}

    @staticmethod
    def _coerce_bool(value: Any, default: bool) -> bool:
        if value is None:
            return default
        if isinstance(value, bool):
            return value
        if isinstance(value, (int, float)):
            return value != 0
        if isinstance(value, str):
            token = value.strip().lower()
            if token in {"1", "true", "yes", "on"}:
                return True
            if token in {"0", "false", "no", "off", ""}:
                return False
        return bool(value)

    def _resolve_bool_option(
        self,
        options: Dict[str, Any],
        keys: List[str],
        default: bool,
    ) -> bool:
        for key in keys:
            if key in options:
                return self._coerce_bool(options.get(key), default)
        return default

    def render(self, document: Document, template_name: str = "ieee") -> "DocxTemplate":
        """Render a template using document context."""
        if not _DOCXTPL_AVAILABLE:
            raise ImportError(
                "docxtpl is not installed. Run: pip install 'docxtpl>=1.0.0'"
            )
        if not document:
            raise ValueError("document must not be None")
        template_name = (template_name or "ieee").strip() or "ieee"
        try:
            template_path = self._resolve_template_path(template_name)
            context = self.build_context(document)
            tpl = DocxTemplate(str(template_path))
            tpl.render(context)
            return tpl
        except Exception as exc:
            logger.error("Failed to render template '%s': %s", template_name, exc)
            raise

    def has_renderable_template(self, template_name: str) -> bool:
        """Return True when the template directory contains a docxtpl-capable source."""
        style = (template_name or "ieee").lower()
        template_dir = self.templates_dir / style
        if (template_dir / "template.jinja2").is_file():
            return True

        candidate = template_dir / "template.docx"
        return candidate.is_file() and self._has_template_markers(candidate)

    def build_context(self, document: Document) -> Dict[str, Any]:
        """Build Jinja2 context from a pipeline document."""
        try:
            metadata = document.metadata
            abstract_text = (metadata.abstract or "").strip()
            if not abstract_text:
                abstract_text = self._first_block_text(document.blocks, "abstract_body")

            keywords = list(metadata.keywords or [])
            if not keywords:
                raw_keywords = self._first_block_text(document.blocks, "keywords_body")
                if raw_keywords:
                    keywords = [item.strip() for item in raw_keywords.split(",") if item.strip()]

            authors = list(metadata.authors or [])
            if not authors:
                authors = self._all_block_text(document.blocks, "author")

            affiliations = list(metadata.affiliations or [])
            if not affiliations:
                affiliations = self._all_block_text(document.blocks, "affiliation")

            title = (metadata.title or "").strip()
            if not title:
                title = self._first_block_text(document.blocks, "title")
            if not title:
                title = document.original_filename or "Untitled Manuscript"

            references = self._collect_references(document)
            sections = self._collect_sections(document.blocks)

            formatting_options = document.formatting_options or {}

            # Templates currently gate title/authors behind `{% if cover_page %}`.
            # Default to rendering front matter so metadata is not dropped when
            # formatting_options does not explicitly provide a flag.
            cover_page = self._resolve_bool_option(
                formatting_options,
                keys=["cover_page", "add_cover_page"],
                default=True,
            )
            toc = self._resolve_bool_option(
                formatting_options,
                keys=["toc", "generate_toc"],
                default=False,
            )
            page_numbers = self._resolve_bool_option(
                formatting_options,
                keys=["page_numbers", "add_page_numbers"],
                default=True,
            )

            return {
                "title": title,
                "authors": authors,
                "affiliations": affiliations,
                "date": datetime.now(timezone.utc).strftime("%B %d, %Y"),
                "abstract": abstract_text,
                "keywords": keywords,
                "sections": sections,
                "references": references,
                "cover_page": cover_page,
                "toc": toc,
                "page_numbers": page_numbers,
                "page_number": formatting_options.get("page_number", "1"),
            }
        except Exception as exc:
            logger.error("Failed to build template context: %s", exc)
            raise

    def _resolve_template_path(self, template_name: str) -> Path:
        style = (template_name or "ieee").lower()
        template_dir = self.templates_dir / style
        jinja_source = template_dir / "template.jinja2"
        if jinja_source.is_file():
            return self._build_template_from_jinja_source(jinja_source)

        candidate = template_dir / "template.docx"
        if candidate.is_file():
            if self._has_template_markers(candidate):
                return candidate
            logger.warning(
                "Template '%s' has no Jinja markers. Using generated fallback template.",
                candidate,
            )
        return self._build_fallback_template()

    def _build_template_from_jinja_source(self, source_path: Path) -> Path:
        """Wrap a plain-text Jinja2 template source in a minimal DOCX container."""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                temp_path = Path(tmp.name)

            doc = WordDocument()
            for line in source_path.read_text(encoding="utf-8").splitlines():
                doc.add_paragraph(line)
            doc.save(str(temp_path))
            return temp_path
        except Exception as exc:
            logger.warning(
                "Failed to build DOCX template from Jinja source '%s': %s. Falling back.",
                source_path,
                exc,
            )
            return self._build_fallback_template()

    def _has_template_markers(self, template_path: Path) -> bool:
        cached = self._template_marker_cache.get(template_path)
        if cached is not None:
            return cached

        has_markers = False
        try:
            with ZipFile(template_path) as zf:
                xml_entries = [
                    name for name in zf.namelist()
                    if name.startswith("word/") and name.endswith(".xml")
                ]

                for xml_name in xml_entries:
                    xml = zf.read(xml_name).decode("utf-8", errors="ignore")
                    # Some templates split markers across <w:t> runs; strip tags to
                    # inspect the underlying text stream as well.
                    xml_text = re.sub(r"<[^>]+>", "", xml)
                    if ("{{" in xml) or ("{%" in xml) or ("{{" in xml_text) or ("{%" in xml_text):
                        has_markers = True
                        break
        except Exception as exc:
            logger.warning(
                "Unable to inspect template '%s' for Jinja markers: %s",
                template_path,
                exc,
            )
            has_markers = False

        self._template_marker_cache[template_path] = has_markers
        return has_markers

    def _build_fallback_template(self) -> Path:
        """Create a temporary fallback DOCX template if none exists."""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                temp_path = Path(tmp.name)
            doc = WordDocument()
            doc.add_paragraph("{{ title }}", style="Title")
            doc.add_paragraph("{% for author in authors %}{{ author }}{% if not loop.last %}, {% endif %}{% endfor %}")
            doc.add_paragraph("{% for affiliation in affiliations %}{{ affiliation }}{% if not loop.last %}; {% endif %}{% endfor %}")
            doc.add_paragraph("{% if abstract %}Abstract{% endif %}")
            doc.add_paragraph("{{ abstract }}")
            doc.add_paragraph("Keywords: {% for keyword in keywords %}{{ keyword }}{% if not loop.last %}, {% endif %}{% endfor %}")
            doc.add_paragraph("{% for section in sections %}")
            doc.add_paragraph("{{ section.heading }}")
            doc.add_paragraph("{% for paragraph in section.paragraphs %}{{ paragraph }}{% endfor %}")
            doc.add_paragraph("{% endfor %}")
            doc.add_paragraph("{% if references %}References{% endif %}")
            doc.add_paragraph("{% for reference in references %}{{ reference }}{% endfor %}")
            doc.save(str(temp_path))
            logger.warning("Using fallback DOCX template at '%s'", temp_path)
            return temp_path
        except Exception as exc:
            logger.error("Failed to build fallback template: %s", exc)
            raise

    def _collect_references(self, document: Document) -> List[str]:
        if document.references:
            ordered = sorted(document.references, key=lambda ref: ref.index)
            refs = [
                (ref.formatted_text or ref.raw_text or "").strip()
                for ref in ordered
                if (ref.formatted_text or ref.raw_text or "").strip()
            ]
            if refs:
                return refs

        ref_blocks = [
            block.text.strip()
            for block in sorted(document.blocks, key=lambda b: b.index)
            if str(block.block_type).lower() == "reference_entry" and block.text.strip()
        ]
        return ref_blocks

    def _collect_sections(self, blocks: List[Block]) -> List[Dict[str, Any]]:
        """Group non-reference content into heading sections."""
        sections: List[Dict[str, Any]] = []
        current_heading = "Body"
        current_paragraphs: List[str] = []

        skip_types = {
            "title",
            "author",
            "affiliation",
            "abstract_heading",
            "abstract_body",
            "keywords_heading",
            "keywords_body",
            "references_heading",
            "reference_entry",
            "figure_caption",
            "table_caption",
        }

        for block in sorted(blocks, key=lambda b: b.index):
            block_type = str(block.block_type).lower()
            text = (block.text or "").strip()
            if not text or block_type in skip_types:
                continue

            if block_type.startswith("heading_"):
                if current_paragraphs:
                    sections.append({"heading": current_heading, "paragraphs": current_paragraphs})
                    current_paragraphs = []
                current_heading = text
                continue

            current_paragraphs.append(text)

        if current_paragraphs:
            sections.append({"heading": current_heading, "paragraphs": current_paragraphs})

        return sections

    def _first_block_text(self, blocks: List[Block], block_type: str) -> str:
        for block in sorted(blocks, key=lambda b: b.index):
            if str(block.block_type).lower() == block_type:
                value = (block.text or "").strip()
                if value:
                    return value
        return ""

    def _all_block_text(self, blocks: List[Block], block_type: str) -> List[str]:
        values: List[str] = []
        for block in sorted(blocks, key=lambda b: b.index):
            if str(block.block_type).lower() == block_type:
                value = (block.text or "").strip()
                if value:
                    values.append(value)
        return values
