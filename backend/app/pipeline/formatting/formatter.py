"""
Formatter Module - Applies structure and styles to create a Word document.

All rendering helpers are wrapped in @safe_function so that a failure in
one block/figure/equation does not abort the entire document generation.
"""

import logging
import os
import re
import yaml
from copy import deepcopy
from xml.etree import ElementTree as ET
from zipfile import ZIP_DEFLATED, ZipFile
from typing import Optional, Any
from docx import Document as WordDocument
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt
from io import BytesIO
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

from app.models import PipelineDocument as Document, BlockType, Figure
from app.pipeline.contracts.loader import ContractLoader
from app.pipeline.formatting.style_mapper import StyleMapper
from app.pipeline.formatting.numbering import NumberingEngine
from app.pipeline.formatting.reference_formatter import ReferenceFormatter
from app.pipeline.formatting.template_renderer import TemplateRenderer
from app.pipeline.tables.renderer import TableRenderer
from app.pipeline.safety.safe_execution import safe_function, safe_execution

class Formatter:
    """
    Formats the validated Document into a python-docx object based on a template.
    """
    
    def __init__(self, templates_dir: str = "app/templates", contracts_dir: str = "app/pipeline/contracts"):
        self.templates_dir = templates_dir
        self.contract_loader = ContractLoader(contracts_dir=contracts_dir)
        self.style_mapper = StyleMapper(self.contract_loader)
        self.numbering_engine = NumberingEngine(self.contract_loader)
        self.reference_formatter = ReferenceFormatter(self.contract_loader)
        self.template_renderer = TemplateRenderer(templates_dir=templates_dir)
        self.table_renderer = TableRenderer()
        
    def process(self, document: Document) -> Document:
        """Standard pipeline stage entry point."""
        template_name = (
            document.template.template_name
            if document.template and hasattr(document.template, "template_name")
            else "none"
        )
        # We store the resulting Word object in a transient field
        document.generated_doc = self.format(document, template_name)
        return document

    @safe_function(fallback_value=None, error_message="Formatter.format failed")
    def format(self, document: Document, template_name: str = "IEEE") -> Optional[Any]:
        """
        Apply formatting using contract-driven modular components.
        """
        if not template_name:
            template_name = "none" # No default template - use neutral formatting
        options = document.formatting_options or {}
        add_cover_page = self._resolve_bool_option(
            options,
            "cover_page",
            aliases=("add_cover_page",),
            default=False,
        )
        add_toc = self._resolve_bool_option(
            options,
            "toc",
            aliases=("generate_toc",),
            default=False,
        )
        add_page_numbers = self._resolve_bool_option(
            options,
            "page_numbers",
            aliases=("add_page_numbers",),
            default=True,
        )
        add_borders = self._resolve_bool_option(
            options,
            "borders",
            aliases=("add_borders",),
            default=False,
        )
        add_line_numbers = self._resolve_bool_option(
            options,
            "line_numbers",
            aliases=("add_line_numbers",),
            default=False,
        )
        template_key = template_name.lower()
        renderer_mode = str(options.get("template_engine", "auto")).strip().lower()
        footnote_lookup = self._build_footnote_lookup(document)
        use_template_renderer = (
            renderer_mode != "legacy"
            and template_key != "none"
            and self.template_renderer.has_renderable_template(template_name)
        )
            
        # 1. Apply rules to model before rendering
        document = self.numbering_engine.apply_numbering(document, template_name)
        self._prepare_references(document, template_name)

        # 2. Primary path: docxtpl/Jinja2 template rendering.
        # Keep legacy python-docx path as a fallback for robustness.
        if use_template_renderer:
            try:
                logger.info("Using docxtpl renderer for template '%s'", template_name)
                rendered = self.template_renderer.render(document, template_name)
                self._post_process_template_render(
                    rendered,
                    document,
                    template_name,
                    options,
                    footnote_lookup=footnote_lookup,
                )
                self._install_post_save_hook(rendered, footnote_lookup)
                return rendered
            except Exception as exc:
                logger.warning(
                    "Falling back to python-docx renderer for template '%s'. Error: %s",
                    template_name, exc
                )
        elif template_key != "none":
            logger.info(
                "Falling back to python-docx renderer for template '%s' because no docxtpl-capable template was found.",
                template_name,
            )
        
        # 2. Load Resources
        # Note: template.docx is still the base for styles
        is_none = template_key == "none"
        template_path = os.path.join(self.templates_dir, template_name.lower(), "template.docx")
        contract_path = os.path.join(self.templates_dir, template_name.lower(), "contract.yaml")
        
        if is_none:
            logger.info("No template specified (General Formatting). Using blank document.")
            word_doc = WordDocument()
        elif not os.path.exists(template_path):
            logger.warning("Template file not found at %s. Using blank.", template_path)
            word_doc = WordDocument()
        else:
            word_doc = WordDocument(template_path)
            
        contract = self._load_contract(contract_path)
        style_map = contract.get("styles", {})
        
        # 2. Add Content
        items_to_insert = []
        
        # Add Blocks
        for block in document.blocks:
            # Skip parser-extracted structural artifacts from main body flow.
            # Footnotes are rendered in a dedicated section later.
            if (
                block.metadata.get("is_header")
                or block.metadata.get("is_footer")
                or block.metadata.get("is_footnote")
                or block.metadata.get("is_endnote")
                or block.block_type == BlockType.FOOTNOTE
            ):
                continue

            # SKIP figure captions - check both Enum and String to be safe
            b_type = str(block.block_type).upper()
            if "FIGURE_CAPTION" in b_type or "TABLE_CAPTION" in b_type:
                continue
            
            # Special handling for references: if it's a reference entry, we might want to re-format it
            if block.block_type == BlockType.REFERENCE_ENTRY:
                # Find matching reference object
                ref = next((r for r in document.references if r.block_id == block.block_id), None)
                if ref:
                    block.text = self.reference_formatter.format_reference(ref, template_name)
            
            items_to_insert.append({
                "type": "block",
                "index": block.index,
                "obj": block
            })
            
        # Add Figures (using index)
        # We need sequential numbering for captions: Figure 1, Figure 2, ...
        # The document.figures list is ordered by extraction.
        for i, fig in enumerate(document.figures):
            # i+1 is the sequential number
            
            # fig.metadata["block_index"] is where it was found in parser.
            b_idx = fig.metadata.get("block_index", -1)
            
            # We insert figure AFTER the block it was attached to
            # Use small offset to ensure it comes after the block
            # HARDENING FIX: Deterministic sub-offset for multiple figures same paragraph
            # If multiple figures share same block_index, add sub-offset based on position
            # Example: 3 figures at index 100 -> 100.1, 100.101, 100.102
            # Maintains base offset (+0.1) and avoids equation collision (+0.2)
            sub_offset = i * 0.001  # Position-based sub-offset
            items_to_insert.append({
                "type": "figure",
                "index": b_idx + 0.1 + sub_offset, 
                "obj": fig,
                "number": i + 1
            })

            
        # Add Equations
        for i, eqn in enumerate(document.equations):
            # FORENSIC FIX: Sort by block_index (position in text) not equation index (creation order)
            # Use small offset (+0.2) to place immediately after parent block
            sort_index = eqn.metadata.get("block_index", eqn.index)
            items_to_insert.append({
                "type": "equation",
                "index": sort_index + 0.2,
                "obj": eqn
            })
            
        # Add Tables
        for i, table in enumerate(document.tables):
             items_to_insert.append({
                "type": "table",
                "index": table.block_index, # Fixed: Tables have global block_index
                "obj": table,
                "number": i + 1  # Sequential number for caption (Table 1, Table 2, ...)
            })
            
        # Sort by index
        items_to_insert.sort(key=lambda x: x["index"])
        
        # 3. Render
        self._apply_initial_layout(word_doc, template_name)
        self._apply_page_size(word_doc, self._resolve_page_size(template_name, options))
        
        # Cover page and TOC must be inserted before main content.
        if add_cover_page:
            self._add_cover_page(word_doc, document)
            
        if add_toc:
            self._add_table_of_contents(word_doc)

        current_columns = None
        
        for item in items_to_insert:
            if item["type"] == "block":
                block = item["obj"]
                
                # Layout logic: Check if we need a column change
                target_cols = self._get_target_columns(block, template_name)
                if current_columns is not None and target_cols != current_columns:
                    # Switch layout: Add new section
                    new_section = word_doc.add_section()
                    self._set_columns(new_section, target_cols)
                
                if current_columns is None:
                    # Initialize first section
                    self._set_columns(word_doc.sections[0], target_cols)
                    
                current_columns = target_cols
                self._render_block(
                    word_doc,
                    block,
                    template_name,
                    footnote_lookup=footnote_lookup,
                )
                
            elif item["type"] == "figure":
                self._render_figure(word_doc, item["obj"], item["number"])
            elif item["type"] == "equation":
                self._render_equation(word_doc, item["obj"])
            elif item["type"] == "table":
                self.table_renderer.render(word_doc, item["obj"], item.get("number"))
                
        # Apply section-level options after all content/sections are created.
        if add_page_numbers:
            self._remove_static_page_number_placeholders(word_doc)
            self._add_page_numbers(word_doc)
        if add_borders:
            self._add_page_borders(word_doc)
        if add_line_numbers:
            self._add_line_numbers(word_doc)
        self._apply_global_line_spacing(word_doc, template_name, options)
        self._install_post_save_hook(word_doc, footnote_lookup)
                
        return word_doc

    def _prepare_references(self, document: Document, template_name: str) -> None:
        """Populate missing reference.formatted_text values before rendering."""
        if not document.references:
            return

        for ref in document.references:
            if ref.formatted_text and ref.formatted_text.strip():
                continue
            try:
                ref.formatted_text = self.reference_formatter.format_reference(ref, template_name)
            except Exception:
                ref.formatted_text = ref.raw_text or ''

    @safe_function(fallback_value=None, error_message="Equation rendering failed")
    def _render_equation(self, doc, equation):
        """Render an equation block."""
        p = doc.add_paragraph()
        p.style = "Normal"
        p.paragraph_format.alignment = 1  # Center

        equation_text = (equation.text or "").strip() or " "

        omath_para = OxmlElement("m:oMathPara")
        omath = OxmlElement("m:oMath")
        run = OxmlElement("m:r")
        text_node = OxmlElement("m:t")
        text_node.text = equation_text
        run.append(text_node)
        omath.append(run)
        omath_para.append(omath)
        p._p.append(omath_para)

        if equation.number:
            p.add_run(f"  ({equation.number})")

    def _apply_initial_layout(self, doc, publisher: str):
        """Set margins and initial properties."""
        contract = self.contract_loader.load(publisher)
        layout = contract.get("layout", {})
        if not layout:
            return

        margins = layout.get("margins", {})
        for section in doc.sections:
            section.top_margin = Inches(margins.get("top", 1.0))
            section.bottom_margin = Inches(margins.get("bottom", 1.0))
            section.left_margin = Inches(margins.get("left", 1.0))
            section.right_margin = Inches(margins.get("right", 1.0))

    def _get_target_columns(self, block, publisher: str) -> int:
        """Determine required column count for a block."""
        contract = self.contract_loader.load(publisher)
        layout = contract.get("layout", {})
        default = layout.get("default_columns", 1)
        
        s_name = (block.section_name or "").lower()
        overrides = layout.get("section_overrides", {})
        
        # Check canonical matching or direct matching
        for key, val in overrides.items():
            if key in s_name:
                return val
        
        return default

    def _apply_page_size(self, doc, size_name: str):
        """Sets the page size for all sections."""
        from docx.shared import Inches, Mm
        
        size_map = {
            "Letter": (Inches(8.5), Inches(11)),
            "A4": (Mm(210), Mm(297)),
            "Legal": (Inches(8.5), Inches(14))
        }
        
        width, height = size_map.get(size_name, size_map["Letter"])
        
        for section in doc.sections:
            section.page_width = width
            section.page_height = height

    def _resolve_page_size(self, template_name: str, options: dict) -> str:
        """Resolve page size from options first, then contract, then default."""
        requested = str(options.get("page_size", "")).strip()
        if requested:
            return requested

        contract = self.contract_loader.load(template_name)
        layout = contract.get("layout", {})
        contract_page_size = str(layout.get("page_size", "")).strip()
        return contract_page_size or "Letter"

    def _resolve_line_spacing(self, template_name: str, options: dict) -> Optional[float]:
        """Resolve global line spacing from options/contract if present."""
        raw_value = options.get("line_spacing", None)
        if raw_value is None:
            raw_value = options.get("add_line_spacing", None)
        if raw_value is None:
            contract = self.contract_loader.load(template_name)
            raw_value = (contract.get("layout", {}) or {}).get("line_spacing")

        if raw_value in (None, "", False):
            return None
        try:
            value = float(raw_value)
            if value <= 0:
                return None
            return value
        except (TypeError, ValueError):
            logger.warning("Invalid line spacing value '%s'. Ignoring.", raw_value)
            return None

    @staticmethod
    def _coerce_bool_option(value: Any, default: bool) -> bool:
        if value is None:
            return default
        if isinstance(value, bool):
            return value
        if isinstance(value, (int, float)):
            return value != 0
        if isinstance(value, str):
            token = value.strip().lower()
            if token in {"1", "true", "yes", "on"}:
                return True
            if token in {"0", "false", "no", "off", ""}:
                return False
        return bool(value)

    def _resolve_bool_option(
        self,
        options: dict,
        primary_key: str,
        aliases=(),
        default: bool = False,
    ) -> bool:
        if not isinstance(options, dict):
            return default
        for key in (primary_key, *aliases):
            if key in options:
                return self._coerce_bool_option(options.get(key), default)
        return default

    def _add_cover_page(self, doc, document_obj):
        """Adds a cover page with title and metadata."""
        # Insert a new paragraph at the very beginning
        # Note: We can't easily prepend in python-docx, so we rely on this being called 
        # BEFORE content addition if we want it first, OR we add a section break.
        # However, typically cover pages are separate sections at the start.
        # Since we call this early in `format()`, we can just add to the empty doc.
        
        p = doc.add_paragraph()
        p.alignment = 1 # Center
        
        # Title
        title = document_obj.metadata.title or document_obj.original_filename or "Untitled Document"
        run = p.add_run(title + "\n\n")
        run.bold = True
        run.font.size = Pt(24)
        
        # Authors
        authors = ", ".join(document_obj.metadata.authors) if document_obj.metadata.authors else "Unknown Author"
        run = p.add_run(authors + "\n")
        run.font.size = Pt(14)
        
        # Date
        from datetime import datetime
        run = p.add_run(datetime.now().strftime("%B %d, %Y"))
        run.font.size = Pt(12)
        
        doc.add_page_break()

    def _add_table_of_contents(self, doc, prepend: bool = False, add_page_break: bool = True):
        """Adds a TOC field code."""
        inserted = []

        p = doc.add_paragraph()
        run = p.add_run("Table of Contents")
        run.bold = True
        run.font.size = Pt(16)
        inserted.append(p)
        
        # XML for TOC field
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar)
        
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = 'TOC \\o "1-3" \\h \\z \\u'
        run._r.append(instr)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run._r.append(fldChar2)
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar3)
        inserted.append(paragraph)

        if add_page_break:
            page_break_para = doc.add_paragraph()
            page_break_para.add_run().add_break(WD_BREAK.PAGE)
            inserted.append(page_break_para)

        if prepend:
            body = doc._body._element
            for para in reversed(inserted):
                body.remove(para._p)
                body.insert(0, para._p)

    def _add_page_numbers(self, doc):
        """Adds simple page numbers to the footer."""
        for section in doc.sections:
            footer = section.footer
            if footer.paragraphs:
                p = footer.paragraphs[0]
            else:
                p = footer.add_paragraph()
            p.alignment = 1  # Center
            if self._paragraph_has_field_code(p, "PAGE"):
                continue

            run = p.add_run()
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar)

            instr = OxmlElement('w:instrText')
            instr.set(qn('xml:space'), 'preserve')
            instr.text = "PAGE"
            run._r.append(instr)

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar2)

    def _add_page_borders(self, doc):
        """Adds page borders via OXML."""
        for section in doc.sections:
            sec_pr = section._sectPr
            existing = sec_pr.xpath('./w:pgBorders')
            for node in existing:
                sec_pr.remove(node)

            pg_borders = OxmlElement('w:pgBorders')
            pg_borders.set(qn('w:offsetFrom'), 'page')

            for border_name in ('top', 'left', 'bottom', 'right'):
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '24')
                border.set(qn('w:color'), 'auto')
                pg_borders.append(border)

            sec_pr.append(pg_borders)

    def _add_line_numbers(self, doc, count_by: int = 1):
        """Enable line numbering for all sections."""
        for section in doc.sections:
            sec_pr = section._sectPr
            existing = sec_pr.xpath('./w:lnNumType')
            ln_num = existing[0] if existing else OxmlElement('w:lnNumType')
            ln_num.set(qn('w:countBy'), str(max(1, int(count_by))))
            ln_num.set(qn('w:start'), '1')
            ln_num.set(qn('w:distance'), '360')
            if not existing:
                sec_pr.append(ln_num)

    def _apply_global_line_spacing(self, doc, template_name: str, options: dict) -> None:
        """Apply global line spacing from options or template contract."""
        line_spacing = self._resolve_line_spacing(template_name, options)
        if line_spacing is None:
            return
        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.line_spacing = line_spacing

    def _paragraph_has_field_code(self, paragraph, field_name: str) -> bool:
        """Check whether paragraph XML already contains a field code token."""
        xml = paragraph._p.xml if paragraph is not None else ""
        return field_name in xml

    def _remove_paragraph(self, paragraph) -> None:
        """Remove a paragraph node from the document body."""
        if paragraph is None:
            return
        p = paragraph._p
        parent = p.getparent()
        if parent is not None:
            parent.remove(p)

    def _prepend_paragraph(self, doc, text: str = "", style: Optional[str] = None, alignment: Optional[int] = None):
        """Create a paragraph and move it to the beginning of the document."""
        try:
            paragraph = doc.add_paragraph(style=style) if style else doc.add_paragraph()
        except Exception:
            paragraph = doc.add_paragraph()

        if text:
            paragraph.add_run(text)
        if alignment is not None:
            paragraph.alignment = alignment

        body = doc._body._element
        body.remove(paragraph._p)
        body.insert(0, paragraph._p)
        return paragraph

    def _document_contains_text(self, doc, text: str) -> bool:
        needle = (text or "").strip().lower()
        if not needle:
            return False
        for paragraph in doc.paragraphs:
            if needle in (paragraph.text or "").strip().lower():
                return True
        return False

    def _prepend_front_matter(self, doc, document_obj: Document, as_cover_page: bool) -> None:
        """Insert title/authors at top when template omitted front-matter markers."""
        title = document_obj.metadata.title or document_obj.original_filename or "Untitled Document"
        authors = ", ".join(document_obj.metadata.authors) if document_obj.metadata.authors else "Unknown Author"
        affiliations = "; ".join(document_obj.metadata.affiliations) if document_obj.metadata.affiliations else ""

        if as_cover_page:
            from datetime import datetime
            page_break = self._prepend_paragraph(doc)
            page_break.add_run().add_break(WD_BREAK.PAGE)

            date_para = self._prepend_paragraph(doc, datetime.now().strftime("%B %d, %Y"), alignment=1)
            if date_para.runs:
                date_para.runs[0].font.size = Pt(12)

            if affiliations:
                aff_para = self._prepend_paragraph(doc, affiliations, alignment=1)
                if aff_para.runs:
                    aff_para.runs[0].italic = True
                    aff_para.runs[0].font.size = Pt(12)

            author_para = self._prepend_paragraph(doc, authors, alignment=1)
            if author_para.runs:
                author_para.runs[0].font.size = Pt(14)

            title_para = self._prepend_paragraph(doc, title, style="Title", alignment=1)
            if title_para.runs:
                title_para.runs[0].bold = True
                title_para.runs[0].font.size = Pt(24)
        else:
            if affiliations:
                aff_para = self._prepend_paragraph(doc, affiliations, alignment=1)
                if aff_para.runs:
                    aff_para.runs[0].italic = True
                    aff_para.runs[0].font.size = Pt(11)

            author_para = self._prepend_paragraph(doc, authors, alignment=1)
            if author_para.runs:
                author_para.runs[0].font.size = Pt(12)

            title_para = self._prepend_paragraph(doc, title, style="Title", alignment=1)
            if title_para.runs:
                title_para.runs[0].bold = True

    def _remove_static_page_number_placeholders(self, doc) -> None:
        """Remove template-injected static page labels before adding PAGE fields."""
        for paragraph in list(doc.paragraphs):
            text = (paragraph.text or "").strip()
            if re.fullmatch(r"Page\s+\d+", text, flags=re.IGNORECASE):
                self._remove_paragraph(paragraph)

    def _remove_static_toc_block(self, doc) -> None:
        """Remove text-only TOC placeholder blocks generated by template loops."""
        paragraphs = list(doc.paragraphs)
        for idx, paragraph in enumerate(paragraphs):
            if (paragraph.text or "").strip().lower() != "table of contents":
                continue
            self._remove_paragraph(paragraph)
            scan = idx + 1
            while scan < len(paragraphs):
                candidate = paragraphs[scan]
                candidate_text = (candidate.text or "").strip()
                if re.fullmatch(r"\d+\.\s+.+", candidate_text):
                    self._remove_paragraph(candidate)
                    scan += 1
                    continue
                if not candidate_text:
                    self._remove_paragraph(candidate)
                    scan += 1
                    continue
                break
            break

    def _ensure_dynamic_toc(self, doc) -> None:
        """Ensure a dynamic Word TOC field is present."""
        if 'TOC \\o "1-3" \\h \\z \\u' in doc._body._element.xml:
            return
        self._add_table_of_contents(doc, prepend=True, add_page_break=True)

    def _post_process_template_render(
        self,
        rendered,
        source_document: Document,
        template_name: str,
        options: dict,
        footnote_lookup: Optional[dict] = None,
    ) -> None:
        """Apply backend layout options to docxtpl-rendered templates as well."""
        word_doc = getattr(rendered, "docx", None)
        if word_doc is None:
            logger.warning("Template render returned object without docx payload; skipping post-process.")
            return

        self._apply_initial_layout(word_doc, template_name)
        self._apply_page_size(word_doc, self._resolve_page_size(template_name, options))

        title = source_document.metadata.title or source_document.original_filename or ""
        cover_enabled = self._resolve_bool_option(
            options,
            "cover_page",
            aliases=("add_cover_page",),
            default=True,
        )
        if title and not self._document_contains_text(word_doc, title):
            self._prepend_front_matter(word_doc, source_document, as_cover_page=cover_enabled)

        add_toc = self._resolve_bool_option(
            options,
            "toc",
            aliases=("generate_toc",),
            default=False,
        )
        add_page_numbers = self._resolve_bool_option(
            options,
            "page_numbers",
            aliases=("add_page_numbers",),
            default=True,
        )
        add_borders = self._resolve_bool_option(
            options,
            "borders",
            aliases=("add_borders",),
            default=False,
        )
        add_line_numbers = self._resolve_bool_option(
            options,
            "line_numbers",
            aliases=("add_line_numbers",),
            default=False,
        )

        if add_toc:
            self._remove_static_toc_block(word_doc)
            self._ensure_dynamic_toc(word_doc)
        else:
            self._remove_static_toc_block(word_doc)

        self._remove_static_page_number_placeholders(word_doc)
        if add_page_numbers:
            self._add_page_numbers(word_doc)
        if add_borders:
            self._add_page_borders(word_doc)
        if add_line_numbers:
            self._add_line_numbers(word_doc)
        self._rehydrate_template_render(
            word_doc,
            source_document,
            template_name,
            footnote_lookup=footnote_lookup or {},
        )
        self._apply_global_line_spacing(word_doc, template_name, options)

    def _build_footnote_lookup(self, document: Document) -> dict:
        """Map source footnote identifiers to valid Word footnote IDs."""
        lookup = {}
        next_word_id = 1

        for block in sorted(document.blocks, key=lambda item: item.index):
            if not (
                block.block_type == BlockType.FOOTNOTE
                or block.metadata.get("is_footnote")
            ):
                continue

            footnote_text = (block.text or "").strip()
            if not footnote_text:
                continue

            raw_id = str(
                block.metadata.get("footnote_id")
                or block.metadata.get("endnote_id")
                or next_word_id
            )
            if raw_id in lookup:
                continue

            lookup[raw_id] = {"word_id": next_word_id, "text": footnote_text}
            next_word_id += 1

        return lookup

    def _rehydrate_template_render(
        self,
        doc,
        source_document: Document,
        template_name: str,
        footnote_lookup: dict,
    ) -> None:
        """Upgrade docxtpl output with inline hyperlinks, footnotes, and missing rich content."""
        used_paragraphs = set()

        for block in sorted(source_document.blocks, key=lambda item: item.index):
            if (
                block.metadata.get("is_header")
                or block.metadata.get("is_footer")
                or block.metadata.get("is_footnote")
                or block.metadata.get("is_endnote")
                or block.block_type == BlockType.FOOTNOTE
            ):
                continue

            clean_text = (block.text or "").strip()
            has_inline_artifacts = bool(
                block.metadata.get("hyperlinks") or block.metadata.get("footnote_refs")
            )
            target_paragraph = self._find_matching_paragraph(doc, clean_text, used_paragraphs)
            if target_paragraph is not None:
                used_paragraphs.add(id(target_paragraph))
                if has_inline_artifacts:
                    self._replace_paragraph_inline_content(
                        target_paragraph,
                        block,
                        footnote_lookup=footnote_lookup,
                    )
                continue

            if block.block_type == BlockType.REFERENCE_ENTRY or has_inline_artifacts:
                self._render_block(
                    doc,
                    deepcopy(block),
                    template_name,
                    footnote_lookup=footnote_lookup,
                )

        if source_document.figures:
            for number, figure in enumerate(source_document.figures, start=1):
                self._render_figure(doc, figure, number)
        if source_document.equations:
            for equation in source_document.equations:
                self._render_equation(doc, equation)
        if source_document.tables:
            for number, table in enumerate(source_document.tables, start=1):
                self.table_renderer.render(doc, table, number)

    def _find_matching_paragraph(self, doc, text: str, used_paragraphs: set):
        needle = " ".join((text or "").split())
        if not needle:
            return None

        for paragraph in doc.paragraphs:
            if id(paragraph) in used_paragraphs:
                continue
            haystack = " ".join((paragraph.text or "").split())
            if haystack == needle:
                return paragraph

        for paragraph in doc.paragraphs:
            if id(paragraph) in used_paragraphs:
                continue
            haystack = " ".join((paragraph.text or "").split())
            if needle and needle in haystack:
                return paragraph

        return None

    def _replace_paragraph_inline_content(self, paragraph, block, footnote_lookup: dict) -> None:
        """Replace a template-rendered paragraph with hyperlink-aware content."""
        self._clear_paragraph_content(paragraph)
        self._write_inline_content(
            paragraph,
            (block.text or "").strip(),
            block.metadata.get("hyperlinks", []),
            block.metadata.get("footnote_refs", []),
            footnote_lookup,
        )

    def _clear_paragraph_content(self, paragraph) -> None:
        for child in list(paragraph._p):
            if child.tag != qn("w:pPr"):
                paragraph._p.remove(child)

    def _write_inline_content(
        self,
        paragraph,
        text: str,
        hyperlinks: list,
        footnote_refs: list,
        footnote_lookup: dict,
    ) -> None:
        plain_text = text or ""
        cursor = 0
        wrote_content = False

        for hyperlink in hyperlinks or []:
            label = str(hyperlink.get("text") or "").strip()
            url = str(hyperlink.get("url") or "").strip()
            if not label or not url:
                continue

            match_index = plain_text.find(label, cursor)
            if match_index < 0:
                continue

            if match_index > cursor:
                paragraph.add_run(plain_text[cursor:match_index])
                wrote_content = True

            self._add_hyperlink(paragraph, label, url)
            wrote_content = True
            cursor = match_index + len(label)

        remaining = plain_text[cursor:]
        if remaining:
            paragraph.add_run(remaining)
            wrote_content = True
        elif plain_text and not wrote_content:
            paragraph.add_run(plain_text)
            wrote_content = True

        if not wrote_content:
            paragraph.add_run("")

        for raw_ref in footnote_refs or []:
            entry = footnote_lookup.get(str(raw_ref))
            if entry:
                self._append_footnote_reference(paragraph, entry["word_id"])

    def _add_hyperlink(self, paragraph, text: str, url: str) -> None:
        """Append a true external Word hyperlink to a paragraph."""
        relation_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), relation_id)

        run = OxmlElement("w:r")
        run_properties = OxmlElement("w:rPr")
        run_style = OxmlElement("w:rStyle")
        run_style.set(qn("w:val"), "Hyperlink")
        run_properties.append(run_style)
        run.append(run_properties)

        text_node = OxmlElement("w:t")
        text_node.text = text
        run.append(text_node)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    def _append_footnote_reference(self, paragraph, word_footnote_id: int) -> None:
        """Append a footnote reference marker to a paragraph."""
        run = OxmlElement("w:r")
        run_properties = OxmlElement("w:rPr")

        run_style = OxmlElement("w:rStyle")
        run_style.set(qn("w:val"), "FootnoteReference")
        run_properties.append(run_style)

        vert_align = OxmlElement("w:vertAlign")
        vert_align.set(qn("w:val"), "superscript")
        run_properties.append(vert_align)

        run.append(run_properties)

        reference = OxmlElement("w:footnoteReference")
        reference.set(qn("w:id"), str(word_footnote_id))
        run.append(reference)
        paragraph._p.append(run)

    def _install_post_save_hook(self, rendered_obj, footnote_lookup: dict) -> None:
        """Patch saved DOCX files so Word has a proper footnotes part."""
        if not footnote_lookup or getattr(rendered_obj, "_scholarform_save_hook_installed", False):
            return

        original_save = rendered_obj.save

        def _save_with_footnotes(target, *args, **kwargs):
            result = original_save(target, *args, **kwargs)
            try:
                self._patch_saved_docx_with_footnotes(target, footnote_lookup)
            except Exception as exc:
                logger.warning("Failed to patch saved DOCX footnotes: %s", exc)
            return result

        rendered_obj.save = _save_with_footnotes
        rendered_obj._scholarform_save_hook_installed = True

    def _patch_saved_docx_with_footnotes(self, target, footnote_lookup: dict) -> None:
        if hasattr(target, "read") and hasattr(target, "write") and hasattr(target, "seek"):
            current_position = target.tell()
            target.seek(0)
            payload = target.read()
            if not payload:
                return
            patched_payload = self._patch_docx_payload(payload, footnote_lookup)
            target.seek(0)
            target.truncate()
            target.write(patched_payload)
            target.seek(current_position)
            return

        target_path = os.fspath(target)
        with open(target_path, "rb") as handle:
            payload = handle.read()

        patched_payload = self._patch_docx_payload(payload, footnote_lookup)
        with open(target_path, "wb") as handle:
            handle.write(patched_payload)

    def _patch_docx_payload(self, payload: bytes, footnote_lookup: dict) -> bytes:
        with ZipFile(BytesIO(payload), "r") as source_zip:
            parts = {name: source_zip.read(name) for name in source_zip.namelist()}

        if b"w:footnoteReference" not in parts.get("word/document.xml", b""):
            return payload

        parts["word/footnotes.xml"] = self._build_footnotes_part(footnote_lookup)
        parts["[Content_Types].xml"] = self._patch_content_types(parts["[Content_Types].xml"])
        parts["word/_rels/document.xml.rels"] = self._patch_document_relationships(
            parts.get("word/_rels/document.xml.rels", b""),
        )
        parts["word/settings.xml"] = self._patch_settings_xml(parts.get("word/settings.xml", b""))

        output_stream = BytesIO()
        with ZipFile(output_stream, "w", compression=ZIP_DEFLATED) as target_zip:
            for name, content in parts.items():
                target_zip.writestr(name, content)
        return output_stream.getvalue()

    def _build_footnotes_part(self, footnote_lookup: dict) -> bytes:
        footnotes_root = ET.Element(qn("w:footnotes"))

        separator = ET.SubElement(
            footnotes_root,
            qn("w:footnote"),
            {qn("w:type"): "separator", qn("w:id"): "-1"},
        )
        separator_paragraph = ET.SubElement(separator, qn("w:p"))
        separator_run = ET.SubElement(separator_paragraph, qn("w:r"))
        ET.SubElement(separator_run, qn("w:separator"))

        continuation = ET.SubElement(
            footnotes_root,
            qn("w:footnote"),
            {qn("w:type"): "continuationSeparator", qn("w:id"): "0"},
        )
        continuation_paragraph = ET.SubElement(continuation, qn("w:p"))
        continuation_run = ET.SubElement(continuation_paragraph, qn("w:r"))
        ET.SubElement(continuation_run, qn("w:continuationSeparator"))

        sorted_entries = sorted(footnote_lookup.values(), key=lambda item: int(item["word_id"]))
        for entry in sorted_entries:
            footnote = ET.SubElement(
                footnotes_root,
                qn("w:footnote"),
                {qn("w:id"): str(entry["word_id"])},
            )
            paragraph = ET.SubElement(footnote, qn("w:p"))
            reference_run = ET.SubElement(paragraph, qn("w:r"))
            reference_props = ET.SubElement(reference_run, qn("w:rPr"))
            reference_style = ET.SubElement(reference_props, qn("w:rStyle"))
            reference_style.set(qn("w:val"), "FootnoteReference")
            ET.SubElement(reference_run, qn("w:footnoteRef"))

            text_run = ET.SubElement(paragraph, qn("w:r"))
            text_node = ET.SubElement(text_run, qn("w:t"))
            text_node.set(qn("xml:space"), "preserve")
            text_node.text = f" {entry['text']}"

        return ET.tostring(footnotes_root, encoding="utf-8", xml_declaration=True)

    def _patch_content_types(self, content_types_xml: bytes) -> bytes:
        root = ET.fromstring(content_types_xml)
        existing = root.findall(f"./{{*}}Override[@PartName='/word/footnotes.xml']")
        if not existing:
            override = ET.Element(
                f"{{{root.tag.partition('}')[0].strip('{')}}}Override",
                {
                    "PartName": "/word/footnotes.xml",
                    "ContentType": "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
                },
            )
            root.append(override)
        return ET.tostring(root, encoding="utf-8", xml_declaration=True)

    def _patch_document_relationships(self, relationships_xml: bytes) -> bytes:
        if relationships_xml:
            root = ET.fromstring(relationships_xml)
        else:
            root = ET.Element("Relationships")

        for relationship in root.findall("./{*}Relationship"):
            if relationship.get("Type") == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes":
                return ET.tostring(root, encoding="utf-8", xml_declaration=True)

        existing_ids = {relationship.get("Id") for relationship in root.findall("./{*}Relationship")}
        next_id = 1
        while f"rId{next_id}" in existing_ids:
            next_id += 1

        new_relationship = ET.Element(
            "Relationship",
            {
                "Id": f"rId{next_id}",
                "Type": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes",
                "Target": "footnotes.xml",
            },
        )
        root.append(new_relationship)
        return ET.tostring(root, encoding="utf-8", xml_declaration=True)

    def _patch_settings_xml(self, settings_xml: bytes) -> bytes:
        if not settings_xml:
            return settings_xml

        root = ET.fromstring(settings_xml)
        existing = root.find(f"./{{*}}footnotePr")
        if existing is None:
            footnote_properties = ET.Element(qn("w:footnotePr"))
            num_format = ET.SubElement(footnote_properties, qn("w:numFmt"))
            num_format.set(qn("w:val"), "decimal")
            root.append(footnote_properties)
        return ET.tostring(root, encoding="utf-8", xml_declaration=True)

    def _set_columns(self, section, count: int):
        """Helper to set column count on a python-docx section."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        
        # Access underlying XML
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')
        if not cols:
            cols = OxmlElement('w:cols')
            sectPr.append(cols)
        else:
            cols = cols[0]
            
        cols.set(qn('w:num'), str(count))
        # Ensure space between columns if > 1
        if count > 1:
            cols.set(qn('w:space'), '720') # 0.5 inch (720 twips)

    def _load_contract(self, path: str) -> dict:
        if not os.path.exists(path):
            return {}
        try:
            with open(path, 'r') as f:
                return yaml.safe_load(f) or {}
        except Exception as e:
            logger.warning("Failed to load contract %s: %s", path, e)
            return {}


    @safe_function(fallback_value=None, error_message="Block rendering failed")
    def _render_block(self, doc, block, template_name, footnote_lookup: Optional[dict] = None):
        """Render a block with contract-driven spacing, formatting, and dynamic list detection."""
        # Skip rendering empty anchor blocks (preserve in pipeline)
        if block.text.strip() == "" and block.metadata.get("has_figure", False):
            return  # Block remains in pipeline for anchor stability
        
        if block.text.strip() == "" and block.metadata.get("has_equation", False):
            return  # Block remains in pipeline for anchor stability
        
        # DYNAMIC LIST DETECTION
        clean_text = block.text.strip()
        if not clean_text:
            return  # Skip empty blocks
        
        try:
            # Check if this is a list item
            if self._is_bullet_list_item(clean_text):
                # Render as bullet list
                clean_text = self._clean_list_text(clean_text)
                p = doc.add_paragraph(style="List Bullet")
            elif self._is_numbered_list_item(clean_text):
                # Render as numbered list
                clean_text = self._clean_list_text(clean_text)
                p = doc.add_paragraph(style="List Number")
            else:
                # Normal paragraph rendering
                word_style = self.style_mapper.get_style_name(block, template_name)
                p = doc.add_paragraph(style=word_style)
            
            self._write_inline_content(
                p,
                clean_text,
                block.metadata.get("hyperlinks", []),
                block.metadata.get("footnote_refs", []),
                footnote_lookup or {},
            )
            # Apply contract-driven spacing
            self._apply_spacing_from_contract(p, block, template_name)
            
        except Exception:
            # Fallback if style missing
            if clean_text:
                p = doc.add_paragraph()
                self._write_inline_content(
                    p,
                    clean_text,
                    block.metadata.get("hyperlinks", []),
                    block.metadata.get("footnote_refs", []),
                    footnote_lookup or {},
                )
                self._apply_spacing_from_contract(p, block, template_name)
            else:
                return

        return p


    def _apply_spacing_from_contract(self, paragraph, block, template_name):
        """Apply spacing rules from contract to paragraph."""
        contract = self.contract_loader.load(template_name)
        layout = contract.get("layout", {})
        spacing_rules = layout.get("spacing", {})
        
        if not spacing_rules:
            return  # No spacing rules in contract
        
        # Determine block type and get appropriate spacing
        if hasattr(block, 'is_heading') and block.is_heading():
            spacing = spacing_rules.get("heading", {})
        elif str(block.block_type).upper() in ["FIGURE_CAPTION", "TABLE_CAPTION"]:
            # Use figure or table spacing for captions
            if "FIGURE" in str(block.block_type).upper():
                spacing = spacing_rules.get("figure", {})
            else:
                spacing = spacing_rules.get("table", {})
        elif "REFERENCES" in str(block.block_type).upper():
            # Special spacing for references section
            spacing = spacing_rules.get("references", spacing_rules.get("heading", {}))
        else:
            spacing = spacing_rules.get("paragraph", {})
        
        # Apply spacing if defined
        if spacing:
            before = spacing.get("before", 0)
            after = spacing.get("after", 0)
            paragraph.paragraph_format.space_before = Pt(before)
            paragraph.paragraph_format.space_after = Pt(after)

        line_spacing = layout.get("line_spacing")
        try:
            if line_spacing:
                paragraph.paragraph_format.line_spacing = float(line_spacing)
        except (TypeError, ValueError):
            pass

    @safe_function(fallback_value=None, error_message="Image sizing failed")
    def _calculate_image_size(self, figure: Figure):
        """
        Calculate optimal image size based on actual dimensions and page constraints.
        
        Strategy:
        1. Use figure.width/height if available (from Figure model)
        2. Constrain to page width (6.5 inches for standard letter with 1" margins)
        3. Maintain aspect ratio
        4. Set reasonable min/max bounds
        """
        # Page constraints (standard letter: 8.5" wide, 1" margins each side = 6.5" usable)
        MAX_WIDTH = Inches(6.5)
        MAX_HEIGHT = Inches(9.0)  # Standard letter: 11" tall, 1" margins = 9" usable
        MIN_WIDTH = Inches(2.0)   # Minimum for readability
        DEFAULT_WIDTH = Inches(5.0)  # Good default for most images
        
        # If figure has dimensions, use them
        if figure.width and figure.height:
            # Convert to inches (assuming pixels at 96 DPI)
            img_width_inches = Inches(figure.width / 96.0)
            img_height_inches = Inches(figure.height / 96.0)
            
            # Calculate aspect ratio
            aspect_ratio = figure.width / figure.height
            
            # Scale to fit within page width
            if img_width_inches > MAX_WIDTH:
                final_width = MAX_WIDTH
                final_height = Inches(MAX_WIDTH.inches / aspect_ratio)
            elif img_width_inches < MIN_WIDTH:
                # Small images: scale up to minimum width
                final_width = MIN_WIDTH
                final_height = Inches(MIN_WIDTH.inches / aspect_ratio)
            else:
                # Image fits naturally
                final_width = img_width_inches
                final_height = img_height_inches
            
            # Ensure height doesn't exceed page
            if final_height > MAX_HEIGHT:
                final_height = MAX_HEIGHT
                final_width = Inches(MAX_HEIGHT.inches * aspect_ratio)
            
            return final_width, final_height
        else:
            # No dimensions available: use smart default
            # Default to 5 inches wide (good for most academic figures)
            return DEFAULT_WIDTH, None  # Let python-docx maintain aspect ratio

    @safe_function(fallback_value=None, error_message="Figure rendering failed")
    def _render_figure(self, doc, figure: Figure, number: int):
        """Render a figure with dynamic sizing based on image dimensions."""
        # 1. Add image with dynamic sizing
        if figure.export_path and os.path.exists(figure.export_path):
            try:
                # Calculate optimal size based on image dimensions
                width, height = self._calculate_image_size(figure)
                # Add image and get the paragraph containing it
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(figure.export_path, width=width, height=height)
                paragraph.alignment = 1  # Center the image
            except Exception as e:
                logger.warning("Failed to render figure from export_path: %s", e)
                # Fallback: add placeholder text if image fails
                p = doc.add_paragraph(f"[Image: {figure.export_path}]")
                p.alignment = 1  # Center
        elif figure.image_data:
            # If we have image_data, use BytesIO (more reliable than temp files)
            try:
                from io import BytesIO
                image_stream = BytesIO(figure.image_data)
                # Calculate optimal size
                width, height = self._calculate_image_size(figure)
                # Add image to a paragraph and center it
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                if height:
                    run.add_picture(image_stream, width=width, height=height)
                else:
                    run.add_picture(image_stream, width=width)
                paragraph.alignment = 1  # Center the image
                logger.info("Rendered figure %d from image_data (%d bytes)", number, len(figure.image_data))
            except Exception as e:
                logger.warning("Failed to render figure from image_data: %s", e)
                # Fallback: add placeholder
                p = doc.add_paragraph(f"[Figure {number} - Image rendering failed: {str(e)[:50]}]")
                p.alignment = 1  # Center
        else:
            # If no image data or path, add a placeholder
            p = doc.add_paragraph(f"[Figure {number} Placeholder - No image data]")
            p.alignment = 1  # Center the placeholder too
        
        # 2. Add caption with bold prefix (check if prefix already exists)
        if figure.caption_text:
            caption_p = doc.add_paragraph(style="Caption")
            # Check if caption already starts with "Figure N:" to avoid duplication
            caption_lower = figure.caption_text.lower().strip()
            if caption_lower.startswith(f"figure {number}:"):
                # Caption already has prefix, just add it as-is with bold prefix
                run = caption_p.add_run(f"Figure {number}: ")
                run.bold = True
                # Add the rest after the prefix
                rest_text = figure.caption_text[len(f"Figure {number}:"):].strip()
                caption_p.add_run(rest_text)
            else:
                # Caption doesn't have prefix, add it
                run = caption_p.add_run(f"Figure {number}: ")
                run.bold = True
                caption_p.add_run(figure.caption_text)
            caption_p.alignment = 1  # Center

    def _is_bullet_list_item(self, text: str) -> bool:
        """Dynamically detect if text is a bullet list item."""
        if not text:
            return False
        stripped = text.lstrip()
        # Check for common bullet markers
        return stripped.startswith(('\u2022', '-', '*', '\u00b7', '\u25e6', '\u25aa', '\u25ab'))
    
    def _is_numbered_list_item(self, text: str) -> bool:
        """Dynamically detect if text is a numbered list item."""
        if not text:
            return False
        import re
        # Match patterns like "1. ", "1) ", "a. ", "a) ", "i. ", "i) "
        return bool(re.match(r'^\s*([0-9]+|[a-z]|[ivxlcdm]+)[\.)\s]\s+', text, re.IGNORECASE))
    
    def _clean_list_text(self, text: str) -> str:
        """Remove list markers from text for proper Word list rendering."""
        import re
        # Remove bullet markers
        text = re.sub(r'^\s*[\u2022\-\*\u00b7\u25e6\u25aa\u25ab]\s+', '', text)
        # Remove numbered markers
        text = re.sub(r'^\s*([0-9]+|[a-z]|[ivxlcdm]+)[\.)\s]\s+', '', text, flags=re.IGNORECASE)
        return text

