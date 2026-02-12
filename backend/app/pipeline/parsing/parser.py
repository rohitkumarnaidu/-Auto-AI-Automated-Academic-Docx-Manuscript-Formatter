"""
DOCX Parser - Extract raw content from DOCX documents.

This is the first stage of the pipeline. It reads a DOCX file and extracts:
- Paragraphs with text and style information
- Tables with cell structure
- Images/figures with binary data

Output: Document model with all raw content preserved in original order.

IMPORTANT:
- All blocks are marked as BlockType.UNKNOWN (classification happens later)
- No normalization, structure detection, or formatting
- Never drop content silently
- Preserve exact order from source document
"""

import io
import os
from typing import List, Optional, Tuple
from datetime import datetime
from pathlib import Path

from docx import Document as DocxDocument
from docx.document import Document as DocxDocumentType
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shape import InlineShape
from docx.oxml.ns import qn

from app.pipeline.tables.extractor import TableExtractor

from app.models import (
    PipelineDocument as Document,
    DocumentMetadata,
    Block,
    BlockType,
    TextStyle,
    Figure,
    ImageFormat,
    Table,
    TableCell,
    Equation,
)
from app.utils.id_generator import (
    generate_block_id, 
    generate_figure_id, 
    generate_table_id,
    generate_equation_id
)


class DocxParser:
    """
    Parses DOCX files into Document model instances.
    
    This parser extracts raw content without interpreting semantic meaning.
    Structure detection and classification happen in later pipeline stages.
    """
    
    def __init__(self):
        """Initialize the parser."""
        self.block_counter = 0
        self.element_counter = 0  # Global sequential index
        self.figure_counter = 0
        self.table_counter = 0
        self.equation_counter = 0
        self.current_page = None  # DOCX doesn't provide reliable page numbers
    
    def parse(self, docx_path: str, document_id: str) -> Document:
        """
        Parse a DOCX file into a Document model.
        
        Args:
            docx_path: Path to the .docx file
            document_id: Unique identifier for this document (e.g., job ID)
        
        Returns:
            Document instance with all extracted content
        
        Raises:
            FileNotFoundError: If DOCX file doesn't exist
            ValueError: If file is not a valid DOCX
        """
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"DOCX file not found: {docx_path}")
        
        # Reset counters for this parse
        self.block_counter = 0
        self.element_counter = 0
        self.figure_counter = 0
        self.table_counter = 0
        self.equation_counter = 0
        
        # Open DOCX document
        try:
            docx = DocxDocument(docx_path)
        except Exception as e:
            raise ValueError(f"Failed to open DOCX file: {e}")
        
        # Initialize document
        # FIX: Explicitly cast document_id to string to avoid Pydantic ValidationError
        # The orchestrator might pass a UUID object, but PipelineDocument expects a str.
        if not isinstance(document_id, str):
            document_id = str(document_id)

        document = Document(
            document_id=document_id,
            original_filename=Path(docx_path).name,
            source_path=docx_path,
            created_at=datetime.utcnow(),
            updated_at=datetime.utcnow()
        )
        
        # Extract core properties
        document.metadata = self._extract_core_properties(docx)
        
        # Extract document content in order
        # DOCX structure: paragraphs and tables are interspersed
        blocks, figures, tables, equations = self._extract_body_content(docx)

        # SAFE EXTENSION: Headers and Footers
        # Extract and append at the very end to avoid index corruption
        header_footer_blocks = self._extract_headers_and_footers(docx)
        if header_footer_blocks:
            blocks.extend(header_footer_blocks)

        # SAFE EXTENSION: Footnotes and Endnotes
        # We append them after the main body to preserve main index order logic
        # and prevent breaking any body-relative logic.
        note_blocks = self._extract_footnotes_and_endnotes(docx)
        if note_blocks:
            blocks.extend(note_blocks)
            # Update history message to reflect notes (optional, but good for visibility)
        
        document.blocks = blocks
        document.figures = figures
        document.tables = tables
        document.equations = equations
        
        # Add processing history
        msg = f"Parsed {len(blocks)} blocks, {len(figures)} figures, {len(tables)} tables, {len(equations)} equations"
        if note_blocks or header_footer_blocks:
            msg += f" (incl. {len(note_blocks)} notes and {len(header_footer_blocks)} header/footers)"

        document.add_processing_stage(
            stage_name="parsing",
            status="success",
            message=msg
        )
        
        return document
    
    def _extract_footnotes_and_endnotes(self, docx: DocxDocumentType) -> List[Block]:
        """
        Extract content from Footnotes and Endnotes parts.
        
        Args:
            docx: python-docx Document object
            
        Returns:
            List of Block objects marked with is_footnote/is_endnote
        """
        note_blocks = []
        
        # 1. Footnotes
        try:
            part = None
            if hasattr(docx, 'part'):
                if hasattr(docx.part, 'footnotes_part'):
                    part = docx.part.footnotes_part
            
            if part:
                # Iterate XML elements directly as python-docx high-level API is limited
                from docx.oxml.ns import qn
                root = part.element
                for i, fn in enumerate(root.findall(qn('w:footnote'))):
                    # Skip separator/continuation footnotes (usually ids < 0 or specific types)
                    # We just try to extract text from paragraphs inside
                    fn_id = fn.get(qn('w:id'))
                    
                    # Extract paragraphs within footnote
                    for p_element in fn.findall(qn('w:p')):
                        # Create a temp paragraph to use existing extraction logic is tricky
                        # because _extract_paragraph expects a docx parent.
                        # We will do manual extraction to be safe and simple.
                        text_chunks = []
                        for r in p_element.findall(qn('w:r')):
                            t = r.find(qn('w:t'))
                            if t is not None and t.text:
                                text_chunks.append(t.text)
                        
                        text = "".join(text_chunks).strip()
                        if text:
                            # Create Block
                            block_id = generate_block_id(self.block_counter)
                            self.block_counter += 1
                            
                            # Fake style (Footnote Text usually)
                            style = TextStyle(font_size=10.0) # Assumption
                            
                            block = Block(
                                block_id=block_id,
                                text=text,
                                index=self.block_counter - 1, # Continue sequence
                                block_type=BlockType.UNKNOWN,
                                style=style
                            )
                            block.metadata["is_footnote"] = True
                            block.metadata["footnote_id"] = fn_id
                            note_blocks.append(block)
        except Exception as e:
            # Silent failure as per requirements ("No crashes allowed")
            print(f"Warning: Footnote extraction skipped: {e}")

        # 2. Endnotes
        try:
            part = None
            if hasattr(docx, 'part'):
                if hasattr(docx.part, 'endnotes_part'):
                    part = docx.part.endnotes_part
            
            if part:
                from docx.oxml.ns import qn
                root = part.element
                for i, en in enumerate(root.findall(qn('w:endnote'))):
                    en_id = en.get(qn('w:id'))
                    for p_element in en.findall(qn('w:p')):
                        text_chunks = []
                        for r in p_element.findall(qn('w:r')):
                            t = r.find(qn('w:t'))
                            if t is not None and t.text:
                                text_chunks.append(t.text)
                        
                        text = "".join(text_chunks).strip()
                        if text:
                            block_id = generate_block_id(self.block_counter)
                            self.block_counter += 1
                            
                            style = TextStyle(font_size=10.0)
                            
                            block = Block(
                                block_id=block_id,
                                text=text,
                                index=self.block_counter - 1,
                                block_type=BlockType.UNKNOWN,
                                style=style
                            )
                            block.metadata["is_endnote"] = True
                            block.metadata["endnote_id"] = en_id
                            note_blocks.append(block)
        except Exception:
            pass

        return note_blocks

    def _extract_headers_and_footers(self, docx: DocxDocumentType) -> List[Block]:
        """
        Extract content from Headers and Footers of all sections.
        """
        hf_blocks = []
        try:
            for i, section in enumerate(docx.sections):
                # Process Headers
                if section.header:
                    for p in section.header.paragraphs:
                        block = self._extract_paragraph(p)
                        if block:
                            block.metadata["is_header"] = True
                            block.metadata["section_index"] = i
                            hf_blocks.append(block)
                
                # Process Footers
                if section.footer:
                    for p in section.footer.paragraphs:
                        block = self._extract_paragraph(p)
                        if block:
                            block.metadata["is_footer"] = True
                            block.metadata["section_index"] = i
                            hf_blocks.append(block)
        except Exception as e:
            print(f"Warning: Header/Footer extraction skipped: {e}")
            
        return hf_blocks

    def _extract_core_properties(self, docx: DocxDocumentType) -> DocumentMetadata:
        """
        Extract document metadata from DOCX core properties.
        
        Args:
            docx: python-docx Document object
        
        Returns:
            DocumentMetadata instance
        """
        core_props = docx.core_properties
        
        metadata = DocumentMetadata()
        
        # Extract available metadata
        if core_props.title:
            metadata.title = core_props.title
        
        if core_props.author:
            # Author may be a single string or semicolon-separated
            authors = [a.strip() for a in core_props.author.split(';')]
            metadata.authors = authors
        
        if core_props.subject:
            metadata.abstract = core_props.subject
        
        if core_props.keywords:
            # Keywords are usually comma or semicolon separated
            keywords = [k.strip() for k in core_props.keywords.replace(';', ',').split(',')]
            metadata.keywords = [k for k in keywords if k]
        
        if core_props.created:
            metadata.publication_date = core_props.created
        
        return metadata
    
    def _extract_body_content(
        self, docx: DocxDocumentType
    ) -> Tuple[List[Block], List[Figure], List[Table], List[Equation]]:
        """
        Extract all content from document body in original order.
        
        DOCX documents contain a mix of paragraphs and tables.
        We need to preserve their original sequential order.
        
        Args:
            docx: python-docx Document object
        
        Returns:
            Tuple of (blocks, figures, tables)
        """
        blocks: List[Block] = []
        figures: List[Figure] = []
        tables: List[Table] = []
        equations: List[Equation] = []
        
        # Iterate through body elements in order
        # This preserves the exact sequence of content
        for element in docx.element.body:
            if isinstance(element, CT_P):
                # Paragraph element
                paragraph = DocxParagraph(element, docx)
                
                # Extract text block
                block = self._extract_paragraph(paragraph)
                if block:  # Only add if not None (skip certain empties)
                    block.index = self.element_counter
                    self.element_counter += 1
                    blocks.append(block)
                
                # Check for inline images
                inline_figures = self._extract_inline_images(paragraph)
                
                # Attach block index to figures for proximity matching
                if block:
                    for figure in inline_figures:
                        figure.metadata["block_index"] = block.index
                    # FORENSIC FIX: Mark block as container to prevent Normalizer dropping it
                    if inline_figures:
                        block.metadata["has_figure"] = True
                
                figures.extend(inline_figures)

                # Check for equations in paragraph
                paragraph_equations = self._extract_equations(paragraph)
                for eqn in paragraph_equations:
                    if block:
                        eqn.block_id = block.block_id
                        # FORENSIC FIX: Anchor equation to block index for correct sorting
                        eqn.metadata["block_index"] = block.index
                        # FORENSIC FIX: Mark block as container
                        block.metadata["has_equation"] = True
                equations.extend(paragraph_equations)
            
            elif isinstance(element, CT_Tbl):
                # Table element
                table = DocxTable(element, docx)
                extracted_table = self._extract_table(table, self.element_counter)
                self.element_counter += 1
                tables.append(extracted_table)
        
        return blocks, figures, tables, equations
    
    def _extract_paragraph(self, paragraph: DocxParagraph) -> Optional[Block]:
        """
        Extract a paragraph as a Block.
        
        Args:
            paragraph: python-docx Paragraph object
        
        Returns:
            Block instance or None if paragraph should be skipped
        """
        # Get text content
        text = paragraph.text
        
        # Edge case: completely empty paragraphs
        # We preserve them as they might indicate spacing/structure
        # but mark them with empty text
        
        # Extract style information
        style = self._extract_paragraph_style(paragraph)
        
        # Generate unique ID
        block_id = generate_block_id(self.block_counter)
        self.block_counter += 1
        
        # Create block
        # Note: block_type is UNKNOWN - classification happens later
        block = Block(
            block_id=block_id,
            text=text,
            index=self.block_counter - 1,
            block_type=BlockType.UNKNOWN,
            style=style,
            page_number=None,  # DOCX doesn't provide direct page numbers
        )
        
        # Store paragraph style name in metadata for later structure detection
        if paragraph.style and paragraph.style.name:
            block.metadata["style_name"] = paragraph.style.name
        
        # Store alignment
        if paragraph.alignment is not None:
            block.metadata["alignment"] = str(paragraph.alignment)
        
        # SAFE EXTENSION: Hyperlink Extraction
        hyperlinks = self._extract_hyperlinks(paragraph)
        if hyperlinks:
            block.metadata["hyperlinks"] = hyperlinks

        # SAFE EXTENSION: Nested List Depth
        list_info = self._get_list_info(paragraph)
        if list_info:
            block.metadata.update(list_info)
        
        return block

    def _extract_hyperlinks(self, paragraph: DocxParagraph) -> List[Dict[str, str]]:
        """Extract URLs from w:hyperlink elements."""
        links = []
        try:
            # Hyperlinks are at the paragraph element level in the XML
            for hyperlink in paragraph._element.findall(qn('w:hyperlink')):
                # Get RId
                r_id = hyperlink.get(qn('r:id'))
                if r_id:
                    # Resolve URL from relationship
                    try:
                        url = paragraph.part.rels[r_id].target_ref
                        # Get Text
                        text = "".join([t.text for t in hyperlink.findall(qn('w:r')) if t.text])
                        if text:
                            links.append({"text": text, "url": url})
                    except (KeyError, AttributeError):
                        continue
        except Exception:
            pass
        return links

    def _get_list_info(self, paragraph: DocxParagraph) -> Optional[Dict[str, Any]]:
        """Detect list level and status using w:ilvl."""
        try:
            pPr = paragraph._element.find(qn('w:pPr'))
            if pPr is not None:
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    ilvl = numPr.find(qn('w:ilvl'))
                    if ilvl is not None:
                        depth = int(ilvl.get(qn('w:val'), 0))
                        return {
                            "list_level": depth,
                            "is_list_item": True
                        }
                    # If numPr exists but no ilvl, it's still a list item (level 0)
                    return {
                        "list_level": 0,
                        "is_list_item": True
                    }
        except Exception:
            pass
        return None
    
    def _extract_paragraph_style(self, paragraph: DocxParagraph) -> TextStyle:
        """
        Extract text style from paragraph.
        
        For paragraphs with mixed formatting (runs), we extract the style
        of the first non-empty run as a representative sample.
        
        Args:
            paragraph: python-docx Paragraph object
        
        Returns:
            TextStyle instance
        """
        # Default style
        bold = False
        italic = False
        underline = False
        font_name = None
        font_size = None
        
        # Check runs for formatting
        # Edge case: paragraph may have multiple runs with different styles
        # We take the first run with actual content as representative
        for run in paragraph.runs:
            if run.text.strip():  # First non-empty run
                if run.bold is not None:
                    bold = run.bold
                if run.italic is not None:
                    italic = run.italic
                if run.underline is not None:
                    underline = True
                if run.font.name:
                    font_name = run.font.name
                if run.font.size:
                    font_size = run.font.size.pt  # Convert to points
                break
        
        # Alternative: check paragraph-level formatting if no runs
        if not paragraph.runs and paragraph.style:
            try:
                if paragraph.style.font.bold:
                    bold = True
                if paragraph.style.font.italic:
                    italic = True
                if paragraph.style.font.name:
                    font_name = paragraph.style.font.name
                if paragraph.style.font.size:
                    font_size = paragraph.style.font.size.pt
            except AttributeError:
                pass  # Style doesn't have font properties
        
        return TextStyle(
            bold=bold,
            italic=italic,
            underline=underline,
            font_name=font_name,
            font_size=font_size
        )
    
    def _extract_inline_images(self, paragraph: DocxParagraph) -> List[Figure]:
        """
        Extract inline images from a paragraph.
        
        Images in DOCX can be embedded inline within paragraphs.
        
        Args:
            paragraph: python-docx Paragraph object
        
        Returns:
            List of Figure instances
        """
        figures = []
        
        for run in paragraph.runs:
            # Check for inline shapes (images)
            if hasattr(run, '_element'):
                for drawing in run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                    # Found an image
                    try:
                        # Get image part
                        inline_shapes = run._element.findall(
                            './/{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline'
                        )
                        
                        for inline in inline_shapes:
                            # Safely get part from run (support different python-docx versions)
                            # This prevents 'AttributeError: 'Run' object has no attribute '_part''
                            part = None
                            try:
                                part = getattr(run, 'part', getattr(run, '_part', None))
                            except AttributeError:
                                # Certain environments might still throw if even getattr is restricted (rare)
                                pass
                                
                            if part:
                                figure = self._extract_image_from_inline(inline, part)
                                if figure:
                                    figures.append(figure)
                    except Exception as e:
                        # If image extraction fails, log but continue
                        # Don't let image errors stop paragraph processing
                        print(f"Warning: Failed to extract inline image: {e}")
        
        return figures
    
    def _extract_image_from_inline(self, inline_element, part) -> Optional[Figure]:
        """
        Extract image data from inline shape element.
        
        Args:
            inline_element: XML element for inline shape
            part: Document part containing image relationships
        
        Returns:
            Figure instance or None if extraction fails
        """
        try:
            # Find the image reference
            blip = inline_element.find(
                './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
            )
            
            if blip is None:
                return None
            
            # Get relationship ID
            embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            
            if not embed_id:
                return None
            
            # Get image part from relationship
            image_part = part.related_parts[embed_id]
            image_data = image_part.blob
            
            # Determine image format from content type
            content_type = image_part.content_type
            image_format = self._get_image_format(content_type)
            
            # Get dimensions if available
            extent = inline_element.find(
                './/{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}extent'
            )
            width = None
            height = None
            if extent is not None:
                # EMUs (English Metric Units) - convert to pixels (approx)
                cx = extent.get('cx')
                cy = extent.get('cy')
                if cx:
                    width = int(cx) / 9525  # EMU to pixels (approx)
                if cy:
                    height = int(cy) / 9525
            
            # Create figure
            figure_id = generate_figure_id(self.figure_counter)
            self.figure_counter += 1
            
            figure = Figure(
                figure_id=figure_id,
                index=self.figure_counter - 1,
                image_data=image_data,
                image_format=image_format,
                width=width,
                height=height
            )
            
            return figure
        
        except Exception as e:
            print(f"Warning: Failed to extract image: {e}")
            return None
    
    def _get_image_format(self, content_type: str) -> ImageFormat:
        """
        Map MIME content type to ImageFormat enum.
        
        Args:
            content_type: MIME type (e.g., 'image/png')
        
        Returns:
            ImageFormat enum value
        """
        format_map = {
            'image/png': ImageFormat.PNG,
            'image/jpeg': ImageFormat.JPEG,
            'image/jpg': ImageFormat.JPG,
            'image/gif': ImageFormat.GIF,
            'image/bmp': ImageFormat.BMP,
            'image/tiff': ImageFormat.TIFF,
            'image/svg+xml': ImageFormat.SVG,
            'image/x-emf': ImageFormat.EMF,
            'image/x-wmf': ImageFormat.WMF,
        }
        
        return format_map.get(content_type, ImageFormat.UNKNOWN)
    
    def _extract_table(self, table: DocxTable, block_index: int) -> Table:
        """
        Extract a table using the specialized TableExtractor.
        
        Args:
            table: python-docx Table object
            block_index: Global sequential index in document
        
        Returns:
            Table instance
        """
        extractor = TableExtractor()
        extracted_table = extractor.extract(
            table, 
            generate_table_id(self.table_counter),
            self.table_counter,
            block_index
        )
        self.table_counter += 1
        return extracted_table
    


    def _extract_equations(self, paragraph: DocxParagraph) -> List[Equation]:
        """
        Extract equations (OMML) from a paragraph.
        """
        equations = []
        
        # Word Math namespaces
        # m:oMathPara (Block equation wrapper)
        # m:oMath (The actual equation)
        
        # Look for oMathPara (Block Equations)
        for om_para in paragraph._element.findall(qn('m:oMathPara')):
            for om in om_para.findall(qn('m:oMath')):
                eqn = self._extract_math_element(om, is_block=True)
                if eqn:
                    equations.append(eqn)
                    
        # Look for inline oMath (not inside oMathPara)
        # This is a bit tricky with findall if they are nested differently.
        # Most inline equations are direct children of w:p or inside w:r.
        for om in paragraph._element.findall(qn('m:oMath')):
            # Check if already processed as block
            if any(om is b_om for b_om in paragraph._element.findall(f".//{qn('m:oMathPara')}/{qn('m:oMath')}")):
                continue
                
            eqn = self._extract_math_element(om, is_block=False)
            if eqn:
                equations.append(eqn)
                
        return equations

    def _extract_math_element(self, om_element, is_block: bool) -> Optional[Equation]:
        """Extract data from an oMath element."""
        try:
            # Get raw XML (OMML)
            from lxml import etree
            omml_str = etree.tostring(om_element, encoding='unicode')
            
            # Simple text extraction (heuristic)
            # oMath contains w:t or m:t elements
            math_text = "".join(t.text for t in om_element.findall(f".//{qn('m:t')}"))
            if not math_text:
                math_text = "".join(t.text for t in om_element.findall(f".//{qn('w:t')}"))
            
            # Create Equation
            eqn_id = generate_equation_id(self.equation_counter)
            self.equation_counter += 1
            
            return Equation(
                equation_id=eqn_id,
                index=self.equation_counter - 1,
                text=math_text,
                omml=omml_str,
                is_block=is_block
            )
        except Exception as e:
            print(f"Warning: Failed to extract equation: {e}")
            return None


# Convenience function
def parse_docx(docx_path: str, document_id: str) -> Document:
    """
    Parse a DOCX file into a Document model.
    
    This is a convenience wrapper around DocxParser.
    
    Args:
        docx_path: Path to the .docx file
        document_id: Unique identifier for this document
    
    Returns:
        Document instance with extracted content
    
    Example:
        >>> doc = parse_docx("manuscript.docx", "job_123")
        >>> print(f"Extracted {len(doc.blocks)} blocks")
    """
    parser = DocxParser()
    return parser.parse(docx_path, document_id)


# Helper functions for unique IDs
def generate_figure_id(counter: int) -> str:
    """Generate a unique ID for a figure (e.g., 'fig_001')."""
    return f"fig_{counter:03d}"


def generate_table_id(counter: int) -> str:
    """Generate a unique ID for a table (e.g., 'tbl_001')."""
    return f"tbl_{counter:03d}"


def generate_equation_id(counter: int) -> str:
    """Generate a unique ID for an equation (e.g., 'eqn_001')."""
    return f"eqn_{counter:03d}"
