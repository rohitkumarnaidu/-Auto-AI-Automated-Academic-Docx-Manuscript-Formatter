"""
Table Renderer - Renders Table models into python-docx tables.
"""

import logging

from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models import Table

logger = logging.getLogger(__name__)

class TableRenderer:
    """
    Renders Table objects into a Word document.
    """
    
    def __init__(self):
        pass
        
    def render(self, doc, table_model: Table, number: int = None):
        """
        Render a Table model into the document with caption.
        """
        if not table_model or not table_model.rows:
            return

        try:
            # 1. ADD CAPTION FIRST (if exists) - Academic standard: caption ABOVE table
            if table_model.caption_text:
                caption_p = doc.add_paragraph(style="Caption")
                table_num = number if number is not None else (table_model.index + 1)
                
                caption_lower = table_model.caption_text.lower().strip()
                if caption_lower.startswith(f"table {table_num}:"):
                    run = caption_p.add_run(f"Table {table_num}: ")
                    run.bold = True
                    rest_text = table_model.caption_text[len(f"Table {table_num}:"):].strip()
                    caption_p.add_run(rest_text)
                else:
                    run = caption_p.add_run(f"Table {table_num}: ")
                    run.bold = True
                    caption_p.add_run(table_model.caption_text)
                caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 2. CREATE TABLE
            rows = len(table_model.rows)
            cols = len(table_model.rows[0]) if rows > 0 else 0
            
            if cols == 0:
                return
                
            word_table = doc.add_table(rows=rows, cols=cols)
            
            # 3. APPLY STYLE SAFELY
            try:
                if 'Table Grid' in doc.styles:
                    word_table.style = 'Table Grid'
            except Exception as exc:
                logger.debug("Could not apply 'Table Grid' style: %s", exc)
            
            # 4. POPULATE CELLS
            for cell_model in table_model.cells:
                try:
                    r, c = cell_model.row, cell_model.col
                    if r < len(word_table.rows) and c < len(word_table.rows[r].cells):
                        word_cell = word_table.rows[r].cells[c]
                        word_cell.text = cell_model.text if cell_model.text else ""
                        
                        # RECURSIVE RENDERING: Check for nested tables
                        nested_tables = cell_model.metadata.get("nested_tables", [])
                        for nested_tbl in nested_tables:
                            try:
                                self.render(word_cell, nested_tbl)
                            except Exception as exc:
                                word_cell.add_paragraph(f"[Nested Table: {nested_tbl.table_id}]")
                                logger.warning("Nested table rendering failed for '%s': %s", nested_tbl.table_id, exc)
                except Exception as exc:
                    logger.warning("Failed to populate cell (%d,%d): %s", cell_model.row, cell_model.col, exc)
        except Exception as exc:
            logger.error("Table rendering failed for table '%s': %s", getattr(table_model, 'table_id', '?'), exc)
            raise

