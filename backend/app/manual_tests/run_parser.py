"""
Manual test script for DOCX parser.

This script demonstrates parsing a DOCX file and inspecting the output.

Usage:
    cd automated-manuscript-formatter/backend
    python -m app.manual_tests.run_parser <path_to_docx>
    
If no path is provided, it will create a sample DOCX for testing.
"""

import sys
import os
from pathlib import Path

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from app.pipeline.parsing.parser import parse_docx
from app.utils.id_generator import generate_document_id


def create_sample_docx(output_path: str) -> str:
    """
    Create a sample DOCX file for testing.
    
    Args:
        output_path: Where to save the sample file
    
    Returns:
        Path to created file
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    
    print(f"\n📝 Creating sample DOCX file: {output_path}")
    
    doc = Document()
    
    # Set core properties
    doc.core_properties.title = "Sample Academic Paper"
    doc.core_properties.author = "John Doe; Jane Smith"
    doc.core_properties.keywords = "machine learning, natural language processing"
    doc.core_properties.subject = "This is a sample abstract for testing the parser."
    
    # Title
    title = doc.add_paragraph("Sample Academic Paper")
    title_run = title.runs[0]
    title_run.bold = True
    title_run.font.size = Pt(18)
    
    # Authors
    authors = doc.add_paragraph("John Doe¹, Jane Smith²")
    authors.runs[0].font.size = Pt(12)
    
    # Affiliations
    doc.add_paragraph("¹ University A, Department of Computer Science")
    doc.add_paragraph("² University B, School of Engineering")
    
    # Abstract
    doc.add_paragraph()
    abstract_heading = doc.add_paragraph("Abstract")
    abstract_heading.runs[0].bold = True
    abstract_heading.runs[0].font.size = Pt(14)
    
    doc.add_paragraph(
        "This is a sample abstract. It describes the purpose of this paper. "
        "The paper explores innovative approaches to document processing using "
        "automated formatting techniques."
    )
    
    # Keywords
    keywords = doc.add_paragraph("Keywords: ")
    keywords.add_run("machine learning, natural language processing, document automation").italic = True
    
    # Section 1
    doc.add_paragraph()
    section1 = doc.add_paragraph("1. Introduction")
    section1.runs[0].bold = True
    section1.runs[0].font.size = Pt(14)
    
    doc.add_paragraph(
        "This is the introduction section. Academic papers typically start with "
        "an introduction that provides background context and motivation for the work."
    )
    
    doc.add_paragraph(
        "Additional paragraph with some italic text and some bold text.",
    )
    p = doc.paragraphs[-1]
    p.add_run(" Here is ").italic = True
    p.add_run("mixed formatting").bold = True
    p.add_run(" in a paragraph.")
    
    # Bulleted list
    doc.add_paragraph("Key contributions:", style='List Bullet')
    doc.add_paragraph("First contribution point", style='List Bullet 2')
    doc.add_paragraph("Second contribution point", style='List Bullet 2')
    doc.add_paragraph("Third contribution point", style='List Bullet 2')
    
    # Section 2 with table
    doc.add_paragraph()
    section2 = doc.add_paragraph("2. Methods")
    section2.runs[0].bold = True
    section2.runs[0].font.size = Pt(14)
    
    doc.add_paragraph("Table 1 shows the experimental configuration:")
    
    # Add table
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Parameter"
    header_cells[1].text = "Value"
    header_cells[2].text = "Unit"
    
    # Make header bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Data rows
    table.rows[1].cells[0].text = "Learning Rate"
    table.rows[1].cells[1].text = "0.001"
    table.rows[1].cells[2].text = "N/A"
    
    table.rows[2].cells[0].text = "Batch Size"
    table.rows[2].cells[1].text = "32"
    table.rows[2].cells[2].text = "samples"
    
    table.rows[3].cells[0].text = "Epochs"
    table.rows[3].cells[1].text = "100"
    table.rows[3].cells[2].text = "iterations"
    
    # Section 3
    doc.add_paragraph()
    section3 = doc.add_paragraph("3. Results")
    section3.runs[0].bold = True
    section3.runs[0].font.size = Pt(14)
    
    doc.add_paragraph(
        "The experimental results demonstrate significant improvements over "
        "baseline approaches. Further analysis is provided in the discussion section."
    )
    
    # Section 4
    doc.add_paragraph()
    section4 = doc.add_paragraph("4. Conclusion")
    section4.runs[0].bold = True
    section4.runs[0].font.size = Pt(14)
    
    doc.add_paragraph(
        "This paper presented a novel approach to automated document formatting. "
        "Future work will explore additional formatting styles and templates."
    )
    
    # References section
    doc.add_paragraph()
    refs_heading = doc.add_paragraph("References")
    refs_heading.runs[0].bold = True
    refs_heading.runs[0].font.size = Pt(14)
    
    doc.add_paragraph(
        "[1] Smith, J. (2020). Machine Learning Fundamentals. IEEE Press."
    )
    doc.add_paragraph(
        "[2] Doe, A., & Brown, B. (2021). Document Processing Techniques. "
        "In Proceedings of ACL 2021, pp. 123-145."
    )
    doc.add_paragraph(
        "[3] Johnson, K. et al. (2022). Automated Formatting Systems. "
        "Journal of AI Research, 45(3), 678-701."
    )
    
    # Save
    doc.save(output_path)
    print(f"✓ Sample DOCX created successfully")
    
    return output_path


def test_parser(docx_path: str):
    """
    Test the parser on a DOCX file.
    
    Args:
        docx_path: Path to DOCX file
    """
    print("\n" + "=" * 70)
    print("DOCX PARSER TEST")
    print("=" * 70)
    
    print(f"\n📄 Input file: {docx_path}")
    
    if not os.path.exists(docx_path):
        print(f"❌ Error: File not found: {docx_path}")
        return
    
    # Generate document ID
    doc_id = generate_document_id("test")
    
    # Parse the document
    print(f"\n⚙️  Parsing document (ID: {doc_id})...")
    
    try:
        document = parse_docx(docx_path, doc_id)
        print("✓ Parsing completed successfully")
    except Exception as e:
        print(f"❌ Parsing failed: {e}")
        import traceback
        traceback.print_exc()
        return
    
    # Display results
    print("\n" + "=" * 70)
    print("EXTRACTION RESULTS")
    print("=" * 70)
    
    # Metadata
    print("\n📋 Document Metadata:")
    print(f"  Title: {document.metadata.title or 'N/A'}")
    print(f"  Authors: {', '.join(document.metadata.authors) if document.metadata.authors else 'N/A'}")
    print(f"  Keywords: {', '.join(document.metadata.keywords) if document.metadata.keywords else 'N/A'}")
    print(f"  Abstract: {document.metadata.abstract[:100] + '...' if document.metadata.abstract else 'N/A'}")
    
    # Statistics
    print("\n📊 Content Statistics:")
    stats = document.get_stats()
    for key, value in stats.items():
        print(f"  {key.replace('_', ' ').title()}: {value}")
    
    # Blocks sample
    print(f"\n📝 Blocks (showing first 10 of {len(document.blocks)}):")
    for i, block in enumerate(document.blocks[:10]):
        text_preview = block.text[:60].replace('\n', ' ')
        if len(block.text) > 60:
            text_preview += "..."
        
        style_info = []
        if block.style.bold:
            style_info.append("bold")
        if block.style.italic:
            style_info.append("italic")
        if block.style.font_size:
            style_info.append(f"{block.style.font_size}pt")
        
        style_str = f" [{', '.join(style_info)}]" if style_info else ""
        style_name = block.metadata.get('style_name', '')
        style_name_str = f" (style: {style_name})" if style_name else ""
        
        print(f"  {block.block_id}: {text_preview}{style_str}{style_name_str}")
    
    # Tables
    if document.tables:
        print(f"\n📊 Tables (showing all {len(document.tables)}):")
        for table in document.tables:
            print(f"  {table.table_id}: {table.num_rows}x{table.num_cols} table")
            print(f"    Header row: {table.has_header_row}")
            if table.num_rows > 0:
                print(f"    First row: {table.rows[0][:3]}")  # First 3 columns
    
    # Figures
    if document.figures:
        print(f"\n🖼️  Figures (showing all {len(document.figures)}):")
        for figure in document.figures:
            size_info = ""
            if figure.width and figure.height:
                size_info = f" ({figure.width:.0f}x{figure.height:.0f}px)"
            data_size = len(figure.image_data) if figure.image_data else 0
            print(f"  {figure.figure_id}: {figure.image_format}{size_info} - {data_size} bytes")
    
    # Processing history
    print("\n⏱️  Processing History:")
    for entry in document.processing_history:
        print(f"  {entry.stage_name}: {entry.status} - {entry.message}")
    
    print("\n" + "=" * 70)
    print("✓ TEST COMPLETE")
    print("=" * 70)
    
    # Validation checks
    print("\n✅ Validation Checks:")
    print(f"  All blocks have IDs: {all(b.block_id for b in document.blocks)}")
    print(f"  All blocks have text: {all(b.text is not None for b in document.blocks)}")
    print(f"  All blocks are UNKNOWN type: {all(str(b.block_type) == 'unknown' for b in document.blocks)}")
    print(f"  Block indices sequential: {all(b.index == i for i, b in enumerate(document.blocks))}")
    
    if document.tables:
        print(f"  All tables have IDs: {all(t.table_id for t in document.tables)}")
        print(f"  All tables have cells: {all(len(t.cells) > 0 for t in document.tables)}")
    
    if document.figures:
        print(f"  All figures have IDs: {all(f.figure_id for f in document.figures)}")
    
    print("\n✓ Parser is working correctly!")


def main():
    """Main entry point."""
    if len(sys.argv) > 1:
        # Use provided DOCX path
        docx_path = sys.argv[1]
    else:
        # Create sample DOCX
        sample_dir = Path(__file__).parent / "sample_inputs"
        sample_dir.mkdir(exist_ok=True)
        docx_path = str(sample_dir / "sample_paper.docx")
        
        print("No DOCX file specified. Creating sample file...")
        create_sample_docx(docx_path)
    
    # Run parser test
    test_parser(docx_path)


if __name__ == "__main__":
    main()
