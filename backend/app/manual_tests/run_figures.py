"""
Manual test script for figure-caption linking.

Tests detecting captions and linking them to figures.
"""

from typing import List
import sys
import os
from pathlib import Path

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from app.models import Document, Block, Figure, BlockType
from app.pipeline.parsing.parser import parse_docx
from app.pipeline.figures.caption_matcher import link_figures
from app.utils.id_generator import generate_document_id

def create_figure_test_docx(path: str):
    """Create a DOCX with figures and captions for testing."""
    from docx import Document as DocxDocument
    from docx.shared import Inches
    
    doc = DocxDocument()
    doc.add_heading("Figure Test", 0)
    
    doc.add_paragraph("Top paragraph.")
    
    # Figure 1
    p = doc.add_paragraph()
    run = p.add_run()
    # Add a small dummy image
    # Note: python-docx requires a real image file.
    # We will skip adding real binary image content and rely on text structure
    # for this mock test if possible, OR utilize the 'parse' logic.
    # Actually, the parser needs real inline shapes to detect figures.
    # Generating a real image is complex in this script.
    
    # Alternative: We can mock the PARSED document structure manually 
    # instead of parsing a real file, to test the linking logic specifically.
    pass

def test_mock_linking():
    """Test linking logic using manually constructed Document objects."""
    print("\n" + "=" * 70)
    print("MOCK LINKING TEST")
    print("=" * 70)
    
    from app.models import Document, Block, Figure, BlockType
    
    # Create mock document
    doc = Document(document_id="test_doc", original_filename="test.docx")
    
    # Scenario: 
    # Block 0: Text
    # Block 1: [Figure 1 here]
    # Block 2: Caption 1
    # Block 3: Text
    # Block 4: [Figure 2 here]
    # Block 5: Caption 2
    
    # Create Blocks
    doc.blocks = [
        Block(block_id="b0", text="Intro text", index=0),
        Block(block_id="b1", text="", index=1), # Figure 1 holder
        Block(block_id="b2", text="Figure 1: Analysis Results", index=2, block_type=BlockType.BODY),
        Block(block_id="b3", text="Spacer text", index=3),
        Block(block_id="b4", text="", index=4), # Figure 2 holder
        Block(block_id="b5", text="Fig 2. System Architecture", index=5, block_type=BlockType.BODY),
    ]
    
    # Create Figures (with metadata from parser)
    doc.figures = [
        Figure(figure_id="f1", index=0, metadata={"block_index": 1}),
        Figure(figure_id="f2", index=1, metadata={"block_index": 4}),
    ]
    
    print(f"Input: {len(doc.figures)} figures, {len(doc.blocks)} blocks")
    
    # Run Linking
    print("Running link_figures()...")
    doc = link_figures(doc)
    
    # Assertions
    print("\nResults:")
    for fig in doc.figures:
        print(f"  {fig.figure_id} -> Caption: '{fig.caption_text}'")
        
    f1 = doc.figures[0]
    f2 = doc.figures[1]
    
    success = True
    if f1.caption_text != "Figure 1: Analysis Results":
        print("❌ Figure 1 mismatch")
        success = False
    if f2.caption_text != "Fig 2. System Architecture":
        print("❌ Figure 2 mismatch")
        success = False
        
    # Check Metadata on Caption Blocks
    b2 = doc.blocks[2]
    if not b2.metadata.get("is_figure_caption"):
        print("❌ Block 2 not marked as caption")
        success = False
        
    if success:
        print("\n✅ Mock Test Passed!")
    else:
        print("\n❌ Mock Test Failed")

def main():
    test_mock_linking()

if __name__ == "__main__":
    main()
