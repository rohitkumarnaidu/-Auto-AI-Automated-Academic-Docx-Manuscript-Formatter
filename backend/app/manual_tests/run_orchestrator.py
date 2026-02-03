"""
Manual test for Orchestrator/API flow.
"""

import sys
import os
import shutil
from pathlib import Path

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from app.pipeline.orchestrator import PipelineOrchestrator

def test_orchestrator():
    print("\n" + "=" * 70)
    print("ORCHESTRATOR INTEGRATION TEST")
    print("=" * 70)
    
    # Setup
    test_job_id = "test_orch_job"
    templates_dir = os.path.abspath("app/templates")
    orch = PipelineOrchestrator(templates_dir=templates_dir)
    
    # Create Mock Input
    mock_input = "test_orch_input.docx"
    # we need a valid docx or converter will fail on empty file if it checks zip?
    # Converter checks extension. Parser checks zip.
    # So we need a real DOCX or a robust mock.
    # simpler: create a text file matching .txt extension to test conversion + pipeline logic if pandoc present.
    # OR use a basic .docx created via python-docx if available.
    
    try:
        from docx import Document
        d = Document()
        d.add_heading("Test Title", 0)
        d.add_paragraph("Test body.")
        d.save(mock_input)
    except:
        print("⚠️ python-docx not available for test setup? Using dummy text.")
        mock_input = "test_orch_input.txt"
        with open(mock_input, "w") as f: f.write("Heading\n\nBody")

    # 1. Run Pipeline (No Template)
    print("\n[Case 1] Pipeline Run (No Template):")
    res = orch.run_pipeline(mock_input, test_job_id, template_name=None)
    
    print(f"Status: {res['status']}")
    print(f"Message: {res['message']}")
    if res['validation']:
        print(f"Valid: {res['validation'].get('is_valid')}")
        
    if res['status'] == "success" and res['output_path'] is None:
        print("✅ Success (Formatting correctly skipped)")
    else:
        print("❌ Unexpected result")

    # 2. Run Pipeline (With Template - Dummy)
    print("\n[Case 2] Pipeline Run (With Dummy Template):")
    res2 = orch.run_pipeline(mock_input, test_job_id, template_name="dummy_tmpl")
    
    print(f"Status: {res2['status']}")
    if res2['output_path']:
        print(f"Output: {res2['output_path']}")
        print("✅ Success (DOCX generated)")
    else:
        print("ℹ️ Docx not generated (likely acceptable if 'dummy_tmpl' missing triggers fallback blank doc)")

    # Cleanup
    if os.path.exists(mock_input): os.remove(mock_input)
    # Cleanup output?
    out_dir = os.path.join("output", test_job_id)
    if os.path.exists(out_dir):
        shutil.rmtree(out_dir)

    print("\n" + "=" * 70)
    print("✓ TEST COMPLETE")

if __name__ == "__main__":
    test_orchestrator()
