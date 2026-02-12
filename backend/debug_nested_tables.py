
import os
import sys
from docx import Document
from pathlib import Path

# Add backend to path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../')))

from app.pipeline.parsing.parser import DocxParser

def create_test_docx(path):
    doc = Document()
    doc.add_paragraph("Top Level Paragraph")
    
    # Add main table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Cell 0,0"
    
    # Add nested table to Cell 0,1
    cell = table.cell(0, 1)
    cell.text = "Parent Cell Text"
    # Note: in python-docx, to add table to cell:
    nested_table = cell.add_table(rows=2, cols=2)
    nested_table.cell(0, 0).text = "Nested 0,0"
    nested_table.cell(1, 1).text = "Nested 1,1"
    
    doc.save(path)
    print(f"Created test document at {path}")

def verify_extraction(path):
    parser = DocxParser()
    doc_obj = parser.parse(str(path), "test_nested")
    
    print(f"Parsed {len(doc_obj.tables)} top-level tables")
    
    found_nested = False
    for table in doc_obj.tables:
        for cell in table.cells:
            nested = cell.metadata.get("nested_tables", [])
            if nested:
                found_nested = True
                print(f"✅ Found {len(nested)} nested tables in Table {table.table_id} Cell ({cell.row}, {cell.col})")
                for nt in nested:
                    print(f"   - Nested Table ID: {nt.table_id}, Dims: {nt.num_rows}x{nt.num_cols}")
                    print(f"     Data Sample: {nt.data[0][0]}")
    
    if not found_nested:
        print("❌ FAILED: No nested tables found in metadata")
        sys.exit(1)
    else:
        print("✅ SUCCESS: Nested table extraction verified")

if __name__ == "__main__":
    test_path = Path("test_nested_gen.docx")
    try:
        create_test_docx(test_path)
        verify_extraction(test_path)
    finally:
        if test_path.exists():
            os.remove(test_path)
