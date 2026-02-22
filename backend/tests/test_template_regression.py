import pytest
import os
from pathlib import Path
from docx import Document as WordDocument
from app.models import (
    Block,
    BlockType,
    DocumentMetadata,
    PipelineDocument,
    Reference,
)
from app.pipeline.formatting.template_renderer import TemplateRenderer

def _build_test_document(title="Regression Test Title") -> PipelineDocument:
    doc = PipelineDocument(
        document_id="doc-reg-001",
        metadata=DocumentMetadata(
            title=title,
            authors=["John Doe", "Jane Smith"],
            affiliations=["University A"],
            abstract="This is a regression test abstract.",
            keywords=["testing", "regression"]
        ),
    )
    doc.blocks = [
        Block(block_id="b1", index=1, block_type=BlockType.HEADING_1, text="Introduction"),
        Block(block_id="b2", index=2, block_type=BlockType.BODY, text="Some body text with citations."),
    ]
    doc.references = [
        Reference(
            reference_id="r1",
            citation_key="Doe2024",
            title="Formatting Science",
            authors=["Doe, J."],
            year=2024,
            formatted_text="Doe, J. (2024). Formatting Science.",
            index=1
        )
    ]
    return doc

def _get_doc_paragraphs(path) -> list:
    doc = WordDocument(str(path))
    return [p.text for p in doc.paragraphs if p.text]

@pytest.mark.regression
@pytest.mark.parametrize("template", ["ieee", "apa", "nature", "springer", "numeric"])
def test_template_regression_basics(template, tmp_path):
    """
    Regression test to ensure basic document structure is rendered across major templates.
    """
    renderer = TemplateRenderer(templates_dir="app/templates")
    doc_obj = _build_test_document(f"Title for {template}")
    
    # Note: lowercase template name
    rendered = renderer.render(doc_obj, template_name=template.lower())
    out_path = tmp_path / f"reg_{template}.docx"
    rendered.save(str(out_path))
    
    paragraphs = _get_doc_paragraphs(out_path)
    full_text = "\n".join(paragraphs)
    
    # 1. Title Presence
    assert f"Title for {template}" in full_text
    
    # 2. Section Heading Presence
    assert any("Introduction" in p for p in paragraphs)
    
    # 3. Authors Presence
    assert any("John Doe" in p for p in paragraphs)
    
    # 4. References Section Presence (usually triggered by {% if references %})
    if template != "none":
        assert any("References" in p for p in paragraphs) or any("BIBLIOGRAPHY" in p.upper() for p in paragraphs)
    
    # 5. Keywords Presence
    assert "testing" in full_text

@pytest.mark.regression
def test_numeric_template_custom_rendering(tmp_path):
    """Specific check for the new 'numeric' template logic."""
    renderer = TemplateRenderer(templates_dir="app/templates")
    doc_obj = _build_test_document("Numeric Test")
    
    rendered = renderer.render(doc_obj, template_name="numeric")
    out_path = tmp_path / "reg_numeric_specific.docx"
    rendered.save(str(out_path))
    
    full_text = "\n".join(_get_doc_paragraphs(out_path))
    assert "Numeric Test" in full_text
    assert "Doe, J. (2024)" in full_text
