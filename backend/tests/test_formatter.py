"""
Tests for Formatter — ensures it produces a valid .docx artifact.
"""
from __future__ import annotations

import io
from zipfile import ZipFile
import pytest
from docx import Document as WordDocument
from docxtpl import DocxTemplate
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.models import Block, BlockType
from app.pipeline.formatting.formatter import Formatter
from app.pipeline.formatting.template_renderer import TemplateRenderer
from app.pipeline.parsing.parser import DocxParser


def _add_word_hyperlink(paragraph, text: str, url: str) -> None:
    relation_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    run_style = OxmlElement("w:rStyle")
    run_style.set(qn("w:val"), "Hyperlink")
    run_properties.append(run_style)
    run.append(run_properties)

    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


class TestFormatterOutput:

    def test_formatter_process_sets_generated_doc(self, minimal_doc):
        """Formatter.process() sets document.generated_doc."""
        formatter = Formatter(templates_dir="app/templates", contracts_dir="app/pipeline/contracts")
        result = formatter.process(minimal_doc)
        assert result.generated_doc is not None, "generated_doc must not be None after formatting"

    def test_formatter_docxtpl_path_returns_docxtemplate(self, full_doc):
        """When a template.docx exists, Formatter.format() returns DocxTemplate."""
        formatter = Formatter(templates_dir="app/templates", contracts_dir="app/pipeline/contracts")
        rendered = formatter.format(full_doc, template_name="ieee")
        # Primary path: DocxTemplate; fallback: WordDocument
        assert isinstance(rendered, (DocxTemplate, WordDocument)), (
            f"Expected DocxTemplate or WordDocument, got {type(rendered)}"
        )

    def test_formatter_none_template_falls_back_to_word_doc(self, minimal_doc):
        """template_name='none' uses blank WordDocument fallback path."""
        formatter = Formatter(templates_dir="app/templates", contracts_dir="app/pipeline/contracts")
        rendered = formatter.format(minimal_doc, template_name="none")
        # 'none' has no Jinja2 template → may be DocxTemplate (fallback template) or WordDocument
        assert rendered is not None

    def test_formatter_output_is_saveable(self, full_doc, tmp_path):
        """The generated_doc can be saved to a real .docx file."""
        formatter = Formatter(templates_dir="app/templates", contracts_dir="app/pipeline/contracts")
        formatter.process(full_doc)
        out = tmp_path / "output.docx"
        full_doc.generated_doc.save(str(out))
        assert out.exists(), "Output .docx file must be written to disk"
        assert out.stat().st_size > 0, "Output .docx must not be empty"

    def test_formatter_output_is_valid_zip(self, full_doc, tmp_path):
        """A .docx saved by the formatter is a valid ZIP/OOXML file."""
        import zipfile
        formatter = Formatter(templates_dir="app/templates", contracts_dir="app/pipeline/contracts")
        formatter.process(full_doc)
        out = tmp_path / "output.docx"
        full_doc.generated_doc.save(str(out))
        assert zipfile.is_zipfile(str(out)), ".docx must be a valid ZIP archive"

    def test_formatter_docxtpl_preserves_title_when_cover_disabled(self, full_doc, tmp_path):
        """Even without cover page, title metadata should remain visible in output."""
        formatter = Formatter(templates_dir="app/templates", contracts_dir="app/pipeline/contracts")
        full_doc.formatting_options = {"cover_page": False, "toc": False, "page_numbers": False}

        rendered = formatter.format(full_doc, template_name="ieee")
        out = tmp_path / "cover_disabled.docx"
        rendered.save(str(out))

        text = "\n".join(p.text for p in WordDocument(str(out)).paragraphs if p.text.strip())
        assert "Test Manuscript" in text

    def test_formatter_docxtpl_replaces_static_page_tokens_with_word_field(self, full_doc):
        """Template placeholders like 'Page 1' are converted to real PAGE field codes."""
        formatter = Formatter(templates_dir="app/templates", contracts_dir="app/pipeline/contracts")
        full_doc.formatting_options = {"page_numbers": True}

        rendered = formatter.format(full_doc, template_name="ieee")
        buffer = io.BytesIO()
        rendered.save(buffer)
        buffer.seek(0)

        with ZipFile(buffer, "r") as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
            footer_xml = archive.read("word/footer1.xml").decode("utf-8", errors="ignore")

        assert "Page 1" not in document_xml
        assert "PAGE" in footer_xml

    def test_formatter_legacy_applies_line_numbers_and_borders_to_all_sections(self, minimal_doc):
        """Legacy mode should stamp section-level options on every generated section."""
        formatter = Formatter(templates_dir="app/templates", contracts_dir="app/pipeline/contracts")
        minimal_doc.blocks = [
            Block(block_id="b1", index=1, block_type=BlockType.TITLE, text="Paper Title", section_name="abstract"),
            Block(block_id="b2", index=2, block_type=BlockType.BODY, text="Abstract paragraph.", section_name="abstract"),
            Block(block_id="b3", index=3, block_type=BlockType.HEADING_1, text="Introduction", section_name="introduction"),
            Block(block_id="b4", index=4, block_type=BlockType.BODY, text="Body paragraph.", section_name="introduction"),
        ]
        minimal_doc.formatting_options = {
            "template_engine": "legacy",
            "line_numbers": True,
            "borders": True,
            "page_numbers": True,
        }

        rendered = formatter.format(minimal_doc, template_name="ieee")
        assert rendered.__class__.__name__ == "Document"
        assert len(rendered.sections) >= 2

        buffer = io.BytesIO()
        rendered.save(buffer)
        buffer.seek(0)
        xml = ZipFile(buffer, "r").read("word/document.xml").decode("utf-8", errors="ignore")

        assert xml.count("<w:lnNumType") >= len(rendered.sections)
        assert xml.count("<w:pgBorders") >= len(rendered.sections)

    def test_formatter_accepts_upload_option_aliases_in_template_mode(self, full_doc):
        """Alias keys from upload APIs should map to formatting behavior."""
        formatter = Formatter(templates_dir="app/templates", contracts_dir="app/pipeline/contracts")
        full_doc.formatting_options = {
            "add_page_numbers": True,
            "add_borders": True,
            "add_line_numbers": True,
            "generate_toc": False,
            "add_cover_page": False,
        }

        rendered = formatter.format(full_doc, template_name="ieee")
        buffer = io.BytesIO()
        rendered.save(buffer)
        buffer.seek(0)

        with ZipFile(buffer, "r") as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
            footer_xml = archive.read("word/footer1.xml").decode("utf-8", errors="ignore")

        assert "PAGE" in footer_xml
        assert "<w:pgBorders" in document_xml
        assert "<w:lnNumType" in document_xml

    def test_formatter_preserves_docx_hyperlinks(self, tmp_path):
        """DOCX hyperlinks should remain true Word hyperlinks after formatting."""
        source_doc = WordDocument()
        paragraph = source_doc.add_paragraph("Visit ")
        _add_word_hyperlink(paragraph, "ScholarForm AI", "https://scholarform.ai")
        source_path = tmp_path / "hyperlinks.docx"
        source_doc.save(str(source_path))

        parsed = DocxParser().parse(str(source_path), "hyper-doc")
        for block in parsed.blocks:
            if block.text.strip():
                block.block_type = BlockType.BODY

        formatter = Formatter(templates_dir="app/templates", contracts_dir="app/pipeline/contracts")
        rendered = formatter.format(parsed, template_name="ieee")
        output_path = tmp_path / "hyperlinks_out.docx"
        rendered.save(str(output_path))

        with ZipFile(output_path, "r") as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
            relationships_xml = archive.read("word/_rels/document.xml.rels").decode("utf-8", errors="ignore")

        assert "w:hyperlink" in document_xml
        assert "https://scholarform.ai" in relationships_xml

    def test_formatter_writes_word_footnotes(self, minimal_doc, tmp_path):
        """Footnotes should be emitted as a Word footnotes part, not appended endnotes."""
        minimal_doc.blocks = [
            Block(block_id="b1", index=1, block_type=BlockType.HEADING_1, text="Introduction"),
            Block(
                block_id="b2",
                index=2,
                block_type=BlockType.BODY,
                text="Benchmark-driven formatting is safer.",
                metadata={"footnote_refs": ["1"]},
            ),
            Block(
                block_id="b3",
                index=3,
                block_type=BlockType.FOOTNOTE,
                text="Footnotes should live in word/footnotes.xml.",
                metadata={"footnote_id": "1"},
            ),
        ]

        formatter = Formatter(templates_dir="app/templates", contracts_dir="app/pipeline/contracts")
        rendered = formatter.format(minimal_doc, template_name="ieee")
        output_path = tmp_path / "footnotes_out.docx"
        rendered.save(str(output_path))

        with ZipFile(output_path, "r") as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
            footnotes_xml = archive.read("word/footnotes.xml").decode("utf-8", errors="ignore")

        assert "w:footnoteReference" in document_xml
        assert "Footnotes should live in word/footnotes.xml." in footnotes_xml
        assert "FOOTNOTES" not in document_xml
