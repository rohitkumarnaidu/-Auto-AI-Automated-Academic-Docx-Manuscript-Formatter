"""
Production Stress Test - Final Baseline Validation
Tests "none" template professional baseline format with real documents.

PRODUCTION HARDENING MODE - Verification Only
Executable with: pytest backend/tests/stress/test_production_stress.py -v
"""

import sys
import os
import time
import uuid
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from unittest.mock import patch, MagicMock

BACKEND_ROOT = Path(__file__).resolve().parents[2]
if str(BACKEND_ROOT) not in sys.path:
    sys.path.insert(0, str(BACKEND_ROOT))


def _orchestrator_available():
    """Check if the orchestrator can be instantiated with mocked dependencies."""
    try:
        with (
            patch("app.pipeline.orchestrator.get_reasoning_engine", return_value=None),
            patch("app.pipeline.orchestrator.get_rag_engine", return_value=None),
            patch("app.pipeline.orchestrator.GROBIDClient"),
            patch("app.pipeline.orchestrator.DoclingClient"),
            patch("app.pipeline.orchestrator.get_supabase_client", return_value=None),
        ):
            from app.pipeline.orchestrator import PipelineOrchestrator
            PipelineOrchestrator(
                templates_dir="app/templates",
                temp_dir="temp_stress_check"
            )
            return True
    except Exception:
        return False
    finally:
        import shutil
        if os.path.exists("temp_stress_check"):
            shutil.rmtree("temp_stress_check", ignore_errors=True)


@pytest.fixture
def orchestrator():
    """Create a mocked pipeline orchestrator for stress testing."""
    with (
        patch("app.pipeline.orchestrator.get_reasoning_engine", return_value=None),
        patch("app.pipeline.orchestrator.get_rag_engine", return_value=None),
        patch("app.pipeline.orchestrator.GROBIDClient") as mock_grobid,
        patch("app.pipeline.orchestrator.DoclingClient") as mock_docling,
        patch("app.pipeline.orchestrator.get_supabase_client", return_value=None),
    ):
        mock_grobid.return_value = MagicMock()
        mock_docling.return_value = MagicMock()

        from app.pipeline.orchestrator import PipelineOrchestrator

        orchestrator = PipelineOrchestrator(
            templates_dir="app/templates",
            temp_dir="temp_stress_test"
        )
        yield orchestrator

    import shutil
    if os.path.exists("temp_stress_test"):
        shutil.rmtree("temp_stress_test", ignore_errors=True)


@pytest.fixture
def complex_hierarchy_docx(tmp_path):
    """DOCX with 4+ heading levels."""
    doc = DocxDocument()
    doc.add_heading("Level 1 Heading", level=1)
    doc.add_paragraph("Introduction content.")
    doc.add_heading("Level 2 Heading", level=2)
    doc.add_paragraph("Section content.")
    doc.add_heading("Level 3 Heading", level=3)
    doc.add_paragraph("Subsection content.")
    doc.add_heading("Level 4 Heading", level=4)
    doc.add_paragraph("Deep nested content.")
    doc.add_heading("Another Level 2", level=2)
    doc.add_paragraph("Sibling section.")

    path = tmp_path / "complex_hierarchy.docx"
    doc.save(str(path))
    return str(path)


@pytest.fixture
def multi_media_docx(tmp_path):
    """DOCX with 3+ figures and 3+ tables."""
    doc = DocxDocument()
    doc.add_heading("Multi-Media Document", level=1)

    for i in range(4):
        doc.add_paragraph(f"Figure paragraph {i}")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = f"Table {i} Cell 0,0"
        table.cell(0, 1).text = f"Table {i} Cell 0,1"
        table.cell(1, 0).text = f"Table {i} Cell 1,0"
        table.cell(1, 1).text = f"Table {i} Cell 1,1"
        doc.add_paragraph(f"Content after table {i}")

    path = tmp_path / "multi_media.docx"
    doc.save(str(path))
    return str(path)


@pytest.fixture
def reference_heavy_docx(tmp_path):
    """DOCX with 15+ references."""
    doc = DocxDocument()
    doc.add_heading("Reference Heavy Document", level=1)
    doc.add_paragraph("Main content with citations [1]-[20].")

    doc.add_heading("References", level=1)
    for i in range(1, 21):
        doc.add_paragraph(f"[{i}] Author {i}, Title {i}, Journal {i}, {2020 + i}.")

    path = tmp_path / "reference_heavy.docx"
    doc.save(str(path))
    return str(path)


@pytest.fixture
def minimal_docx(tmp_path):
    """DOCX with title + 2 paragraphs."""
    doc = DocxDocument()
    doc.add_heading("Minimal Document", level=1)
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")

    path = tmp_path / "minimal.docx"
    doc.save(str(path))
    return str(path)


@pytest.fixture
def edge_case_docx(tmp_path):
    """DOCX with empty sections and sparse content."""
    doc = DocxDocument()
    doc.add_heading("Edge Case Document", level=1)
    doc.add_heading("Empty Section", level=2)
    doc.add_heading("Another Empty", level=2)
    doc.add_paragraph("Only paragraph in entire document.")
    doc.add_heading("Trailing Empty", level=2)

    path = tmp_path / "edge_case.docx"
    doc.save(str(path))
    return str(path)


class TestProductionStressValidation:
    """Validates professional baseline format with real documents."""

    @pytest.fixture(autouse=True)
    def skip_without_orchestrator(self):
        """Skip all stress tests if orchestrator cannot be instantiated."""
        if not _orchestrator_available():
            pytest.skip("PipelineOrchestrator unavailable (missing dependencies)")

    @pytest.mark.slow
    @pytest.mark.performance
    def test_complex_hierarchy(self, orchestrator, complex_hierarchy_docx, tmp_path):
        """Test complex heading hierarchy (4+ levels)."""
        job_id = str(uuid.uuid4())
        start = time.time()
        result = orchestrator.run_pipeline(
            input_path=complex_hierarchy_docx,
            job_id=job_id,
            template_name="none",
        )
        elapsed = time.time() - start

        assert result is not None
        assert elapsed < 60, f"Pipeline took too long: {elapsed:.2f}s"

    @pytest.mark.slow
    @pytest.mark.performance
    def test_multi_media(self, orchestrator, multi_media_docx, tmp_path):
        """Test document with multiple figures and tables."""
        job_id = str(uuid.uuid4())
        start = time.time()
        result = orchestrator.run_pipeline(
            input_path=multi_media_docx,
            job_id=job_id,
            template_name="none",
        )
        elapsed = time.time() - start

        assert result is not None
        assert elapsed < 60, f"Pipeline took too long: {elapsed:.2f}s"

    @pytest.mark.slow
    @pytest.mark.performance
    def test_reference_heavy(self, orchestrator, reference_heavy_docx, tmp_path):
        """Test document with 15+ references."""
        job_id = str(uuid.uuid4())
        start = time.time()
        result = orchestrator.run_pipeline(
            input_path=reference_heavy_docx,
            job_id=job_id,
            template_name="none",
        )
        elapsed = time.time() - start

        assert result is not None
        assert elapsed < 60, f"Pipeline took too long: {elapsed:.2f}s"

    @pytest.mark.slow
    @pytest.mark.performance
    def test_minimal(self, orchestrator, minimal_docx, tmp_path):
        """Test minimal document (title + 2 paragraphs)."""
        job_id = str(uuid.uuid4())
        start = time.time()
        result = orchestrator.run_pipeline(
            input_path=minimal_docx,
            job_id=job_id,
            template_name="none",
        )
        elapsed = time.time() - start

        assert result is not None
        assert elapsed < 30, f"Minimal doc took too long: {elapsed:.2f}s"

    @pytest.mark.slow
    @pytest.mark.performance
    def test_edge_case(self, orchestrator, edge_case_docx, tmp_path):
        """Test edge case document with empty sections."""
        job_id = str(uuid.uuid4())
        start = time.time()
        result = orchestrator.run_pipeline(
            input_path=edge_case_docx,
            job_id=job_id,
            template_name="none",
        )
        elapsed = time.time() - start

        assert result is not None
        assert elapsed < 30, f"Edge case took too long: {elapsed:.2f}s"

    @pytest.mark.slow
    @pytest.mark.performance
    def test_all_templates_with_minimal(self, orchestrator, minimal_docx, tmp_path):
        """Test minimal document against all available templates."""
        templates = [
            "none", "ieee", "apa", "mla", "chicago", "harvard",
            "vancouver", "nature", "springer", "elsevier", "acm",
            "numeric",
        ]

        results = {}
        for template in templates:
            job_id = str(uuid.uuid4())
            start = time.time()
            result = orchestrator.run_pipeline(
                input_path=minimal_docx,
                job_id=job_id,
                template_name=template,
            )
            elapsed = time.time() - start
            results[template] = {
                "status": "PASS" if result else "FAIL",
                "time": round(elapsed, 2),
            }

        all_pass = all(r["status"] == "PASS" for r in results.values())
        assert all_pass, f"Some templates failed: {results}"

    @pytest.mark.performance
    def test_validation_report_structure(self, orchestrator, minimal_docx, tmp_path):
        """Test that validation results have expected structure."""
        job_id = str(uuid.uuid4())
        result = orchestrator.run_pipeline(
            input_path=minimal_docx,
            job_id=job_id,
            template_name="none",
        )

        assert result is not None
        assert isinstance(result, dict)
