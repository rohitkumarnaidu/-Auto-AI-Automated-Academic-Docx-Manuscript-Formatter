
import os
import sys
from docx import Document
from pathlib import Path

# Add backend to path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../')))

from app.pipeline.parsing.parser import DocxParser
from app.pipeline.tables.renderer import TableRenderer

def verify_rendering():
    # 1. Create a dummy table model with nesting
    from app.models.table import Table, TableCell
    
    nested = Table(
        table_id="nested_1",
        index=1,
        block_index=10,
        num_rows=1,
        num_cols=1,
        cells=[TableCell(row=0, col=0, text="Deep Content")],
        data=[["Deep Content"]],
        rows=[["Deep Content"]]
    )
    
    parent_cell = TableCell(row=0, col=0, text="Parent Cell")
    parent_cell.metadata["nested_tables"] = [nested]
    
    parent = Table(
        table_id="parent",
        index=0,
        block_index=0,
        num_rows=1,
        num_cols=1,
        cells=[parent_cell],
        data=[["Parent Cell"]],
        rows=[["Parent Cell"]]
    )
    
    # 2. Render
    doc = Document()
    renderer = TableRenderer()
    renderer.render(doc, parent)
    
    # 3. Save and Verify
    out_path = Path("debug_rendered_nested.docx")
    doc.save(out_path)
    print(f"Rendered document saved to {out_path}")
    
    # 4. Re-parse to verify structure
    parser = DocxParser()
    reparsed = parser.parse(str(out_path), "verify_render")
    
    found_nested = False
    for t in reparsed.tables:
        for c in t.cells:
            if c.metadata.get("nested_tables"):
                found_nested = True
                print("✅ VERIFIED: Nested table survived rendering and re-parsing")
                break
    
    if not found_nested:
        print("❌ FAILED: Nested table not found in rendered document")
        sys.exit(1)

if __name__ == "__main__":
    try:
        verify_rendering()
    finally:
        p = Path("debug_rendered_nested.docx")
        if p.exists():
            os.remove(p)
