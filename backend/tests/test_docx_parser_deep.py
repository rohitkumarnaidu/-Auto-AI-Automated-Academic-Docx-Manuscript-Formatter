"""Deep unit tests for DocxParser — covers every internal method and edge case."""

from __future__ import annotations

from datetime import datetime, timezone
from typing import Any, Dict, List, Optional
from unittest.mock import MagicMock, PropertyMock, patch, call
import os
import io
import uuid

import pytest
from docx import Document as DocxDocument
from docx.shared import Inches, Pt

from app.pipeline.parsing.parser import (
    DocxParser,
    parse_docx,
    generate_figure_id,
    generate_table_id,
    generate_equation_id,
)
from lxml import etree

from app.models import (
    Block,
    BlockType,
    TextStyle,
    ImageFormat,
    Figure,
    Equation,
    DocumentMetadata,
    PipelineDocument as Document,
)
from app.pipeline.parsing.base_parser import BaseParser


class TestHelperFunctions:
    def test_generate_figure_id(self):
        assert generate_figure_id(0) == "fig_000"
        assert generate_figure_id(1) == "fig_001"
        assert generate_figure_id(100) == "fig_100"

    def test_generate_table_id(self):
        assert generate_table_id(0) == "tbl_000"
        assert generate_table_id(1) == "tbl_001"

    def test_generate_equation_id(self):
        assert generate_equation_id(0) == "eqn_000"
        assert generate_equation_id(42) == "eqn_042"

    def test_parse_docx_convenience(self, tmp_path):
        doc = DocxDocument()
        doc.add_paragraph("Hello")
        path = tmp_path / "c.docx"
        doc.save(str(path))
        result = parse_docx(str(path), "test-conv")
        assert result is not None
        assert len(result.blocks) > 0

    def test_parse_docx_not_found(self):
        with pytest.raises(FileNotFoundError):
            parse_docx("/nonexistent/path.docx", "test")


class TestSupportsFormat:
    @pytest.fixture
    def parser(self):
        return DocxParser()

    def test_supports_docx(self, parser):
        assert parser.supports_format(".docx") is True

    def test_supports_doc_upper(self, parser):
        assert parser.supports_format(".DOC") is True

    def test_supports_docx_upper(self, parser):
        assert parser.supports_format(".DOCX") is True

    def test_supports_doc_lower(self, parser):
        assert parser.supports_format(".doc") is True

    def test_rejects_pdf(self, parser):
        assert parser.supports_format(".pdf") is False

    def test_rejects_txt(self, parser):
        assert parser.supports_format(".txt") is False

    def test_rejects_empty(self, parser):
        assert parser.supports_format("") is False

    def test_rejects_none_extension(self, parser):
        assert parser.supports_format("makefile") is False


class TestParseEdgeCases:
    @pytest.fixture
    def parser(self):
        return DocxParser()

    @patch("app.pipeline.parsing.parser.DocxDocument")
    def test_parse_with_uuid_document_id(self, mock_docx_cls, parser, tmp_path):
        mock_docx = MagicMock()
        mock_docx.element.body = []
        mock_docx_cls.return_value = mock_docx
        mock_docx.sections = []
        path = tmp_path / "d.docx"
        path.write_text("")
        uid = uuid.UUID("12345678-1234-5678-1234-567812345678")
        result = parser.parse(str(path), uid)
        assert isinstance(result.document_id, str)

    @patch("app.pipeline.parsing.parser.DocxDocument")
    def test_parse_invalid_docx_raises_value_error(self, mock_docx_cls, parser, tmp_path):
        mock_docx_cls.side_effect = RuntimeError("bad zip")
        path = tmp_path / "bad.docx"
        path.write_text("garbage")
        with pytest.raises(ValueError, match="Failed to open DOCX file"):
            parser.parse(str(path), "test-id")

    def test_parse_file_not_found(self, parser):
        with pytest.raises(FileNotFoundError, match="DOCX file not found"):
            parser.parse("/no/such/file.docx", "x")

    @patch("app.pipeline.parsing.parser.DocxDocument")
    def test_parse_empty_body(self, mock_docx_cls, parser, tmp_path):
        mock_docx = MagicMock()
        mock_docx.element.body = []
        mock_docx.sections = []
        mock_docx_cls.return_value = mock_docx
        path = tmp_path / "empty.docx"
        path.write_text("")
        result = parser.parse(str(path), "empty")
        assert result.blocks == []
        assert result.figures == []
        assert result.tables == []

    @patch("app.pipeline.parsing.parser.DocxDocument")
    def test_parse_processing_stage_added(self, mock_docx_cls, parser, tmp_path):
        mock_docx = MagicMock()
        mock_docx.element.body = []
        mock_docx.sections = []
        mock_docx_cls.return_value = mock_docx
        path = tmp_path / "p.docx"
        path.write_text("")
        result = parser.parse(str(path), "proc")
        assert len(result.processing_history) > 0
        stage = result.processing_history[0]
        assert stage.stage_name == "parsing"
        assert stage.status == "success"


class TestCoreProperties:
    @pytest.fixture
    def parser(self):
        return DocxParser()

    def _make_mock_docx(self, **kwargs):
        docx = MagicMock()
        props = MagicMock()
        for k, v in kwargs.items():
            setattr(props, k, v)
        docx.core_properties = props
        return docx

    def test_empty_metadata(self, parser):
        docx = self._make_mock_docx(title=None, author=None, subject=None, keywords=None, created=None)
        meta = parser._extract_core_properties(docx)
        assert meta.title is None or meta.title == ""
        assert meta.abstract is None or meta.abstract == ""
        assert meta.keywords is None or meta.keywords == []

    def test_title_extracted(self, parser):
        docx = self._make_mock_docx(title="My Paper Title")
        meta = parser._extract_core_properties(docx)
        assert meta.title == "My Paper Title"

    def test_authors_semicolon_separated(self, parser):
        docx = self._make_mock_docx(author="Alice; Bob; Charlie")
        meta = parser._extract_core_properties(docx)
        assert meta.authors == ["Alice", "Bob", "Charlie"]

    def test_authors_single(self, parser):
        docx = self._make_mock_docx(author="Alice")
        meta = parser._extract_core_properties(docx)
        assert meta.authors == ["Alice"]

    def test_subject_as_abstract(self, parser):
        docx = self._make_mock_docx(subject="This is the abstract")
        meta = parser._extract_core_properties(docx)
        assert meta.abstract == "This is the abstract"

    def test_keywords_comma_separated(self, parser):
        docx = self._make_mock_docx(keywords="machine learning, nlp, transformers")
        meta = parser._extract_core_properties(docx)
        assert "machine learning" in meta.keywords
        assert "nlp" in meta.keywords
        assert "transformers" in meta.keywords

    def test_keywords_semicolon_separated(self, parser):
        docx = self._make_mock_docx(keywords="ml; ai; nlp")
        meta = parser._extract_core_properties(docx)
        assert "ml" in meta.keywords
        assert "ai" in meta.keywords

    def test_keywords_mixed_separators(self, parser):
        docx = self._make_mock_docx(keywords="a,b;c,d")
        meta = parser._extract_core_properties(docx)
        assert "a" in meta.keywords
        assert "b" in meta.keywords
        assert "c" in meta.keywords
        assert "d" in meta.keywords

    def test_created_date(self, parser):
        from datetime import datetime
        dt = datetime(2024, 6, 15)
        docx = self._make_mock_docx(created=dt)
        meta = parser._extract_core_properties(docx)
        assert meta.publication_date == dt

    def test_no_keywords_returns_empty(self, parser):
        docx = self._make_mock_docx(keywords="", title="T")
        meta = parser._extract_core_properties(docx)
        assert meta.keywords is None or meta.keywords == []


class TestParagraphStyleExtraction:
    @pytest.fixture
    def parser(self):
        return DocxParser()

    def _make_paragraph(self, runs_data: List[Dict], style_name=None, no_runs=False):
        """Helper to create a mocked paragraph."""
        para = MagicMock(spec=["runs", "style", "_element", "text", "alignment"])
        style = MagicMock()
        style.name = style_name
        font = MagicMock()
        font.name = None
        font.size = None
        font.bold = None
        font.italic = None
        style.font = font
        para.style = style

        if no_runs:
            para.runs = []
        else:
            runs = []
            for rd in runs_data:
                run = MagicMock()
                run.text = rd.get("text", "")
                run.bold = rd.get("bold")
                run.italic = rd.get("italic")
                run.underline = rd.get("underline")
                run_font = MagicMock()
                run_font.name = rd.get("font_name")
                run_font.size = rd.get("font_size")
                run.font = run_font
                runs.append(run)
            para.runs = runs

        return para

    def test_default_style(self, parser):
        para = self._make_paragraph([{"text": ""}])
        style = parser._extract_paragraph_style(para)
        assert style.bold is False
        assert style.italic is False
        assert style.underline is False
        assert style.font_name is None
        assert style.font_size is None

    def test_bold_italic(self, parser):
        para = self._make_paragraph([{"text": "hello", "bold": True, "italic": True}])
        style = parser._extract_paragraph_style(para)
        assert style.bold is True
        assert style.italic is True

    def test_underline(self, parser):
        para = self._make_paragraph([{"text": "hello", "underline": True}])
        style = parser._extract_paragraph_style(para)
        assert style.underline is True

    def test_font_name_size(self, parser):
        para = self._make_paragraph([{"text": "hello", "font_name": "Arial", "font_size": Pt(12)}])
        style = parser._extract_paragraph_style(para)
        assert style.font_name == "Arial"
        assert style.font_size == 12.0

    def test_skips_empty_runs(self, parser):
        para = self._make_paragraph([
            {"text": "  ", "bold": None, "italic": None},
            {"text": "actual", "bold": True},
        ])
        style = parser._extract_paragraph_style(para)
        assert style.bold is True

    def test_paragraph_style_fallback(self, parser):
        para = self._make_paragraph([], no_runs=True)
        para.style.font.bold = True
        para.style.font.italic = True
        para.style.font.name = "Times New Roman"
        para.style.font.size = Pt(10)
        style = parser._extract_paragraph_style(para)
        assert style.bold is True
        assert style.italic is True
        assert style.font_name == "Times New Roman"
        assert style.font_size == 10.0

    def test_paragraph_style_fallback_attribute_error(self, parser):
        para = self._make_paragraph([], no_runs=True)
        style_mock = MagicMock(spec=["font"])
        font_mock = MagicMock(spec=["bold", "italic", "name", "size"])
        font_mock.bold = True
        font_mock.italic = True
        font_mock.name = "Arial"
        font_mock.size = Pt(11)
        style_mock.font = font_mock
        para.style = style_mock
        style = parser._extract_paragraph_style(para)
        assert style.font_size == 11.0

    def test_paragraph_style_font_size_none(self, parser):
        para = self._make_paragraph([{"text": "x", "font_size": None}])
        style = parser._extract_paragraph_style(para)
        assert style.font_size is None


class TestHyperlinkExtraction:
    @pytest.fixture
    def parser(self):
        return DocxParser()

    def _make_paragraph_with_hyperlinks(self, links_data, raise_on_findall=False):
        para = MagicMock()
        hyperlink_elements = []

        for ld in links_data:
            hl = MagicMock()
            hl.get.return_value = ld.get("r_id")
            hl.findall.return_value = []
            if ld.get("text_chunks"):
                run_elem = MagicMock()
                tn = MagicMock()
                tn.text = ld["text_chunks"]
                run_elem.findall.return_value = [tn]
                hl.findall.return_value = [run_elem]
            hyperlink_elements.append(hl)

        element = MagicMock()
        if raise_on_findall:
            element.findall.side_effect = Exception("XML parse error")
        else:
            element.findall.return_value = hyperlink_elements
        para._element = element

        rels = {}
        for ld in links_data:
            if ld.get("r_id") and ld.get("url"):
                rel = MagicMock()
                rel.target_ref = ld["url"]
                rels[ld["r_id"]] = rel
        para.part.rels = rels

        return para

    def test_no_hyperlinks(self, parser):
        para = MagicMock()
        para._element.findall.return_value = []
        result = parser._extract_hyperlinks(para)
        assert result == []

    def test_single_hyperlink(self, parser):
        para = self._make_paragraph_with_hyperlinks([
            {"r_id": "rId1", "url": "https://example.com", "text_chunks": "Example"},
        ])
        result = parser._extract_hyperlinks(para)
        assert len(result) == 1
        assert result[0]["text"] == "Example"
        assert result[0]["url"] == "https://example.com"

    def test_multiple_hyperlinks(self, parser):
        para = self._make_paragraph_with_hyperlinks([
            {"r_id": "rId1", "url": "https://a.com", "text_chunks": "Link A"},
            {"r_id": "rId2", "url": "https://b.com", "text_chunks": "Link B"},
        ])
        result = parser._extract_hyperlinks(para)
        assert len(result) == 2

    def test_hyperlink_no_text_skipped(self, parser):
        para = self._make_paragraph_with_hyperlinks([
            {"r_id": "rId1", "url": "https://a.com", "text_chunks": ""},
        ])
        result = parser._extract_hyperlinks(para)
        assert result == []

    def test_hyperlink_missing_rid_skipped(self, parser):
        para = self._make_paragraph_with_hyperlinks([
            {"r_id": None, "url": "https://a.com", "text_chunks": "Link"},
        ])
        result = parser._extract_hyperlinks(para)
        assert result == []

    def test_hyperlink_bad_relation_skipped(self, parser):
        para = self._make_paragraph_with_hyperlinks([
            {"r_id": "rId99", "url": None, "text_chunks": "Link"},
        ])
        result = parser._extract_hyperlinks(para)
        assert result == []

    def test_hyperlink_extraction_failure_swallowed(self, parser):
        para = self._make_paragraph_with_hyperlinks([], raise_on_findall=True)
        result = parser._extract_hyperlinks(para)
        assert result == []


class TestNoteReferences:
    @pytest.fixture
    def parser(self):
        return DocxParser()

    def _make_paragraph(self, ref_ids: List[str], raise_error=False):
        para = MagicMock()
        element = MagicMock()
        if raise_error:
            element.findall.side_effect = Exception("XML error")
        else:
            ref_elements = []
            for rid in ref_ids:
                ref = MagicMock()
                ref.get.return_value = rid
                ref_elements.append(ref)
            element.findall.return_value = ref_elements
        para._element = element
        return para

    def test_no_refs(self, parser):
        para = self._make_paragraph([])
        result = parser._extract_note_references(para, "w:footnoteReference")
        assert result == []

    def test_footnote_refs(self, parser):
        para = self._make_paragraph(["1", "2", "3"])
        result = parser._extract_note_references(para, "w:footnoteReference")
        assert result == ["1", "2", "3"]

    def test_endnote_refs(self, parser):
        para = self._make_paragraph(["4", "5"])
        result = parser._extract_note_references(para, "w:endnoteReference")
        assert result == ["4", "5"]

    def test_error_handling(self, parser):
        para = self._make_paragraph([], raise_error=True)
        result = parser._extract_note_references(para, "w:footnoteReference")
        assert result == []


class TestListInfo:
    @pytest.fixture
    def parser(self):
        return DocxParser()

    def _make_paragraph(self, numpr_data=None, style_val=None, ppr_none=False):
        para = MagicMock()
        element = MagicMock()

        if ppr_none:
            element.find.return_value = None
            para._element = element
            return para

        ppr = MagicMock()

        def ppr_find(path):
            from docx.oxml.ns import qn
            if "numPr" in path:
                if numpr_data is not None:
                    numpr = MagicMock()
                    def numpr_find(sub):
                        if "ilvl" in sub:
                            if "ilvl" in numpr_data:
                                ilvl = MagicMock()
                                ilvl.get.return_value = numpr_data["ilvl"]
                                return ilvl
                            return None
                        if "numId" in sub:
                            if "numId" in numpr_data:
                                numId = MagicMock()
                                numId.get.return_value = numpr_data["numId"]
                                return numId
                            return None
                        return None
                    numpr.find.side_effect = numpr_find
                    return numpr
                return None
            if "pStyle" in path:
                if style_val:
                    pstyle = MagicMock()
                    pstyle.get.return_value = style_val
                    return pstyle
                return None
            return None

        ppr.find.side_effect = ppr_find

        element.find.return_value = ppr
        para._element = element
        return para

    def test_no_list(self, parser):
        para = self._make_paragraph(ppr_none=True)
        result = parser._get_list_info(para)
        assert result is None

    def test_numbered_list(self, parser):
        para = self._make_paragraph(numpr_data={"ilvl": "0", "numId": "1"})
        result = parser._get_list_info(para)
        assert result is not None
        assert result["is_list_item"] is True
        assert result["list_level"] == 0
        assert result["list_id"] == "1"

    def test_nested_list(self, parser):
        para = self._make_paragraph(numpr_data={"ilvl": "2", "numId": "1"})
        result = parser._get_list_info(para)
        assert result["list_level"] == 2

    def test_no_ilvl_defaults_to_zero(self, parser):
        para = self._make_paragraph(numpr_data={"numId": "5"})
        result = parser._get_list_info(para)
        assert result["list_level"] == 0

    def test_style_fallback_list(self, parser):
        para = self._make_paragraph(numpr_data=None, style_val="ListBullet")
        result = parser._get_list_info(para)
        assert result is not None
        assert result["is_list_item"] is True
        assert result["list_id"] == "style_listbullet"

    def test_style_fallback_list_number_parsed(self, parser):
        para = self._make_paragraph(numpr_data=None, style_val="ListNumber2")
        result = parser._get_list_info(para)
        assert result["list_level"] == 1

    def test_style_fallback_not_list(self, parser):
        para = self._make_paragraph(numpr_data=None, style_val="Normal")
        result = parser._get_list_info(para)
        assert result is None

    def test_error_handling(self, parser):
        para = MagicMock()
        element = MagicMock()
        element.find.side_effect = Exception("error")
        para._element = element
        result = parser._get_list_info(para)
        assert result is None


class TestImageFormat:
    @pytest.fixture
    def parser(self):
        return DocxParser()

    def test_png(self, parser):
        assert parser._get_image_format("image/png") == ImageFormat.PNG

    def test_jpeg(self, parser):
        assert parser._get_image_format("image/jpeg") == ImageFormat.JPEG

    def test_jpg(self, parser):
        assert parser._get_image_format("image/jpg") == ImageFormat.JPG

    def test_gif(self, parser):
        assert parser._get_image_format("image/gif") == ImageFormat.GIF

    def test_bmp(self, parser):
        assert parser._get_image_format("image/bmp") == ImageFormat.BMP

    def test_tiff(self, parser):
        assert parser._get_image_format("image/tiff") == ImageFormat.TIFF

    def test_svg(self, parser):
        assert parser._get_image_format("image/svg+xml") == ImageFormat.SVG

    def test_emf(self, parser):
        assert parser._get_image_format("image/x-emf") == ImageFormat.EMF

    def test_wmf(self, parser):
        assert parser._get_image_format("image/x-wmf") == ImageFormat.WMF

    def test_unknown_format(self, parser):
        assert parser._get_image_format("image/webp") == ImageFormat.UNKNOWN

    def test_empty_content_type(self, parser):
        assert parser._get_image_format("") == ImageFormat.UNKNOWN


class TestInlineImageExtraction:
    @pytest.fixture
    def parser(self):
        return DocxParser()

    def _make_run(self, has_inline=False, has_anchor=False, inline_element=None,
                  anchor_element=None, part=None):
        run = MagicMock()
        element = MagicMock()

        if has_inline:
            element.findall.return_value = [inline_element or MagicMock()]
        elif has_anchor:
            def findall_side_effect(ns):
                if "inline" in ns:
                    return []
                if "anchor" in ns:
                    return [anchor_element or MagicMock()]
                return []
            element.findall.side_effect = findall_side_effect
        else:
            element.findall.return_value = []

        run._element = element

        if part:
            try:
                run.part = part
            except AttributeError:
                pass

        return run

    def test_no_images_in_run(self, parser):
        run = self._make_run()
        para = MagicMock()
        para.runs = [run]
        figures = parser._extract_inline_images(para)
        assert figures == []

    def test_run_no_hasattr_element(self, parser):
        run = MagicMock()
        del run._element
        para = MagicMock()
        para.runs = [run]
        figures = parser._extract_inline_images(para)
        assert figures == []

    def test_part_getattr_fallback(self, parser):
        run = MagicMock()
        element = MagicMock()
        inline = MagicMock()
        element.findall.return_value = [inline]
        run._element = element
        run.part = MagicMock()
        para = MagicMock()
        para.runs = [run]
        with patch.object(parser, "_extract_image_from_inline", return_value=None):
            figures = parser._extract_inline_images(para)
            assert figures == []

    def test_image_extraction_failure_swallowed(self, parser):
        run = MagicMock()
        run._element.findall.side_effect = Exception("bad image")
        para = MagicMock()
        para.runs = [run]
        figures = parser._extract_inline_images(para)
        assert figures == []

    def test_anchor_shapes_also_extracted(self, parser):
        run = MagicMock()
        element = MagicMock()

        def findall_side_effect(ns):
            if "inline" in ns:
                return []
            if "anchor" in ns:
                return [MagicMock()]
            return []
        element.findall.side_effect = findall_side_effect
        run._element = element
        run.part = MagicMock()

        para = MagicMock()
        para.runs = [run]

        with patch.object(parser, "_extract_image_from_inline", return_value=MagicMock(spec=Figure)):
            figures = parser._extract_inline_images(para)
            assert len(figures) > 0


class TestImageFromInline:
    @pytest.fixture
    def parser(self):
        return DocxParser()

    def test_no_blip_element(self, parser):
        inline = MagicMock()
        inline.find.return_value = None
        result = parser._extract_image_from_inline(inline, MagicMock())
        assert result is None

    def test_no_embed_id(self, parser):
        inline = MagicMock()
        blip = MagicMock()
        blip.get.return_value = None
        inline.find.return_value = blip
        result = parser._extract_image_from_inline(inline, MagicMock())
        assert result is None

    def test_successful_extraction(self, parser):
        inline = MagicMock()
        blip = MagicMock()
        blip.get.return_value = "rId1"
        inline.find.return_value = blip

        image_part = MagicMock()
        image_part.blob = b"fakeimagedata"
        image_part.content_type = "image/png"

        part = MagicMock()
        part.related_parts = {"rId1": image_part}

        extent = MagicMock()
        extent.get.side_effect = lambda k: {"cx": "100000", "cy": "50000"}.get(k)
        inline.find.return_value = extent

        def find_side_effect(path):
            if "blip" in path:
                return blip
            if "extent" in path:
                return extent
            return None
        inline.find.side_effect = find_side_effect

        result = parser._extract_image_from_inline(inline, part)
        assert result is not None
        assert isinstance(result, Figure)
        assert result.image_format == ImageFormat.PNG
        assert result.image_data == b"fakeimagedata"

    def test_extraction_exception(self, parser):
        inline = MagicMock()
        inline.find.side_effect = Exception("fail")
        result = parser._extract_image_from_inline(inline, MagicMock())
        assert result is None


class TestExtractParagraph:
    @pytest.fixture
    def parser(self):
        return DocxParser()

    def _make_paragraph(self, text="", style_name=None, alignment=None):
        para = MagicMock(spec=["runs", "style", "_element", "text", "alignment", "part"])
        para.text = text
        para.alignment = alignment
        style = MagicMock()
        style.name = style_name if style_name else None
        font = MagicMock()
        font.name = None
        font.size = None
        font.bold = None
        font.italic = None
        style.font = font
        para.style = style

        element = MagicMock()
        element.find.return_value = None
        element.findall.return_value = []
        para._element = element
        para.runs = []

        return para

    def test_empty_paragraph_creates_block(self, parser):
        para = self._make_paragraph(text="")
        block = parser._extract_paragraph(para)
        assert block is not None
        assert block.text == ""

    def test_paragraph_with_style_name(self, parser):
        para = self._make_paragraph(text="Hello", style_name="Heading 1")
        block = parser._extract_paragraph(para)
        assert block.metadata.get("style_name") == "Heading 1"

    def test_paragraph_with_alignment(self, parser):
        para = self._make_paragraph(text="Center", alignment=1)
        block = parser._extract_paragraph(para)
        assert "alignment" in block.metadata

    def test_paragraph_no_style(self, parser):
        para = self._make_paragraph(text="No style")
        para.style.name = None
        block = parser._extract_paragraph(para)
        assert "style_name" not in block.metadata or block.metadata.get("style_name") is None

    def test_paragraph_no_alignment(self, parser):
        para = self._make_paragraph(text="No align", alignment=None)
        block = parser._extract_paragraph(para)
        assert "alignment" not in block.metadata


class TestHeadersAndFooters:
    @pytest.fixture
    def parser(self):
        return DocxParser()

    def _make_block_mock(self):
        block = MagicMock()
        block.metadata = {}
        return block

    def _make_section(self, has_header=True, has_footer=True,
                      header_texts=None, footer_texts=None):
        section = MagicMock()

        if has_header:
            header = MagicMock()
            h_paragraphs = []
            for t in (header_texts or []):
                p = MagicMock()
                p.text = t
                p.alignment = None
                h_paragraphs.append(p)
            header.paragraphs = h_paragraphs
            section.header = header
        else:
            section.header = None

        if has_footer:
            footer = MagicMock()
            f_paragraphs = []
            for t in (footer_texts or []):
                p = MagicMock()
                p.text = t
                p.alignment = None
                f_paragraphs.append(p)
            footer.paragraphs = f_paragraphs
            section.footer = footer
        else:
            section.footer = None

        return section

    def test_no_sections(self, parser):
        docx = MagicMock()
        docx.sections = []
        result = parser._extract_headers_and_footers(docx)
        assert result == []

    def test_header_extracted(self, parser):
        parser._extract_paragraph = MagicMock(return_value=self._make_block_mock())
        docx = MagicMock()
        docx.sections = [self._make_section(has_header=True, header_texts=["Page Title"])]
        result = parser._extract_headers_and_footers(docx)
        assert len(result) > 0

    def test_footer_extracted(self, parser):
        parser._extract_paragraph = MagicMock(return_value=self._make_block_mock())
        docx = MagicMock()
        docx.sections = [self._make_section(has_footer=True, header_texts=[], footer_texts=["Page 1"])]
        result = parser._extract_headers_and_footers(docx)
        assert len(result) > 0

    def test_header_footer_metadata_tags(self, parser):
        parser._extract_paragraph = MagicMock(return_value=self._make_block_mock())
        docx = MagicMock()
        docx.sections = [self._make_section(header_texts=["H"], footer_texts=["F"])]
        result = parser._extract_headers_and_footers(docx)
        header_blocks = [b for b in result if b.metadata.get("is_header")]
        footer_blocks = [b for b in result if b.metadata.get("is_footer")]
        assert len(header_blocks) > 0
        assert len(footer_blocks) > 0

    def test_exception_swallowed(self, parser):
        docx = MagicMock()
        docx.sections.side_effect = Exception("No sections")
        result = parser._extract_headers_and_footers(docx)
        assert result == []

    def test_no_header_no_footer(self, parser):
        parser._extract_paragraph = MagicMock(return_value=None)
        docx = MagicMock()
        section = MagicMock()
        section.header = None
        section.footer = None
        docx.sections = [section]
        result = parser._extract_headers_and_footers(docx)
        assert result == []


class TestFootnotesAndEndnotes:
    @pytest.fixture
    def parser(self):
        return DocxParser()

    def _make_docx_with_notes(self, footnote_texts=None, endnote_texts=None):
        docx = MagicMock()

        if footnote_texts:
            fn_part = MagicMock()
            fn_root = MagicMock()
            fn_elements = []
            for i, ft in enumerate(footnote_texts):
                fn_elem = MagicMock()
                fn_elem.get.return_value = str(i + 1)
                p_elem = MagicMock()
                r_elem = MagicMock()
                t_elem = MagicMock()
                t_elem.text = ft
                r_elem.find.return_value = t_elem
                p_elem.findall.return_value = [r_elem]
                fn_elem.findall.return_value = [p_elem]
                fn_elements.append(fn_elem)
            fn_root.findall.return_value = fn_elements
            fn_part.element = fn_root
            docx.part.footnotes_part = fn_part
        else:
            docx.part.footnotes_part = MagicMock()
            docx.part.footnotes_part.element.findall.return_value = []

        if endnote_texts:
            en_part = MagicMock()
            en_root = MagicMock()
            en_elements = []
            for i, et in enumerate(endnote_texts):
                en_elem = MagicMock()
                en_elem.get.return_value = str(i + 1)
                p_elem = MagicMock()
                r_elem = MagicMock()
                t_elem = MagicMock()
                t_elem.text = et
                r_elem.find.return_value = t_elem
                p_elem.findall.return_value = [r_elem]
                en_elem.findall.return_value = [p_elem]
                en_elements.append(en_elem)
            en_root.findall.return_value = en_elements
            en_part.element = en_root
            docx.part.endnotes_part = en_part
        else:
            docx.part.endnotes_part = MagicMock()
            docx.part.endnotes_part.element.findall.return_value = []

        return docx

    def test_no_notes(self, parser):
        docx = MagicMock()
        docx.part.footnotes_part = None
        docx.part.endnotes_part = None
        result = parser._extract_footnotes_and_endnotes(docx)
        assert result == []

    def test_footnotes_extracted(self, parser):
        docx = self._make_docx_with_notes(footnote_texts=["First footnote", "Second footnote"])
        result = parser._extract_footnotes_and_endnotes(docx)
        footnotes = [b for b in result if b.metadata.get("is_footnote")]
        assert len(footnotes) > 0

    def test_endnotes_extracted(self, parser):
        docx = self._make_docx_with_notes(endnote_texts=["An endnote"])
        result = parser._extract_footnotes_and_endnotes(docx)
        endnotes = [b for b in result if b.metadata.get("is_endnote")]
        assert len(endnotes) > 0

    def test_footnote_metadata(self, parser):
        docx = self._make_docx_with_notes(footnote_texts=["Note text"])
        result = parser._extract_footnotes_and_endnotes(docx)
        if result:
            b = result[0]
            assert b.metadata.get("is_footnote") is True
            assert b.metadata.get("footnote_id") is not None

    def test_endnote_metadata(self, parser):
        docx = self._make_docx_with_notes(endnote_texts=["EN text"])
        result = parser._extract_footnotes_and_endnotes(docx)
        if result:
            b = result[0]
            assert b.metadata.get("is_endnote") is True
            assert b.metadata.get("endnote_id") is not None

    def test_empty_text_skipped(self, parser):
        docx = self._make_docx_with_notes(footnote_texts=[""])
        result = parser._extract_footnotes_and_endnotes(docx)
        footnotes = [b for b in result if b.metadata.get("is_footnote")]
        assert len(footnotes) == 0

    def test_exception_swallowed_footnotes(self, parser):
        docx = MagicMock()
        docx.part.footnotes_part.element.findall.side_effect = Exception("fail")
        docx.part.endnotes_part.element.findall.side_effect = Exception("fail")
        result = parser._extract_footnotes_and_endnotes(docx)
        assert result == []

    def test_exception_swallowed_no_part(self, parser):
        docx = MagicMock()
        docx.part.footnotes_part = None
        docx.part.endnotes_part = None
        result = parser._extract_footnotes_and_endnotes(docx)
        assert result == []


class TestEquationExtraction:
    @pytest.fixture
    def parser(self):
        return DocxParser()

    def _make_paragraph_with_math(self, omath_paras=None, omath_direct=None, omath_inside_para=None):
        para = MagicMock()
        element = MagicMock()

        ns_math = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"
        nested_path = f".//{ns_math}oMathPara/{ns_math}oMath"

        def findall_side_effect(path):
            if path == nested_path:
                return omath_inside_para or []
            if path.endswith("}oMathPara"):
                return omath_paras or []
            if path.endswith("}oMath"):
                return omath_direct or []
            return []

        element.findall.side_effect = findall_side_effect
        para._element = element
        return para

    def test_no_equations(self, parser):
        para = self._make_paragraph_with_math()
        result = parser._extract_equations(para)
        assert result == []

    def test_block_equation(self, parser):
        om = MagicMock()
        om_para = MagicMock()
        om_para.findall.return_value = [om]
        para = self._make_paragraph_with_math(omath_paras=[om_para])

        with patch.object(parser, "_extract_math_element", return_value=MagicMock(spec=Equation)):
            result = parser._extract_equations(para)
            assert len(result) == 1

    def test_inline_equation_skips_already_processed(self, parser):
        om = MagicMock()
        om_para = MagicMock()
        om_para.findall.return_value = [om]

        ns_math = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"
        nested_path = f".//{ns_math}oMathPara/{ns_math}oMath"
        para = MagicMock()
        element = MagicMock()
        call_log = []

        def findall_fn(path):
            call_log.append(path)
            if path == nested_path:
                return [om]
            if path.endswith("}oMathPara"):
                return [om_para]
            if path.endswith("}oMath"):
                return [om]
            return []

        element.findall.side_effect = findall_fn
        para._element = element

        with patch.object(parser, "_extract_math_element", return_value=MagicMock(spec=Equation)) as mock_extract:
            result = parser._extract_equations(para)
            assert len(result) == 1, f"Expected 1, got {len(result)}. findall calls: {call_log}"

    def test_inline_equation_direct(self, parser):
        om1 = MagicMock()
        om2 = MagicMock()
        para = self._make_paragraph_with_math(omath_direct=[om1, om2])

        with patch.object(parser, "_extract_math_element", return_value=MagicMock(spec=Equation)):
            result = parser._extract_equations(para)
            assert len(result) == 2

    def test_inline_equation_direct(self, parser):
        om1 = MagicMock()
        om2 = MagicMock()
        para = self._make_paragraph_with_math(omath_direct=[om1, om2])

        with patch.object(parser, "_extract_math_element", return_value=MagicMock(spec=Equation)):
            result = parser._extract_equations(para)
            assert len(result) == 2

    def test_no_equations(self, parser):
        para = self._make_paragraph_with_math()
        result = parser._extract_equations(para)
        assert result == []

    def test_block_equation(self, parser):
        om = MagicMock()
        om_para = MagicMock()
        om_para.findall.return_value = [om]
        para = self._make_paragraph_with_math(omath_paras=[om_para])

        with patch.object(parser, "_extract_math_element", return_value=MagicMock(spec=Equation)):
            result = parser._extract_equations(para)
            assert len(result) == 1

    def test_inline_equation_skips_already_processed(self, parser):
        om = MagicMock()
        om_para = MagicMock()
        om_para.findall.return_value = [om]
        para = self._make_paragraph_with_math(
            omath_paras=[om_para],
            omath_direct=[om],
            omath_inside_para=[om],
        )

        with patch.object(parser, "_extract_math_element", return_value=MagicMock(spec=Equation)):
            result = parser._extract_equations(para)
            assert len(result) == 1

    def test_inline_equation_direct(self, parser):
        om1 = MagicMock()
        om2 = MagicMock()
        para = self._make_paragraph_with_math(omath_direct=[om1, om2])

        with patch.object(parser, "_extract_math_element", return_value=MagicMock(spec=Equation)):
            result = parser._extract_equations(para)
            assert len(result) == 2


class TestMathElementExtraction:
    @pytest.fixture
    def parser(self):
        return DocxParser()

    def _make_real_om_element(self, m_t_texts=None):
        nsmap = {
            "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        }
        om = etree.Element("{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath", nsmap=nsmap)
        if m_t_texts:
            for t in m_t_texts:
                t_elem = etree.SubElement(om, "{http://schemas.openxmlformats.org/officeDocument/2006/math}t")
                t_elem.text = t
        return om

    def test_math_text_extraction_m_t(self, parser):
        om = self._make_real_om_element(m_t_texts=["E = mc", "\u00b2"])
        eqn = parser._extract_math_element(om, is_block=True)
        assert eqn is not None
        assert "E = mc" in eqn.text

    def test_math_text_fallback_w_t(self, parser):
        nsmap = {
            "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
            "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        }
        om = etree.Element("{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath", nsmap=nsmap)
        t_elem = etree.SubElement(om, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
        t_elem.text = "F = ma"

        eqn = parser._extract_math_element(om, is_block=False)
        assert eqn is not None
        assert eqn.text == "F = ma"

    def test_math_no_text(self, parser):
        om = self._make_real_om_element()
        eqn = parser._extract_math_element(om, is_block=True)
        assert eqn is not None
        assert eqn.text == ""

    def test_math_is_block_flag(self, parser):
        om = self._make_real_om_element(m_t_texts=["x=1"])
        eqn = parser._extract_math_element(om, is_block=True)
        assert eqn.is_block is True
        eqn2 = parser._extract_math_element(om, is_block=False)
        assert eqn2.is_block is False

    def test_math_omml_present(self, parser):
        om = self._make_real_om_element(m_t_texts=["a+b"])
        eqn = parser._extract_math_element(om, is_block=True)
        assert eqn.omml is not None

    def test_extraction_exception(self, parser):
        om = MagicMock()
        om.findall.side_effect = Exception("math error")
        eqn = parser._extract_math_element(om, is_block=True)
        assert eqn is None


class TestBlockCounter:
    @pytest.fixture
    def parser(self):
        return DocxParser()

    def test_counter_increments_across_parses(self, parser):
        assert parser.block_counter == 0
        para = MagicMock(spec=["runs", "style", "_element", "text", "alignment", "part"])
        para.text = "text"
        para.alignment = None
        style = MagicMock()
        style.name = None
        font = MagicMock()
        font.name = None
        font.size = None
        font.bold = None
        font.italic = None
        style.font = font
        para.style = style
        para._element.findall.return_value = []
        para._element.find.return_value = None
        para.runs = []

        parser._extract_paragraph(para)
        assert parser.block_counter == 1
        parser._extract_paragraph(para)
        assert parser.block_counter == 2


class TestParseDocumentIdCoercion:
    @pytest.fixture
    def parser(self):
        return DocxParser()

    @patch("app.pipeline.parsing.parser.DocxDocument")
    def test_uuid_coerced_to_string(self, mock_docx_cls, parser, tmp_path):
        mock_docx = MagicMock()
        mock_docx.element.body = []
        mock_docx.sections = []
        mock_docx_cls.return_value = mock_docx
        path = tmp_path / "u.docx"
        path.write_text("")
        result = parser.parse(str(path), uuid.uuid4())
        assert isinstance(result.document_id, str)

    @patch("app.pipeline.parsing.parser.DocxDocument")
    def test_int_coerced_to_string(self, mock_docx_cls, parser, tmp_path):
        mock_docx = MagicMock()
        mock_docx.element.body = []
        mock_docx.sections = []
        mock_docx_cls.return_value = mock_docx
        path = tmp_path / "i.docx"
        path.write_text("")
        result = parser.parse(str(path), 42)
        assert isinstance(result.document_id, str)
        assert result.document_id == "42"


class TestInheritance:
    def test_extends_base_parser(self):
        assert issubclass(DocxParser, BaseParser)
