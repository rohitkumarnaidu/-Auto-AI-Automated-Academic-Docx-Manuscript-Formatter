"""
Pipeline Integration Tests
Tests end-to-end pipeline execution and stage integration.
"""

import os
import pytest
from unittest.mock import Mock, patch, MagicMock
from pathlib import Path
from docx import Document as DocxDocument


class TestPipeline:
    """Test suite for pipeline integration."""

    @pytest.mark.pipeline
    def test_orchestrator_initialization(self):
        """Test pipeline orchestrator can be initialized."""
        with patch('app.pipeline.orchestrator.get_reasoning_engine'):
            with patch('app.pipeline.orchestrator.get_rag_engine'):
                from app.pipeline.orchestrator import PipelineOrchestrator

                orchestrator = PipelineOrchestrator()
                assert orchestrator is not None

    @pytest.mark.pipeline
    def test_contract_loading(self):
        """Test contract loading for templates."""
        from app.pipeline.contracts.loader import ContractLoader

        loader = ContractLoader()
        contract = loader.load("none")

        assert contract is not None
        assert "spacing" in contract or "publisher" in contract

    @pytest.mark.pipeline
    def test_contract_validation(self):
        """Test contract has required fields."""
        from app.pipeline.contracts.loader import ContractLoader

        loader = ContractLoader()
        contract = loader.load("none")

        assert contract is not None

    @pytest.mark.slow
    @pytest.mark.pipeline
    def test_pipeline_error_handling(self):
        """Test pipeline handles errors gracefully."""
        with patch('app.pipeline.orchestrator.get_reasoning_engine') as mock_reasoning:
            mock_reasoning.return_value.generate_instruction_set.side_effect = Exception("Test error")

            from app.pipeline.orchestrator import PipelineOrchestrator

            orchestrator = PipelineOrchestrator()
            assert orchestrator is not None

    @pytest.mark.pipeline
    def test_pipeline_run_with_valid_docx(self, tmp_path):
        """Test pipeline execution with a valid DOCX file."""
        doc = DocxDocument()
        doc.add_heading("Test Document", level=1)
        doc.add_paragraph("This is a test paragraph.")
        doc_path = tmp_path / "test_pipeline.docx"
        doc.save(str(doc_path))

        with (
            patch('app.pipeline.orchestrator.get_reasoning_engine') as mock_reasoning,
            patch('app.pipeline.orchestrator.get_rag_engine') as mock_rag,
            patch('app.pipeline.orchestrator.GROBIDClient') as mock_grobid,
            patch('app.pipeline.orchestrator.DoclingClient') as mock_docling,
            patch('app.pipeline.orchestrator.get_supabase_client', return_value=None),
        ):
            mock_reasoning.return_value = MagicMock()
            mock_rag.return_value = MagicMock()
            mock_grobid.return_value = MagicMock()
            mock_docling.return_value = MagicMock()

            from app.pipeline.orchestrator import PipelineOrchestrator

            orchestrator = PipelineOrchestrator(temp_dir=str(tmp_path / "temp"))
            output_path = str(tmp_path / "output.docx")

            result = orchestrator.run_pipeline(
                input_path=str(doc_path),
                job_id="test-job-id-valid-docx",
                template_name="none",
            )
            assert result is not None

    @pytest.mark.pipeline
    def test_pipeline_run_with_missing_file(self, tmp_path):
        """Test pipeline handles missing input file."""
        with (
            patch('app.pipeline.orchestrator.get_reasoning_engine') as mock_reasoning,
            patch('app.pipeline.orchestrator.get_rag_engine') as mock_rag,
            patch('app.pipeline.orchestrator.GROBIDClient') as mock_grobid,
            patch('app.pipeline.orchestrator.DoclingClient') as mock_docling,
            patch('app.pipeline.orchestrator.get_supabase_client', return_value=None),
        ):
            mock_reasoning.return_value = MagicMock()
            mock_rag.return_value = MagicMock()
            mock_grobid.return_value = MagicMock()
            mock_docling.return_value = MagicMock()

            from app.pipeline.orchestrator import PipelineOrchestrator

            orchestrator = PipelineOrchestrator(temp_dir=str(tmp_path / "temp"))
            result = orchestrator.run_pipeline(
                input_path="/nonexistent/file.docx",
                job_id="test-job-id-missing",
                template_name="none",
            )
            assert result is not None
            assert "error" in result or result.get("status") == "failed" or "document" not in result

    @pytest.mark.pipeline
    def test_pipeline_stage_interface_check(self):
        """Test stage interface validation."""
        with (
            patch('app.pipeline.orchestrator.get_reasoning_engine'),
            patch('app.pipeline.orchestrator.get_rag_engine'),
            patch('app.pipeline.orchestrator.GROBIDClient'),
            patch('app.pipeline.orchestrator.DoclingClient'),
            patch('app.pipeline.orchestrator.get_supabase_client', return_value=None),
        ):
            from app.pipeline.orchestrator import PipelineOrchestrator

            orchestrator = PipelineOrchestrator()

            class ValidStage:
                def process(self, doc):
                    return doc

            orchestrator._check_stage_interface(ValidStage(), "process", "ValidStage")

            class InvalidStage:
                pass

            with pytest.raises(RuntimeError, match="does not implement required method"):
                orchestrator._check_stage_interface(InvalidStage(), "process", "InvalidStage")

    @pytest.mark.pipeline
    def test_pipeline_with_corrupted_docx(self, tmp_path):
        """Test pipeline handles corrupted DOCX gracefully."""
        corrupted = tmp_path / "corrupted.docx"
        corrupted.write_bytes(b"\x00\x01\x02CORRUPTED")

        with (
            patch('app.pipeline.orchestrator.get_reasoning_engine'),
            patch('app.pipeline.orchestrator.get_rag_engine'),
            patch('app.pipeline.orchestrator.GROBIDClient'),
            patch('app.pipeline.orchestrator.DoclingClient'),
            patch('app.pipeline.orchestrator.get_supabase_client', return_value=None),
        ):
            from app.pipeline.orchestrator import PipelineOrchestrator

            orchestrator = PipelineOrchestrator(temp_dir=str(tmp_path / "temp"))
            result = orchestrator.run_pipeline(
                input_path=str(corrupted),
                job_id="test-job-id-corrupted",
                template_name="none",
            )
            assert result is not None


if __name__ == "__main__":
    pytest.main([__file__, "-v", "-m", "pipeline"])
