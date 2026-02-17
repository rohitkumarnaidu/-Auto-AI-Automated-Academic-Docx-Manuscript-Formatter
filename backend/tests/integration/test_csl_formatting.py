"""
Integration tests for CSL formatting and template rendering.
"""

from __future__ import annotations

from zipfile import ZipFile

import pytest
from docx import Document as WordDocument

from app.models import (
    Block,
    BlockType,
    DocumentMetadata,
    PipelineDocument,
    Reference,
    ReferenceType,
)
from app.pipeline.contracts.loader import ContractLoader
from app.pipeline.formatting.template_renderer import TemplateRenderer
from app.pipeline.references.formatter_engine import ReferenceFormatterEngine
from app.pipeline.services.csl_engine import CSLEngine


def _build_document() -> PipelineDocument:
    doc = PipelineDocument(
        document_id="integration-csl-doc",
        metadata=DocumentMetadata(
            title="CSL Integration Validation",
            authors=["Alice Smith", "Bob Lee"],
            affiliations=["University X", "Institute Y"],
            abstract="Integration test for CSL and Jinja2 templates.",
            keywords=["csl", "docxtpl", "integration"],
        ),
        formatting_options={
            "cover_page": True,
            "toc": True,
            "page_numbers": True,
            "page_number": "1",
        },
    )

    doc.blocks = [
        Block(block_id="h1", index=0, block_type=BlockType.HEADING_1, text="Introduction"),
        Block(
            block_id="b1",
            index=1,
            block_type=BlockType.BODY,
            text="This manuscript validates CSL reference formatting.",
        ),
    ]

    doc.references = [
        Reference(
            reference_id="ref_001",
            citation_key="ref_1",
            raw_text="[1] A. Smith, B. Lee, Robust Pipelines, 2024.",
            reference_type=ReferenceType.JOURNAL_ARTICLE,
            authors=["Smith, Alice", "Lee, Bob"],
            title="Robust Pipelines",
            journal="Journal of Document Automation",
            year=2024,
            volume="12",
            issue="3",
            pages="101-110",
            doi="10.1000/jda.2024.001",
            index=0,
        )
    ]
    return doc


def _read_visible_text(path) -> str:
    doc = WordDocument(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _find_unresolved_jinja_tag(path):
    with ZipFile(path, "r") as archive:
        for name in archive.namelist():
            if not name.startswith("word/") or not name.endswith(".xml"):
                continue
            xml = archive.read(name).decode("utf-8", errors="ignore")
            for token in ("{{", "{%", "{#"):
                if token in xml:
                    return token, name
    return None, None


@pytest.mark.integration
def test_csl_ieee_formatting_integrates_with_template_renderer(tmp_path):
    loader = ContractLoader(contracts_dir="app/pipeline/contracts")
    ref_engine = ReferenceFormatterEngine(contract_loader=loader, csl_engine=CSLEngine())
    renderer = TemplateRenderer(templates_dir="app/templates")
    document = _build_document()

    document.references = ref_engine.format_all(document.references, publisher="ieee")
    rendered = renderer.render(document, template_name="ieee")

    out = tmp_path / "ieee_integration.docx"
    rendered.save(str(out))

    text = _read_visible_text(out)
    assert "CSL Integration Validation" in text
    assert "Robust Pipelines" in text
    assert "References" in text

    token, xml_name = _find_unresolved_jinja_tag(out)
    assert token is None, f"Found unresolved Jinja tag token '{token}' in {xml_name}"


@pytest.mark.integration
def test_csl_apa_formatting_integrates_with_template_renderer(tmp_path):
    loader = ContractLoader(contracts_dir="app/pipeline/contracts")
    ref_engine = ReferenceFormatterEngine(contract_loader=loader, csl_engine=CSLEngine())
    renderer = TemplateRenderer(templates_dir="app/templates")
    document = _build_document()

    document.references = ref_engine.format_all(document.references, publisher="apa")
    rendered = renderer.render(document, template_name="apa")

    out = tmp_path / "apa_integration.docx"
    rendered.save(str(out))

    text = _read_visible_text(out)
    assert "CSL Integration Validation" in text
    assert "Robust Pipelines" in text
    assert "References" in text

    token, xml_name = _find_unresolved_jinja_tag(out)
    assert token is None, f"Found unresolved Jinja tag token '{token}' in {xml_name}"


@pytest.mark.integration
def test_template_validation_all_styles_render_all_tags(tmp_path):
    renderer = TemplateRenderer(templates_dir="app/templates")

    for style in ("ieee", "apa", "none"):
        document = _build_document()
        rendered = renderer.render(document, template_name=style)
        out = tmp_path / f"{style}_validation.docx"
        rendered.save(str(out))

        visible_text = _read_visible_text(out)
        assert "CSL Integration Validation" in visible_text
        assert "Table of Contents" in visible_text
        assert "Page 1" in visible_text

        token, xml_name = _find_unresolved_jinja_tag(out)
        assert token is None, f"Style '{style}' has unresolved token '{token}' in {xml_name}"
