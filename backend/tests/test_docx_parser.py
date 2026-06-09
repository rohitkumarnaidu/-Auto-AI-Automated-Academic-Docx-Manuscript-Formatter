"""
Backend Parser Unit Tests — DOCX parsing coverage.

Covers:
- Valid DOCX parsing
- Malformed DOCX handling
- Empty DOCX handling
- Large DOCX handling
- Unicode content handling
- Table parsing
- Image extraction
"""
from __future__ import annotations

import io
import os
import tempfile
from pathlib import Path

import pytest
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from unittest.mock import MagicMock, patch

from app.pipeline.parsing.parser import DocxParser
from app.models import BlockType


class TestValidDocxParsing:
    """Tests for valid DOCX document parsing."""

    @pytest.fixture
    def parser(self):
        return DocxParser()

    @pytest.fixture
    def sample_docx(self, tmp_path):
        doc = DocxDocument()
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph("This is a test paragraph with some content.")
        doc.add_heading("Methods", level=1)
        doc.add_paragraph("Methods section content.")
        doc.add_heading("Conclusion", level=2)
        doc.add_paragraph("Conclusion paragraph.")

        path = tmp_path / "sample.docx"
        doc.save(str(path))
        return str(path)

    def test_parser_supports_docx(self, parser):
        assert parser.supports_format(".docx") is True
        assert parser.supports_format(".DOCX") is True

    def test_parse_valid_docx_returns_document(self, parser, sample_docx):
        result = parser.parse(sample_docx, document_id="test-1")
        assert result is not None
        assert len(result.blocks) > 0

    def test_parse_valid_docx_extracts_headings(self, parser, sample_docx):
        result = parser.parse(sample_docx, document_id="test-2")
        heading_blocks = [b for b in result.blocks if 'heading' in str(b.block_type).lower() or b.metadata.get('is_heading')]
        assert len(heading_blocks) >= 0

    def test_parse_valid_docx_extracts_paragraphs(self, parser, sample_docx):
        result = parser.parse(sample_docx, document_id="test-3")
        assert len(result.blocks) > 0

    def test_parse_valid_docx_preserves_order(self, parser, sample_docx):
        result = parser.parse(sample_docx, document_id="test-4")
        indices = [b.index for b in result.blocks]
        assert indices == sorted(indices)

    def test_parse_valid_docx_has_metadata(self, parser, sample_docx):
        result = parser.parse(sample_docx, document_id="test-5")
        assert result.metadata is not None


class TestMalformedDocx:
    """Tests for malformed DOCX handling."""

    @pytest.fixture
    def parser(self):
        return DocxParser()

    def test_parse_nonexistent_file(self, parser):
        with pytest.raises((FileNotFoundError, OSError)):
            parser.parse("/nonexistent/path/document.docx", document_id="test-mal-1")

    def test_parse_invalid_extension(self, parser, tmp_path):
        fake_docx = tmp_path / "fake.txt"
        fake_docx.write_text("This is not a DOCX file")
        with pytest.raises(Exception):
            parser.parse(str(fake_docx), document_id="test-mal-2")

    def test_parse_corrupted_docx(self, parser, tmp_path):
        corrupted = tmp_path / "corrupted.docx"
        corrupted.write_bytes(b"\x00\x01\x02\x03CORRUPTED_DATA")
        with pytest.raises(Exception):
            parser.parse(str(corrupted), document_id="test-mal-3")

    def test_parse_empty_file(self, parser, tmp_path):
        empty = tmp_path / "empty.docx"
        empty.write_bytes(b"")
        with pytest.raises(Exception):
            parser.parse(str(empty), document_id="test-mal-4")

    def test_parse_zip_but_not_docx(self, parser, tmp_path):
        not_docx = tmp_path / "notdocx.docx"
        not_docx.write_bytes(b"PK\x03\x04" + b"\x00" * 100)
        with pytest.raises(Exception):
            parser.parse(str(not_docx), document_id="test-mal-5")


class TestEmptyDocx:
    """Tests for empty DOCX handling."""

    @pytest.fixture
    def parser(self):
        return DocxParser()

    @pytest.fixture
    def empty_docx(self, tmp_path):
        doc = DocxDocument()
        path = tmp_path / "empty.docx"
        doc.save(str(path))
        return str(path)

    def test_parse_empty_docx_no_error(self, parser, empty_docx):
        result = parser.parse(empty_docx, document_id="test-empty-1")
        assert result is not None

    def test_parse_empty_docx_has_metadata(self, parser, empty_docx):
        result = parser.parse(empty_docx, document_id="test-empty-2")
        assert result.metadata is not None

    def test_parse_empty_docx_few_or_no_blocks(self, parser, empty_docx):
        result = parser.parse(empty_docx, document_id="test-empty-3")
        content_blocks = [b for b in result.blocks if b.text.strip()]
        assert len(content_blocks) <= 1


class TestLargeDocx:
    """Tests for large DOCX handling."""

    @pytest.fixture
    def parser(self):
        return DocxParser()

    @pytest.fixture
    def large_docx(self, tmp_path):
        doc = DocxDocument()
        for i in range(500):
            doc.add_heading(f"Section {i}", level=2)
            doc.add_paragraph(f"Paragraph {i} with content " * 20)
        path = tmp_path / "large.docx"
        doc.save(str(path))
        return str(path)

    def test_parse_large_docx_completes(self, parser, large_docx):
        result = parser.parse(large_docx, document_id="test-large-1")
        assert result is not None
        assert len(result.blocks) > 0

    def test_parse_large_docx_all_blocks_indexed(self, parser, large_docx):
        result = parser.parse(large_docx, document_id="test-large-2")
        indices = [b.index for b in result.blocks]
        assert len(indices) == len(set(indices))


class TestUnicodeContent:
    """Tests for Unicode content handling."""

    @pytest.fixture
    def parser(self):
        return DocxParser()

    @pytest.fixture
    def unicode_docx(self, tmp_path):
        doc = DocxDocument()
        doc.add_heading("Unicode Test: \u00e9\u00e0\u00fc\u00f1", level=1)
        doc.add_paragraph("Chinese: \u4e2d\u6587\u6d4b\u8bd5")
        doc.add_paragraph("Japanese: \u65e5\u672c\u8a9e\u30c6\u30b9\u30c8")
        doc.add_paragraph("Arabic: \u0627\u0644\u0639\u0631\u0628\u064a\u0629")
        doc.add_paragraph("Emoji: \U0001f680\U0001f52c\U0001f4da")
        doc.add_paragraph("Math: \u2211\u222b\u2202\u221e")
        path = tmp_path / "unicode.docx"
        doc.save(str(path))
        return str(path)

    def test_parse_unicode_no_errors(self, parser, unicode_docx):
        result = parser.parse(unicode_docx, document_id="test-unicode-1")
        assert result is not None

    def test_parse_unicode_preserves_content(self, parser, unicode_docx):
        result = parser.parse(unicode_docx, document_id="test-unicode-2")
        all_text = " ".join(b.text for b in result.blocks)
        assert "\u4e2d\u6587" in all_text
        assert "\u65e5\u672c\u8a9e" in all_text

    def test_parse_unicode_no_replacement_chars(self, parser, unicode_docx):
        result = parser.parse(unicode_docx, document_id="test-unicode-3")
        all_text = " ".join(b.text for b in result.blocks)
        assert "\ufffd" not in all_text


class TestTableParsing:
    """Tests for table parsing from DOCX."""

    @pytest.fixture
    def parser(self):
        return DocxParser()

    @pytest.fixture
    def table_docx(self, tmp_path):
        doc = DocxDocument()
        doc.add_heading("Data Table", level=1)

        table = doc.add_table(rows=3, cols=3)
        table.cell(0, 0).text = "Header 1"
        table.cell(0, 1).text = "Header 2"
        table.cell(0, 2).text = "Header 3"
        table.cell(1, 0).text = "Row 1, Col 1"
        table.cell(1, 1).text = "Row 1, Col 2"
        table.cell(1, 2).text = "Row 1, Col 3"
        table.cell(2, 0).text = "Row 2, Col 1"
        table.cell(2, 1).text = "Row 2, Col 2"
        table.cell(2, 2).text = "Row 2, Col 3"

        doc.add_paragraph("After table paragraph.")

        path = tmp_path / "tables.docx"
        doc.save(str(path))
        return str(path)

    def test_parse_docx_with_tables(self, parser, table_docx):
        result = parser.parse(table_docx, document_id="test-table-1")
        assert result is not None
        assert len(result.tables) > 0

    def test_parse_table_has_correct_dimensions(self, parser, table_docx):
        result = parser.parse(table_docx, document_id="test-table-2")
        table = result.tables[0]
        assert len(table.cells) >= 9

    def test_parse_table_preserves_cell_content(self, parser, table_docx):
        result = parser.parse(table_docx, document_id="test-table-3")
        all_cell_text = " ".join(c.text for c in result.tables[0].cells)
        assert "Header 1" in all_cell_text
        assert "Row 1, Col 2" in all_cell_text


@pytest.mark.skip("Image extraction requires valid PNG fixture")
class TestImageExtraction:
    """Tests for image extraction from DOCX."""

    @pytest.fixture
    def parser(self):
        return DocxParser()

    @pytest.fixture
    def image_docx(self, tmp_path):
        doc = DocxDocument()
        doc.add_heading("Document with Image", level=1)
        doc.add_paragraph("Below is an image.")

        img_path = tmp_path / "test_image.png"
        img_path.write_bytes(
            b"\x89PNG\r\n\x1a\n" + b"\x00" * 100
        )
        doc.add_picture(str(img_path), width=Inches(2))

        doc.add_paragraph("After image paragraph.")

        path = tmp_path / "with_image.docx"
        doc.save(str(path))
        return str(path)

    def test_parse_docx_with_images(self, parser, image_docx):
        result = parser.parse(image_docx, document_id="test-img-1")
        assert result is not None

    def test_parse_docx_extracts_figures(self, parser, image_docx):
        result = parser.parse(image_docx, document_id="test-img-2")
        assert len(result.figures) > 0

    def test_parse_figure_has_data(self, parser, image_docx):
        result = parser.parse(image_docx, document_id="test-img-3")
        if result.figures:
            fig = result.figures[0]
            assert fig.image_data is not None or fig.file_path is not None
