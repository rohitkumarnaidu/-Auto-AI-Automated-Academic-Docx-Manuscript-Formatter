from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any, Dict, List, Tuple
from xml.etree import ElementTree as ET
from zipfile import ZipFile

import yaml
from docx import Document as WordDocument

from app.models import (
    Block,
    BlockType,
    DocumentMetadata,
    Equation,
    PipelineDocument,
    Reference,
    ReferenceType,
    Table,
    TableCell,
    TemplateInfo,
)
from app.pipeline.formatting.formatter import Formatter

BASE_DIR = Path(__file__).parent / "golden_files"
INPUT_DIR = BASE_DIR / "inputs"
GOLDEN_DIR = BASE_DIR / "goldens"
FORMATTER = Formatter(templates_dir="app/templates", contracts_dir="app/pipeline/contracts")


def _split_frontmatter(raw_text: str) -> Tuple[Dict[str, Any], str]:
    if not raw_text.startswith("---\n"):
        return {}, raw_text

    _, frontmatter, body = raw_text.split("---\n", 2)
    return yaml.safe_load(frontmatter) or {}, body


def _extract_markdown_hyperlinks(text: str) -> Tuple[str, List[Dict[str, str]]]:
    links: List[Dict[str, str]] = []

    def _replace(match: re.Match[str]) -> str:
        anchor = match.group(1).strip()
        url = match.group(2).strip()
        if anchor and url:
            links.append({"text": anchor, "url": url})
        return anchor

    cleaned = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", _replace, text)
    return cleaned, links


def _extract_footnote_refs(text: str) -> Tuple[str, List[str]]:
    refs = re.findall(r"\[\^([^\]]+)\]", text)
    cleaned = re.sub(r"\[\^[^\]]+\]", "", text)
    return cleaned.strip(), refs


def _build_pipeline_document(sample_path: Path, template_name: str) -> PipelineDocument:
    raw_text = sample_path.read_text(encoding="utf-8")
    frontmatter, body = _split_frontmatter(raw_text)

    metadata = DocumentMetadata(
        title=frontmatter.get("title"),
        authors=list(frontmatter.get("authors") or []),
        affiliations=list(frontmatter.get("affiliations") or []),
        keywords=list(frontmatter.get("keywords") or []),
    )
    document = PipelineDocument(
        document_id=f"golden-{template_name}",
        original_filename=sample_path.name,
        source_path=str(sample_path),
        metadata=metadata,
        template=TemplateInfo(template_name=template_name),
        formatting_options={
            "cover_page": False,
            "toc": False,
            "page_numbers": False,
            "borders": False,
            "line_numbers": False,
        },
    )

    blocks: List[Block] = []
    references: List[Reference] = []
    current_mode = "body"
    current_section_name = ""
    block_index = 100
    reference_index = 0

    for raw_line in body.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue

        footnote_match = re.match(r"^\[\^([^\]]+)\]:\s*(.+)$", line.strip())
        if footnote_match:
            blocks.append(
                Block(
                    block_id=f"{template_name}-block-{len(blocks) + 1}",
                    index=block_index,
                    block_type=BlockType.FOOTNOTE,
                    text=footnote_match.group(2).strip(),
                    section_name="footnotes",
                    metadata={"footnote_id": footnote_match.group(1).strip()},
                )
            )
            block_index += 100
            continue

        heading_match = re.match(r"^(#{1,4})\s+(.+?)\s*$", line)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            lowered = heading_text.lower()

            if lowered == "abstract":
                block_type = BlockType.ABSTRACT_HEADING
                current_mode = "abstract"
                current_section_name = "abstract"
            elif lowered == "keywords":
                block_type = BlockType.KEYWORDS_HEADING
                current_mode = "keywords"
                current_section_name = "keywords"
            elif lowered == "references":
                block_type = BlockType.REFERENCES_HEADING
                current_mode = "references"
                current_section_name = "references"
            else:
                block_type = {
                    1: BlockType.HEADING_1,
                    2: BlockType.HEADING_2,
                    3: BlockType.HEADING_3,
                    4: BlockType.HEADING_4,
                }[min(level, 4)]
                current_mode = "body"
                current_section_name = lowered

            blocks.append(
                Block(
                    block_id=f"{template_name}-block-{len(blocks) + 1}",
                    index=block_index,
                    block_type=block_type,
                    text=heading_text,
                    level=level,
                    section_name=current_section_name,
                )
            )
            block_index += 100
            continue

        cleaned_text, hyperlinks = _extract_markdown_hyperlinks(line.strip())
        cleaned_text, footnote_refs = _extract_footnote_refs(cleaned_text)
        block_metadata: Dict[str, Any] = {}
        if hyperlinks:
            block_metadata["hyperlinks"] = hyperlinks
        if footnote_refs:
            block_metadata["footnote_refs"] = footnote_refs

        if current_mode == "abstract":
            metadata.abstract = cleaned_text
            block_type = BlockType.ABSTRACT_BODY
        elif current_mode == "keywords":
            if not metadata.keywords:
                metadata.keywords = [item.strip() for item in cleaned_text.split(",") if item.strip()]
            block_type = BlockType.KEYWORDS_BODY
        elif current_mode == "references":
            block_type = BlockType.REFERENCE_ENTRY
            references.append(
                Reference(
                    reference_id=f"{template_name}-ref-{reference_index + 1}",
                    citation_key=f"{template_name}_ref_{reference_index + 1}",
                    raw_text=cleaned_text,
                    reference_type=ReferenceType.JOURNAL_ARTICLE,
                    index=reference_index,
                    formatted_text=cleaned_text,
                )
            )
            reference_index += 1
        else:
            block_type = BlockType.BODY

        blocks.append(
            Block(
                block_id=f"{template_name}-block-{len(blocks) + 1}",
                index=block_index,
                block_type=block_type,
                text=cleaned_text,
                section_name=current_section_name,
                metadata=block_metadata,
            )
        )
        block_index += 100

    document.blocks = blocks
    document.references = references
    return document


def _build_structural_summary(output_path: Path, golden: Dict[str, Any]) -> Dict[str, Any]:
    rendered = WordDocument(str(output_path))
    paragraphs = [paragraph.text.strip() for paragraph in rendered.paragraphs if paragraph.text.strip()]
    full_text = "\n".join(paragraphs)

    actual_hierarchy = [
        heading
        for heading in golden["heading_hierarchy"]
        if heading["text"] in full_text
    ]
    reference_count = sum(
        1 for reference_text in golden.get("reference_entries", []) if reference_text in full_text
    )

    return {
        "section_count": len(actual_hierarchy),
        "heading_hierarchy": actual_hierarchy,
        "template_metadata": {
            "template_name": golden["template_metadata"]["template_name"],
            "title_present": golden["template_metadata"]["title"] in full_text,
        },
        "reference_count": reference_count,
    }


def _load_golden(template_name: str) -> Dict[str, Any]:
    return json.loads((GOLDEN_DIR / f"{template_name}.json").read_text(encoding="utf-8"))


def _run_fixture(sample_name: str, tmp_path: Path) -> Dict[str, Any]:
    sample_path = INPUT_DIR / f"{sample_name}.md"
    golden = _load_golden(sample_name)
    document = _build_pipeline_document(sample_path, sample_name)

    rendered = FORMATTER.format(document, template_name=sample_name)
    output_path = tmp_path / f"{sample_name}.docx"
    rendered.save(str(output_path))
    return _build_structural_summary(output_path, golden)


def test_formatter_golden_files(tmp_path: Path):
    sample_names = ["ieee", "apa", "acm", "nature", "resume"]

    for sample_name in sample_names:
        golden = _load_golden(sample_name)
        actual = _run_fixture(sample_name, tmp_path)
        expected = {
            "section_count": golden["section_count"],
            "heading_hierarchy": golden["heading_hierarchy"],
            "template_metadata": {
                "template_name": golden["template_metadata"]["template_name"],
                "title_present": True,
            },
            "reference_count": golden["reference_count"],
        }
        assert actual == expected, sample_name


def _build_rich_content_document() -> PipelineDocument:
    document = PipelineDocument(
        document_id="golden-rich-content",
        original_filename="rich-content.md",
        source_path="rich-content.md",
        metadata=DocumentMetadata(
            title="Rich Content Regression",
            authors=["Ada Lovelace"],
            affiliations=["Analytical Engine Institute"],
        ),
        template=TemplateInfo(template_name="ieee"),
        formatting_options={"cover_page": False, "toc": False, "page_numbers": False},
    )
    document.blocks = [
        Block(block_id="title-1", index=0, block_type=BlockType.TITLE, text="Rich Content Regression"),
        Block(block_id="h1", index=100, block_type=BlockType.HEADING_1, text="Results", level=1),
        Block(
            block_id="body-1",
            index=200,
            block_type=BlockType.BODY,
            text="Visit OpenAI for richer context",
            metadata={
                "hyperlinks": [{"text": "OpenAI", "url": "https://openai.com"}],
                "footnote_refs": ["fn-1"],
            },
        ),
        Block(
            block_id="fn-1-block",
            index=900,
            block_type=BlockType.FOOTNOTE,
            text="Footnote text survives the DOCX patching step.",
            metadata={"footnote_id": "fn-1"},
        ),
    ]
    document.equations = [
        Equation(
            equation_id="eq-1",
            index=300,
            block_id="body-1",
            text="E = mc^2",
            metadata={"block_index": 300},
        )
    ]
    document.tables = [
        Table(
            table_id="tbl-1",
            index=0,
            block_index=400,
            num_rows=2,
            num_cols=2,
            rows=[["Metric", "Value"], ["Accuracy", "0.98"]],
            data=[["Metric", "Value"], ["Accuracy", "0.98"]],
            cells=[
                TableCell(row=0, col=0, text="Metric", is_header=True, bold=True),
                TableCell(row=0, col=1, text="Value", is_header=True, bold=True),
                TableCell(row=1, col=0, text="Accuracy"),
                TableCell(row=1, col=1, text="0.98"),
            ],
            has_header=True,
            has_header_row=True,
            header_rows=1,
            caption_text="Benchmark Summary",
        )
    ]
    return document


def _read_docx_part(output_path: Path, part_name: str) -> bytes:
    with ZipFile(output_path) as archive:
        return archive.read(part_name)


def test_formatter_golden_rich_content_regression(tmp_path: Path):
    document = _build_rich_content_document()
    rendered = FORMATTER.format(document, template_name="ieee")
    output_path = tmp_path / "rich-content.docx"
    rendered.save(str(output_path))

    rendered_doc = WordDocument(str(output_path))
    full_text = "\n".join(paragraph.text for paragraph in rendered_doc.paragraphs if paragraph.text)

    assert any("Benchmark Summary" in paragraph.text for paragraph in rendered_doc.paragraphs)
    assert len(rendered_doc.tables) == 1
    assert rendered_doc.tables[0].cell(0, 0).text == "Metric"
    assert rendered_doc.tables[0].cell(1, 1).text == "0.98"
    assert "Footnote text survives the DOCX patching step." not in full_text

    document_xml = _read_docx_part(output_path, "word/document.xml").decode("utf-8")
    relationships_xml = _read_docx_part(output_path, "word/_rels/document.xml.rels").decode("utf-8")
    footnotes_xml = _read_docx_part(output_path, "word/footnotes.xml")

    assert "w:hyperlink" in document_xml
    assert "https://openai.com" in relationships_xml
    assert "m:oMathPara" in document_xml
    assert "E = mc^2" in document_xml

    footnotes_root = ET.fromstring(footnotes_xml)
    footnote_text = "".join(node.text or "" for node in footnotes_root.iter() if node.text)
    assert "Footnote text survives the DOCX patching step." in footnote_text
    assert "footnoteReference" in document_xml
