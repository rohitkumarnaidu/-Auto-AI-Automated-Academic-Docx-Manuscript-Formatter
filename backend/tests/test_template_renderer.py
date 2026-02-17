"""
Tests for docxtpl/Jinja2 template rendering.
"""

from docx import Document as WordDocument

from app.models import (
    Block,
    BlockType,
    DocumentMetadata,
    PipelineDocument,
    Reference,
    ReferenceType,
)
from app.pipeline.formatting.formatter import Formatter
from app.pipeline.formatting.template_renderer import TemplateRenderer


def _build_sample_document() -> PipelineDocument:
    doc = PipelineDocument(
        document_id="doc-template-001",
        metadata=DocumentMetadata(
            title="Neural Rendering for Scientific Manuscripts",
            authors=["Alice Smith", "Bob Lee"],
            affiliations=["University X", "Institute Y"],
            abstract="This study proposes a robust formatting workflow.",
            keywords=["formatting", "docxtpl"],
        ),
        formatting_options={"cover_page": True},
    )

    doc.blocks = [
        Block(
            block_id="b1",
            index=1,
            block_type=BlockType.HEADING_1,
            text="Introduction",
        ),
        Block(
            block_id="b2",
            index=2,
            block_type=BlockType.BODY,
            text="Template-driven rendering improves consistency.",
        ),
    ]

    doc.references = [
        Reference(
            reference_id="ref_001",
            citation_key="ref_1",
            raw_text="[1] A. Smith, Rendering, 2024.",
            reference_type=ReferenceType.JOURNAL_ARTICLE,
            authors=["Smith, A."],
            title="Rendering",
            journal="Journal of Formatting",
            year=2024,
            index=0,
            formatted_text="[1] A. Smith, \"Rendering,\" Journal of Formatting, 2024.",
        )
    ]
    return doc


def _read_docx_text(path) -> str:
    doc = WordDocument(str(path))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text)


def test_template_renderer_renders_jinja_context(tmp_path):
    renderer = TemplateRenderer(templates_dir="app/templates")
    pipeline_doc = _build_sample_document()

    rendered = renderer.render(pipeline_doc, template_name="ieee")
    out_path = tmp_path / "rendered_ieee.docx"
    rendered.save(str(out_path))

    text = _read_docx_text(out_path)
    assert "Neural Rendering for Scientific Manuscripts" in text
    assert "Alice Smith" in text
    assert "Bob Lee" in text
    assert "Introduction" in text
    assert "{{ title }}" not in text
    assert "{% for author in authors %}" not in text


def test_formatter_uses_docxtpl_and_avoids_duplicate_cover_title(tmp_path):
    formatter = Formatter(templates_dir="app/templates", contracts_dir="app/pipeline/contracts")
    pipeline_doc = _build_sample_document()

    rendered = formatter.format(pipeline_doc, template_name="ieee")
    out_path = tmp_path / "formatted_ieee.docx"
    rendered.save(str(out_path))

    text = _read_docx_text(out_path)
    title = "Neural Rendering for Scientific Manuscripts"
    assert rendered.__class__.__name__ == "DocxTemplate"
    assert text.count(title) == 1


def test_template_renderer_renders_apa_template(tmp_path):
    renderer = TemplateRenderer(templates_dir="app/templates")
    pipeline_doc = _build_sample_document()

    rendered = renderer.render(pipeline_doc, template_name="apa")
    out_path = tmp_path / "rendered_apa.docx"
    rendered.save(str(out_path))

    text = _read_docx_text(out_path)
    assert "Neural Rendering for Scientific Manuscripts" in text
    assert "References" in text
    assert "Journal of Formatting" in text
