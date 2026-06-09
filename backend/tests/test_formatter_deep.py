from __future__ import annotations

import io
import logging
import os
import re
from unittest.mock import MagicMock, PropertyMock, call, patch, ANY
from xml.etree import ElementTree as ET
from zipfile import ZipFile

import pytest
import yaml
from docx import Document as WordDocument
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.models import (
    Block,
    BlockType,
    DocumentMetadata,
    Equation,
    Figure,
    PipelineDocument,
    Reference,
    ReferenceType,
    Table,
    TemplateInfo,
)
from app.pipeline.formatting.formatter import Formatter


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

@pytest.fixture
def mock_sub_components():
    """Return mocks for the 6 sub-components so no real filesystem is needed."""
    with (
        patch("app.pipeline.formatting.formatter.ContractLoader") as mock_cl,
        patch("app.pipeline.formatting.formatter.StyleMapper") as mock_sm,
        patch("app.pipeline.formatting.formatter.NumberingEngine") as mock_ne,
        patch("app.pipeline.formatting.formatter.ReferenceFormatter") as mock_rf,
        patch("app.pipeline.formatting.formatter.TemplateRenderer") as mock_tr,
        patch("app.pipeline.formatting.formatter.TableRenderer") as mock_tbl,
    ):
        contract_loader = mock_cl.return_value
        contract_loader.load.return_value = {}
        style_mapper = mock_sm.return_value
        numbering_engine = mock_ne.return_value
        reference_formatter = mock_rf.return_value
        template_renderer = mock_tr.return_value
        table_renderer = mock_tbl.return_value
        yield {
            "contract_loader": contract_loader,
            "style_mapper": style_mapper,
            "numbering_engine": numbering_engine,
            "reference_formatter": reference_formatter,
            "template_renderer": template_renderer,
            "table_renderer": table_renderer,
        }


@pytest.fixture
def formatter(mock_sub_components):
    return Formatter(templates_dir="/fake/templates", contracts_dir="/fake/contracts")


@pytest.fixture
def minimal_doc():
    return PipelineDocument(
        document_id="doc-1",
        metadata=DocumentMetadata(title="Test Title", authors=["Alice"]),
        blocks=[
            Block(block_id="b1", index=0, block_type=BlockType.TITLE, text="Test Title"),
            Block(block_id="b2", index=100, block_type=BlockType.BODY, text="Hello world."),
        ],
        template=TemplateInfo(template_name="ieee"),
    )


@pytest.fixture
def word_doc():
    return WordDocument()


# ---------------------------------------------------------------------------
# Initialization & process
# ---------------------------------------------------------------------------

class TestInit:
    def test_initializes_sub_components(self, mock_sub_components):
        f = Formatter(templates_dir="/t", contracts_dir="/c")
        assert f.templates_dir == "/t"

    def test_process_sets_generated_doc(self, formatter, minimal_doc):
        with patch.object(formatter, "format", return_value="fake_doc"):
            result = formatter.process(minimal_doc)
        assert result.generated_doc == "fake_doc"

    def test_process_falls_back_to_none(self, formatter):
        doc = PipelineDocument(document_id="d", template=None)
        with patch.object(formatter, "format", return_value="fake_doc"):
            result = formatter.process(doc)
        assert result.generated_doc == "fake_doc"

    def test_process_no_template_attr(self, formatter):
        doc = PipelineDocument(document_id="d", template=TemplateInfo(template_name="none"))
        with patch.object(formatter, "format", return_value="fake_doc"):
            result = formatter.process(doc)
        assert result.generated_doc == "fake_doc"


# ---------------------------------------------------------------------------
# _coerce_bool_option  (static, no mocks)
# ---------------------------------------------------------------------------

class TestCoerceBoolOption:
    def test_none_returns_default(self):
        assert Formatter._coerce_bool_option(None, True) is True
        assert Formatter._coerce_bool_option(None, False) is False

    def test_bool_passthrough(self):
        assert Formatter._coerce_bool_option(True, False) is True
        assert Formatter._coerce_bool_option(False, True) is False

    def test_int_zero_is_false(self):
        assert Formatter._coerce_bool_option(0, True) is False

    def test_int_nonzero_is_true(self):
        assert Formatter._coerce_bool_option(1, False) is True
        assert Formatter._coerce_bool_option(42, False) is True

    def test_float(self):
        assert Formatter._coerce_bool_option(0.0, True) is False
        assert Formatter._coerce_bool_option(1.0, False) is True

    def test_string_true_tokens(self):
        for token in ("1", "true", "True", "yes", "on"):
            assert Formatter._coerce_bool_option(token, False) is True

    def test_string_false_tokens(self):
        for token in ("0", "false", "no", "off", ""):
            assert Formatter._coerce_bool_option(token, True) is False

    def test_string_unknown_returns_bool(self):
        assert Formatter._coerce_bool_option("garbage", False) is True

    def test_list_nonempty(self):
        assert Formatter._coerce_bool_option([1], False) is True

    def test_empty_list(self):
        assert Formatter._coerce_bool_option([], True) is False


# ---------------------------------------------------------------------------
# _resolve_bool_option
# ---------------------------------------------------------------------------

class TestResolveBoolOption:
    def test_none_options_returns_default(self, formatter):
        assert formatter._resolve_bool_option(None, "x", default=True) is True

    def test_primary_key_found(self, formatter):
        assert formatter._resolve_bool_option({"x": True}, "x") is True

    def test_alias_found(self, formatter):
        assert formatter._resolve_bool_option({"y": True}, "x", aliases=("y",)) is True

    def test_not_found_returns_default(self, formatter):
        assert formatter._resolve_bool_option({"z": True}, "x", default=False) is False

    def test_primary_key_takes_precedence(self, formatter):
        opts = {"x": True, "y": False}
        assert formatter._resolve_bool_option(opts, "x", aliases=("y",)) is True


# ---------------------------------------------------------------------------
# _resolve_page_size
# ---------------------------------------------------------------------------

class TestResolvePageSize:
    def test_from_options(self, formatter):
        assert formatter._resolve_page_size("ieee", {"page_size": "A4"}) == "A4"

    def test_from_contract(self, formatter):
        formatter.contract_loader.load.return_value = {"layout": {"page_size": "Legal"}}
        assert formatter._resolve_page_size("ieee", {}) == "Legal"

    def test_default_to_letter(self, formatter):
        formatter.contract_loader.load.return_value = {}
        assert formatter._resolve_page_size("ieee", {}) == "Letter"

    def test_empty_string_in_options(self, formatter):
        assert formatter._resolve_page_size("ieee", {"page_size": ""}) != ""


# ---------------------------------------------------------------------------
# _resolve_line_spacing
# ---------------------------------------------------------------------------

class TestResolveLineSpacing:
    def test_from_options(self, formatter):
        assert formatter._resolve_line_spacing("ieee", {"line_spacing": 2.0}) == 2.0

    def test_from_alias(self, formatter):
        assert formatter._resolve_line_spacing("ieee", {"add_line_spacing": 1.5}) == 1.5

    def test_from_contract(self, formatter):
        formatter.contract_loader.load.return_value = {"layout": {"line_spacing": 1.15}}
        assert formatter._resolve_line_spacing("ieee", {}) == 1.15

    def test_none_value_returns_none(self, formatter):
        assert formatter._resolve_line_spacing("ieee", {}) is None

    def test_empty_string_returns_none(self, formatter):
        assert formatter._resolve_line_spacing("ieee", {"line_spacing": ""}) is None

    def test_false_returns_none(self, formatter):
        assert formatter._resolve_line_spacing("ieee", {"line_spacing": False}) is None

    def test_invalid_string_returns_none(self, formatter):
        result = formatter._resolve_line_spacing("ieee", {"line_spacing": "not-a-number"})
        assert result is None

    def test_negative_returns_none(self, formatter):
        assert formatter._resolve_line_spacing("ieee", {"line_spacing": -1}) is None

    def test_zero_returns_none(self, formatter):
        assert formatter._resolve_line_spacing("ieee", {"line_spacing": 0}) is None


# ---------------------------------------------------------------------------
# _prepare_references
# ---------------------------------------------------------------------------

class TestPrepareReferences:
    def test_no_references_noop(self, formatter):
        doc = PipelineDocument(document_id="d", blocks=[], references=[])
        formatter._prepare_references(doc, "ieee")
        formatter.reference_formatter.format_reference.assert_not_called()

    def test_already_formatted_skipped(self, formatter):
        ref = Reference(
            reference_id="r1", block_id="b1", formatted_text="Already formatted.",
            raw_text="raw", citation_key="ck", index=0,
        )
        doc = PipelineDocument(document_id="d", blocks=[], references=[ref])
        formatter._prepare_references(doc, "ieee")
        formatter.reference_formatter.format_reference.assert_not_called()

    def test_empty_formatted_text_replaced(self, formatter):
        formatter.reference_formatter.format_reference.return_value = "Formatted Ref"
        ref = Reference(reference_id="r1", block_id="b1", formatted_text="", raw_text="raw",
                        citation_key="ck", index=0)
        doc = PipelineDocument(document_id="d", blocks=[], references=[ref])
        formatter._prepare_references(doc, "ieee")
        assert ref.formatted_text == "Formatted Ref"

    def test_whitespace_only_skipped(self, formatter):
        ref = Reference(reference_id="r1", block_id="b1", formatted_text="   ", raw_text="raw",
                        citation_key="ck", index=0)
        doc = PipelineDocument(document_id="d", blocks=[], references=[ref])
        formatter._prepare_references(doc, "ieee")
        formatter.reference_formatter.format_reference.assert_called_once()

    def test_exception_falls_back_to_raw(self, formatter):
        formatter.reference_formatter.format_reference.side_effect = Exception("boom")
        ref = Reference(reference_id="r1", block_id="b1", formatted_text="", raw_text="fallback",
                        citation_key="ck", index=0)
        doc = PipelineDocument(document_id="d", blocks=[], references=[ref])
        formatter._prepare_references(doc, "ieee")
        assert ref.formatted_text == "fallback"


# ---------------------------------------------------------------------------
# _render_equation
# ---------------------------------------------------------------------------

class TestRenderEquation:
    def test_creates_omath_para(self, formatter, word_doc):
        eq = Equation(equation_id="e1", index=0, text="E=mc^2")
        formatter._render_equation(word_doc, eq)
        xml = word_doc.paragraphs[0]._p.xml
        assert "oMathPara" in xml
        assert "E=mc^2" in xml

    def test_center_aligned(self, formatter, word_doc):
        eq = Equation(equation_id="e1", index=0, text="x=y")
        formatter._render_equation(word_doc, eq)
        assert word_doc.paragraphs[0].paragraph_format.alignment == 1

    def test_empty_text_uses_space(self, formatter, word_doc):
        eq = Equation(equation_id="e1", index=0, text="")
        formatter._render_equation(word_doc, eq)
        xml = word_doc.paragraphs[0]._p.xml
        assert xml is not None

    def test_with_number(self, formatter, word_doc):
        eq = Equation(equation_id="e1", index=0, text="a+b", number="1")
        formatter._render_equation(word_doc, eq)
        assert "(1)" in word_doc.paragraphs[0].text

    def test_equation_number_none(self, formatter, word_doc):
        eq = Equation(equation_id="e1", index=0, text="c=d", number=None)
        formatter._render_equation(word_doc, eq)
        assert "(" not in word_doc.paragraphs[0].text

    def test_error_swallowed(self, formatter, word_doc):
        eq = MagicMock()
        eq.text = "test"
        eq.number = None
        with patch.object(formatter, "_render_equation", wraps=formatter._render_equation):
            result = formatter._render_equation(word_doc, eq)
            assert result is None  # safe_function returns fallback_value


# ---------------------------------------------------------------------------
# _apply_initial_layout
# ---------------------------------------------------------------------------

class TestApplyInitialLayout:
    def test_sets_margins(self, formatter, word_doc):
        formatter.contract_loader.load.return_value = {
            "layout": {"margins": {"top": 2.0, "bottom": 1.5, "left": 1.25, "right": 1.0}}
        }
        formatter._apply_initial_layout(word_doc, "ieee")
        s = word_doc.sections[0]
        assert s.top_margin == Inches(2.0)
        assert s.bottom_margin == Inches(1.5)
        assert s.left_margin == Inches(1.25)
        assert s.right_margin == Inches(1.0)

    def test_empty_layout_returns_early(self, formatter, word_doc):
        formatter.contract_loader.load.return_value = {}
        formatter._apply_initial_layout(word_doc, "ieee")
        s = word_doc.sections[0]
        assert s.top_margin == Inches(1.0)

    def test_no_layout_key(self, formatter, word_doc):
        formatter.contract_loader.load.return_value = {"styles": {}}
        formatter._apply_initial_layout(word_doc, "ieee")
        s = word_doc.sections[0]
        assert s.top_margin == Inches(1.0)


# ---------------------------------------------------------------------------
# _get_target_columns
# ---------------------------------------------------------------------------

class TestGetTargetColumns:
    def test_default_columns(self, formatter):
        formatter.contract_loader.load.return_value = {"layout": {"default_columns": 1}}
        block = Block(block_id="b1", index=0, block_type=BlockType.BODY, text="body")
        assert formatter._get_target_columns(block, "ieee") == 1

    def test_section_override(self, formatter):
        formatter.contract_loader.load.return_value = {
            "layout": {"default_columns": 1, "section_overrides": {"abstract": 2}}
        }
        block = Block(block_id="b1", index=0, block_type=BlockType.BODY, text="body", section_name="abstract")
        assert formatter._get_target_columns(block, "ieee") == 2

    def test_section_override_substring_match(self, formatter):
        formatter.contract_loader.load.return_value = {
            "layout": {"default_columns": 1, "section_overrides": {"intro": 2}}
        }
        block = Block(block_id="b1", index=0, block_type=BlockType.BODY, text="body", section_name="introduction")
        assert formatter._get_target_columns(block, "ieee") == 2

    def test_no_layout_returns_1(self, formatter):
        formatter.contract_loader.load.return_value = {}
        block = Block(block_id="b1", index=0, block_type=BlockType.BODY, text="body")
        assert formatter._get_target_columns(block, "ieee") == 1


# ---------------------------------------------------------------------------
# _apply_page_size
# ---------------------------------------------------------------------------

class TestApplyPageSize:
    def test_letter(self, formatter, word_doc):
        formatter._apply_page_size(word_doc, "Letter")
        s = word_doc.sections[0]
        assert s.page_width == Inches(8.5)
        assert s.page_height == Inches(11)

    def test_a4(self, formatter, word_doc):
        from docx.shared import Mm
        formatter._apply_page_size(word_doc, "A4")
        s = word_doc.sections[0]
        assert abs(s.page_width - Mm(210)) < 10000  # Allow emu rounding
        assert abs(s.page_height - Mm(297)) < 10000

    def test_legal(self, formatter, word_doc):
        formatter._apply_page_size(word_doc, "Legal")
        s = word_doc.sections[0]
        assert s.page_width == Inches(8.5)
        assert s.page_height == Inches(14)

    def test_unknown_falls_back_to_letter(self, formatter, word_doc):
        formatter._apply_page_size(word_doc, "Unknown")
        s = word_doc.sections[0]
        assert s.page_width == Inches(8.5)


# ---------------------------------------------------------------------------
# _add_cover_page
# ---------------------------------------------------------------------------

class TestAddCoverPage:
    def test_adds_title_author_date(self, formatter, word_doc, minimal_doc):
        formatter._add_cover_page(word_doc, minimal_doc)
        text = " ".join(p.text for p in word_doc.paragraphs)
        assert "Test Title" in text
        assert "Alice" in text
        assert "20" in text  # year in date

    def test_center_aligned(self, formatter, word_doc, minimal_doc):
        formatter._add_cover_page(word_doc, minimal_doc)
        assert word_doc.paragraphs[0].alignment == 1

    def test_unknown_author_fallback(self, formatter, word_doc):
        doc = PipelineDocument(document_id="d", metadata=DocumentMetadata(title="T"))
        formatter._add_cover_page(word_doc, doc)
        assert "Unknown Author" in word_doc.paragraphs[0].text

    def test_no_title_uses_filename(self, formatter, word_doc):
        doc = PipelineDocument(
            document_id="d", metadata=DocumentMetadata(), original_filename="paper.docx"
        )
        formatter._add_cover_page(word_doc, doc)
        assert "paper.docx" in word_doc.paragraphs[0].text

    def test_page_break_after_cover(self, formatter, word_doc, minimal_doc):
        formatter._add_cover_page(word_doc, minimal_doc)
        assert len(word_doc.paragraphs) >= 2


# ---------------------------------------------------------------------------
# _add_table_of_contents
# ---------------------------------------------------------------------------

class TestAddTableOfContents:
    def test_adds_toc_field(self, formatter, word_doc):
        formatter._add_table_of_contents(word_doc)
        xml = word_doc._body._element.xml
        assert 'TOC \\o "1-3" \\h \\z \\u' in xml

    def test_toc_heading_added(self, formatter, word_doc):
        formatter._add_table_of_contents(word_doc)
        assert "Table of Contents" in word_doc.paragraphs[0].text

    def test_prepend_moves_to_front(self, formatter, word_doc):
        word_doc.add_paragraph("Existing")
        formatter._add_table_of_contents(word_doc, prepend=True)
        assert word_doc.paragraphs[0].text == "Table of Contents"

    def test_without_page_break(self, formatter, word_doc):
        formatter._add_table_of_contents(word_doc, add_page_break=False)
        xml = word_doc._body._element.xml
        assert "Table of Contents" in xml

    def test_prepend_removes_original_location(self, formatter, word_doc):
        word_doc.add_paragraph("First")
        formatter._add_table_of_contents(word_doc, prepend=True)
        # Should only have 1 TOC heading now
        toc_count = sum(1 for p in word_doc.paragraphs if p.text == "Table of Contents")
        assert toc_count == 1


# ---------------------------------------------------------------------------
# _add_page_numbers
# ---------------------------------------------------------------------------

class TestAddPageNumbers:
    def test_adds_page_field_to_footer(self, formatter, word_doc):
        formatter._add_page_numbers(word_doc)
        footer = word_doc.sections[0].footer
        xml = footer._element.xml
        assert "PAGE" in xml

    def test_skips_if_already_has_page_field(self, formatter, word_doc):
        footer = word_doc.sections[0].footer
        p = footer.paragraphs[0]
        run = p.add_run()
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "begin")
        run._r.append(fld)
        instr = OxmlElement("w:instrText")
        instr.text = "PAGE"
        run._r.append(instr)
        formatter._add_page_numbers(word_doc)
        xml = footer._element.xml
        assert xml.count("PAGE") == 1

    def test_center_aligned(self, formatter, word_doc):
        formatter._add_page_numbers(word_doc)
        assert word_doc.sections[0].footer.paragraphs[0].alignment == 1


# ---------------------------------------------------------------------------
# _add_page_borders
# ---------------------------------------------------------------------------

class TestAddPageBorders:
    def test_adds_pgborders(self, formatter, word_doc):
        formatter._add_page_borders(word_doc)
        xml = word_doc.sections[0]._sectPr.xml
        assert "w:pgBorders" in xml

    def test_removes_existing_borders(self, formatter, word_doc):
        sec_pr = word_doc.sections[0]._sectPr
        existing = OxmlElement("w:pgBorders")
        sec_pr.append(existing)
        formatter._add_page_borders(word_doc)
        pg_borders = word_doc.sections[0]._sectPr.xpath("./w:pgBorders")
        assert len(pg_borders) == 1  # replaced

    def test_offset_from_page(self, formatter, word_doc):
        formatter._add_page_borders(word_doc)
        xml = word_doc.sections[0]._sectPr.xml
        assert 'offsetFrom="page"' in xml


# ---------------------------------------------------------------------------
# _add_line_numbers
# ---------------------------------------------------------------------------

class TestAddLineNumbers:
    def test_adds_ln_num_type(self, formatter, word_doc):
        formatter._add_line_numbers(word_doc)
        xml = word_doc.sections[0]._sectPr.xml
        assert "w:lnNumType" in xml

    def test_reuses_existing_element(self, formatter, word_doc):
        sec_pr = word_doc.sections[0]._sectPr
        existing = OxmlElement("w:lnNumType")
        sec_pr.append(existing)
        formatter._add_line_numbers(word_doc, count_by=5)
        xml = word_doc.sections[0]._sectPr.xml
        assert xml.count("w:lnNumType") == 1
        assert 'countBy="5' in xml

    def test_count_by_forced_positive(self, formatter, word_doc):
        formatter._add_line_numbers(word_doc, count_by=-3)
        xml = word_doc.sections[0]._sectPr.xml
        assert 'countBy="1"' in xml


# ---------------------------------------------------------------------------
# _paragraph_has_field_code
# ---------------------------------------------------------------------------

class TestParagraphHasFieldCode:
    def test_field_present(self, formatter, word_doc):
        p = word_doc.add_paragraph()
        run = p.add_run()
        instr = OxmlElement("w:instrText")
        instr.text = "PAGE"
        run._r.append(instr)
        assert formatter._paragraph_has_field_code(p, "PAGE") is True

    def test_field_absent(self, formatter, word_doc):
        p = word_doc.add_paragraph("Hello")
        assert formatter._paragraph_has_field_code(p, "PAGE") is False

    def test_none_paragraph(self, formatter):
        assert formatter._paragraph_has_field_code(None, "PAGE") is False


# ---------------------------------------------------------------------------
# _remove_paragraph & _prepend_paragraph
# ---------------------------------------------------------------------------

class TestRemoveParagraph:
    def test_removes_from_body(self, formatter, word_doc):
        p = word_doc.add_paragraph("Remove me")
        assert len(word_doc.paragraphs) == 1
        formatter._remove_paragraph(p)
        assert len(word_doc.paragraphs) == 0

    def test_none_noop(self, formatter):
        formatter._remove_paragraph(None)  # should not raise


class TestPrependParagraph:
    def test_prepends_to_empty_doc(self, formatter, word_doc):
        p = formatter._prepend_paragraph(word_doc, "Hello")
        assert word_doc.paragraphs[0].text == "Hello"

    def test_prepends_before_existing(self, formatter, word_doc):
        word_doc.add_paragraph("Second")
        formatter._prepend_paragraph(word_doc, "First")
        assert word_doc.paragraphs[0].text == "First"

    def test_with_alignment(self, formatter, word_doc):
        p = formatter._prepend_paragraph(word_doc, "Center", alignment=1)
        assert p.alignment == 1

    def test_with_style(self, formatter, word_doc):
        p = formatter._prepend_paragraph(word_doc, "Title", style="Title")
        assert p.text == "Title"

    def test_style_exception_fallback(self, formatter, word_doc):
        with patch.object(word_doc, "add_paragraph", side_effect=[Exception("no style"), word_doc.add_paragraph()]):
            p = formatter._prepend_paragraph(word_doc, "T", style="Missing")
            assert p.text == "T"


# ---------------------------------------------------------------------------
# _document_contains_text
# ---------------------------------------------------------------------------

class TestDocumentContainsText:
    def test_text_found(self, formatter, word_doc):
        word_doc.add_paragraph("Hello World")
        assert formatter._document_contains_text(word_doc, "Hello") is True

    def test_text_not_found(self, formatter, word_doc):
        word_doc.add_paragraph("Goodbye")
        assert formatter._document_contains_text(word_doc, "Hello") is False

    def test_empty_needle(self, formatter, word_doc):
        assert formatter._document_contains_text(word_doc, "") is False

    def test_case_insensitive(self, formatter, word_doc):
        word_doc.add_paragraph("Hello World")
        assert formatter._document_contains_text(word_doc, "hello") is True


# ---------------------------------------------------------------------------
# _build_footnote_lookup
# ---------------------------------------------------------------------------

class TestBuildFootnoteLookup:
    def test_no_footnote_blocks(self, formatter):
        doc = PipelineDocument(document_id="d", blocks=[
            Block(block_id="b1", index=0, block_type=BlockType.BODY, text="Hello"),
        ])
        assert formatter._build_footnote_lookup(doc) == {}

    def test_footnote_block_included(self, formatter):
        doc = PipelineDocument(document_id="d", blocks=[
            Block(block_id="b1", index=0, block_type=BlockType.FOOTNOTE, text="A note",
                  metadata={"footnote_id": "fn1"}),
        ])
        lookup = formatter._build_footnote_lookup(doc)
        assert "fn1" in lookup
        assert lookup["fn1"]["text"] == "A note"
        assert lookup["fn1"]["word_id"] == 1

    def test_is_footnote_metadata(self, formatter):
        doc = PipelineDocument(document_id="d", blocks=[
            Block(block_id="b1", index=0, block_type=BlockType.BODY, text="Note text",
                  metadata={"is_footnote": True, "footnote_id": "f1"}),
        ])
        lookup = formatter._build_footnote_lookup(doc)
        assert "f1" in lookup

    def test_empty_text_skipped(self, formatter):
        doc = PipelineDocument(document_id="d", blocks=[
            Block(block_id="b1", index=0, block_type=BlockType.FOOTNOTE, text="",
                  metadata={"footnote_id": "fn1"}),
        ])
        assert formatter._build_footnote_lookup(doc) == {}

    def test_whitespace_only_skipped(self, formatter):
        doc = PipelineDocument(document_id="d", blocks=[
            Block(block_id="b1", index=0, block_type=BlockType.FOOTNOTE, text="   ",
                  metadata={"footnote_id": "fn1"}),
        ])
        assert formatter._build_footnote_lookup(doc) == {}

    def test_duplicate_raw_id_skipped(self, formatter):
        doc = PipelineDocument(document_id="d", blocks=[
            Block(block_id="b1", index=0, block_type=BlockType.FOOTNOTE, text="First",
                  metadata={"footnote_id": "dup"}),
            Block(block_id="b2", index=1, block_type=BlockType.FOOTNOTE, text="Second",
                  metadata={"footnote_id": "dup"}),
        ])
        lookup = formatter._build_footnote_lookup(doc)
        assert len(lookup) == 1
        assert lookup["dup"]["text"] == "First"

    def test_uses_endnote_id_when_footnote_id_missing(self, formatter):
        doc = PipelineDocument(document_id="d", blocks=[
            Block(block_id="b1", index=0, block_type=BlockType.FOOTNOTE, text="Note",
                  metadata={"endnote_id": "en1"}),
        ])
        lookup = formatter._build_footnote_lookup(doc)
        assert "en1" in lookup

    def test_sequential_word_ids(self, formatter):
        doc = PipelineDocument(document_id="d", blocks=[
            Block(block_id="b1", index=0, block_type=BlockType.FOOTNOTE, text="First",
                  metadata={"footnote_id": "f1"}),
            Block(block_id="b2", index=2, block_type=BlockType.FOOTNOTE, text="Second",
                  metadata={"footnote_id": "f2"}),
            Block(block_id="b3", index=1, block_type=BlockType.BODY, text="ignore"),
        ])
        lookup = formatter._build_footnote_lookup(doc)
        assert lookup["f1"]["word_id"] == 1
        assert lookup["f2"]["word_id"] == 2

    def test_no_id_falls_back_to_counter(self, formatter):
        doc = PipelineDocument(document_id="d", blocks=[
            Block(block_id="b1", index=0, block_type=BlockType.FOOTNOTE, text="Note",
                  metadata={"is_footnote": True}),
        ])
        lookup = formatter._build_footnote_lookup(doc)
        key = list(lookup.keys())[0]
        assert lookup[key]["word_id"] == 1


# ---------------------------------------------------------------------------
# _add_hyperlink
# ---------------------------------------------------------------------------

class TestAddHyperlink:
    def test_adds_hyperlink_element(self, formatter, word_doc):
        p = word_doc.add_paragraph()
        with patch.object(p.part, "relate_to", return_value="rId99"):
            formatter._add_hyperlink(p, "Click Here", "https://example.com")
        xml = p._p.xml
        assert "w:hyperlink" in xml
        assert "Click Here" in xml

    def test_uses_relate_to(self, formatter, word_doc):
        p = word_doc.add_paragraph()
        with patch.object(p.part, "relate_to", return_value="rId99") as mock_relate:
            formatter._add_hyperlink(p, "Link", "https://example.com")
        mock_relate.assert_called_with("https://example.com", ANY, is_external=True)


# ---------------------------------------------------------------------------
# _append_footnote_reference
# ---------------------------------------------------------------------------

class TestAppendFootnoteReference:
    def test_appends_footnote_ref(self, formatter, word_doc):
        p = word_doc.add_paragraph()
        formatter._append_footnote_reference(p, 1)
        xml = p._p.xml
        assert "w:footnoteReference" in xml
        assert 'w:id="1"' in xml

    def test_run_has_superscript(self, formatter, word_doc):
        p = word_doc.add_paragraph()
        formatter._append_footnote_reference(p, 2)
        xml = p._p.xml
        assert 'w:val="superscript"' in xml


# ---------------------------------------------------------------------------
# _write_inline_content
# ---------------------------------------------------------------------------

class TestWriteInlineContent:
    def test_writes_plain_text(self, formatter, word_doc):
        p = word_doc.add_paragraph()
        formatter._write_inline_content(p, "Hello", [], [], {})
        assert p.text == "Hello"

    def test_writes_with_hyperlink(self, formatter, word_doc):
        p = word_doc.add_paragraph()
        with patch.object(formatter, "_add_hyperlink"):
            formatter._write_inline_content(
                p, "Visit Example", [{"text": "Example", "url": "https://example.com"}], [], {}
            )
        assert "Visit " in p.text

    def test_hyperlink_no_label_skipped(self, formatter, word_doc):
        p = word_doc.add_paragraph()
        with patch.object(formatter, "_add_hyperlink") as mock_add:
            formatter._write_inline_content(
                p, "text", [{"text": "", "url": "https://example.com"}], [], {}
            )
        mock_add.assert_not_called()

    def test_hyperlink_no_url_skipped(self, formatter, word_doc):
        p = word_doc.add_paragraph()
        with patch.object(formatter, "_add_hyperlink") as mock_add:
            formatter._write_inline_content(
                p, "text", [{"text": "Label", "url": ""}], [], {}
            )
        mock_add.assert_not_called()

    def test_footnote_ref_appended(self, formatter, word_doc):
        p = word_doc.add_paragraph()
        footnote_lookup = {"1": {"word_id": 5, "text": "Note"}}
        with patch.object(formatter, "_append_footnote_reference") as mock_append:
            formatter._write_inline_content(p, "text", [], ["1"], footnote_lookup)
        mock_append.assert_called_with(p, 5)

    def test_unknown_footnote_ref_skipped(self, formatter, word_doc):
        p = word_doc.add_paragraph()
        with patch.object(formatter, "_append_footnote_reference") as mock_append:
            formatter._write_inline_content(p, "text", [], ["unknown"], {})
        mock_append.assert_not_called()

    def test_empty_text_still_writes_run(self, formatter, word_doc):
        p = word_doc.add_paragraph()
        formatter._write_inline_content(p, "", [], [], {})
        assert p.text == ""

    def test_hyperlink_after_text(self, formatter, word_doc):
        p = word_doc.add_paragraph()
        with patch.object(formatter, "_add_hyperlink"):
            formatter._write_inline_content(
                p, "prefixLink", [{"text": "Link", "url": "https://x.com"}], [], {}
            )
        assert "prefix" in p.text

    def test_remaining_text_after_hyperlink(self, formatter, word_doc):
        p = word_doc.add_paragraph()
        with patch.object(formatter, "_add_hyperlink"):
            formatter._write_inline_content(
                p, "Click Here now", [{"text": "Here", "url": "https://x.com"}], [], {}
            )
        assert " now" in p.text


# ---------------------------------------------------------------------------
# _clear_paragraph_content
# ---------------------------------------------------------------------------

class TestClearParagraphContent:
    def test_removes_all_but_pPr(self, formatter, word_doc):
        p = word_doc.add_paragraph("Hello")
        initial_children = list(p._p)
        assert len(initial_children) >= 1
        formatter._clear_paragraph_content(p)
        remaining = list(p._p)
        for child in remaining:
            assert child.tag == qn("w:pPr")


# ---------------------------------------------------------------------------
# _find_matching_paragraph
# ---------------------------------------------------------------------------

class TestFindMatchingParagraph:
    def test_exact_match(self, formatter, word_doc):
        word_doc.add_paragraph("Hello World")
        result = formatter._find_matching_paragraph(word_doc, "Hello World", set())
        assert result is not None
        assert result.text == "Hello World"

    def test_partial_match(self, formatter, word_doc):
        word_doc.add_paragraph("Hello World Here")
        result = formatter._find_matching_paragraph(word_doc, "Hello", set())
        assert result is not None

    def test_used_paragraph_skipped(self, formatter, word_doc):
        p = word_doc.add_paragraph("Hello World")
        word_doc.add_paragraph("Hello World")  # second copy not in used set
        used = {id(p)}
        result = formatter._find_matching_paragraph(word_doc, "Hello World", used)
        assert result is not None
        assert id(result) != id(p)

    def test_empty_needle_returns_none(self, formatter, word_doc):
        assert formatter._find_matching_paragraph(word_doc, "", set()) is None

    def test_whitespace_only_normalized(self, formatter, word_doc):
        word_doc.add_paragraph("Hello  World")
        result = formatter._find_matching_paragraph(word_doc, "Hello World", set())
        assert result is not None


# ---------------------------------------------------------------------------
# _replace_paragraph_inline_content
# ---------------------------------------------------------------------------

class TestReplaceParagraphInlineContent:
    def test_clears_and_writes(self, formatter, word_doc):
        p = word_doc.add_paragraph("Old text")
        block = Block(block_id="b1", index=0, block_type=BlockType.BODY, text="New text")
        with patch.object(formatter, "_clear_paragraph_content") as mock_clear:
            with patch.object(formatter, "_write_inline_content") as mock_write:
                formatter._replace_paragraph_inline_content(p, block, {})
        mock_clear.assert_called_with(p)
        mock_write.assert_called_with(p, "New text", [], [], {})


# ---------------------------------------------------------------------------
# _remove_static_page_number_placeholders
# ---------------------------------------------------------------------------

class TestRemoveStaticPageNumberPlaceholders:
    def test_removes_page_n_pattern(self, formatter, word_doc):
        word_doc.add_paragraph("Page 1")
        word_doc.add_paragraph("Normal text")
        formatter._remove_static_page_number_placeholders(word_doc)
        texts = [p.text for p in word_doc.paragraphs]
        assert "Page 1" not in texts
        assert "Normal text" in texts

    def test_removes_page_123(self, formatter, word_doc):
        word_doc.add_paragraph("Page 123")
        formatter._remove_static_page_number_placeholders(word_doc)
        assert len(word_doc.paragraphs) == 0

    def test_case_insensitive(self, formatter, word_doc):
        word_doc.add_paragraph("PAGE 1")
        formatter._remove_static_page_number_placeholders(word_doc)
        assert len(word_doc.paragraphs) == 0

    def test_no_match_preserved(self, formatter, word_doc):
        word_doc.add_paragraph("Page One")
        word_doc.add_paragraph("Keep me")
        formatter._remove_static_page_number_placeholders(word_doc)
        assert len(word_doc.paragraphs) == 2


# ---------------------------------------------------------------------------
# _remove_static_toc_block
# ---------------------------------------------------------------------------

class TestRemoveStaticTocBlock:
    def test_removes_toc_text_only(self, formatter, word_doc):
        word_doc.add_paragraph("Table of Contents")
        formatter._remove_static_toc_block(word_doc)
        assert len(word_doc.paragraphs) == 0

    def test_removes_toc_and_entries(self, formatter, word_doc):
        word_doc.add_paragraph("Table of Contents")
        word_doc.add_paragraph("1. Introduction")
        word_doc.add_paragraph("2. Methods")
        formatter._remove_static_toc_block(word_doc)
        assert len(word_doc.paragraphs) == 0

    def test_stops_at_non_entry(self, formatter, word_doc):
        word_doc.add_paragraph("Table of Contents")
        word_doc.add_paragraph("1. Intro")
        word_doc.add_paragraph("Real content")
        formatter._remove_static_toc_block(word_doc)
        assert len(word_doc.paragraphs) == 1

    def test_empty_paragraphs_between_entries_removed(self, formatter, word_doc):
        word_doc.add_paragraph("Table of Contents")
        word_doc.add_paragraph("")
        word_doc.add_paragraph("1. Intro")
        formatter._remove_static_toc_block(word_doc)
        assert len(word_doc.paragraphs) == 0

    def test_no_toc_noop(self, formatter, word_doc):
        word_doc.add_paragraph("Hello")
        formatter._remove_static_toc_block(word_doc)
        assert len(word_doc.paragraphs) == 1

    def test_case_sensitive(self, formatter, word_doc):
        word_doc.add_paragraph("table of contents")
        formatter._remove_static_toc_block(word_doc)
        assert len(word_doc.paragraphs) == 0


# ---------------------------------------------------------------------------
# _ensure_dynamic_toc
# ---------------------------------------------------------------------------

class TestEnsureDynamicToc:
    def test_adds_if_missing(self, formatter, word_doc):
        formatter._ensure_dynamic_toc(word_doc)
        xml = word_doc._body._element.xml
        assert 'TOC \\o "1-3" \\h \\z \\u' in xml

    def test_skips_if_present(self, formatter, word_doc):
        formatter._add_table_of_contents(word_doc)
        with patch.object(formatter, "_add_table_of_contents") as mock_add:
            formatter._ensure_dynamic_toc(word_doc)
        mock_add.assert_not_called()


# ---------------------------------------------------------------------------
# _install_post_save_hook
# ---------------------------------------------------------------------------

class TestInstallPostSaveHook:
    def test_installs_hook(self, formatter, word_doc):
        lookup = {"1": {"word_id": 1, "text": "Note"}}
        formatter._install_post_save_hook(word_doc, lookup)
        assert hasattr(word_doc, "_scholarform_save_hook_installed")
        assert word_doc._scholarform_save_hook_installed is True

    def test_empty_lookup_skips(self, formatter, word_doc):
        formatter._install_post_save_hook(word_doc, {})
        assert not hasattr(word_doc, "_scholarform_save_hook_installed")

    def test_already_installed_skips(self, formatter, word_doc):
        lookup = {"1": {"word_id": 1, "text": "N"}}
        formatter._install_post_save_hook(word_doc, lookup)
        original_save = word_doc.save
        formatter._install_post_save_hook(word_doc, lookup)
        assert word_doc.save is original_save

    def test_save_calls_patch(self, formatter, word_doc, tmp_path):
        lookup = {"1": {"word_id": 1, "text": "Note"}}
        formatter._install_post_save_hook(word_doc, lookup)
        out = tmp_path / "test.docx"
        with patch.object(formatter, "_patch_saved_docx_with_footnotes") as mock_patch:
            word_doc.save(str(out))
        mock_patch.assert_called_once()


# ---------------------------------------------------------------------------
# _patch_docx_payload
# ---------------------------------------------------------------------------

class TestPatchDocxPayload:
    def _make_docx_bytes(self, has_footnote_ref=True):
        buf = io.BytesIO()
        with ZipFile(buf, "w") as zf:
            doc_xml = (
                '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:body><w:p></w:p></w:body></w:document>'
            )
            if has_footnote_ref:
                doc_xml = doc_xml.replace("</w:p>", '<w:footnoteReference w:id="1"/></w:p>')
            zf.writestr("word/document.xml", doc_xml.encode())
            zf.writestr("[Content_Types].xml", b'<?xml version="1.0"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"></Types>')
            zf.writestr("word/_rels/document.xml.rels", b'<?xml version="1.0"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"></Relationships>')
            zf.writestr("word/settings.xml", b'<?xml version="1.0"?><w:settings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"></w:settings>')
        return buf.getvalue()

    def test_no_footnote_ref_returns_unchanged(self, formatter):
        payload = self._make_docx_bytes(has_footnote_ref=False)
        result = formatter._patch_docx_payload(payload, {"1": {"word_id": 1, "text": "N"}})
        assert result == payload

    def test_footnote_ref_adds_footnotes_part(self, formatter):
        payload = self._make_docx_bytes(has_footnote_ref=True)
        result = formatter._patch_docx_payload(payload, {"1": {"word_id": 1, "text": "Note"}})
        with ZipFile(io.BytesIO(result), "r") as zf:
            assert "word/footnotes.xml" in zf.namelist()


# ---------------------------------------------------------------------------
# _build_footnotes_part
# ---------------------------------------------------------------------------

class TestBuildFootnotesPart:
    def test_creates_footnotes_xml(self, formatter):
        lookup = {"1": {"word_id": 1, "text": "First note"}}
        xml_bytes = formatter._build_footnotes_part(lookup)
        root = ET.fromstring(xml_bytes)
        notes = root.findall("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}footnote")
        texts = [node.text or "" for node in root.iter() if node.text]
        assert any("First note" in t for t in texts)
        assert len(notes) >= 3  # separator + continuation + at least 1 content

    def test_includes_separator_and_continuation(self, formatter):
        xml_bytes = formatter._build_footnotes_part({"1": {"word_id": 1, "text": "N"}})
        root = ET.fromstring(xml_bytes)
        types = [fn.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type") or ""
                 for fn in root]
        assert "separator" in types
        assert "continuationSeparator" in types

    def test_sorted_by_word_id(self, formatter):
        lookup = {
            "2": {"word_id": 3, "text": "Third"},
            "1": {"word_id": 1, "text": "First"},
        }
        xml_bytes = formatter._build_footnotes_part(lookup)
        xml_str = xml_bytes.decode("utf-8")
        first_pos = xml_str.index("First")
        third_pos = xml_str.index("Third")
        assert first_pos < third_pos


# ---------------------------------------------------------------------------
# _patch_content_types
# ---------------------------------------------------------------------------

class TestPatchContentTypes:
    def test_adds_override(self, formatter):
        xml = b'<?xml version="1.0"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"></Types>'
        result = formatter._patch_content_types(xml)
        assert "/word/footnotes.xml" in result.decode("utf-8")

    def test_skips_if_already_present(self, formatter):
        xml = (
            b'<?xml version="1.0"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            b'<Override PartName="/word/footnotes.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>'
            b"</Types>"
        )
        result = formatter._patch_content_types(xml)
        assert result.decode("utf-8").count("Override") == 1


# ---------------------------------------------------------------------------
# _patch_document_relationships
# ---------------------------------------------------------------------------

class TestPatchDocumentRelationships:
    def test_adds_footnotes_relationship(self, formatter):
        xml = b'<?xml version="1.0"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"></Relationships>'
        result = formatter._patch_document_relationships(xml)
        text = result.decode("utf-8")
        assert "footnotes" in text
        assert "footnotes.xml" in text

    def test_skips_if_already_present(self, formatter):
        xml = (
            b'<?xml version="1.0"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            b'<Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" Target="footnotes.xml"/>'
            b"</Relationships>"
        )
        result = formatter._patch_document_relationships(xml)
        root = ET.fromstring(result)
        rels = root.findall("./{*}Relationship")
        assert len(rels) == 1

    def test_empty_input_creates_root(self, formatter):
        result = formatter._patch_document_relationships(b"")
        assert b"Relationships" in result

    def test_uses_next_available_id(self, formatter):
        xml = (
            b'<?xml version="1.0"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            b'<Relationship Id="rId1" Type="other" Target="other.xml"/>'
            b"</Relationships>"
        )
        result = formatter._patch_document_relationships(xml)
        assert b'rId2' in result


# ---------------------------------------------------------------------------
# _patch_settings_xml
# ---------------------------------------------------------------------------

class TestPatchSettingsXml:
    def test_adds_footnote_pr(self, formatter):
        xml = b'<?xml version="1.0"?><w:settings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"></w:settings>'
        result = formatter._patch_settings_xml(xml)
        text = result.decode("utf-8")
        assert "footnotePr" in text
        assert "decimal" in text

    def test_skips_if_already_present(self, formatter):
        xml = (
            b'<?xml version="1.0"?><w:settings xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            b"<w:footnotePr><w:numFmt w:val=\"decimal\"/></w:footnotePr>"
            b"</w:settings>"
        )
        result = formatter._patch_settings_xml(xml)
        root = ET.fromstring(result)
        fps = root.findall("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}footnotePr")
        assert len(fps) == 1

    def test_empty_input_returns_as_is(self, formatter):
        assert formatter._patch_settings_xml(b"") == b""


# ---------------------------------------------------------------------------
# _set_columns
# ---------------------------------------------------------------------------

class TestSetColumns:
    def test_sets_single_column(self, formatter, word_doc):
        formatter._set_columns(word_doc.sections[0], 1)
        cols = word_doc.sections[0]._sectPr.xpath("./w:cols")
        assert cols[0].get(qn("w:num")) == "1"

    def test_two_columns_adds_space(self, formatter, word_doc):
        formatter._set_columns(word_doc.sections[0], 2)
        cols = word_doc.sections[0]._sectPr.xpath("./w:cols")
        assert cols[0].get(qn("w:num")) == "2"
        assert cols[0].get(qn("w:space")) == "720"

    def test_reuses_existing_cols_element(self, formatter, word_doc):
        sect_pr = word_doc.sections[0]._sectPr
        existing = sect_pr.xpath("./w:cols")
        if existing:
            existing[0].set(qn("w:num"), "3")
        else:
            cols = OxmlElement("w:cols")
            cols.set(qn("w:num"), "3")
            sect_pr.append(cols)
        formatter._set_columns(word_doc.sections[0], 2)
        all_cols = word_doc.sections[0]._sectPr.xpath("./w:cols")
        assert len(all_cols) == 1
        assert all_cols[0].get(qn("w:num")) == "2"


# ---------------------------------------------------------------------------
# _load_contract
# ---------------------------------------------------------------------------

class TestLoadContract:
    def test_missing_file_returns_empty(self, formatter):
        with patch("os.path.exists", return_value=False):
            assert formatter._load_contract("/nonexistent") == {}

    def test_loads_yaml(self, formatter, tmp_path):
        contract_file = tmp_path / "contract.yaml"
        contract_file.write_text("layout:\n  margins:\n    top: 2\n")
        yaml_content = contract_file.read_bytes()
        mock_file = MagicMock()
        mock_file.__enter__.return_value = io.StringIO(yaml_content.decode())
        with patch("os.path.exists", return_value=True):
            with patch("builtins.open", return_value=mock_file):
                result = formatter._load_contract(str(contract_file))
        assert result["layout"]["margins"]["top"] == 2

    def test_exception_returns_empty(self, formatter):
        with patch("os.path.exists", return_value=True), \
             patch("builtins.open", side_effect=PermissionError("denied")):
            result = formatter._load_contract("/path/contract.yaml")
        assert result == {}


# ---------------------------------------------------------------------------
# _is_bullet_list_item / _is_numbered_list_item / _clean_list_text
# ---------------------------------------------------------------------------

class TestListDetection:
    @pytest.mark.parametrize("marker", ["\u2022", "-", "*", "\u00b7", "\u25e6", "\u25aa", "\u25ab"])
    def test_bullet_detected(self, formatter, marker):
        assert formatter._is_bullet_list_item(f"{marker} item") is True

    def test_bullet_with_leading_space(self, formatter):
        assert formatter._is_bullet_list_item("  - item") is True

    def test_not_bullet(self, formatter):
        assert formatter._is_bullet_list_item("Regular text") is False

    def test_empty_not_bullet(self, formatter):
        assert formatter._is_bullet_list_item("") is False

    @pytest.mark.parametrize("prefix", ["1. ", "1) ", "a. ", "a) ", "i. ", "I. "])
    def test_numbered_detected(self, formatter, prefix):
        assert formatter._is_numbered_list_item(f"{prefix}item") is True

    def test_roman_numeral(self, formatter):
        assert formatter._is_numbered_list_item("iii. item") is True

    def test_not_numbered(self, formatter):
        assert formatter._is_numbered_list_item("Regular text") is False

    def test_empty_not_numbered(self, formatter):
        assert formatter._is_numbered_list_item("") is False

    def test_clean_bullet(self, formatter):
        assert formatter._clean_list_text("- Hello") == "Hello"

    def test_clean_numbered(self, formatter):
        assert formatter._clean_list_text("1. Hello") == "Hello"

    def test_clean_numbered_paren(self, formatter):
        assert formatter._clean_list_text("1) Hello") == "Hello"


# ---------------------------------------------------------------------------
# _calculate_image_size
# ---------------------------------------------------------------------------

class TestCalculateImageSize:
    def test_fits_naturally(self, formatter):
        fig = Figure(figure_id="f1", index=0, width=480, height=320)
        w, h = formatter._calculate_image_size(fig)
        assert w is not None

    def test_scales_down_when_too_wide(self, formatter):
        fig = Figure(figure_id="f1", index=0, width=10000, height=500)
        w, h = formatter._calculate_image_size(fig)
        assert w <= Inches(6.5)

    def test_scales_up_when_too_small(self, formatter):
        fig = Figure(figure_id="f1", index=0, width=50, height=25)
        w, h = formatter._calculate_image_size(fig)
        assert w >= Inches(2.0)

    def test_no_dimensions_returns_default(self, formatter):
        fig = Figure(figure_id="f1", index=0, width=None, height=None)
        w, h = formatter._calculate_image_size(fig)
        assert w == Inches(5.0)
        assert h is None

    def test_height_constrained(self, formatter):
        fig = Figure(figure_id="f1", index=0, width=100, height=5000)
        w, h = formatter._calculate_image_size(fig)
        assert h <= Inches(9.0)

    def test_zero_width_does_not_crash(self, formatter):
        fig = Figure(figure_id="f1", index=0, width=0, height=100)
        w, h = formatter._calculate_image_size(fig)
        assert w is not None


# ---------------------------------------------------------------------------
# _render_figure
# ---------------------------------------------------------------------------

class TestRenderFigure:
    def test_export_path_image(self, formatter, word_doc, tmp_path):
        img_path = tmp_path / "test.png"
        img_path.write_bytes(b"fake_png_bytes")
        fig = Figure(
            figure_id="f1", index=0, number=1, export_path=str(img_path),
            width=96, height=96, caption_text="Test Caption",
        )
        with patch.object(formatter, "_calculate_image_size", return_value=(Inches(3), Inches(3))):
            formatter._render_figure(word_doc, fig, 1)
        assert any("Test Caption" in p.text for p in word_doc.paragraphs)

    def test_image_data_path(self, formatter, word_doc):
        fig = Figure(
            figure_id="f1", index=0, number=1, image_data=b"fake_png_bytes",
            width=96, height=96, caption_text="Data Caption",
        )
        with patch.object(formatter, "_calculate_image_size", return_value=(Inches(3), Inches(3))):
            formatter._render_figure(word_doc, fig, 1)
        assert any("Data Caption" in p.text for p in word_doc.paragraphs)

    def test_no_image_data_placeholder(self, formatter, word_doc):
        fig = Figure(figure_id="f1", index=0, number=1, caption_text="No image")
        formatter._render_figure(word_doc, fig, 1)
        text = " ".join(p.text for p in word_doc.paragraphs)
        assert "Placeholder" in text

    def test_caption_with_existing_prefix(self, formatter, word_doc):
        fig = Figure(
            figure_id="f1", index=0, number=1, image_data=b"data",
            width=96, height=96, caption_text="Figure 1: My Fig",
        )
        with patch.object(formatter, "_calculate_image_size", return_value=(Inches(3), Inches(3))):
            formatter._render_figure(word_doc, fig, 1)
        text = " ".join(p.text for p in word_doc.paragraphs)
        assert "My Fig" in text

    def test_caption_without_prefix(self, formatter, word_doc):
        fig = Figure(
            figure_id="f1", index=0, number=1, image_data=b"data",
            width=96, height=96, caption_text="My Fig",
        )
        with patch.object(formatter, "_calculate_image_size", return_value=(Inches(3), Inches(3))):
            formatter._render_figure(word_doc, fig, 1)
        text = " ".join(p.text for p in word_doc.paragraphs)
        assert "Figure 1:" in text

    def test_export_path_failure_falls_back(self, formatter, word_doc):
        fig = Figure(
            figure_id="f1", index=0, export_path="/nonexistent/image.png",
            caption_text="Fallback",
        )
        formatter._render_figure(word_doc, fig, 1)
        texts = [p.text for p in word_doc.paragraphs]
        combined = " ".join(texts)
        assert "Fallback" in combined or "failed" in combined or "Placeholder" in combined


# ---------------------------------------------------------------------------
# _render_block
# ---------------------------------------------------------------------------

class TestRenderBlock:
    def test_renders_normal_paragraph(self, formatter, word_doc):
        formatter.style_mapper.get_style_name.return_value = "Normal"
        block = Block(block_id="b1", index=0, block_type=BlockType.BODY, text="Hello world")
        formatter._render_block(word_doc, block, "ieee")
        assert len(word_doc.paragraphs) >= 1

    def test_skips_empty_block_with_figure(self, formatter, word_doc):
        block = Block(block_id="b1", index=0, block_type=BlockType.BODY, text="",
                      metadata={"has_figure": True})
        formatter._render_block(word_doc, block, "ieee")
        assert len(word_doc.paragraphs) == 0

    def test_skips_empty_block_with_equation(self, formatter, word_doc):
        block = Block(block_id="b1", index=0, block_type=BlockType.BODY, text="",
                      metadata={"has_equation": True})
        formatter._render_block(word_doc, block, "ieee")
        assert len(word_doc.paragraphs) == 0

    def test_skips_truly_empty_block(self, formatter, word_doc):
        block = Block(block_id="b1", index=0, block_type=BlockType.BODY, text="")
        formatter._render_block(word_doc, block, "ieee")
        assert len(word_doc.paragraphs) == 0

    def test_bullet_list_renders_with_style(self, formatter, word_doc):
        block = Block(block_id="b1", index=0, block_type=BlockType.BODY, text="- Item")
        with patch.object(formatter, "_is_bullet_list_item", return_value=True):
            with patch.object(formatter, "_clean_list_text", return_value="Item"):
                formatter._render_block(word_doc, block, "ieee")
        assert len(word_doc.paragraphs) == 1

    def test_numbered_list_renders_with_style(self, formatter, word_doc):
        block = Block(block_id="b1", index=0, block_type=BlockType.BODY, text="1. Item")
        with patch.object(formatter, "_is_numbered_list_item", return_value=True):
            with patch.object(formatter, "_clean_list_text", return_value="Item"):
                formatter._render_block(word_doc, block, "ieee")
        assert len(word_doc.paragraphs) == 1

    def test_exception_falls_back(self, formatter, word_doc):
        formatter.style_mapper.get_style_name.side_effect = Exception("no style")
        block = Block(block_id="b1", index=0, block_type=BlockType.BODY, text="Hello")
        formatter._render_block(word_doc, block, "ieee")
        assert len(word_doc.paragraphs) == 1

    def test_exception_with_empty_text_returns_none(self, formatter, word_doc):
        formatter.style_mapper.get_style_name.side_effect = Exception("no style")
        block = Block(block_id="b1", index=0, block_type=BlockType.BODY, text="")
        result = formatter._render_block(word_doc, block, "ieee")
        assert result is None


# ---------------------------------------------------------------------------
# _apply_spacing_from_contract
# ---------------------------------------------------------------------------

class TestApplySpacingFromContract:
    def test_heading_spacing_applied(self, formatter, word_doc):
        formatter.contract_loader.load.return_value = {
            "layout": {"spacing": {"heading": {"before": 12, "after": 6}}}
        }
        block = Block(block_id="b1", index=0, block_type=BlockType.HEADING_1, text="H1")
        p = word_doc.add_paragraph()
        formatter._apply_spacing_from_contract(p, block, "ieee")
        assert p.paragraph_format.space_before == Pt(12)
        assert p.paragraph_format.space_after == Pt(6)

    def test_paragraph_spacing_applied(self, formatter, word_doc):
        formatter.contract_loader.load.return_value = {
            "layout": {"spacing": {"paragraph": {"before": 3, "after": 6}}}
        }
        block = Block(block_id="b1", index=0, block_type=BlockType.BODY, text="Body")
        p = word_doc.add_paragraph()
        formatter._apply_spacing_from_contract(p, block, "ieee")
        assert p.paragraph_format.space_before == Pt(3)

    def test_no_spacing_rules(self, formatter, word_doc):
        formatter.contract_loader.load.return_value = {"layout": {"spacing": {}}}
        block = Block(block_id="b1", index=0, block_type=BlockType.BODY, text="Body")
        p = word_doc.add_paragraph()
        formatter._apply_spacing_from_contract(p, block, "ieee")
        before = p.paragraph_format.space_before
        assert before is None or before.pt == 0

    def test_line_spacing_from_contract(self, formatter, word_doc):
        formatter.contract_loader.load.return_value = {
            "layout": {"spacing": {"paragraph": {"before": 0, "after": 0}}, "line_spacing": 2.0}
        }
        block = Block(block_id="b1", index=0, block_type=BlockType.BODY, text="Body")
        p = word_doc.add_paragraph()
        formatter._apply_spacing_from_contract(p, block, "ieee")
        assert p.paragraph_format.line_spacing == 2.0


# ---------------------------------------------------------------------------
# _apply_global_line_spacing
# ---------------------------------------------------------------------------

class TestApplyGlobalLineSpacing:
    def test_applies_to_all_paragraphs(self, formatter, word_doc):
        word_doc.add_paragraph("One")
        word_doc.add_paragraph("Two")
        with patch.object(formatter, "_resolve_line_spacing", return_value=1.5):
            formatter._apply_global_line_spacing(word_doc, "ieee", {})
        for p in word_doc.paragraphs:
            assert p.paragraph_format.line_spacing == 1.5

    def test_none_skipped(self, formatter, word_doc):
        word_doc.add_paragraph("One")
        with patch.object(formatter, "_resolve_line_spacing", return_value=None):
            formatter._apply_global_line_spacing(word_doc, "ieee", {})
        for p in word_doc.paragraphs:
            assert p.paragraph_format.line_spacing is None


# ---------------------------------------------------------------------------
# _prepend_front_matter
# ---------------------------------------------------------------------------

class TestPrependFrontMatter:
    def test_as_cover_page(self, formatter, word_doc):
        doc = PipelineDocument(
            document_id="d",
            metadata=DocumentMetadata(title="My Paper", authors=["Alice", "Bob"],
                                       affiliations=["MIT"]),
        )
        formatter._prepend_front_matter(word_doc, doc, as_cover_page=True)
        text = " ".join(p.text for p in word_doc.paragraphs)
        assert "My Paper" in text
        assert "Alice, Bob" in text
        assert "MIT" in text

    def test_inline_mode(self, formatter, word_doc):
        doc = PipelineDocument(
            document_id="d",
            metadata=DocumentMetadata(title="Inline", authors=["Author"],
                                       affiliations=["Org"]),
        )
        formatter._prepend_front_matter(word_doc, doc, as_cover_page=False)
        text = " ".join(p.text for p in word_doc.paragraphs)
        assert "Inline" in text
        assert "Author" in text

    def test_no_affiliations(self, formatter, word_doc):
        doc = PipelineDocument(
            document_id="d",
            metadata=DocumentMetadata(title="T", authors=["A"]),
        )
        formatter._prepend_front_matter(word_doc, doc, as_cover_page=True)
        pagination_runs = any("August" in p.text or "202" in p.text for p in word_doc.paragraphs)
        assert pagination_runs

    def test_unknown_author_fallback(self, formatter, word_doc):
        doc = PipelineDocument(
            document_id="d",
            metadata=DocumentMetadata(title="T"),
        )
        formatter._prepend_front_matter(word_doc, doc, as_cover_page=False)
        assert "Unknown Author" in " ".join(p.text for p in word_doc.paragraphs)


# ---------------------------------------------------------------------------
# _post_process_template_render
# ---------------------------------------------------------------------------

class TestPostProcessTemplateRender:
    def test_no_docx_returns_early(self, formatter):
        rendered = MagicMock()
        rendered.docx = None
        formatter._post_process_template_render(rendered, MagicMock(), "ieee", {})
        # Should complete without error; no docx calls to check

    def test_applies_layout_and_page_size(self, formatter):
        doc = MagicMock()
        body_mock = MagicMock()
        body_mock.element.xml = ""
        doc.docx = MagicMock(spec=WordDocument)
        doc.docx.paragraphs = []
        doc.docx._body = body_mock
        doc.docx.sections = []
        source_doc = MagicMock()
        source_doc.metadata.title = ""
        source_doc.original_filename = ""
        source_doc.blocks = []
        source_doc.figures = []
        source_doc.equations = []
        source_doc.tables = []

        with patch.object(formatter, "_apply_initial_layout") as mock_layout:
            with patch.object(formatter, "_apply_page_size") as mock_size:
                formatter._post_process_template_render(doc, source_doc, "ieee", {},
                                                         footnote_lookup={})
        mock_layout.assert_called_once()
        mock_size.assert_called_once()

    def test_adds_front_matter_when_title_missing(self, formatter):
        doc = MagicMock()
        body_mock = MagicMock()
        body_mock.element.xml = ""
        doc.docx = MagicMock(spec=WordDocument)
        doc.docx.paragraphs = []
        doc.docx._body = body_mock
        doc.docx.sections = []
        source_doc = MagicMock()
        source_doc.metadata.title = "Missing Title"
        source_doc.original_filename = "paper"
        source_doc.blocks = []
        source_doc.figures = []
        source_doc.equations = []
        source_doc.tables = []

        with patch.object(formatter, "_document_contains_text", return_value=False):
            with patch.object(formatter, "_prepend_front_matter") as mock_front:
                with patch.object(formatter, "_prepend_paragraph") as mock_pp:
                    formatter._post_process_template_render(doc, source_doc, "ieee", {},
                                                             footnote_lookup={})
        mock_front.assert_called_once()
