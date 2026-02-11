import os
import sys
import json
import shutil
import argparse
from pathlib import Path

# Add backend to path (Depth 3: manual_tests/normal/phase1_identification)
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../../../')))

# We need these for sample creation if requested
from docx import Document as DocxDocument
from docx.shared import Pt

from app.pipeline.input_conversion.converter import InputConverter, ConversionError
from app.pipeline.parsing.parser import DocxParser

def create_sample_docx(output_path: str) -> str:
    """Legacy utility from app/manual_tests/run_parser.py"""
    print(f"\n📝 Creating sample DOCX file: {output_path}")
    doc = DocxDocument()
    
    # Title
    title = doc.add_paragraph("Sample Academic Paper")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)
    
    # Abstract
    doc.add_paragraph()
    abstract_heading = doc.add_paragraph("Abstract")
    abstract_heading.runs[0].bold = True
    
    doc.add_paragraph("This is a sample abstract for document processing tests.")
    
    # Section 1
    doc.add_paragraph()
    doc.add_paragraph("1. Introduction").runs[0].bold = True
    doc.add_paragraph("This is the introduction section text.")
    
    # Table 1
    doc.add_paragraph("Table 1: Test Data")
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = "Key"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Test"
    table.cell(1, 1).text = "Pass"
    
    doc.save(output_path)
    print(f"✓ Sample DOCX created successfully")
    return output_path

def main():
    parser = argparse.ArgumentParser(description="ScholarForm AI Input Conversion Test")
    parser.add_argument("input_path", nargs='?', help="Path to input file")
    parser.add_argument("--create-sample", action="store_true", help="Create a sample DOCX")
    args = parser.parse_args()

    print(f"\n🚀 PHASE 1: INPUT CONVERSION & PARSING")
    
    # 1. Handle Sample Creation
    if args.create_sample or not args.input_path:
        sample_path = "manual_tests/sample_inputs/sample_paper.docx"
        os.makedirs("manual_tests/sample_inputs", exist_ok=True)
        input_path = create_sample_docx(sample_path)
    else:
        input_path = args.input_path

    if not os.path.exists(input_path):
        print(f"❌ ERROR: File not found: {input_path}")
        return

    # 2. Conversion Phase (Legacy run_conversion.py logic)
    converter = InputConverter()
    job_id = "test_conversion_job"
    
    print(f"\n[1] Testing Conversion for: {Path(input_path).suffix}...")
    try:
        docx_path = converter.convert_to_docx(input_path, job_id)
        print(f"✅ Conversion successful: {docx_path}")
    except ConversionError as e:
        print(f"❌ Conversion failed (expected for non-docx/md/pdf): {e}")
        return
    except Exception as e:
        print(f"❌ Unexpected Error: {e}")
        return

    # 3. Parsing Phase (Legacy run_parser.py logic)
    print(f"\n[2] Testing Initial DOCX Parsing...")
    parser = DocxParser()
    try:
        blocks = parser.parse_docx(docx_path)
        print(f"✅ Parsed {len(blocks)} blocks.")
    except Exception as e:
        print(f"❌ Parsing Error: {e}")
        return

    # 4. Save Output
    output_dir = Path("manual_tests/outputs")
    output_dir.mkdir(parents=True, exist_ok=True)
    output_file = output_dir / "01_blocks.json"
    
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump({
            "summary": {
                "input_file": str(input_path),
                "total_blocks": len(blocks),
                "parsing_status": "success"
            },
            "blocks": [b.model_dump() for b in blocks]
        }, f, indent=2)
    
    print(f"\n--- Analysis Summary ---")
    print(f"Total Blocks: {len(blocks)}")
    print(f"------------------------")
    print(f"\n✅ SUCCESS: Result saved to {output_file}")

if __name__ == "__main__":
    main()
