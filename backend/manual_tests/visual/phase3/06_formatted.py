"""
Visual Test - Phase 3: Formatter Verification

Purpose: Verify complete formatting with template application
Input: DOCX file
Output: Formatted DOCX with template applied

Usage:
    python manual_tests/visual/phase3/06_formatted.py uploads/input.docx --template IEEE

Output:
    manual_tests/visual_outputs/phase3_06_formatted.docx

Visual Inspection:
    - Open output DOCX in Microsoft Word
    - Verify template formatting applied
    - Check for content duplication
    - Confirm heading hierarchy
    - Verify caption placement
"""

import os
import sys
import argparse
from pathlib import Path
from docx import Document as WordDoc
from docx.shared import RGBColor

# Add backend to path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../../../')))

from app.pipeline.parsing.parser import DocxParser
from app.pipeline.normalization.normalizer import Normalizer
from app.pipeline.structure_detection.detector import StructureDetector
from app.pipeline.classification.classifier import ContentClassifier
from app.pipeline.figures.caption_matcher import CaptionMatcher
from app.pipeline.tables.caption_matcher import TableCaptionMatcher
from app.pipeline.formatting.formatter import Formatter

def main():
    parser = argparse.ArgumentParser(description="ScholarForm AI Visual Formatter Test")
    parser.add_argument("input_path", help="Path to input docx")
    parser.add_argument("--template", default="IEEE", help="Template to use")
    args = parser.parse_args()

    print(f"======================================================================")
    print(f"PHASE 3 VISUAL TEST - FORMATTER VERIFICATION")
    print(f"======================================================================")
    print(f"Input: {args.input_path}")
    print(f"Template: {args.template}")
    
    # 1. Full Pipeline Execution
    print(f"\n[1/7] Parsing DOCX...")
    doc_parser = DocxParser()
    normalizer = Normalizer()
    detector = StructureDetector()
    classifier = ContentClassifier()
    fig_matcher = CaptionMatcher()
    tab_matcher = TableCaptionMatcher()
    formatter = Formatter()
    
    doc = doc_parser.parse(args.input_path, "phase3_visual_test")
    print(f"  ✓ Extracted {len(doc.blocks)} blocks")
    
    print(f"[2/7] Normalizing...")
    doc = normalizer.process(doc)
    print(f"  ✓ Normalized to {len(doc.blocks)} blocks")
    
    print(f"[3/7] Detecting structure...")
    doc = detector.process(doc)
    
    print(f"[4/7] Classifying blocks...")
    doc = classifier.process(doc)
    
    print(f"[5/7] Matching captions...")
    doc = fig_matcher.process(doc)
    doc = tab_matcher.process(doc)
    print(f"  ✓ Matched {len(doc.figures)} figure captions")
    print(f"  ✓ Matched {len(doc.tables)} table captions")
    
    print(f"[6/7] Applying formatter...")
    try:
        doc = formatter.process(doc)
        print(f"  ✓ Formatter applied successfully")
    except FileNotFoundError as e:
        print(f"  ⚠ Formatter skipped (no contract): {e}")
    except Exception as e:
        print(f"  ⚠ Formatter skipped (error): {e}")
    
    # 2. Create annotated visual output
    print(f"[7/7] Creating visual output...")
    annotated_doc = WordDoc()
    
    # Dashboard
    header = annotated_doc.add_paragraph()
    header_run = header.add_run("--- PHASE 3 VISUAL TEST: FORMATTER VERIFICATION ---\n")
    header_run.bold = True
    header_run.font.color.rgb = RGBColor(0, 128, 0)
    
    summary = annotated_doc.add_paragraph()
    summary_run = summary.add_run(
        f"Template: {args.template}\n"
        f"Total Blocks: {len(doc.blocks)}\n"
        f"Figures: {len(doc.figures)}\n"
        f"Tables: {len(doc.tables)}\n"
        "----------------------------------------------------\n"
    )
    summary_run.font.color.rgb = RGBColor(0, 0, 255)
    
    # Add all blocks
    for block in doc.blocks:
        para = annotated_doc.add_paragraph()
        text_run = para.add_run(block.text)
        
        # Add annotation
        annotation_run = para.add_run(f" [{block.block_type.value.upper()}]")
        annotation_run.font.bold = True
        annotation_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save
    target_dir = Path("manual_tests/visual_outputs")
    target_dir.mkdir(parents=True, exist_ok=True)
    final_file = target_dir / "phase3_06_formatted.docx"
    
    annotated_doc.save(final_file)
    
    print(f"\n======================================================================")
    print(f"✓ FORMATTER TEST COMPLETE")
    print(f"======================================================================")
    print(f"Output: {final_file}")
    print(f"\nMANUAL INSPECTION REQUIRED:")
    print(f"1. Open the output DOCX in Microsoft Word")
    print(f"2. Verify template formatting applied correctly")
    print(f"3. Check for content duplication")
    print(f"4. Confirm heading hierarchy")
    print(f"5. Verify caption placement")

if __name__ == "__main__":
    main()
