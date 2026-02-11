import os
import sys
from pathlib import Path
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

# Add backend to path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../../../')))

from app.pipeline.parsing.parser import DocxParser
from app.pipeline.normalization.normalizer import Normalizer
from app.pipeline.structure_detection.detector import StructureDetector
from app.pipeline.classification.classifier import ContentClassifier
from app.pipeline.figures.caption_matcher import CaptionMatcher
from app.models.block import BlockType

def annotate_visual(input_path, output_path):
    print(f"Annotating {input_path}")
    
    # 1. Pipeline Execution
    parser = DocxParser()
    normalizer = Normalizer()
    detector = StructureDetector()
    classifier = ContentClassifier()
    matcher = CaptionMatcher()
    
    doc = parser.parse(input_path, "visual_test")
    doc = normalizer.process(doc)
    doc = detector.process(doc)
    doc = classifier.process(doc)
    doc = matcher.process(doc)
    
    blocks = doc.blocks
    figures = doc.figures
    
    # Apply numbering
    for i, fig in enumerate(figures, 1):
        fig.number = i
        fig.metadata["figure_number"] = i
    
    # Find insertion points (simple heuristic: after caption)
    anchors_found = 0
    for fig in figures:
        if fig.caption_block_id:
            caption_block = next((b for b in blocks if b.block_id == fig.caption_block_id), None)
            if caption_block:
                fig.metadata["anchor_index"] = caption_block.index + 1
                anchors_found += 1
    
    # 2. Annotate DOCX
    annotated_doc = Document(input_path)
    
    # Dashboard insertion
    if annotated_doc.paragraphs:
        first_para = annotated_doc.paragraphs[0]
        
        first_para.insert_paragraph_before("------------------------------------------------\n")
        first_para.insert_paragraph_before(f"Insertion Anchors Found: {anchors_found}")
        first_para.insert_paragraph_before(f"Figures: {len(figures)}")
        first_para.insert_paragraph_before(f"Total Blocks: {len(blocks)}")
        
        header_p = first_para.insert_paragraph_before(
            "--- QA VISUAL DASHBOARD: PHASE 1 (FIGURE INSERTION) ---"
        )
        if header_p.runs:
            header_p.runs[0].bold = True
    
    # Highlight figure captions and mark insertion points
    for fig in figures:
        if fig.caption_block_id:
            caption_block = next((b for b in blocks if b.block_id == fig.caption_block_id), None)
            if caption_block:
                idx = caption_block.index
                if 0 <= idx < len(annotated_doc.paragraphs):
                    para = annotated_doc.paragraphs[idx]
                    for run in para.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                    anchor_idx = fig.metadata.get("anchor_index", "?")
                    para.add_run(f" [FIG {fig.number} → INSERT AT {anchor_idx}]").font.bold = True
                    
    annotated_doc.save(output_path)
    return output_path

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python 05c_figure_insertion.py <docx_path>")
        sys.exit(1)
        
    input_docx = sys.argv[1]
    output_docx = "manual_tests/visual_outputs/05c_figure_insertion_annotated.docx"
    
    os.makedirs("manual_tests/visual_outputs", exist_ok=True)
    annotate_visual(input_docx, output_docx)
    print(f"Done. Visual test saved to {output_docx}")
