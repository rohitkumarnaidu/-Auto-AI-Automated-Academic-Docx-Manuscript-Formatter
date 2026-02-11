"""
Visual Test - Stage 3: Figures and Tables
Purpose: Verify figure and table caption matching visually
Input: DOCX file
Output: Annotated DOCX with figure and table captions highlighted
"""

import sys
import os
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

# Add backend to path (Depth 3)
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../../../')))

from app.pipeline.parsing.parser import DocxParser
from app.pipeline.normalization.normalizer import Normalizer as TextNormalizer
from app.pipeline.structure_detection.detector import StructureDetector
from app.pipeline.classification.classifier import ContentClassifier
from app.pipeline.figures.caption_matcher import CaptionMatcher
from app.pipeline.tables.caption_matcher import TableCaptionMatcher
from app.models.block import BlockType

def add_comment_to_paragraph(paragraph, comment_text):
    """Add inline note."""
    run = paragraph.add_run(f" [{comment_text}]")
    run.font.color.rgb = RGBColor(0, 0, 255) # Blue

def main():
    if len(sys.argv) < 2:
        print("Usage: python 03_figures_tables.py <input.docx>")
        sys.exit(1)
        
    input_path = sys.argv[1]
    output_dir = Path(__file__).parent.parent.parent / "visual_outputs"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "03_figures_tables_annotated.docx"
    
    print("=" * 70)
    print("VISUAL TEST - STAGE 3: FIGURES AND TABLES")
    print("=" * 70)
    print(f"Input: {input_path}")
    print(f"Output: {output_path}")
    
    # Execution
    print("[1/5] Parsing & Normalizing...")
    parser = DocxParser()
    normalizer = TextNormalizer()
    doc_obj = parser.parse(input_path, "visual_test_ft")
    doc_obj = normalizer.process(doc_obj)
    
    print("[2/5] Structure & Classification...")
    detector = StructureDetector()
    classifier = ContentClassifier()
    doc_obj = detector.process(doc_obj)
    doc_obj = classifier.process(doc_obj)
    
    print("[3/5] Caption Matching (Figures & Tables)...")
    fig_matcher = CaptionMatcher()
    tab_matcher = TableCaptionMatcher()
    doc_obj = fig_matcher.process(doc_obj)
    doc_obj = tab_matcher.process(doc_obj)
    
    # Annotation
    print("[4/5] Generating annotated DOCX...")
    annotated_doc = Document()
    
    # Summary Dashboard
    summary = annotated_doc.add_paragraph()
    summary.add_run("=== FIGURES & TABLES DASHBOARD ===\n").bold = True
    summary.add_run(f"Total Blocks: {len(doc_obj.blocks)}\n")
    summary.add_run(f"Figures Extracted: {len(doc_obj.figures)}\n")
    summary.add_run(f"Tables Extracted: {len(doc_obj.tables)}\n")
    summary.add_run("==================================\n")
    summary.runs[0].font.color.rgb = RGBColor(0, 128, 0)

    highlight_count = 0
    for block in doc_obj.blocks:
        para = annotated_doc.add_paragraph()
        run = para.add_run(block.text)
        
        if block.block_type == BlockType.FIGURE_CAPTION:
            run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE
            add_comment_to_paragraph(para, "FIGURE CAPTION")
            highlight_count += 1
        elif block.block_type == BlockType.TABLE_CAPTION:
            run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE
            add_comment_to_paragraph(para, "TABLE CAPTION")
            highlight_count += 1

    annotated_doc.save(str(output_path))
    print(f"  ? Highlighted {highlight_count} captions")
    print(f"\n✅ SUCCESS: Result saved to {output_path}")

if __name__ == "__main__":
    main()
