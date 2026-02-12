"""
Visual Test - Stage 1: Parse and Structure Detection

Purpose: Verify parsing and heading detection visually
Input: DOCX file
Output: Annotated DOCX with headings highlighted

Usage: python manual_tests/visual/phase1/01_parse_and_structure.py uploads/input.docx
Output: manual_tests/visual_outputs/01_structure_annotated.docx

Visual Inspection:
- Open output DOCX in Microsoft Word
- Yellow highlights = Detected headings
- Comments show heading levels
- Red highlights = Duplicate blocks (if any)
"""

import sys
import os
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn

# Add backend to path (Depth 3: manual_tests/visual/phase1)
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../../../')))

from app.pipeline.parsing.parser import DocxParser
from app.pipeline.normalization.normalizer import Normalizer as TextNormalizer
from app.pipeline.structure_detection.detector import StructureDetector

def add_highlight(paragraph, color_rgb):
    """Add highlight to paragraph."""
    # Assuming color_rgb is WD_COLOR_INDEX (int)
    for run in paragraph.runs:
        run.font.highlight_color = color_rgb

def add_comment_to_paragraph(paragraph, comment_text):
    """Add comment text as a note (simulated via text)."""
    # Note: python-docx doesn't support true comments, so we'll add inline notes
    run = paragraph.add_run(f" [{comment_text}]")
    run.font.color.rgb = RGBColor(0, 0, 255) # Blue
    run.font.size = run.font.size // 2 if run.font.size else None

def main():
    if len(sys.argv) < 2:
        print("Usage: python 01_parse_and_structure.py <input.docx>")
        sys.exit(1)
        
    input_path = sys.argv[1]
    output_dir = Path(__file__).parent.parent.parent / "visual_outputs" # backend/manual_tests/visual_outputs
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "01_structure_annotated.docx"
    
    print("=" * 70)
    print("VISUAL TEST - STAGE 1: PARSE AND STRUCTURE DETECTION")
    print("=" * 70)
    print(f"Input: {input_path}")
    print(f"Output: {output_path}")
    print()
    
    # Stage 1: Parse
    print("[1/4] Parsing DOCX...")
    parser = DocxParser()
    doc_obj = parser.parse(input_path, document_id="visual_test_stage1")
    print(f"  ? Extracted {len(doc_obj.blocks)} blocks")
    
    # Stage 2: Normalize
    print("[2/4] Normalizing text...")
    normalizer = TextNormalizer()
    doc_obj = normalizer.process(doc_obj)
    print(f"  ? Normalized and split blocks into {len(doc_obj.blocks)}")
    
    # Stage 3: Structure Detection
    print("[3/4] Detecting structure...")
    detector = StructureDetector()
    doc_obj = detector.process(doc_obj)
    
    # Count headings
    heading_count = sum(1 for b in doc_obj.blocks if b.metadata.get('is_heading_candidate', False))
    print(f"  ? Detected {heading_count} headings")
    
    # Stage 4: Create annotated DOCX
    print("[4/4] Creating annotated DOCX...")
    
    # Create NEW document to write blocks
    annotated_doc = Document()
    
    # Add summary at the beginning
    summary_para = annotated_doc.add_paragraph()
    summary_para.add_run("=== STRUCTURE DETECTION SUMMARY ===\n").bold = True
    summary_para.add_run(f"Total Blocks: {len(doc_obj.blocks)}\n")
    summary_para.add_run(f"Headings Detected: {heading_count}\n")
    block_ids_seen = set()
    dup_count = 0
    merged_count = 0
    
    for block in doc_obj.blocks:
        is_split = block.metadata.get('split_from_original', False)
        
        # Determine if it's a "true" duplicate (same ID but not a split part)
        # Note: If normalizer splits, it might reuse ID or append suffix. 
        # Usually split parts share ID, so we check split flag.
        is_duplicate = (block.block_id in block_ids_seen) and (not is_split)
        
        if is_duplicate:
            dup_count += 1
        if is_split:
            merged_count += 1 # Tracking splits/merges roughly
            
        block_ids_seen.add(block.block_id)
    
    summary_para.add_run(f"Duplicates: {dup_count}\n")
    summary_para.add_run(f"Split/Merged Blocks: {merged_count}\n")
    summary_para.add_run("===================================\n")
    summary_para.runs[0].font.color.rgb = RGBColor(0, 128, 0)

    # Reset seen set for actual writing
    block_ids_seen = set()

    # Annotate paragraphs based on normalized blocks
    for block in doc_obj.blocks:
        para = annotated_doc.add_paragraph()
        run = para.add_run(block.text)
        
        is_split = block.metadata.get('split_from_original', False)
        is_duplicate = (block.block_id in block_ids_seen) and (not is_split)
        
        if is_duplicate:
            # Highlight duplicate in RED (WD_COLOR_INDEX.RED is 6)
            run.font.highlight_color = 6 
            add_comment_to_paragraph(para, f"DUPLICATE: {block.block_id}")
            
        elif block.metadata.get('is_heading_candidate', False):
            # Highlight headings in YELLOW (WD_COLOR_INDEX.YELLOW is 7)
            run.font.highlight_color = 7 
            level = block.metadata.get('level', '?')
            add_comment_to_paragraph(para, f"HEADING Level {level}")
            
        elif block.metadata.get('is_footnote', False):
            # Highlight footnotes in GRAY (WD_COLOR_INDEX.GRAY_25 is 16, let's use GRAY_25)
            run.font.highlight_color = 16 
            fn_id = block.metadata.get('footnote_id', '?')
            add_comment_to_paragraph(para, f"FOOTNOTE (ID: {fn_id})")
            
        elif is_split:
             # Just note it's split, maybe green highlight or just note? user says "Expected", so maybe no highlight or just note
             # User requirement: "Acceptance criteria: '1 IntroductionAcademic...' appears split correctly"
             # No specific highlight requested for split, just correct text.
             pass

        block_ids_seen.add(block.block_id)

    # Save annotated document
    annotated_doc.save(str(output_path))
    
    print()
    print("=" * 70)
    print("? STAGE 1 COMPLETE")
    print("=" * 70)
    print(f"Output saved to: {output_path}")
    print()
    print("NEXT STEP:")
    print("1. Open the output DOCX in Microsoft Word")
    print("2. Look for:")
    print(" - Yellow highlights = Detected headings")
    print(" - Blue annotations = Heading levels")
    print(" - Red text = Duplicate blocks")
    print("3. Verify headings are correctly identified")
    print("4. Report any issues before proceeding to Stage 2")
    print()

if __name__ == "__main__":
    main()
