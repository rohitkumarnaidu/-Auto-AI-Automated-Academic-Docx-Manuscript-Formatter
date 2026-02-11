"""
Visual Test - Stage 5b: Figure Numbering
Purpose: Verify sequential figure numbering visually
Input: DOCX file
Output: Annotated DOCX with figure numbers in captions
"""

import sys
import os
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

# Add backend to path (Depth 3)
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../../../')))

from app.pipeline.parsing.parser import DocxParser
from app.pipeline.normalization.normalizer import Normalizer as TextNormalizer
from app.pipeline.structure_detection.detector import StructureDetector
from app.pipeline.classification.classifier import ContentClassifier
from app.pipeline.figures.caption_matcher import CaptionMatcher

def add_comment_to_paragraph(paragraph, comment_text):
    """Add inline note."""
    run = paragraph.add_run(f" [{comment_text}]")
    run.font.color.rgb = RGBColor(0, 0, 255) # Blue

def main():
    if len(sys.argv) < 2:
        print("Usage: python 05b_figure_numbering.py <input.docx>")
        sys.exit(1)
        
    input_path = sys.argv[1]
    output_dir = Path(__file__).parent.parent.parent / "visual_outputs"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "05b_figure_numbering_annotated.docx"
    
    print("=" * 70)
    print("VISUAL TEST - STAGE 5b: FIGURE NUMBERING")
    print("=" * 70)
    print(f"Input: {input_path}")
    print(f"Output: {output_path}")
    
    # Execution
    print("[1/3] Running Pipeline stages...")
    parser = DocxParser()
    normalizer = TextNormalizer()
    detector = StructureDetector()
    classifier = ContentClassifier()
    matcher = CaptionMatcher()
    
    doc_obj = parser.parse(input_path, "visual_test_fignum")
    doc_obj = normalizer.process(doc_obj)
    doc_obj = detector.process(doc_obj)
    doc_obj = classifier.process(doc_obj)
    doc_obj = matcher.process(doc_obj)
    
    # 2. Sequential Numbering
    for i, fig in enumerate(doc_obj.figures, 1):
        fig.number = i
    
    # Annotation
    print("[2/3] Generating annotated DOCX...")
    annotated_doc = Document()
    
    summary = annotated_doc.add_paragraph()
    summary.add_run("=== FIGURE NUMBERING SUMMARY ===\n").bold = True
    summary.add_run(f"Total Blocks: {len(doc_obj.blocks)}\n")
    summary.add_run(f"Figures Numbered: {len(doc_obj.figures)}\n")
    summary.add_run("===============================\n")
    summary.runs[0].font.color.rgb = RGBColor(0, 128, 0)

    for block in doc_obj.blocks:
        para = annotated_doc.add_paragraph()
        run = para.add_run(block.text)
        
        # Check if this block is a matched caption
        fig = next((f for f in doc_obj.figures if f.caption_block_id == block.block_id), None)
        if fig:
            run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE
            add_comment_to_paragraph(para, f"FIGURE {fig.number}")

    annotated_doc.save(str(output_path))
    print(f"\n✅ SUCCESS: Result saved to {output_path}")

if __name__ == "__main__":
    main()
