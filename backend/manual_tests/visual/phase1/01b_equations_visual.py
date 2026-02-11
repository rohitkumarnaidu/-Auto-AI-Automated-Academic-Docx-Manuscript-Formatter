"""
Visual Test - Stage 1b: Equation Detection
Purpose: Verify equation detection visually
Input: DOCX file
Output: Annotated DOCX with equations highlighted in Turquoise
"""

import sys
import os
from pathlib import Path
from docx import Document
from docx.shared import RGBColor

# Add backend to path (Depth 3)
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../../../')))

from app.pipeline.parsing.parser import DocxParser
from app.pipeline.normalization.normalizer import Normalizer as TextNormalizer

def add_comment_to_paragraph(paragraph, comment_text):
    """Add inline note."""
    run = paragraph.add_run(f" [{comment_text}]")
    run.font.color.rgb = RGBColor(0, 0, 255) # Blue

def main():
    if len(sys.argv) < 2:
        print("Usage: python 01b_equations_visual.py <input.docx>")
        sys.exit(1)
        
    input_path = sys.argv[1]
    output_dir = Path(__file__).parent.parent.parent / "visual_outputs"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "01b_equations_annotated.docx"
    
    print("=" * 70)
    print("VISUAL TEST - STAGE 1b: EQUATION DETECTION")
    print("=" * 70)
    print(f"Input: {input_path}")
    print(f"Output: {output_path}")
    
    # Execution
    print("[1/2] Parsing & Normalizing...")
    parser = DocxParser()
    doc_obj = parser.parse(input_path, "visual_test_eq")
    normalizer = TextNormalizer()
    doc_obj = normalizer.process(doc_obj)
    
    # Count equations found by parser
    eq_count = len(doc_obj.equations)
    print(f"  ? Parser found {eq_count} equation objects")
    
    # Annotation
    print("[2/2] Generating annotated DOCX...")
    annotated_doc = Document()
    summary = annotated_doc.add_paragraph()
    summary.add_run("=== EQUATION DETECTION SUMMARY ===\n").bold = True
    summary.add_run(f"Total Blocks: {len(doc_obj.blocks)}\n")
    summary.add_run(f"Equation Objects: {eq_count}\n")
    summary.add_run("==================================\n")
    
    highlight_count = 0
    for block in doc_obj.blocks:
        para = annotated_doc.add_paragraph()
        run = para.add_run(block.text)
        
        # Correct detection: Match by block_id
        is_equation = any(eq.block_id == block.block_id for eq in doc_obj.equations)
            
        # Hard check on typical symbols if it's an equation block
        if not is_equation and (('=' in block.text or '+' in block.text) and len(block.text) < 100):
            # This is a bit weak, but for visual tests it helps see what MIGHT be one
             pass

        if is_equation:
            run.font.highlight_color = 3 # TURQUOISE (WD_COLOR_INDEX.TURQUOISE is 3)
            add_comment_to_paragraph(para, "EQUATION")
            highlight_count += 1

    annotated_doc.save(str(output_path))
    print(f"  ? Highlighted {highlight_count} potential equation blocks")
    print(f"\n✅ SUCCESS: Result saved to {output_path}")

if __name__ == "__main__":
    main()
