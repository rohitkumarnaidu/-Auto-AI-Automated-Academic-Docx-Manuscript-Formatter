import os
import sys
from pathlib import Path
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

# Add backend to path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../../../')))

from app.pipeline.parsing.parser import DocxParser
from app.pipeline.normalization.normalizer import Normalizer
from app.pipeline.structure_detection.detector import StructureDetector
from app.pipeline.classification.classifier import ContentClassifier
from app.pipeline.figures.caption_matcher import CaptionMatcher
from app.pipeline.tables.caption_matcher import TableCaptionMatcher
from app.pipeline.references.parser import ReferenceParser
from app.pipeline.validation.validator import DocumentValidator
from app.pipeline.formatting.formatter import Formatter
from app.models.block import BlockType

def annotate_visual(input_path, output_path):
    print(f"Annotating {input_path}")
    
    # 1. Full Pipeline Execution (Including Formatting)
    parser = DocxParser()
    normalizer = Normalizer()
    detector = StructureDetector()
    classifier = ContentClassifier()
    fig_matcher = CaptionMatcher()
    tab_matcher = TableCaptionMatcher()
    ref_parser = ReferenceParser()
    validator = DocumentValidator()
    formatter = Formatter()
    
    doc = parser.parse(input_path, "visual_test")
    doc = normalizer.process(doc)
    doc = detector.process(doc)
    doc = classifier.process(doc)
    doc = fig_matcher.process(doc)
    doc = tab_matcher.process(doc)
    doc = ref_parser.process(doc)
    doc = validator.process(doc)
    doc = formatter.process(doc)
    
    blocks = doc.blocks
    figures = doc.figures
    tables = doc.tables
    has_formatted = hasattr(doc, 'generated_doc') and doc.generated_doc is not None
    
    # 2. Annotate DOCX
    annotated_doc = Document(input_path)
    
    # Comprehensive Dashboard
    if annotated_doc.paragraphs:
        first_para = annotated_doc.paragraphs[0]
        
        first_para.insert_paragraph_before("------------------------------------------------\n")
        first_para.insert_paragraph_before(f"Formatted Document Generated: {'YES' if has_formatted else 'NO'}")
        first_para.insert_paragraph_before(f"Template: {doc.template.template_name if doc.template else 'None'}")
        first_para.insert_paragraph_before(f"Validation: {'PASS' if doc.is_valid else 'FAIL'}")
        first_para.insert_paragraph_before(f"Tables: {len(tables)}")
        first_para.insert_paragraph_before(f"Figures: {len(figures)}")
        first_para.insert_paragraph_before(f"Total Blocks: {len(blocks)}")
        
        header_p = first_para.insert_paragraph_before(
            "--- QA VISUAL DASHBOARD: FORMATTED OUTPUT (PHASE 3) ---"
        )
        if header_p.runs:
            header_p.runs[0].bold = True
    
    # Color mapping
    color_map = {
        BlockType.TITLE: WD_COLOR_INDEX.PINK,
        BlockType.AUTHOR: WD_COLOR_INDEX.BLUE,
        BlockType.AFFILIATION: WD_COLOR_INDEX.TEAL,
        BlockType.ABSTRACT_HEADING: WD_COLOR_INDEX.BRIGHT_GREEN,
        BlockType.ABSTRACT_BODY: WD_COLOR_INDEX.GREEN,
        BlockType.HEADING_1: WD_COLOR_INDEX.YELLOW,
        BlockType.HEADING_2: WD_COLOR_INDEX.YELLOW,
        BlockType.HEADING_3: WD_COLOR_INDEX.YELLOW,
        BlockType.HEADING_4: WD_COLOR_INDEX.YELLOW,
        BlockType.REFERENCES_HEADING: WD_COLOR_INDEX.YELLOW,
        BlockType.REFERENCE_ENTRY: WD_COLOR_INDEX.GRAY_25,
        BlockType.FIGURE_CAPTION: WD_COLOR_INDEX.TURQUOISE,
        BlockType.TABLE_CAPTION: WD_COLOR_INDEX.TURQUOISE,
    }
    
    # Highlight all blocks
    for block in blocks:
        idx = block.index
        if 0 <= idx < len(annotated_doc.paragraphs):
            para = annotated_doc.paragraphs[idx]
            
            if block.block_type in color_map:
                for run in para.runs:
                    run.font.highlight_color = color_map[block.block_type]
                para.add_run(f" [{block.block_type.value.upper()}]").font.bold = True
                    
    annotated_doc.save(output_path)
    return output_path

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python 06_formatted.py <docx_path>")
        sys.exit(1)
        
    input_docx = sys.argv[1]
    output_docx = "manual_tests/visual_outputs/06_formatted_annotated.docx"
    
    os.makedirs("manual_tests/visual_outputs", exist_ok=True)
    annotate_visual(input_docx, output_docx)
    print(f"Done. Visual test saved to {output_docx}")
