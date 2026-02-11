"""
Visual Test - Stage 4: References
Purpose: Verify reference parsing visually
Input: DOCX file
Output: Annotated DOCX with references highlighted
"""

import sys
import os
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

# Add backend to path (Depth 3)
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../../../')))

from app.pipeline.parsing.parser import DocxParser
from app.pipeline.normalization.normalizer import Normalizer as TextNormalizer
from app.pipeline.structure_detection.detector import StructureDetector
from app.pipeline.classification.classifier import ContentClassifier
from app.pipeline.references.parser import ReferenceParser
from app.models.block import BlockType

def add_comment_to_paragraph(paragraph, comment_text):
    """Add inline note."""
    run = paragraph.add_run(f" [{comment_text}]")
    run.font.color.rgb = RGBColor(0, 0, 255) # Blue

def main():
    if len(sys.argv) < 2:
        print("Usage: python 04_references.py <input.docx>")
        sys.exit(1)
        
    input_path = sys.argv[1]
    output_dir = Path(__file__).parent.parent.parent / "visual_outputs"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "04_references_annotated.docx"
    
    print("=" * 70)
    print("VISUAL TEST - STAGE 4: REFERENCES")
    print("=" * 70)
    print(f"Input: {input_path}")
    print(f"Output: {output_path}")
    
    # Execution
    print("[1/5] Parsing & Normalizing...")
    parser = DocxParser()
    normalizer = TextNormalizer()
    doc_obj = parser.parse(input_path, "visual_test_ref")
    doc_obj = normalizer.process(doc_obj)
    
    print("[2/5] Structure & Classification...")
    detector = StructureDetector()
    classifier = ContentClassifier()
    doc_obj = detector.process(doc_obj)
    doc_obj = classifier.process(doc_obj)
    
    print("[3/5] Reference Parsing...")
    ref_parser = ReferenceParser()
    doc_obj = ref_parser.process(doc_obj)
    
    # Count references
    ref_entries = [b for b in doc_obj.blocks if b.block_type == BlockType.REFERENCE_ENTRY]
    
    # Annotation
    print("[4/5] Generating annotated DOCX...")
    annotated_doc = Document()
    
    # Summary Dashboard
    summary = annotated_doc.add_paragraph()
    summary.add_run("=== REFERENCES DASHBOARD ===\n").bold = True
    summary.add_run(f"Total Blocks: {len(doc_obj.blocks)}\n")
    summary.add_run(f"Reference Entries: {len(ref_entries)}\n")
    summary.add_run("=============================\n")
    summary.runs[0].font.color.rgb = RGBColor(0, 128, 0)

    for block in doc_obj.blocks:
        para = annotated_doc.add_paragraph()
        run = para.add_run(block.text)
        
        if block.block_type == BlockType.REFERENCES_HEADING:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            add_comment_to_paragraph(para, "REF HEADING")
        elif block.block_type == BlockType.REFERENCE_ENTRY:
            run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
            add_comment_to_paragraph(para, "REF ENTRY")

    annotated_doc.save(str(output_path))
    print(f"\n✅ SUCCESS: Result saved to {output_path}")

if __name__ == "__main__":
    main()
