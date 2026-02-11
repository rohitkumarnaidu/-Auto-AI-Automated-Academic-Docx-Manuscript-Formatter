"""
Visual Test - Stage 6b: Table Numbering
Purpose: Verify sequential table numbering visually
Input: DOCX file
Output: Annotated DOCX with table numbers in captions
"""

import sys
import os
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

# Add backend to path (Depth 3)
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../../../')))

from app.pipeline.parsing.parser import DocxParser
from app.pipeline.normalization.normalizer import Normalizer as TextNormalizer
from app.pipeline.structure_detection.detector import StructureDetector
from app.pipeline.classification.classifier import ContentClassifier
from app.pipeline.tables.caption_matcher import TableCaptionMatcher

def add_comment_to_paragraph(paragraph, comment_text):
    """Add inline note."""
    run = paragraph.add_run(f" [{comment_text}]")
    run.font.color.rgb = RGBColor(0, 0, 255) # Blue

def main():
    if len(sys.argv) < 2:
        print("Usage: python 06b_table_numbering.py <input.docx>")
        sys.exit(1)
        
    input_path = sys.argv[1]
    output_dir = Path(__file__).parent.parent.parent / "visual_outputs"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "06b_table_numbering_annotated.docx"
    
    print("=" * 70)
    print("VISUAL TEST - STAGE 6b: TABLE NUMBERING")
    print("=" * 70)
    print(f"Input: {input_path}")
    print(f"Output: {output_path}")
    
    # Execution
    print("[1/3] Running Pipeline stages...")
    parser = DocxParser()
    normalizer = TextNormalizer()
    detector = StructureDetector()
    classifier = ContentClassifier()
    matcher = TableCaptionMatcher()
    
    doc_obj = parser.parse(input_path, "visual_test_tabnum")
    doc_obj = normalizer.process(doc_obj)
    doc_obj = detector.process(doc_obj)
    doc_obj = classifier.process(doc_obj)
    doc_obj = matcher.process(doc_obj)
    
    # 2. Sequential Numbering
    for i, tab in enumerate(doc_obj.tables, 1):
        tab.number = i
    
    # Annotation
    print("[2/3] Generating annotated DOCX...")
    annotated_doc = Document()
    
    summary = annotated_doc.add_paragraph()
    summary.add_run("=== TABLE NUMBERING SUMMARY ===\n").bold = True
    summary.add_run(f"Total Blocks: {len(doc_obj.blocks)}\n")
    summary.add_run(f"Tables Numbered: {len(doc_obj.tables)}\n")
    summary.add_run("===============================\n")
    summary.runs[0].font.color.rgb = RGBColor(0, 128, 0)

    for block in doc_obj.blocks:
        para = annotated_doc.add_paragraph()
        run = para.add_run(block.text)
        
        # Check if this block is a matched caption
        tab = next((t for t in doc_obj.tables if t.caption_block_id == block.block_id), None)
        if tab:
            run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE
            add_comment_to_paragraph(para, f"TABLE {tab.number}")

    annotated_doc.save(str(output_path))
    print(f"\n✅ SUCCESS: Result saved to {output_path}")

if __name__ == "__main__":
    main()
