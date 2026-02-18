"""
Visual Test - Stage 5: Full Pipeline
Purpose: Verify the entire pipeline visually from parsing to validation
Input: DOCX file
Output: Annotated DOCX with all markers (Headings, Classification, Figs/Tabs, Refs)
"""

import sys
import os
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX

# Add backend to path (Depth 3)
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../../../')))

from app.pipeline.parsing.parser import DocxParser
from app.pipeline.normalization.normalizer import Normalizer as TextNormalizer
from app.pipeline.structure_detection.detector import StructureDetector
from app.pipeline.classification.classifier import ContentClassifier
from app.pipeline.figures.caption_matcher import CaptionMatcher
from app.pipeline.tables.caption_matcher import TableCaptionMatcher
from app.pipeline.references.parser import ReferenceParser
from app.pipeline.validation import DocumentValidator
from app.models.block import BlockType

# Comprehensive Color mapping
COLOR_MAP = {
    BlockType.TITLE: WD_COLOR_INDEX.PINK,
    BlockType.AUTHOR: WD_COLOR_INDEX.BLUE,
    BlockType.AFFILIATION: WD_COLOR_INDEX.TEAL,
    BlockType.ABSTRACT_HEADING: WD_COLOR_INDEX.BRIGHT_GREEN,
    BlockType.ABSTRACT_BODY: WD_COLOR_INDEX.GREEN,
    BlockType.HEADING_1: WD_COLOR_INDEX.YELLOW,
    BlockType.HEADING_2: WD_COLOR_INDEX.YELLOW,
    BlockType.HEADING_3: WD_COLOR_INDEX.YELLOW,
    BlockType.HEADING_4: WD_COLOR_INDEX.YELLOW,
    BlockType.REFERENCES_HEADING: WD_COLOR_INDEX.YELLOW,
    BlockType.REFERENCE_ENTRY: WD_COLOR_INDEX.GRAY_25,
    BlockType.FIGURE_CAPTION: WD_COLOR_INDEX.TURQUOISE,
    BlockType.TABLE_CAPTION: WD_COLOR_INDEX.TURQUOISE,
}

def add_comment_to_paragraph(paragraph, comment_text):
    """Add inline note."""
    run = paragraph.add_run(f" [{comment_text}]")
    run.font.color.rgb = RGBColor(0, 0, 255) # Blue

def main():
    if len(sys.argv) < 2:
        print("Usage: python 05_full_pipeline.py <input.docx>")
        sys.exit(1)
        
    input_path = sys.argv[1]
    output_dir = Path(__file__).parent.parent.parent / "visual_outputs"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "05_full_pipeline_annotated.docx"
    
    print("=" * 70)
    print("VISUAL TEST - STAGE 5: FULL PIPELINE")
    print("=" * 70)
    print(f"Input: {input_path}")
    print(f"Output: {output_path}")
    
    # Execution
    print("[1/3] Running All Pipeline Stages...")
    parser = DocxParser()
    normalizer = TextNormalizer()
    detector = StructureDetector()
    classifier = ContentClassifier()
    fig_matcher = CaptionMatcher()
    tab_matcher = TableCaptionMatcher()
    ref_parser = ReferenceParser()
    validator = DocumentValidator()
    
    doc_obj = parser.parse(input_path, "visual_test_full")
    doc_obj = normalizer.process(doc_obj)
    doc_obj = detector.process(doc_obj)
    doc_obj = classifier.process(doc_obj)
    doc_obj = fig_matcher.process(doc_obj)
    doc_obj = tab_matcher.process(doc_obj)
    doc_obj = ref_parser.process(doc_obj)
    doc_obj = validator.process(doc_obj)
    
    # Annotation
    print("[2/3] Generating annotated DOCX...")
    annotated_doc = Document()
    
    # Summary Dashboard
    summary = annotated_doc.add_paragraph()
    summary.add_run("=== FULL PIPELINE DASHBOARD ===\n").bold = True
    summary.add_run(f"Validation: {'PASS' if doc_obj.is_valid else 'FAIL'}\n")
    summary.add_run(f"Total Blocks: {len(doc_obj.blocks)}\n")
    summary.add_run(f"Figures: {len(doc_obj.figures)}\n")
    summary.add_run(f"Tables: {len(doc_obj.tables)}\n")
    summary.add_run(f"Ref Entries: {len([b for b in doc_obj.blocks if b.block_type == BlockType.REFERENCE_ENTRY])}\n")
    summary.add_run(f"Errors: {len(doc_obj.validation_errors)}\n")
    summary.add_run(f"Warnings: {len(doc_obj.validation_warnings)}\n")
    summary.add_run("===============================\n")
    summary.runs[0].font.color.rgb = RGBColor(0, 128, 0)

    for block in doc_obj.blocks:
        para = annotated_doc.add_paragraph()
        run = para.add_run(block.text)
        
        # Apply color if mapped
        if block.block_type in COLOR_MAP:
            run.font.highlight_color = COLOR_MAP[block.block_type]
            add_comment_to_paragraph(para, block.block_type.value.upper())
        elif block.block_type != BlockType.UNKNOWN:
             add_comment_to_paragraph(para, block.block_type.value.upper())

    annotated_doc.save(str(output_path))
    print(f"\n✅ SUCCESS: Result saved to {output_path}")

if __name__ == "__main__":
    main()
