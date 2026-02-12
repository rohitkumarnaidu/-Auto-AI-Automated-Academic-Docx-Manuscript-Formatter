"""
Visual Test - Phase 2: Full Pipeline Assembly

Purpose: Verify complete pipeline assembly with visual annotations
Input: DOCX file
Output: Annotated DOCX showing full pipeline results

Usage:
    python manual_tests/visual/phase2/05_full_pipeline.py uploads/input.docx

Output:
    manual_tests/visual_outputs/phase2_05_full_pipeline_annotated.docx

Visual Inspection:
    - Open output DOCX in Microsoft Word
    - Verify all blocks are correctly processed
    - Check for RED duplicate warnings
    - Confirm no data loss or corruption
"""

import os
import sys
from pathlib import Path
from docx import Document as WordDoc
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor

# Add backend to path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "../../../")))

from app.pipeline.parsing.parser import DocxParser
from app.pipeline.normalization.normalizer import Normalizer
from app.pipeline.structure_detection.detector import StructureDetector
from app.pipeline.classification.classifier import ContentClassifier
from app.pipeline.figures.caption_matcher import CaptionMatcher
from app.pipeline.tables.caption_matcher import TableCaptionMatcher
from app.models.block import BlockType

def annotate_visual_full(input_path, output_path):
    print(f"Annotating {input_path}")
    
    # 1. Full Pipeline Execution
    parser = DocxParser()
    normalizer = Normalizer()
    detector = StructureDetector()
    classifier = ContentClassifier()
    fig_matcher = CaptionMatcher()
    tab_matcher = TableCaptionMatcher()
    
    doc = parser.parse(input_path, "phase2_visual_test")
    doc = normalizer.process(doc)
    doc = detector.process(doc)
    doc = classifier.process(doc)
    doc = fig_matcher.process(doc)
    doc = tab_matcher.process(doc)
    
    blocks = doc.blocks
    figures = doc.figures
    tables = doc.tables
    
    # 2. Create annotated DOCX
    annotated_doc = WordDoc()
    
    # Dashboard
    header = annotated_doc.add_paragraph()
    header_run = header.add_run("--- QA VISUAL DASHBOARD: PHASE 2 (FULL ASSEMBLY) ---\n")
    header_run.bold = True
    header_run.font.color.rgb = RGBColor(0, 128, 0)
    
    summary = annotated_doc.add_paragraph()
    summary_run = summary.add_run(
        f"Total Blocks: {len(blocks)}\n"
        f"Figures: {len(figures)}\n"
        f"Tables: {len(tables)}\n"
        f"Block Types: {len(set(b.block_type for b in blocks))}\n"
        "----------------------------------------------------\n"
    )
    summary_run.font.color.rgb = RGBColor(0, 0, 255)
    
    # Color mapping
    color_map = {
        BlockType.TITLE: WD_COLOR_INDEX.PINK,
        BlockType.AUTHOR: WD_COLOR_INDEX.BRIGHT_GREEN,
        BlockType.AFFILIATION: WD_COLOR_INDEX.TEAL,
        BlockType.ABSTRACT_HEADING: WD_COLOR_INDEX.RED,
        BlockType.ABSTRACT_BODY: WD_COLOR_INDEX.YELLOW,
        BlockType.HEADING_1: WD_COLOR_INDEX.YELLOW,
        BlockType.HEADING_2: WD_COLOR_INDEX.YELLOW,
        BlockType.FIGURE_CAPTION: WD_COLOR_INDEX.TURQUOISE,
        BlockType.TABLE_CAPTION: WD_COLOR_INDEX.BLUE,
        BlockType.REFERENCES_HEADING: WD_COLOR_INDEX.YELLOW,
        BlockType.REFERENCE_ENTRY: WD_COLOR_INDEX.GRAY_25,
    }
    
    # Add annotated blocks
    for block in blocks:
        para = annotated_doc.add_paragraph()
        
        # Add block text
        text_run = para.add_run(block.text)
        
        # Highlight by block type
        if block.block_type in color_map:
            text_run.font.highlight_color = color_map[block.block_type]
        
        # Add annotation
        annotation_run = para.add_run(f" [{block.block_type.value.upper()}]")
        annotation_run.font.bold = True
        annotation_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save
    annotated_doc.save(output_path)
    return output_path

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python 05_full_pipeline.py <docx_path>")
        sys.exit(1)
        
    input_docx = sys.argv[1]
    output_dir = Path("manual_tests/visual_outputs")
    output_dir.mkdir(parents=True, exist_ok=True)
    output_docx = output_dir / "phase2_05_full_pipeline_annotated.docx"
    
    annotate_visual_full(input_docx, str(output_docx))
    print(f"✓ Visual test saved to {output_docx}")

