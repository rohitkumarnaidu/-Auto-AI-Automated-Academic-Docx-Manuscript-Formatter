"""
Visual Test - Phase 2: Classification Verification After Assembly

Purpose: Re-verify classification after assembly
Input: DOCX file
Output: Annotated DOCX with block type labels

Usage:
    python manual_tests/visual/phase2/02_classification.py uploads/input.docx

Output:
    manual_tests/visual_outputs/phase2_02_classification_annotated.docx

Visual Inspection:
    - Open output DOCX in Microsoft Word
    - Color-coded annotations show block types
    - Verify no classification corruption after assembly
"""

import sys
import os
from pathlib import Path
from docx import Document
from docx.shared import RGBColor

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../../..')))

from app.pipeline.parsing.parser import DocxParser
from app.pipeline.normalization.normalizer import Normalizer as TextNormalizer
from app.pipeline.structure_detection.detector import StructureDetector
from app.pipeline.classification.classifier import ContentClassifier

def add_annotation(paragraph, text, color_rgb):
    """Add colored annotation to paragraph."""
    run = paragraph.add_run(f" [{text}]")
    run.font.color.rgb = RGBColor(*color_rgb)
    run.font.bold = True

def main():
    if len(sys.argv) < 2:
        print("Usage: python 02_classification.py <input.docx>")
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_dir = Path(__file__).parent.parent / "visual_outputs"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "phase2_02_classification_annotated.docx"
    
    print("=" * 70)
    print("VISUAL TEST - PHASE 2: CLASSIFICATION VERIFICATION")
    print("=" * 70)
    print(f"Input:  {input_path}")
    print(f"Output: {output_path}")
    print()
    
    # Run full Phase 1 pipeline
    print("[1/4] Parsing DOCX...")
    parser = DocxParser()
    doc_obj = parser.parse(input_path, document_id="phase2_classification_test")
    print(f"  ✓ Extracted {len(doc_obj.blocks)} blocks")
    
    print("[2/4] Normalizing...")
    normalizer = TextNormalizer()
    doc_obj = normalizer.process(doc_obj)
    print(f"  ✓ Normalized to {len(doc_obj.blocks)} blocks")
    
    print("[3/4] Detecting structure...")
    detector = StructureDetector()
    doc_obj = detector.process(doc_obj)
    
    print("[4/4] Classifying blocks...")
    classifier = ContentClassifier()
    doc_obj = classifier.process(doc_obj)
    
    # Count classifications
    type_counts = {}
    for block in doc_obj.blocks:
        block_type = block.block_type.value
        type_counts[block_type] = type_counts.get(block_type, 0) + 1
    
    print(f"  ✓ Classified {len(doc_obj.blocks)} blocks into {len(type_counts)} types")
    
    # Create annotated DOCX
    print()
    print("Creating annotated DOCX...")
    annotated_doc = Document()
    
    # Add summary
    summary_para = annotated_doc.add_paragraph()
    summary_run = summary_para.add_run(
        f"=== PHASE 2 CLASSIFICATION VERIFICATION ===\n"
        f"Total Blocks: {len(doc_obj.blocks)}\n"
        f"Block Types: {len(type_counts)}\n"
    )
    for block_type, count in sorted(type_counts.items()):
        summary_run.add_text(f"{block_type}: {count}\n")
    summary_run.add_text("=" * 43 + "\n")
    summary_run.font.color.rgb = RGBColor(0, 128, 0)
    summary_run.bold = True
    
    # Color map for block types
    color_map = {
        'title': (0, 0, 255),           # Blue
        'author': (0, 128, 128),        # Teal
        'affiliation': (128, 0, 128),   # Purple
        'abstract_heading': (255, 0, 0),  # Red
        'abstract_body': (255, 100, 100), # Light red
        'heading_1': (0, 128, 0),       # Green
        'heading_2': (0, 200, 0),       # Light green
        'body': (0, 0, 0),              # Black
        'figure_caption': (255, 140, 0), # Orange
        'table_caption': (255, 165, 0),  # Light orange
        'references_heading': (128, 0, 0), # Dark red
        'reference_entry': (150, 50, 50),  # Brown
    }
    
    # Add annotated blocks
    for block in doc_obj.blocks:
        para = annotated_doc.add_paragraph()
        
        # Add block text
        text_run = para.add_run(block.text)
        
        # Add annotation
        block_type = block.block_type.value
        color = color_map.get(block_type, (128, 128, 128))  # Gray for unknown
        
        annotation_run = para.add_run(f" [{block_type.upper()}]")
        annotation_run.font.color.rgb = RGBColor(*color)
        annotation_run.bold = True
    
    # Save
    annotated_doc.save(str(output_path))
    
    print()
    print("=" * 70)
    print("✅ PHASE 2 CLASSIFICATION VERIFICATION COMPLETE")
    print("=" * 70)
    print(f"Output saved to: {output_path}")
    print()
    print("VERIFICATION CHECKLIST:")
    print("1. Open the output DOCX in Microsoft Word")
    print("2. Verify block types are correctly classified")
    print("3. Confirm no classification corruption after assembly")
    print("4. Check that all sections are properly labeled")
    print()

if __name__ == "__main__":
    main()
