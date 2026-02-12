
import os
import sys
from docx import Document
from pathlib import Path
from docx.oxml.ns import qn

def dump_xml(path):
    doc = Document(path)
    for i, p in enumerate(doc.paragraphs):
        print(f"Paragraph {i}: '{p.text}'")
        print(p._element.xml)
        if i > 5: break

if __name__ == "__main__":
    from docx import Document as NewDoc
    test_path = Path("test_list_xml_dump.docx")
    doc = NewDoc()
    doc.add_paragraph("List Item", style='List Bullet')
    doc.save(test_path)
    
    dump_xml(test_path)
    if test_path.exists():
        os.remove(test_path)
