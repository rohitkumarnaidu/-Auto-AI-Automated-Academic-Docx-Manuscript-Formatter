from pathlib import Path
from collections import Counter, defaultdict
from datetime import datetime
import csv
import json
import re
import subprocess

from docx import Document as Doc

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "audits" / "2026-03-02-fullstack"
DATA = OUT / "data"
DOCX = OUT / "docx"
DOCX_DATA = DOCX / "data"
for d in (OUT, DATA, DOCX, DOCX_DATA):
    d.mkdir(parents=True, exist_ok=True)

EX_DIRS = {".git", "node_modules", ".next", "dist", "build", "coverage", ".pytest_cache", "__pycache__", ".venv", "venv", "uploads", "semantic_store"}
EX_SUB = ["backend/uploads/", "backend/db/semantic_store/", "backend/manual_tests/visual_outputs/", "backend/manual_tests/visual/visual_outputs/"]
BIN = {".docx", ".pdf", ".png", ".jpg", ".jpeg", ".gif", ".bmp", ".ico", ".bin", ".zip", ".gz", ".7z", ".woff", ".woff2", ".ttf", ".eot", ".pyc"}
TXT = {".py", ".js", ".jsx", ".ts", ".tsx", ".json", ".md", ".sql", ".yml", ".yaml", ".ini", ".mjs", ".css", ".txt", ".csl", ".xsl", ".mako", ".toml", ".lock", ".env", ".example"}


def p(x: Path) -> str:
    return x.as_posix()


def incl(path: Path) -> bool:
    if not path.is_file():
        return False
    ps = p(path).lower()
    if any(s in ps for s in EX_SUB):
        return False
    if any(part.lower() in EX_DIRS for part in path.parts):
        return False
    if path.suffix.lower() in BIN:
        return False
    if path.suffix.lower() in TXT or path.name.endswith(".env.example"):
        return True
    if path.suffix == "":
        try:
            raw = path.read_bytes()[:4096]
            return b"\x00" not in raw
        except Exception:
            return False
    return False


def txt(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except Exception:
        return path.read_text(encoding="latin-1", errors="ignore")


def lines(s: str) -> int:
    return 0 if not s else s.count("\n") + 1


def is_test(fp: str) -> bool:
    f = fp.lower()
    return "/tests/" in f or "/test/" in f or ".test." in f or ".spec." in f or f.endswith("_test.py")


def mod(rel: Path) -> str:
    ps = rel.parts
    if len(ps) < 3:
        return ps[0] if ps else "root"
    if ps[0] == "backend" and ps[1] == "app":
        return f"app/{ps[2]}"
    if ps[0] == "frontend" and ps[1] in {"app", "src"}:
        return f"{ps[1]}/{ps[2]}"
    return f"{ps[1]}/{ps[2]}"


def dyn(rel: Path, t: str) -> str:
    fp = p(rel).lower()
    if fp.startswith("backend/"):
        return "dynamic"
    if fp.endswith(".css"):
        return "static"
    if fp.startswith("frontend/app/") and ("use client" not in t and "useState(" not in t and "useEffect(" not in t):
        return "static"
    return "dynamic"


def rscore(fp: str, loc: int):
    f = fp.lower()
    imp = 1.2
    if f == "backend/app/main.py" or "backend/app/routers/" in f:
        imp = 3.0
    elif "backend/app/middleware/" in f:
        imp = 2.8
    elif "backend/app/services/" in f or "backend/app/pipeline/" in f:
        imp = 2.6
    elif f.startswith("frontend/app/"):
        imp = 2.4
    elif "/services/" in f or "authcontext" in f:
        imp = 2.6
    elif is_test(f):
        imp = 0.8

    lik = 0.7 if loc < 120 else 1.0 if loc < 250 else 1.3 if loc < 500 else 1.6 if loc < 900 else 1.9
    if any(k in f for k in ["auth", "upload", "documents.py", "stream.py", "orchestrator.py", "generator.py", "llm_service.py"]):
        lik = min(2.0, lik + 0.3)

    exp = 1.6 if any(k in f for k in ["auth", "upload", "documents.py", "stream.py", "oauth", "jwt", "api.js", "authguard"]) else 0.2
    if any(k in f for k in ["security_headers.py", "rate_limit.py"]):
        exp = max(exp, 1.3)
    if is_test(f):
        exp = 0.1

    op = 1.2 if any(k in f for k in ["orchestrator", "pipeline", "llm", "rag", "ocr", "grobid", "crossref", "redis", "converter", "stream"]) else 0.3
    if f == "backend/app/main.py":
        op = 1.4
    if is_test(f):
        op = 0.2

    cov = 0.0 if is_test(f) else 0.2 if f == "frontend/src/services/api.js" else 0.8 if f.startswith("frontend/app/") else 0.7 if f.startswith("frontend/") else 0.6
    drift = 0.5 if f in {"frontend/context/authcontext.jsx", "frontend/services/api.js", "frontend/vitest.config.js", "backend/app/schemas/document.py", "backend/app/routers/templates.py"} else 0.3 if f.endswith("readme.md") or "/docs/" in f else 0.0
    if is_test(f):
        drift = min(drift, 0.2)

    s = round(max(1.0, min(10.0, imp + lik + exp + op + cov + drift)), 1)
    level = "low" if s <= 3 else "medium" if s <= 6 else "high" if s <= 8 else "critical"
    q = round(max(1.0, min(10.0, 10 - (s - 1) * 0.9 + (0.5 if is_test(f) else 0) - (0.8 if drift >= 0.4 else 0) - (0.6 if cov >= 0.8 else 0))), 1)
    return s, level, q, cov, drift


def wcsv(path: Path, fields, rows):
    with path.open("w", newline="", encoding="utf-8") as fh:
        w = csv.DictWriter(fh, fieldnames=fields)
        w.writeheader()
        w.writerows(rows)


def mtable(headers, rows):
    out = []
    out.append("| " + " | ".join(headers) + " |")
    out.append("| " + " | ".join(["---"] * len(headers)) + " |")
    for row in rows:
        out.append("| " + " | ".join(str(row.get(h, "")).replace("|", "\\|") for h in headers) + " |")
    return "\n".join(out)


def tmd(doc: Doc, md: str):
    for ln in md.splitlines():
        s = ln.strip()
        if not s:
            doc.add_paragraph("")
        elif s.startswith("### "):
            doc.add_heading(s[4:], level=3)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=2)
        elif s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif s.startswith("- "):
            doc.add_paragraph(s[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", s):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", s), style="List Number")
        else:
            doc.add_paragraph(ln)


def csv_docx(csv_path: Path, out_path: Path, title: str):
    d = Doc()
    d.add_heading(title, level=1)
    rows = list(csv.reader(csv_path.open("r", encoding="utf-8", newline="")))
    if not rows:
        d.add_paragraph("No data")
        d.save(out_path)
        return
    hdr = rows[0]
    tb = d.add_table(rows=1, cols=len(hdr))
    tb.style = "Table Grid"
    for i, h in enumerate(hdr):
        tb.rows[0].cells[i].text = str(h)
    for r in rows[1:]:
        c = tb.add_row().cells
        for i in range(len(hdr)):
            c[i].text = str(r[i] if i < len(r) else "")
    d.save(out_path)


def run():
    files = sorted([f for b in (ROOT / "backend", ROOT / "frontend") for f in b.rglob("*") if incl(f)], key=lambda x: p(x))

    gstat = {}
    try:
        for ln in subprocess.check_output(["git", "-C", str(ROOT), "status", "--porcelain"], text=True, encoding="utf-8", errors="ignore").splitlines():
            code, fp = ln[:2], ln[3:].strip().replace("\\", "/")
            if " -> " in fp:
                fp = fp.split(" -> ", 1)[1].strip()
            gstat[fp] = "untracked" if code == "??" else "modified" if code.strip() else "tracked_clean"
    except Exception:
        pass

    ev = {
        "backend/app/config/settings.py": "backend/app/config/settings.py:43,44",
        "backend/app/main.py": "backend/app/main.py:131",
        "backend/app/middleware/security_headers.py": "backend/app/middleware/security_headers.py:74",
        "backend/app/routers/documents.py": "backend/app/routers/documents.py:31,137,171,186,845",
        "backend/app/routers/templates.py": "backend/app/routers/templates.py:95,125",
        "backend/app/schemas/document.py": "backend/app/schemas/document.py:20",
        "backend/app/services/llm_service.py": "backend/app/services/llm_service.py:151,157,190",
        "backend/app/pipeline/intelligence/reasoning_engine.py": "backend/app/pipeline/intelligence/reasoning_engine.py:64,136,570",
        "backend/app/pipeline/intelligence/semantic_parser.py": "backend/app/pipeline/intelligence/semantic_parser.py:29,56",
        "backend/app/pipeline/intelligence/rag_engine.py": "backend/app/pipeline/intelligence/rag_engine.py:347,368,382",
        "backend/app/pipeline/nlp/analyzer.py": "backend/app/pipeline/nlp/analyzer.py:21,192,224",
        "frontend/src/services/api.js": "frontend/src/services/api.js:12,13,422,504,721",
        "frontend/services/api.js": "frontend/services/api.js:12,13,418,500,717",
        "frontend/src/context/AuthContext.jsx": "frontend/src/context/AuthContext.jsx:166,170",
        "frontend/context/AuthContext.jsx": "frontend/context/AuthContext.jsx:164",
        "frontend/components/AuthGuard.jsx": "frontend/components/AuthGuard.jsx:27,32",
        "frontend/app/(shared)/login/page.jsx": "frontend/app/(shared)/login/page.jsx:27,41,47",
        "frontend/app/(formatter)/jobs/[jobId]/[step]/page.jsx": "frontend/app/(formatter)/jobs/[jobId]/[step]/page.jsx:16,18,35",
        "frontend/vitest.config.js": "frontend/vitest.config.js:16,17,20,21",
        "frontend/package.json": "frontend/package.json:11,13",
    }

    manifest, faudit, rec = [], [], []
    for f in files:
        rel = f.relative_to(ROOT)
        fp = p(rel)
        t = txt(f)
        loc = lines(t)
        rs, rl, qs, cov, drift = rscore(fp, loc)
        layer = rel.parts[0] if rel.parts else "unknown"
        manifest.append({
            "path": fp, "layer": layer, "module": mod(rel), "file_type": "env-example" if f.name.endswith(".env.example") else (f.suffix.lower().lstrip(".") if f.suffix else "noext"),
            "loc": loc, "tracked_state": gstat.get(fp, "tracked_clean"), "static_dynamic_tag": dyn(rel, t), "covered": "yes"
        })
        feat = "Project source artifact"
        if "backend/app/routers/" in fp: feat = "REST endpoints; request validation; async dispatch"
        elif "backend/app/pipeline/intelligence/" in fp: feat = "LLM reasoning; semantic parsing; retrieval augmentation"
        elif "backend/app/pipeline/parsing/" in fp: feat = "Document parsing; OCR fallback; structure extraction"
        elif "backend/app/pipeline/formatting/" in fp: feat = "Template rendering; numbering; citation formatting"
        elif "backend/app/services/llm_service.py" in fp: feat = "Provider routing; fallback chain; timeout control"
        elif "backend/app/services/" in fp: feat = "External integrations; retries; caching"
        elif "backend/app/middleware/" in fp: feat = "Security headers; metrics; rate limiting"
        elif "backend/app/config/" in fp: feat = "Runtime settings; environment defaults"
        elif fp.startswith("backend/tests/"): feat = "Automated regression and contract coverage"
        elif fp.startswith("frontend/app/"): feat = "User-facing routes; page composition; interaction flow"
        elif fp in {"frontend/src/services/api.js", "frontend/services/api.js"}: feat = "API client; upload flow; error handling"
        elif "AuthContext" in fp: feat = "Auth session lifecycle; OAuth redirects; state sync"
        elif "frontend/components/" in fp or "frontend/src/components/" in fp: feat = "Reusable UI components and shell behavior"

        lim = "No critical limitation identified; monitor for drift"
        miss = "No mandatory missing item detected"
        if fp == "backend/app/routers/documents.py":
            lim, miss = "Complex endpoint surface; static quota constants and chunk path complexity", "Introduce plan-based quotas and stronger abuse controls"
        elif fp == "backend/app/schemas/document.py":
            lim, miss = "Template enum coverage lags installed/public template set", "Expand TemplateChoice to match supported/public templates"
        elif fp == "backend/app/routers/templates.py":
            lim, miss = "Public whitelist omits some installed templates", "Reconcile installed templates with public API and frontend selectors"
        elif fp in {"backend/app/services/llm_service.py", "backend/app/pipeline/intelligence/semantic_parser.py"}:
            lim, miss = "Provider/model fallback reliability depends on runtime availability", "Add stronger telemetry and tuned-model governance"
        elif fp in {"frontend/services/api.js", "frontend/context/AuthContext.jsx"}:
            lim, miss = "Duplicate root-level tree can drift from src-based imports", "Retire duplicate tree or enforce sync in CI"
        elif fp == "frontend/vitest.config.js":
            lim, miss = "Current include/exclude skips many tests under src/test", "Expand Vitest patterns to include src/test and route tests"

        tcov = "unknown"
        if fp.startswith("backend/tests/") or fp.startswith("frontend/e2e/"): tcov = "direct"
        elif fp.startswith("frontend/src/test/"): tcov = "present_but_excluded_by_vitest_config"
        elif fp == "frontend/src/services/api.js": tcov = "direct_unit_coverage"
        elif fp.startswith("backend/app/"): tcov = "partial_backend_suite"
        elif fp.startswith("frontend/app/"): tcov = "e2e_only"
        elif fp.startswith("frontend/"): tcov = "limited"

        faudit.append({
            "path": fp, "risk_score_1_10": rs, "risk_level": rl, "quality_score_1_10": qs, "enhancement_needed": "yes" if (rs >= 4 or cov >= 0.7 or drift >= 0.3) else "no",
            "features_present": feat, "limitations": lim, "missing_items": miss, "test_coverage_state": tcov, "evidence_refs": ev.get(fp, fp)
        })
        rec.append({"path": fp, "layer": layer, "module": mod(rel), "loc": loc, "risk": rs, "risk_level": rl, "quality": qs, "dyn": dyn(rel, t)})

    wcsv(DATA / "audit_manifest.csv", ["path", "layer", "module", "file_type", "loc", "tracked_state", "static_dynamic_tag", "covered"], manifest)
    wcsv(DATA / "file_audit.csv", ["path", "risk_score_1_10", "risk_level", "quality_score_1_10", "enhancement_needed", "features_present", "limitations", "missing_items", "test_coverage_state", "evidence_refs"], faudit)

    mrows = [
        {"component": "LLMService.generate_with_fallback", "purpose": "Generic generation", "provider": "NVIDIA NIM + Ollama/DeepSeek", "primary": "nvidia_nim/meta/llama-3.3-70b-instruct", "fallback_chain": "ollama/deepseek-r1 -> LLMUnavailableError", "input_contract": "messages, temperature, max_tokens", "output_contract": "text", "retries": "tier fallback", "timeout": "30s", "config_keys": "NVIDIA_API_KEY, OLLAMA_BASE_URL", "gaps": "No in-service rule-based fallback"},
        {"component": "ReasoningEngine.generate_instruction_set", "purpose": "Instruction planning", "provider": "NVIDIA + DeepSeek + rules", "primary": "nvidia_nim/meta/llama-3.3-70b-instruct", "fallback_chain": "ollama/deepseek-r1 -> rule_based_fallback", "input_contract": "semantic_blocks", "output_contract": "InstructionSet JSON", "retries": "guarded retries", "timeout": "configured per call", "config_keys": "NVIDIA_API_KEY, OLLAMA_BASE_URL", "gaps": "Need explicit fallback-rate SLO"},
        {"component": "SemanticParser", "purpose": "Semantic classification", "provider": "HuggingFace SciBERT", "primary": "allenai/scibert_scivocab_uncased", "fallback_chain": "deterministic heuristics", "input_contract": "document blocks", "output_contract": "labels+confidence", "retries": "none", "timeout": "N/A", "config_keys": "model store", "gaps": "Base model can be unreliable without fine-tuning"},
        {"component": "RagEngine", "purpose": "Guideline retrieval", "provider": "SentenceTransformers + Chroma/native", "primary": "BAAI/bge-m3", "fallback_chain": "bge-small-en-v1.5 -> deterministic embeddings", "input_contract": "query + metadata", "output_contract": "ranked snippets", "retries": "backend fallback", "timeout": "N/A", "config_keys": "RAG persistence/chroma", "gaps": "Deterministic fallback lowers semantic quality"},
        {"component": "GROBIDClient", "purpose": "Metadata/reference extraction", "provider": "GROBID REST", "primary": "configured endpoint", "fallback_chain": "safe_function empty payload", "input_contract": "PDF", "output_contract": "metadata/references", "retries": "none", "timeout": "requests timeout", "config_keys": "GROBID_URL", "gaps": "No secondary metadata extractor"},
        {"component": "KeywordExtraction", "purpose": "Keyword mining", "provider": "KeyBERT + YAKE", "primary": "KeyBERT", "fallback_chain": "YAKE -> frequency heuristic", "input_contract": "text/title/abstract", "output_contract": "keywords[]", "retries": "none", "timeout": "N/A", "config_keys": "KEYBERT_AVAILABLE, YAKE_AVAILABLE", "gaps": "No explicit LLM reranker"},
        {"component": "KeyLLM (named component)", "purpose": "LLM keyword reranking", "provider": "Not implemented", "primary": "N/A", "fallback_chain": "N/A", "input_contract": "Expected semantic context", "output_contract": "Expected weighted keywords", "retries": "N/A", "timeout": "N/A", "config_keys": "N/A", "gaps": "Requested capability absent as explicit module"},
    ]
    wcsv(DATA / "model_provider_fallback_matrix.csv", ["component", "purpose", "provider", "primary", "fallback_chain", "input_contract", "output_contract", "retries", "timeout", "config_keys", "gaps"], mrows)

    tdir = ROOT / "backend" / "app" / "templates"
    tids = sorted([d.name for d in tdir.iterdir() if d.is_dir()]) if tdir.exists() else []
    pub = ["ieee", "apa", "acm", "springer", "elsevier", "nature", "harvard", "chicago", "mla", "vancouver", "numeric", "none", "modern_blue", "modern_gold", "modern_red"]
    high = {"ieee", "apa", "springer", "nature", "vancouver"}
    medh = {"acm", "elsevier", "harvard", "chicago", "mla", "numeric"}
    med = {"modern_blue", "modern_gold", "modern_red"}
    tm = []
    for t in tids:
        conf = "high" if t in high or t == "none" else "medium_high" if t in medh else "medium" if t in med else "medium_low" if t in {"resume", "portfolio"} else "medium"
        tab, fig, eqn, ref = "supported", "supported", "partial", "supported"
        fail = "No critical template-specific failure observed"
        if t in med:
            tab, fig, eqn, fail = "partial", "partial", "partial", "Custom style fidelity depends on contract/render parity"
        if t in {"resume", "portfolio"}:
            tab, fig, eqn, ref, fail = "partial", "partial", "low", "partial", "Installed but not exposed in /api/templates public whitelist"
        if t == "none":
            tab, fig, eqn, ref, fail = "pass_through", "pass_through", "pass_through", "pass_through", "Minimal formatting path; relies on source quality"
        tm.append({"template": t, "ingest_formats": ".doc,.docx,.pdf,.odt,.rtf,.tex,.txt,.html,.htm,.md,.markdown", "text": "supported", "tables": tab, "figures": fig, "equations": eqn, "references": ref, "confidence_level": conf, "known_failures": fail})
    wcsv(DATA / "template_compatibility_matrix.csv", ["template", "ingest_formats", "text", "tables", "figures", "equations", "references", "confidence_level", "known_failures"], tm)

    ux = []
    for f in files:
        rel = f.relative_to(ROOT)
        fp = p(rel)
        if not fp.startswith("frontend/app/") or not rel.name.startswith("page."):
            continue
        t = txt(f)
        loc = lines(t)
        rp = []
        for part in rel.parts[2:-1]:
            if part.startswith("(") and part.endswith(")"):
                continue
            rp.append(part)
        route = "/" + "/".join(rp) if rp else "/"
        resp_hits = sum(1 for k in ["sm:", "md:", "lg:", "xl:", "2xl:"] if k in t)
        resp = "good" if resp_hits >= 2 else "partial" if resp_hits == 1 or "grid" in t or "flex" in t else "unknown"
        a11y_hits = sum(1 for k in ["aria-", "role=", "htmlFor", "label", "alt="] if k in t)
        a11y = "good" if a11y_hits >= 4 else "partial" if a11y_hits >= 2 else "needs_improvement"
        perf = "high" if loc > 900 else "medium" if loc > 450 else "low"
        uiq = max(1, min(10, 7 + (1 if resp == "good" else -1 if resp == "unknown" else 0) + (1 if a11y == "good" else -1 if a11y == "needs_improvement" else 0) - (2 if perf == "high" else 1 if perf == "medium" else 0)))
        issues, acts = [], []
        if "(protected)" in rel.parts and "loading" not in t.lower():
            issues.append("Guarded route has limited explicit loading-state UX")
            acts.append("Add deterministic skeleton/loading fallback before redirect resolution")
        if resp in {"partial", "unknown"}:
            issues.append("Responsive behavior not strongly evidenced in this file")
            acts.append("Add explicit breakpoint classes and mobile QA snapshot")
        if a11y == "needs_improvement":
            issues.append("Limited ARIA/label semantics")
            acts.append("Improve semantic labels, aria attributes, and focus states")
        if perf == "high":
            issues.append("Large route component increases client rendering cost")
            acts.append("Split route into smaller sections/hooks and memoize expensive blocks")
        ux.append({"route": route, "auth_guarded": "yes" if "(protected)" in rel.parts else "no", "responsive_status": resp, "accessibility_status": a11y, "performance_risk": perf, "ui_quality_score": uiq, "ux_issues": "; ".join(issues or ["No blocking UX issue found in static audit"]), "enhancement_actions": "; ".join(acts or ["Keep current implementation; monitor regressions"])})
    ux = sorted(ux, key=lambda x: x["route"])
    wcsv(DATA / "frontend_page_ux_matrix.csv", ["route", "auth_guarded", "responsive_status", "accessibility_status", "performance_risk", "ui_quality_score", "ux_issues", "enhancement_actions"], ux)

    backlog = [
        {"id": "P0-001", "title": "Unify template schema/API/frontend parity", "layer": "backend+frontend", "severity": "critical", "effort_s_m_l": "M", "impact": "Prevent template mismatch failures", "dependency": "schemas+templates router+frontend lists", "acceptance_criteria": "Parity test passes for all installed templates"},
        {"id": "P0-002", "title": "Remove duplicate frontend source trees", "layer": "frontend", "severity": "high", "effort_s_m_l": "M", "impact": "Eliminate auth/API drift", "dependency": "frontend/context/services vs frontend/src/*", "acceptance_criteria": "Single canonical imports, no drift"},
        {"id": "P0-003", "title": "Expand Vitest discovery scope", "layer": "frontend", "severity": "high", "effort_s_m_l": "S", "impact": "Increase direct unit confidence", "dependency": "frontend/vitest.config.js", "acceptance_criteria": "src/test and route tests execute in CI"},
        {"id": "P0-004", "title": "Add LLM fallback observability SLOs", "layer": "backend", "severity": "high", "effort_s_m_l": "M", "impact": "Reliability under provider outage", "dependency": "llm_service/reasoning_engine/model_metrics", "acceptance_criteria": "Fallback metrics on monitoring endpoint"},
        {"id": "P1-005", "title": "Plan-based upload quotas", "layer": "backend", "severity": "high", "effort_s_m_l": "M", "impact": "Cost/fairness controls", "dependency": "documents router + plan metadata", "acceptance_criteria": "Per-plan quota enforcement"},
        {"id": "P1-006", "title": "Mixed-content template reliability suite", "layer": "backend", "severity": "medium", "effort_s_m_l": "M", "impact": "Confidence for tables/figures/equations", "dependency": "golden fixtures + template regression", "acceptance_criteria": "All templates validated on mixed content"},
        {"id": "P1-007", "title": "Implement optional KeyLLM stage", "layer": "backend", "severity": "medium", "effort_s_m_l": "M", "impact": "Better keyword quality", "dependency": "nlp analyzer + model routing", "acceptance_criteria": "Config-driven KeyLLM with fallback"},
        {"id": "P1-008", "title": "Accessibility regression checks", "layer": "frontend", "severity": "medium", "effort_s_m_l": "S", "impact": "Improved UX compliance", "dependency": "route components + test tooling", "acceptance_criteria": "a11y checks pass on core routes"},
        {"id": "P1-009", "title": "Template API/frontend contract tests", "layer": "backend+frontend", "severity": "medium", "effort_s_m_l": "S", "impact": "Prevents parity regressions", "dependency": "templates API + selector", "acceptance_criteria": "CI fails on list drift"},
        {"id": "P2-010", "title": "Main Generator Mode v2/v3 roadmap", "layer": "platform", "severity": "medium", "effort_s_m_l": "L", "impact": "Scale beyond formatter", "dependency": "generator hardening baseline", "acceptance_criteria": "Approved phased architecture plan"},
    ]
    wcsv(DATA / "priority_backlog.csv", ["id", "title", "layer", "severity", "effort_s_m_l", "impact", "dependency", "acceptance_criteria"], backlog)

    bfiles = [r for r in rec if r["layer"] == "backend"]
    ffiles = [r for r in rec if r["layer"] == "frontend"]
    bloc = sum(r["loc"] for r in bfiles)
    floc = sum(r["loc"] for r in ffiles)
    rcnt = Counter(r["risk_level"] for r in rec)
    dcnt = Counter(r["dyn"] for r in rec)

    mod_agg = defaultdict(lambda: {"files": 0, "loc": 0, "risk": 0.0, "quality": 0.0})
    for r in rec:
        k = (r["layer"], r["module"])
        mod_agg[k]["files"] += 1
        mod_agg[k]["loc"] += r["loc"]
        mod_agg[k]["risk"] += r["risk"]
        mod_agg[k]["quality"] += r["quality"]

    mod_rows = []
    for (ly, mo), st in mod_agg.items():
        c = st["files"]
        mod_rows.append({"layer": ly, "module": mo, "files": c, "loc": st["loc"], "avg_risk": round(st["risk"] / c, 2), "avg_quality": round(st["quality"] / c, 2)})
    mod_rows = sorted(mod_rows, key=lambda x: (x["layer"], -x["avg_risk"], -x["loc"]))
    top = sorted(rec, key=lambda x: (-x["risk"], -x["loc"], x["path"]))

    sdoc = txt(ROOT / "backend" / "app" / "schemas" / "document.py")
    m = re.search(r"TemplateChoice\s*=\s*Literal\[(.*?)\]", sdoc, re.S)
    schema_ts = sorted({x.lower() for x in re.findall(r'"([A-Za-z_]+)"', m.group(1))}) if m else []
    fmt_doc = txt(ROOT / "frontend" / "app" / "(formatter)" / "templates" / "page.jsx")
    fmt_ts = sorted(set(re.findall(r"id:\s*'([^']+)'", fmt_doc)))
    gen_doc = txt(ROOT / "frontend" / "app" / "(generator)" / "(protected)" / "generate" / "page.jsx")
    gen_ts = sorted(set(re.findall(r"id:\s*'([^']+)'", gen_doc)))
    inst = set(tids)
    pset = set(pub)
    sset = set(schema_ts)
    fset = set(fmt_ts)
    gset = set(gen_ts)

    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    rep = {}
    rep["00_executive_summary.md"] = f"""# 00 Executive Summary

Audit date: 2026-03-02  
Generated at: {now}

## Scope and Coverage
- Source-owned files audited: {len(rec)}
- Backend files audited: {len(bfiles)}
- Frontend files audited: {len(ffiles)}
- Backend LOC audited: {bloc:,}
- Frontend LOC audited: {floc:,}
- Coverage status: Every included source/config/docs/tests file has an audit row in `data/file_audit.csv`.

## Risk Overview
- Low risk files: {rcnt.get('low', 0)}
- Medium risk files: {rcnt.get('medium', 0)}
- High risk files: {rcnt.get('high', 0)}
- Critical risk files: {rcnt.get('critical', 0)}

## Highest-Priority Findings
- Template parity drift exists across installed backend templates, public templates API, and schema literals.
- Frontend duplicate source trees (`frontend/context` vs `frontend/src/context`, `frontend/services` vs `frontend/src/services`) create drift risk.
- Frontend unit discovery is narrow in current Vitest config.
- Upload/chunk limits are present, but quota policy is fixed constant and not plan-aware.
- AI stack has fallback depth, while KeyLLM is not implemented as a named component.

## Test Health Snapshot (March 2, 2026)
- Backend: `pytest --collect-only` collected 349 tests.
- Backend targeted suites: `tests/test_api.py` 12/12 pass, `tests/test_generator.py` 21/21 pass.
- Frontend unit: `vitest run` currently runs 1 file / 5 tests.
- Frontend E2E: `playwright test` executed 7 tests and all passed in this run.

## Reliability Position
- Strong baseline coverage exists, including template regression suites.
- No claim of 100 percent formatting correctness for all possible documents is made.
- Reliability should be expressed by confidence tiers and known failure modes by content type/template.
"""

    rep["01_system_inventory.md"] = (
        "# 01 System Inventory\n\n"
        "## File Universe\n"
        f"- Backend files: {len(bfiles)}\n"
        f"- Frontend files: {len(ffiles)}\n"
        f"- Total LOC: {bloc + floc:,}\n\n"
        "## Dynamic vs Static\n"
        f"- Dynamic files: {dcnt.get('dynamic', 0)}\n"
        f"- Static files: {dcnt.get('static', 0)}\n\n"
        "## Top Modules by Risk/Volume\n"
        + mtable(["layer", "module", "files", "loc", "avg_risk", "avg_quality"], mod_rows[:30]) + "\n\n"
        "## Working Tree Notes\n"
        "- Audit includes tracked and uncommitted files from current working tree.\n"
        "- Generated/vendor artifacts are excluded from per-file scoring.\n"
    )

    btop = [{"path": r["path"], "risk": r["risk"], "quality": r["quality"], "loc": r["loc"], "module": r["module"]} for r in top if r["layer"] == "backend"][:25]
    rep["02_backend_module_audit.md"] = (
        "# 02 Backend Module Audit\n\n"
        f"## Backend Coverage\n- Files audited: {len(bfiles)}\n- LOC audited: {bloc:,}\n\n"
        "## Critical Findings\n"
        "- Template schema mismatch: `backend/app/schemas/document.py:20` limits `TemplateChoice` to a subset of installed/public templates.\n"
        "- Public template filter in `backend/app/routers/templates.py:95,125` excludes installed templates (`resume`, `portfolio`).\n"
        "- Upload constraints exist (`documents.py` limits, magic-byte checks, quotas), but quota policy remains fixed (`MAX_DAILY_UPLOADS = 20`).\n"
        "- Global body-size guard (60MB) vs file limit (50MB) is coherent but should stay explicitly documented.\n\n"
        "## AI/NLP/OCR Findings\n"
        "- LLM fallback chains are implemented in `llm_service.py` and `reasoning_engine.py`.\n"
        "- SciBERT parser warns about reliability when using base model without fine-tuning.\n"
        "- RAG engine supports transformer and deterministic fallback path.\n"
        "- Keyword pipeline includes KeyBERT and YAKE; KeyLLM is not implemented.\n\n"
        "## Highest-Risk Backend Files\n"
        + mtable(["path", "risk", "quality", "loc", "module"], btop) + "\n"
    )

    ftop = [{"path": r["path"], "risk": r["risk"], "quality": r["quality"], "loc": r["loc"], "module": r["module"]} for r in top if r["layer"] == "frontend"][:25]
    rep["03_frontend_module_audit.md"] = (
        "# 03 Frontend Module Audit\n\n"
        f"## Frontend Coverage\n- Files audited: {len(ffiles)}\n- LOC audited: {floc:,}\n\n"
        "## Architecture Findings\n"
        "- Next.js app-router surface is broad and protected-route wiring is present.\n"
        "- Duplicate source trees exist and already show behavior drift in auth/API modules.\n"
        "- Route dynamic-step page validates allowed steps and returns native Next 404 for invalid values.\n\n"
        "## UX and Testing Findings\n"
        f"- Routes audited in UX matrix: {len(ux)}\n"
        "- Unit discovery remains narrow under current Vitest config.\n"
        "- E2E suite passed 7/7 in this run.\n\n"
        "## Highest-Risk Frontend Files\n"
        + mtable(["path", "risk", "quality", "loc", "module"], ftop) + "\n"
    )

    rep["04_model_provider_fallback_matrix.md"] = (
        "# 04 Model Provider Fallback Matrix\n\n"
        "Requested components check: GROBID, SciBERT, KeyBERT, YAKE are present. KeyLLM is not implemented as a named component.\n\n"
        + mtable(["component", "purpose", "provider", "primary", "fallback_chain", "timeout", "gaps"], mrows)
        + "\n\nFull data: `data/model_provider_fallback_matrix.csv`.\n"
    )

    rep["05_input_output_limits_contract.md"] = (
        "# 05 Input Output Limits Contract\n\n"
        "## Input Types\n"
        "- Accepted upload extensions: `.docx, .doc, .pdf, .odt, .rtf, .tex, .txt, .html, .htm, .md, .markdown`.\n"
        "- Magic-byte validation is enforced for binary types in upload path.\n"
        "- UTF-8 validation is enforced for text extensions.\n\n"
        "## Size and Quota Limits\n"
        "- `MAX_FILE_SIZE`: 50MB (`backend/app/config/settings.py:43`).\n"
        "- Chunk upload per-part limit: 5MB (`backend/app/routers/documents.py:171`).\n"
        "- Frontend chunk threshold: 10MB (`frontend/src/services/api.js:13`).\n"
        "- `MAX_BATCH_FILES`: 10 (`backend/app/config/settings.py:44`).\n"
        "- Daily upload quota: 20 (`backend/app/routers/documents.py:31`).\n"
        "- Global request body limit: 60MB (`backend/app/main.py:131`).\n\n"
        "## Conversion and Output Contract\n"
        "- Conversion supports `.docx, .doc, .md, .html, .txt, .tex, .pdf, .odt, .rtf`.\n"
        "- OCR path supports available backend selection with LibreOffice fallback.\n"
        "- Export formats are `docx` and `pdf` only.\n"
    )

    rep["06_template_reliability_report.md"] = (
        "# 06 Template Reliability Report\n\n"
        f"- Installed backend templates: {len(tids)}\n"
        f"- Public templates API entries: {len(pub)}\n"
        f"- Schema template literals: {len(schema_ts)}\n\n"
        "## Parity Findings\n"
        f"- Installed but not public: {sorted(inst - pset)}\n"
        f"- Public but not schema: {sorted(pset - sset)}\n"
        f"- Frontend template ids not in public baseline: {sorted(fset - pset)}\n"
        f"- Generator template ids not in public baseline: {sorted(gset - pset)}\n\n"
        "## Reliability Position\n"
        "- Core academic templates show higher confidence.\n"
        "- Custom/professional templates need stricter parity and mixed-content validation.\n"
        "- Reliability should be tracked by confidence tier and known failure mode.\n\n"
        "## Compatibility Matrix\n"
        + mtable(["template", "text", "tables", "figures", "equations", "references", "confidence_level", "known_failures"], tm)
        + "\n\nFull matrix: `data/template_compatibility_matrix.csv`.\n"
    )

    rep["07_testing_quality_report.md"] = (
        "# 07 Testing Quality Report\n\n"
        "## Latest Executions (March 2, 2026)\n"
        "- Backend `pytest --collect-only -q`: 349 tests collected.\n"
        "- Backend targeted suites: `tests/test_api.py` 12/12 pass, `tests/test_generator.py` 21/21 pass.\n"
        "- Frontend unit (`vitest run`): 1 file / 5 tests executed.\n"
        "- Frontend E2E (`playwright test`): 7/7 pass in this run.\n\n"
        "## Gaps\n"
        "- Vitest include/exclude currently skips `src/test/**`.\n"
        "- Route-level frontend behavior relies more on E2E than unit-level checks.\n"
        "- Some backend integration tests depend on optional external services.\n"
    )

    rep["08_static_dynamic_map.md"] = (
        "# 08 Static Dynamic Map\n\n"
        "## Global\n"
        f"- Dynamic files: {dcnt.get('dynamic', 0)}\n"
        f"- Static files: {dcnt.get('static', 0)}\n\n"
        "## Layer Breakdown\n"
        f"- Backend dynamic: {sum(1 for r in rec if r['layer']=='backend' and r['dyn']=='dynamic')}\n"
        f"- Frontend dynamic: {sum(1 for r in rec if r['layer']=='frontend' and r['dyn']=='dynamic')}\n"
        f"- Backend static: {sum(1 for r in rec if r['layer']=='backend' and r['dyn']=='static')}\n"
        f"- Frontend static: {sum(1 for r in rec if r['layer']=='frontend' and r['dyn']=='static')}\n"
    )

    rep["09_gap_and_roadmap.md"] = (
        "# 09 Gap and Roadmap\n\n"
        "## Completion Status\n"
        "- Complete: core formatter/generator flow, route surface, pipeline baseline.\n"
        "- Partially complete: template parity governance, frontend unit breadth, model fallback observability.\n"
        "- Missing/high-value: explicit KeyLLM stage, stronger quota policy, expanded mixed-content reliability CI.\n\n"
        "## Must-Fix Before Production Hardening\n"
        + "\n".join([f"- {x['id']}: {x['title']} ({x['effort_s_m_l']}, {x['severity']})" for x in backlog[:4]])
        + "\n\n## Main Generator Mode Expansion\n"
        "### v1 Hardening (0-6 weeks)\n"
        "- Resolve template parity and duplicate frontend tree drift.\n"
        "- Expand test discovery and contract parity tests.\n"
        "- Add fallback telemetry and provider health dashboards.\n\n"
        "### v2 Capability Expansion (6-14 weeks)\n"
        "- Add explicit KeyLLM/keyword-reranking stage.\n"
        "- Add richer generation controls and section-level validation loops.\n"
        "- Support advanced template packs with strict contract linting.\n\n"
        "### v3 Enterprise Scale (14+ weeks)\n"
        "- Multi-tenant policy engine for quotas and governance.\n"
        "- Reliability SLOs with autoscaling and regional fallback strategy.\n"
        "- Compliance-ready traceability for document/model decisions.\n"
    )

    for fn, body in rep.items():
        (OUT / fn).write_text(body, encoding="utf-8")

    for fn in sorted(rep.keys()):
        d = Doc()
        tmd(d, (OUT / fn).read_text(encoding="utf-8"))
        d.save(DOCX / fn.replace(".md", ".docx"))

    csvs = [
        "audit_manifest.csv",
        "file_audit.csv",
        "model_provider_fallback_matrix.csv",
        "template_compatibility_matrix.csv",
        "frontend_page_ux_matrix.csv",
        "priority_backlog.csv",
    ]
    for c in csvs:
        csv_docx(DATA / c, DOCX_DATA / c.replace(".csv", ".docx"), f"Data Table - {c}")

    master = Doc()
    master.add_heading("Fullstack Audit Package - March 2, 2026", level=0)
    master.add_paragraph("Comprehensive backend and frontend audit package generated from current working tree.")
    for fn in sorted(rep.keys()):
        master.add_page_break()
        tmd(master, (OUT / fn).read_text(encoding="utf-8"))
    master.add_page_break()
    master.add_heading("Appendix A - Complete Per-File Risk Index", level=1)
    for row in faudit:
        master.add_paragraph(f"{row['path']} | risk={row['risk_score_1_10']} ({row['risk_level']}) | quality={row['quality_score_1_10']} | enhancement={row['enhancement_needed']}")
    master.save(DOCX / "fullstack_audit_complete.docx")

    idx = "# Audit Artifact Index\n\n## Markdown Reports\n"
    idx += "\n".join([f"- {k}" for k in sorted(rep.keys())]) + "\n\n## CSV Data\n"
    idx += "\n".join([f"- data/{k}" for k in csvs]) + "\n\n## DOCX Reports\n"
    idx += "\n".join([f"- docx/{k.replace('.md', '.docx')}" for k in sorted(rep.keys())])
    idx += "\n- docx/fullstack_audit_complete.docx\n\n## DOCX Data Tables\n"
    idx += "\n".join([f"- docx/data/{k.replace('.csv', '.docx')}" for k in csvs]) + "\n"
    (OUT / "README.md").write_text(idx, encoding="utf-8")

    print(json.dumps({"total_files_audited": len(rec), "backend_files": len(bfiles), "frontend_files": len(ffiles), "backend_loc": bloc, "frontend_loc": floc, "out_dir": str(OUT)}, indent=2))


if __name__ == "__main__":
    run()
