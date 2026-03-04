from __future__ import annotations

import csv
from collections import Counter
from pathlib import Path

from docx import Document


BASE = Path("docs/audits/2026-03-02-fullstack")
OUT = BASE / "docx" / "fullstack_audit_master_single_source_of_truth.docx"


def add_markdown(doc: Document, content: str) -> None:
    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1] == ".":
            doc.add_paragraph(stripped[3:].strip(), style="List Number")
        else:
            doc.add_paragraph(line)


def add_csv_table(doc: Document, csv_path: Path) -> None:
    with csv_path.open("r", encoding="utf-8", newline="") as fh:
        rows = list(csv.reader(fh))

    if not rows:
        doc.add_paragraph("No rows found.")
        return

    headers = rows[0]
    body = rows[1:]

    doc.add_paragraph(f"Rows: {len(body)}")
    doc.add_paragraph("Columns:")
    for h in headers:
        doc.add_paragraph(h, style="List Bullet")

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = str(h)

    for row in body:
        cells = table.add_row().cells
        for i in range(len(headers)):
            cells[i].text = str(row[i] if i < len(row) else "")


def collect_files(base: Path) -> list[Path]:
    files = [p for p in base.rglob("*") if p.is_file()]
    return sorted(files, key=lambda p: p.as_posix())


def size_kb(p: Path) -> str:
    return f"{p.stat().st_size / 1024:.1f}"


def main() -> None:
    if not BASE.exists():
        raise FileNotFoundError(f"Audit folder not found: {BASE}")

    all_files = collect_files(BASE)
    md_files = sorted([p for p in all_files if p.suffix.lower() == ".md"], key=lambda p: p.name)
    csv_files = sorted([p for p in all_files if p.suffix.lower() == ".csv"], key=lambda p: p.name)
    docx_files = sorted([p for p in all_files if p.suffix.lower() == ".docx"], key=lambda p: p.name)

    ext_counter = Counter(p.suffix.lower() for p in all_files)

    doc = Document()
    doc.add_heading("Fullstack Audit - All In One Complete Document", level=0)
    doc.add_paragraph("Scope date baseline: March 2, 2026")
    doc.add_paragraph(
        "This document consolidates every generated audit artifact into one single file: "
        "inventory, markdown reports, full CSV tables, and final recommendations."
    )

    doc.add_heading("Section 0 - Requirement Coverage Matrix", level=1)
    doc.add_paragraph("This table maps your requested audit items to where they are covered in this same single document.")
    req = doc.add_table(rows=1, cols=4)
    req.style = "Table Grid"
    req.rows[0].cells[0].text = "Requested Item"
    req.rows[0].cells[1].text = "Status"
    req.rows[0].cells[2].text = "Where Covered In This Document"
    req.rows[0].cells[3].text = "Primary Source Artifact"
    req_rows = [
        ("Backend + frontend complete audit", "Complete", "Section 2", "00-09 markdown reports"),
        ("Each file risk score 1-10", "Complete", "Section 3", "data/file_audit.csv"),
        ("Features, limitations, missing items", "Complete", "Section 3", "data/file_audit.csv"),
        ("Model | Purpose | Provider | Fallback", "Complete", "Section 2 + Section 3", "04_*.md + model_provider_fallback_matrix.csv"),
        ("KeyLLM, GROBID, SciBERT, KeyBERT, YAKE status", "Complete", "Section 2 + Section 3", "04_*.md + model_provider_fallback_matrix.csv"),
        ("Upload/input/output limits and quotas", "Complete", "Section 2", "05_input_output_limits_contract.md"),
        ("Template reliability across content types", "Complete", "Section 2 + Section 3", "06_*.md + template_compatibility_matrix.csv"),
        ("Can formatting be 100% for all docs", "Covered with realistic confidence model", "Section 2", "00_*.md + 06_*.md"),
        ("Frontend pages UX/UI/responsive checks", "Complete", "Section 3", "data/frontend_page_ux_matrix.csv"),
        ("Static vs dynamic mapping", "Complete", "Section 2", "08_static_dynamic_map.md"),
        ("Testing status and quality", "Complete", "Section 2", "07_testing_quality_report.md"),
        ("Industry-gap roadmap and priorities", "Complete", "Section 2 + Section 3 + Section 4", "09_*.md + priority_backlog.csv"),
        ("Valuable suggestions", "Complete", "Section 4", "Final Valuable Suggestions"),
    ]
    for item, status, where, source in req_rows:
        r = req.add_row().cells
        r[0].text = item
        r[1].text = status
        r[2].text = where
        r[3].text = source

    doc.add_page_break()
    doc.add_heading("Section 1 - All Documents Created", level=1)
    doc.add_paragraph(f"Total files created in audit package: {len(all_files)}")
    doc.add_paragraph(f"Markdown files: {len(md_files)}")
    doc.add_paragraph(f"CSV files: {len(csv_files)}")
    doc.add_paragraph(f"DOCX files: {len(docx_files)}")
    doc.add_paragraph("File types distribution:")
    for ext, count in sorted(ext_counter.items(), key=lambda x: x[0]):
        label = ext if ext else "(no extension)"
        doc.add_paragraph(f"{label}: {count}", style="List Bullet")

    inv = doc.add_table(rows=1, cols=4)
    inv.style = "Table Grid"
    inv.rows[0].cells[0].text = "#"
    inv.rows[0].cells[1].text = "Relative Path"
    inv.rows[0].cells[2].text = "Type"
    inv.rows[0].cells[3].text = "Size KB"
    for idx, file_path in enumerate(all_files, 1):
        rel = file_path.relative_to(BASE).as_posix()
        row = inv.add_row().cells
        row[0].text = str(idx)
        row[1].text = rel
        row[2].text = file_path.suffix.lower().lstrip(".") or "noext"
        row[3].text = size_kb(file_path)

    doc.add_page_break()
    doc.add_heading("Section 2 - Full Content of All Markdown Reports", level=1)
    for md in md_files:
        rel = md.relative_to(BASE).as_posix()
        doc.add_page_break()
        doc.add_heading(rel, level=2)
        add_markdown(doc, md.read_text(encoding="utf-8"))

    doc.add_page_break()
    doc.add_heading("Section 3 - Full Content of All CSV Data Tables", level=1)
    for csv_file in csv_files:
        rel = csv_file.relative_to(BASE).as_posix()
        doc.add_page_break()
        doc.add_heading(rel, level=2)
        add_csv_table(doc, csv_file)

    doc.add_page_break()
    doc.add_heading("Section 4 - Final Valuable Suggestions", level=1)
    suggestions = [
        "Unify template definitions across backend schema, templates API, and frontend selectors with CI parity tests.",
        "Remove duplicate frontend source trees and enforce one canonical import path.",
        "Expand Vitest include scope to execute src/test and route-level component tests.",
        "Keep Playwright E2E stable with retries, trace collection, and deterministic test fixtures.",
        "Add provider fallback SLO dashboards for NVIDIA, DeepSeek/Ollama, OCR, and GROBID dependency health.",
        "Make upload quotas plan-aware and configurable via policy table instead of hardcoded constants.",
        "Add mixed-content reliability benchmarks per template for text, tables, figures, equations, and references.",
        "Introduce explicit KeyLLM stage as optional keyword reranker with graceful fallback to KeyBERT/YAKE.",
        "Add contract tests to ensure backend template exposure always matches installed templates and frontend UI.",
        "Strengthen accessibility audits for each protected/public route with automated checks in CI.",
        "Add architecture drift checks so README/docs stay synchronized with actual code and config defaults.",
        "For main generator mode expansion, execute v1 hardening before v2 feature growth to avoid scaling existing drift.",
    ]
    for s in suggestions:
        doc.add_paragraph(s, style="List Bullet")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(str(OUT))


if __name__ == "__main__":
    main()
